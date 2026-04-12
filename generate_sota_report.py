"""
Generate the SOTA analysis and positioning report as a .docx file.
Run: python3 generate_sota_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def set_cell_shading(cell, color_hex):
    """Set background shading for a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                            font_size=None, space_after=None, space_before=None,
                            alignment=None, font_name=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if font_name:
        run.font.name = font_name
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if alignment:
        p.alignment = alignment
    return p


def add_citation_paragraph(doc, text, citation_parts=None):
    """Add a paragraph where certain parts are superscript citations."""
    p = doc.add_paragraph()
    if citation_parts is None:
        p.add_run(text)
        return p

    remaining = text
    for marker, ref_num in citation_parts:
        if marker in remaining:
            before, after = remaining.split(marker, 1)
            if before:
                p.add_run(before)
            sup_run = p.add_run(ref_num)
            sup_run.font.superscript = True
            remaining = after
        else:
            continue
    if remaining:
        p.add_run(remaining)
    return p


def create_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def build_document():
    doc = Document()

    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Configure heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            h_style.font.size = Pt(14)
            h_style.paragraph_format.space_before = Pt(18)
        elif level == 2:
            h_style.font.size = Pt(12)
            h_style.paragraph_format.space_before = Pt(14)
        else:
            h_style.font.size = Pt(11)
            h_style.paragraph_format.space_before = Pt(10)

    # ================================================================
    # TITLE
    # ================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        'State-of-the-Art Analysis and Positioning of Our Work\n'
        'in the Context of Recent Publications in\n'
        'Dentomaxillofacial Radiology (DMFR)'
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'Internal Working Document — Pre-Experiment Review\n'
        'April 2026'
    )
    run.italic = True
    run.font.size = Pt(11)
    subtitle.paragraph_format.space_after = Pt(24)

    # ================================================================
    # 1. INTRODUCTION
    # ================================================================
    doc.add_heading('1. Introduction and Purpose', level=1)

    doc.add_paragraph(
        'This internal report provides a systematic analysis of the state of the art '
        'in artificial intelligence (AI) and large language model (LLM) research as '
        'published in Dentomaxillofacial Radiology (DMFR), the official journal of the '
        'International Association of Dentomaxillofacial Radiology (IADMFR). DMFR is a '
        'leading venue for dental imaging research, with an impact factor of 4.1 (2024), '
        'ranked 12th of 163 journals in Dentistry, Oral Surgery & Medicine and 32nd of '
        '213 in Radiology. As noted in a recent editorial by the Editor-in-Chief, '
        'approximately 60% of the most-cited DMFR articles over the previous half-decade '
        'were AI-related, reflecting the journal\'s strong editorial appetite for this '
        'topic [1].'
    )

    doc.add_paragraph(
        'The purpose of this report is threefold: (1) to review relevant recent '
        'publications in DMFR that intersect with our study\'s scope; (2) to identify '
        'the strengths, weaknesses, and limitations of our current experimental design '
        'in light of published standards; and (3) to propose concrete methodological '
        'improvements and a preliminary experiment plan to validate our pipeline before '
        'commencing the full-scale experiment.'
    )

    doc.add_paragraph(
        'Our study, titled "Comparative Analysis of Multimodal Large Language Models '
        'and Dental Students in Spatial Proficiency for Radiographic Anatomic Landmark '
        'Identification: A Novel Grid-Based Assessment Model," proposes a standardised '
        'square grid overlay system for evaluating spatial proficiency on dental '
        'radiographs. We compare the performance of frontier multimodal LLMs (GPT-5.4, '
        'Gemini 3.1 Pro) against 40 fourth-year dental students across 200 images and '
        '900 landmark identification queries spanning three radiographic modalities '
        '(panoramic, periapical, cephalometric).'
    )

    # ================================================================
    # 2. STATE OF THE ART
    # ================================================================
    doc.add_heading('2. State of the Art in DMFR', level=1)

    # 2.1
    doc.add_heading('2.1 LLM and Chatbot Evaluations in Dental Radiology', level=2)

    doc.add_paragraph(
        'The evaluation of LLMs in dentomaxillofacial radiology is a rapidly emerging '
        'research area. A comprehensive systematic review by Liu et al. (2025) [2] '
        'searched five databases and identified 19 studies evaluating LLMs in DMFR '
        'applications. The authors categorised these into three application domains: '
        'qualification examinations (8 studies, 33.3%\u201386.1% correctness), '
        'diagnosis and treatment planning (7 studies, 37%\u201392.5% accuracy), '
        'and report generation (4 studies, 70.4%\u201381.3% accuracy). A key finding '
        'was that domain-specific customisation substantially improved performance, '
        'with a customised GPT-4V variant achieving 91% accuracy for supernumerary '
        'tooth detection. However, the review concluded that "current accuracy, '
        'completeness, and consistency remain variable" and called for further '
        'standardisation of evaluation methods.'
    )

    doc.add_paragraph(
        'Jeong et al. (2024) [3] directly compared LLM chatbots against 120 dental '
        'students on 52 oral and maxillofacial radiology examination questions. Dental '
        'students achieved 81.2% accuracy, outperforming ChatGPT Plus (GPT-4) at 65.4%, '
        'Bing Chat at 63.5%, and both ChatGPT (GPT-3.5) and Google Bard at 50.0%. '
        'Critically, all chatbots scored below 35% on image interpretation questions, '
        'highlighting a fundamental weakness in visual analysis capability.'
    )

    doc.add_paragraph(
        'The most directly relevant study to our work is by Jeong et al. (2026) [4], '
        'which represents the first comprehensive evaluation of multimodal LLM chatbots '
        'in oral radiology using both text- and image-based questions. Testing ChatGPT-4o '
        'and Gemini 2.0 Flash on 90 questions from a Korean dental university, the '
        'authors assessed three dimensions: accuracy, consistency (10 repeated outputs '
        'with Fleiss\' kappa), and hallucination rates (5-point Global Quality Score '
        'rated by two oral radiologists). While text-based performance was excellent '
        '(>80%), image-based accuracy was notably limited (<30%), with high response '
        'variability and frequent hallucinations on visual questions. The authors '
        'concluded that general-purpose multimodal LLMs "are not yet suitable for '
        'reliable use in oral radiology."'
    )

    doc.add_paragraph(
        'Najmuddin (2025) [5] extended the model comparison landscape by evaluating '
        'six AI tools (ChatGPT, ChatGPT-4o, Microsoft Copilot, DeepSeek, Gemini, '
        'Meta AI) on 80 multiple-choice questions across five oral radiology domains. '
        'Microsoft Copilot achieved the highest overall accuracy, followed by ChatGPT-4o. '
        'Notably, both achieved 100% accuracy on radiographic safety questions, '
        'suggesting domain-dependent performance variation.'
    )

    # 2.2
    doc.add_heading('2.2 Automated Landmark Detection on Dental Radiographs', level=2)

    doc.add_paragraph(
        'Automated cephalometric landmark detection is one of the most extensively '
        'studied applications of AI in DMFR. Jiang et al. (2023) [6] developed CephNet, '
        'a two-stage cascading CNN trained on 9,870 cephalograms from 20 institutions, '
        'achieving a mean landmark prediction error of 0.94 \u00b1 0.74 mm and an SDR '
        'within 2 mm of 91.73% across 30 landmarks. Weingart et al. (2023) [7] applied '
        'Deep Neural Patchworks (hierarchical U-Net architecture) to 3D CT data with 60 '
        'landmarks, reporting a mean localisation error of 1.94 mm (SD 1.45 mm) and SDR '
        'within 2 mm of 66.4%. These results establish the current performance ceiling '
        'for purpose-built AI systems on cephalometric data.'
    )

    doc.add_paragraph(
        'Indermun et al. (2023) [8] compared human-assisted landmark detection '
        '(Dolphin Imaging) against an AI system (BoneFinder) on 409 cephalograms with '
        '19 landmarks. The study found no statistically significant differences between '
        'approaches for most landmarks (p > 0.05), though absolute Euclidean distances '
        '(6.19\u201311.29 mm) were notably higher than typical literature values. '
        'Menezes et al. (2023) [9] investigated how image brightness and contrast '
        'settings affect AI landmark detection reliability, finding that extreme contrast '
        'adjustments compromised reproducibility, with ICC values dropping below '
        'acceptable thresholds. Both studies reported inter- and intra-rater reliability '
        'metrics (ICC), reinforcing this as a standard requirement in DMFR landmark '
        'publications.'
    )

    doc.add_paragraph(
        'Kazimierczak et al. (2023) [10] evaluated the CephX commercial AI platform '
        'for skeletal facial asymmetry analysis, finding a 16.8% tracing error rate and '
        'near-zero ICC agreement with manual measurements. This study serves as a '
        'cautionary example and underscores the importance of rigorous validation before '
        'clinical deployment of AI tools in landmark-dependent analyses.'
    )

    # 2.3
    doc.add_heading('2.3 Deep Learning for Detection and Segmentation Tasks', level=2)

    doc.add_paragraph(
        'Beyond landmark detection, several recent DMFR publications are relevant to '
        'our work because they evaluate AI on the same anatomical structures we assess. '
        'Semper-Hogg et al. (2025) [11] achieved human-level performance in mandibular '
        'canal segmentation on CBCT using a Deep Neural Patchwork, reporting a Dice '
        'coefficient of 0.77 \u00b1 0.09 and 95th-percentile Hausdorff distance of '
        '1.66 \u00b1 0.86 mm. This is directly comparable to our mandibular canal '
        'area-based landmark, albeit in 3D versus our 2D grid assessment.'
    )

    doc.add_paragraph(
        'For periapical radiograph interpretation, Hamdan et al. (2024) [12] evaluated '
        'a deep learning tool for apical radiolucency detection in a cross-over reader '
        'study with washout period. Although AI did not significantly improve diagnostic '
        'accuracy (AFROC-AUC), it substantially reduced diagnostic time, suggesting that '
        'speed advantages may be as clinically relevant as accuracy gains. Celik et al. '
        '(2023) [13] benchmarked 10 detection frameworks on 454 periapical lesion '
        'instances across 357 panoramic radiographs, with RetinaNet achieving the '
        'highest mAP (0.953) and ATSS the best F1 (0.895).'
    )

    doc.add_paragraph(
        'Fernandes et al. (2024) [14] provided a baseline comparison of CNN, Vision '
        'Transformer (ViT), and gated MLP architectures across four dental radiology '
        'classification tasks including mental foramen detection. ViT outperformed '
        'others for certain tasks (AUC 0.80\u20130.83), though no single architecture '
        'dominated. This task-dependent performance pattern is important context for '
        'interpreting our multi-landmark results.'
    )

    # 2.4
    doc.add_heading('2.4 Methodological Standards Observed in Recent DMFR Publications', level=2)

    doc.add_paragraph(
        'Analysis of the reviewed papers reveals several methodological patterns that '
        'DMFR reviewers and editors expect:'
    )

    items = [
        ('Inter- and intra-rater reliability: ',
         'All landmark detection studies report ICC for point-based agreement and '
         'Cohen\'s or Fleiss\' kappa for categorical agreement. Intra-rater reliability '
         '(re-evaluation of a subset after a washout period) is reported alongside '
         'inter-rater reliability in papers by Indermun et al. [8], Menezes et al. [9], '
         'and Kazimierczak et al. [10].'),
        ('Consistency measurement for LLMs: ',
         'Jeong et al. (2026) [4] established the precedent of repeated runs with '
         'Fleiss\' kappa for LLM evaluation in DMFR. Studies using temperature = 0 for '
         'deterministic output must at minimum verify and report determinism.'),
        ('Physical distance metrics: ',
         'Cephalometric landmark studies universally report errors in millimetres with '
         'SDR at 2 mm and 4 mm thresholds [6\u20139]. Grid-cell or pixel-based metrics '
         'are less common and require conversion factors for clinical interpretability.'),
        ('Power calculations: ',
         'DMFR author guidelines explicitly require power calculations. Sample size '
         'justification is expected in the methods section.'),
        ('Reporting guidelines: ',
         'The journal mandates adherence to applicable reporting checklists. For AI '
         'evaluation studies, TRIPOD+AI [15] and FLAIR [16] are the most relevant. '
         'STARD (Standards for Reporting Diagnostic Accuracy) may also apply.'),
        ('Response quality assessment: ',
         'LLM evaluation papers assess not only accuracy but also hallucination rates, '
         'response consistency, and reasoning quality [2\u20134].'),
        ('Bland-Altman analysis: ',
         'Agreement studies frequently include Bland-Altman plots to visualise '
         'systematic bias between methods [8, 10].'),
    ]

    for bold_part, normal_part in items:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(normal_part)

    # ================================================================
    # 3. POSITIONING
    # ================================================================
    doc.add_heading('3. Positioning of Our Work', level=1)

    doc.add_heading('3.1 Novelty and Contribution', level=2)

    doc.add_paragraph(
        'Our study occupies a genuinely unexplored intersection in the DMFR literature. '
        'While LLM text-based evaluations are well-represented [2\u20135], and '
        'purpose-built CNN/ViT models for landmark detection have been extensively '
        'studied [6\u201310], no published study in DMFR has evaluated multimodal LLMs '
        'on structured spatial landmark identification tasks using actual dental '
        'radiographs. The closest work, Jeong et al. (2026) [4], tested MLLMs on '
        'exam-style questions (including image-based ones) but not on structured '
        'coordinate-based landmark localisation.'
    )

    doc.add_paragraph(
        'Our contributions are fourfold:'
    )

    contributions = [
        ('Novel grid-based assessment framework: ',
         'The standardised square grid overlay discretises continuous spatial '
         'localisation into alphanumeric coordinates, enabling objective comparison '
         'between AI models and human observers using a common coordinate language. '
         'No DMFR publication employs this approach.'),
        ('First MLLM evaluation on radiographic landmark identification: ',
         'We are the first to evaluate frontier MLLMs (GPT-5.4, Gemini 3.1 Pro) on '
         'the specific task of identifying clinically critical anatomic landmarks on '
         'dental radiographs, as opposed to answering examination questions about them.'),
        ('Zero-shot, training-free paradigm: ',
         'Unlike purpose-built CNN systems that require thousands of annotated training '
         'images [6, 13, 14], our approach requires no training data, no fine-tuning, '
         'and no domain-specific model development. This makes it immediately deployable '
         'and reproducible.'),
        ('Multi-modality coverage with human benchmark: ',
         'We cover three radiographic modalities (panoramic, periapical, cephalometric) '
         'in a single study, with 40 dental students as a clinically meaningful human '
         'benchmark\u2014a scale exceeding most comparable studies in DMFR.'),
    ]

    for i, (bold_part, normal_part) in enumerate(contributions):
        p = doc.add_paragraph(style='List Number')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(normal_part)

    doc.add_heading('3.2 Strengths', level=2)

    strengths = [
        'Large, well-structured dataset (200 images, 900 queries, 12 landmark types) '
        'with balanced representation across modalities.',
        'Rigorous ground truth established by two independent OMFR specialists with '
        'consensus resolution, following the standard protocol observed in DMFR '
        'landmark studies [8, 9].',
        'Two prompting strategies (zero-shot baseline and guided with grid explanation) '
        'enabling investigation of whether explicit spatial anchoring improves MLLM '
        'accuracy\u2014a question directly motivated by the spatial reasoning '
        'deficiencies identified by Jeong et al. [4].',
        'Comprehensive metric suite spanning both point-based (ED, NED, SDR) and '
        'area-based (Jaccard, Dice) landmarks, with SDR at four clinically motivated '
        'thresholds.',
        'Statistical analysis plan with normality testing, non-parametric omnibus and '
        'post-hoc tests, effect sizes, and multiple comparison corrections\u2014aligning '
        'with DMFR\'s statistical rigour expectations.',
        'Ethical approval obtained from the Non-Interventional Clinical Research Ethics '
        'Committee of Burdur Mehmet Akif Ersoy University.',
        'Fully automated, reproducible pipeline using batch APIs with deterministic '
        'settings (temperature = 0), enabling exact replication.',
        'Use of the FDI two-digit numbering system for tooth identification, ensuring '
        'international consistency and avoiding ambiguity with the American numbering '
        'system.',
    ]

    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    # ================================================================
    # 4. WEAKNESSES AND LIMITATIONS
    # ================================================================
    doc.add_heading('4. Identified Weaknesses and Limitations', level=1)

    doc.add_heading('4.1 Critical Issues Requiring Action Before Experiments', level=2)

    # Issue 1
    p = doc.add_paragraph()
    run = p.add_run('Issue 1: Absence of Reproducibility/Consistency Verification')
    run.bold = True

    doc.add_paragraph(
        'Jeong et al. (2026) [4] established the methodological standard for LLM '
        'evaluation in DMFR by running each query 10 times and measuring response '
        'consistency with Fleiss\' kappa. Our current design uses temperature = 0 for '
        'deterministic outputs, which theoretically yields identical responses across '
        'runs. However, we do not verify this empirically, nor do we report any '
        'consistency metric. Reviewers familiar with the Jeong et al. methodology will '
        'expect either (a) empirical verification of determinism through repeated runs '
        'on a representative subset, or (b) explicit justification with supporting '
        'evidence that temperature = 0 guarantees identical outputs for these specific '
        'models and task types.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recommendation: ')
    run.bold = True
    p.add_run(
        'Conduct a consistency verification run on a representative subset (e.g., '
        '60 queries covering all modalities and landmark types) with 3 repetitions per '
        'model at temperature = 0. Compute Fleiss\' kappa and report it. If kappa = 1.0, '
        'this confirms determinism and justifies single-run evaluation. Estimated '
        'additional cost: < $2.'
    )

    # Issue 2
    p = doc.add_paragraph()
    run = p.add_run('Issue 2: Grid-Cell Metrics Lack Physical Distance Context')
    run.bold = True

    doc.add_paragraph(
        'The DMFR cephalometric landmark literature universally reports localisation '
        'errors in millimetres, with SDR at 2 mm and 4 mm as standard thresholds '
        '[6\u20139]. Our metrics are expressed in grid-cell units (SDR@0, SDR@1, '
        'SDR@\u221a2, SDR@2 cells), which are not directly interpretable for readers '
        'accustomed to physical distance reporting. Without a conversion factor, '
        'reviewers cannot contextualise our results against the existing body of '
        'landmark detection research.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recommendation: ')
    run.bold = True
    p.add_run(
        'Compute the approximate physical size of one grid cell for each modality '
        'using the known image pixel dimensions and the physical field of view of the '
        'imaging equipment. Report a conversion table (e.g., "1 panoramic grid cell '
        '\u2248 X mm") and provide dual-unit SDR thresholds where possible. If exact '
        'physical dimensions are unavailable, provide the pixel-based conversion '
        'and discuss the limitation.'
    )

    # Issue 3
    p = doc.add_paragraph()
    run = p.add_run('Issue 3: No Power Calculation')
    run.bold = True

    doc.add_paragraph(
        'DMFR author guidelines explicitly require power calculations. Our sample of '
        '900 queries provides substantial statistical power, but this must be formally '
        'computed and reported. Failure to include a power analysis is a common reason '
        'for desk rejection or major revision requests at DMFR.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recommendation: ')
    run.bold = True
    p.add_run(
        'Perform an a priori power analysis using G*Power or equivalent. For example, '
        'with 100 queries per panoramic landmark and a two-sided \u03b1 of 0.05, '
        'we can demonstrate > 99% power to detect a 0.5-cell mean ED difference '
        'between groups (assuming SD = 1.5 cells, Mann-Whitney U test). Report this '
        'calculation in the Methods section.'
    )

    # Issue 4
    p = doc.add_paragraph()
    run = p.add_run('Issue 4: Missing Intra-Rater Reliability Assessment')
    run.bold = True

    doc.add_paragraph(
        'Our ground truth validation plan includes inter-rater reliability (ICC between '
        'two OMFR specialists), which aligns with standard practice. However, all '
        'reviewed DMFR landmark studies also report intra-rater reliability\u2014the '
        'consistency of each specialist when re-evaluating a random subset after a '
        'washout period of at least two weeks [8\u201310]. The absence of intra-rater '
        'data would be noted as a methodological gap by reviewers.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recommendation: ')
    run.bold = True
    p.add_run(
        'If not already collected, arrange for each OMFR specialist to re-evaluate a '
        'random subset of 30 images (10 per modality) after a minimum 2-week washout '
        'period. Compute intra-rater ICC for point landmarks and intra-rater kappa '
        'for area landmarks. If this is infeasible, acknowledge explicitly in the '
        'Limitations section.'
    )

    doc.add_heading('4.2 Important Methodological Improvements', level=2)

    # Issue 5
    p = doc.add_paragraph()
    run = p.add_run('Issue 5: No Response Categorisation Beyond Parse Success/Failure')
    run.bold = True

    doc.add_paragraph(
        'Our pipeline currently classifies each model response as either parseable '
        '(valid grid coordinate extracted) or unparseable. However, unparseable '
        'responses may arise from fundamentally different failure modes: explicit '
        'refusals ("I cannot identify landmarks in radiographs"), verbose explanations '
        'without coordinates, hallucinated but structurally plausible responses, or '
        'format errors. Jeong et al. (2026) [4] addressed this by using a 5-point '
        'Global Quality Score (GQS) rated by specialists to assess hallucination. '
        'Distinguishing these failure modes provides richer insight into model behaviour '
        'and is expected by DMFR reviewers.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recommendation: ')
    run.bold = True
    p.add_run(
        'Implement an automated response categorisation step in the parsing pipeline '
        'that classifies non-coordinate responses into: (a) refusal, (b) verbose but '
        'contains coordinate, (c) ambiguous/multiple answers, (d) complete failure to '
        'engage. Report the distribution of these categories per model and strategy as '
        'a formal result ("instruction compliance rate").'
    )

    # Issue 6
    p = doc.add_paragraph()
    run = p.add_run('Issue 6: Missing Bland-Altman Agreement Analysis')
    run.bold = True

    doc.add_paragraph(
        'Several DMFR landmark studies use Bland-Altman plots to visualise agreement '
        'between measurement methods [8, 10]. These plots reveal systematic biases '
        '(e.g., a model consistently predicting landmarks 1 cell too far to the right) '
        'that summary statistics like mean ED may obscure. For the planned comparison '
        'between student mean responses and AI predictions, Bland-Altman analysis would '
        'provide a complementary perspective to ICC and statistical tests.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recommendation: ')
    run.bold = True
    p.add_run(
        'Add Bland-Altman plot generation to the analysis pipeline. For point landmarks, '
        'plot the difference between AI ED and student mean ED against their average. '
        'This will be particularly valuable when student data become available.'
    )

    # Issue 7
    p = doc.add_paragraph()
    run = p.add_run(
        'Issue 7: Post-Hoc Test Selection (Dunn\'s Test vs. Mann-Whitney U)')
    run.bold = True

    doc.add_paragraph(
        'Our technical report specifies Dunn\'s test with Bonferroni correction as '
        'the post-hoc procedure following a significant Kruskal-Wallis result, which '
        'is methodologically appropriate because Dunn\'s test uses the pooled rank '
        'information from the omnibus test. However, our code implementation uses '
        'pairwise Mann-Whitney U tests with Bonferroni correction, which is an '
        'independent pairwise approach that does not leverage the pooled ranks. Both '
        'approaches are published and accepted, but they may produce slightly different '
        'p-values. The discrepancy between the documented and implemented methods should '
        'be resolved.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recommendation: ')
    run.bold = True
    p.add_run(
        'Either update the code to use Dunn\'s test (available in scikit-posthocs or '
        'pingouin) to match the technical report, or update the technical report and '
        'paper methods to accurately describe the Mann-Whitney U approach. We recommend '
        'using Dunn\'s test for consistency with the stated plan and because it is '
        'the conventional post-hoc companion to Kruskal-Wallis in the DMFR literature.'
    )

    doc.add_heading('4.3 Inherent Limitations to Acknowledge', level=2)

    limitations = [
        ('Single-institution dataset: ',
         'All 200 images originate from one university hospital. External validation on '
         'images from different equipment, populations, or institutions would strengthen '
         'generalisability. This should be acknowledged and positioned as future work.'),
        ('Grid-cell resolution ceiling: ',
         'The grid system inherently discretises spatial information. A one-cell error '
         'on the periapical grid (8\u00d76) represents a proportionally larger spatial '
         'error than on the panoramic grid (16\u00d78). While NED normalisation '
         'mitigates this, the coarser periapical grid may disadvantage fine-grained '
         'landmark tasks.'),
        ('No comparison with purpose-built AI: ',
         'Our study compares MLLMs against students but not against dedicated CNN/ViT '
         'systems trained on dental radiographs. While this is intentional (our focus '
         'is zero-shot capability), reviewers may expect at least a discussion of how '
         'MLLM performance compares to the CNN benchmarks in the literature [6, 7, 13].'),
        ('Student group design: ',
         'The 5-group \u00d7 8-student design provides only 8 responses per query. '
         'While sufficient for computing a group mean with reasonable confidence, this '
         'limits the granularity of student inter-observer variability analysis. The '
         'choice should be justified by the fatigue-mitigation rationale.'),
        ('Temporal model availability: ',
         'LLMs are updated frequently. GPT-5.4 and Gemini 3.1 Pro represent the state '
         'of the art at the time of the experiment (2026), but newer models may become '
         'available before publication. We should clearly state model versions, access '
         'dates, and API versions for reproducibility.'),
    ]

    for bold_part, normal_part in limitations:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(normal_part)

    # ================================================================
    # 5. LITERATURE COMPARISON TABLE
    # ================================================================
    doc.add_heading('5. Comparative Overview: Our Study vs. Key DMFR Publications', level=1)

    doc.add_paragraph(
        'Table 1 summarises how our study compares to the most relevant recent '
        'publications in DMFR across key methodological dimensions.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Table 1. ')
    run.bold = True
    p.add_run('Comparative positioning against recent DMFR publications.')
    p.paragraph_format.space_after = Pt(4)

    headers = ['Dimension', 'Liu et al.\n2025 [2]', 'Jeong et al.\n2024 [3]',
               'Jeong et al.\n2026 [4]', 'Jiang et al.\n2023 [6]', 'Our Study']
    rows = [
        ['Study type', 'Systematic\nreview', 'Cross-\nsectional', 'Cross-\nsectional',
         'AI development\n& validation', 'Cross-\nsectional'],
        ['Task', 'Various\nLLM tasks', 'Exam MCQs\n(text only)', 'Exam MCQs\n(text+image)',
         'Cephalometric\nlandmarks', 'Landmark\nidentification'],
        ['Models tested', '19 studies\nreviewed', 'GPT-3.5/4,\nBard, Bing',
         'ChatGPT-4o,\nGemini 2.0', 'CephNet\n(custom CNN)', 'GPT-5.4,\nGemini 3.1 Pro'],
        ['Image-based?', 'Varied', 'No', 'Yes', 'Yes', 'Yes'],
        ['Human comparator', 'N/A', '120 students', 'Specialist\nGQS ratings',
         'Manual\ntracing', '40 students\n+ 2 OMFRs'],
        ['N (queries)', 'N/A', '52', '90', '9,870 images\n\u00d730 landmarks',
         '900'],
        ['Spatial metric', 'N/A', 'N/A', 'N/A', 'ED (mm),\nSDR@2mm', 'ED, NED,\nSDR (cells)'],
        ['Overlap metric', 'N/A', 'N/A', 'N/A', 'N/A', 'Jaccard, Dice'],
        ['Consistency', 'N/A', 'N/A', 'Fleiss\' \u03ba\n(10 runs)', 'N/A',
         'To be added\n(3 runs)'],
        ['Prompting\nstrategies', 'N/A', 'Single', 'Single', 'N/A', '2 (zero-shot\n+ guided)'],
        ['Grid-based?', 'No', 'No', 'No', 'No', 'Yes (novel)'],
    ]

    create_table(doc, headers, rows)
    doc.add_paragraph('')  # spacing

    # ================================================================
    # 6. SUGGESTED IMPROVEMENTS SUMMARY
    # ================================================================
    doc.add_heading('6. Summary of Recommended Revisions', level=1)

    doc.add_paragraph(
        'Table 2 consolidates all recommended revisions with their priority, estimated '
        'effort, and impact on the study.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Table 2. ')
    run.bold = True
    p.add_run('Prioritised revision plan.')
    p.paragraph_format.space_after = Pt(4)

    headers2 = ['#', 'Revision', 'Priority', 'Effort', 'Impact']
    rows2 = [
        ['1', 'Add consistency verification\n(3 repeated runs on subset)', 'Must',
         '~$2, 1 hr', 'Addresses top reviewer\nconcern (Jeong 2026)'],
        ['2', 'Compute grid-cell-to-mm\nconversion factors', 'Must',
         '30 min', 'Enables comparison\nwith literature'],
        ['3', 'Add power calculation\nto Methods', 'Must',
         '30 min', 'DMFR requirement'],
        ['4', 'Collect intra-rater reliability\ndata (30-image subset)', 'Must',
         '2\u20133 hrs\n(specialist time)', 'Standard for DMFR\nlandmark papers'],
        ['5', 'Implement response\ncategorisation', 'Should',
         '1\u20132 hrs\n(code)', 'Richer failure\nanalysis'],
        ['6', 'Add Bland-Altman\nplot capability', 'Should',
         '1 hr (code)', 'Standard agreement\nvisualisation'],
        ['7', 'Align post-hoc test\n(Dunn\'s vs. MWU)', 'Should',
         '30 min', 'Method consistency'],
        ['8', 'Report instruction compliance\nrate as formal metric', 'Should',
         'Trivial', 'Quantifies model\ncooperativeness'],
        ['9', 'Consider activating Claude\nSonnet 4.6 as 3rd model', 'Could',
         '~$5\u201310', 'Stronger multi-provider\ncomparison'],
    ]

    create_table(doc, headers2, rows2)
    doc.add_paragraph('')

    # ================================================================
    # 7. PRELIMINARY EXPERIMENT PLAN
    # ================================================================
    doc.add_heading('7. Preliminary Experiment Plan', level=1)

    doc.add_paragraph(
        'Before committing to the full 3,600-call experiment, we propose a preliminary '
        'validation run to verify pipeline correctness, assess response quality, '
        'calibrate expectations, and identify any issues requiring methodology adjustment. '
        'This preliminary experiment is designed to answer five specific questions:'
    )

    questions = [
        'Can the models parse and follow the grid coordinate response format reliably?',
        'What is the approximate accuracy range, and does it align with the <30% '
        'image-based accuracy reported by Jeong et al. (2026) [4], or does the grid '
        'system meaningfully improve spatial performance?',
        'Does the guided prompting strategy produce measurably different results from '
        'the zero-shot baseline?',
        'Are there model-specific failure patterns (refusals, hallucinations, format '
        'errors) that require prompt adjustments?',
        'Is the output deterministic at temperature = 0 across repeated runs '
        '(consistency verification)?',
    ]

    for q in questions:
        doc.add_paragraph(q, style='List Number')

    doc.add_heading('7.1 Sample Design', level=2)

    doc.add_paragraph(
        'The preliminary sample is constructed to be representative across all key '
        'dimensions of the full experiment while remaining cost-efficient:'
    )

    headers3 = ['Modality', 'Images', 'Landmarks/Image', 'Queries', 'Selection Method']
    rows3 = [
        ['Panoramic', '5', '6 (3 point + 3 area)', '30',
         'Stratified random: images 1, 20, 40, 60, 80'],
        ['Periapical', '5', '3 (all point)', '15',
         'Stratified random: images 1, 10, 20, 30, 40'],
        ['Cephalometric', '5', '3 (all point)', '15',
         'Stratified random: images 1, 10, 20, 30, 40'],
        ['TOTAL', '15', '\u2014', '60', '\u2014'],
    ]

    create_table(doc, headers3, rows3)
    doc.add_paragraph('')

    doc.add_paragraph(
        'This yields 60 queries per model per strategy. Including the consistency '
        'verification phase, the total API call count is:'
    )

    headers4 = ['Phase', 'Queries', 'Models', 'Strategies', 'Repetitions', 'Total Calls']
    rows4 = [
        ['Main preliminary', '60', '2', '2', '1', '240'],
        ['Consistency check', '60', '2', '1 (zero-shot)', '3', '360'],
        ['TOTAL', '\u2014', '\u2014', '\u2014', '\u2014', '600'],
    ]

    create_table(doc, headers4, rows4)
    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('Estimated cost: ')
    run.bold = True
    p.add_run(
        'Approximately $2\u20134, based on batch API pricing for GPT-5.4 and '
        'Gemini 3.1 Pro.'
    )

    doc.add_heading('7.2 Execution Protocol', level=2)

    steps = [
        ('Step 1 \u2014 Sample Extraction: ',
         'Extract the 60-query subset from the existing query_index.json using '
         'the stratified image IDs specified above. Verify that all 12 landmark '
         'types and both point/area categories are represented.'),
        ('Step 2 \u2014 Main Preliminary Run: ',
         'Submit the 60-query subset to both models (GPT-5.4, Gemini 3.1 Pro) '
         'using both strategies (zero-shot, guided). Total: 240 API calls. '
         'Parse responses and compute all metrics (ED, NED, Jaccard, Dice, SDR). '
         'Categorise any non-parseable responses.'),
        ('Step 3 \u2014 Consistency Verification: ',
         'Re-submit the same 60 queries to both models using the zero-shot strategy '
         'only, for 3 additional repetitions (total 4 runs including the main run). '
         'Compute Fleiss\' kappa across the 4 runs for each model. If kappa < 1.0, '
         'investigate which queries produced variable responses and assess whether '
         'temperature = 0 is reliably deterministic.'),
        ('Step 4 \u2014 Analysis: ',
         'Generate summary statistics, visualisations, and a comparison table. '
         'Evaluate the five questions listed above. Document any anomalies or '
         'unexpected patterns.'),
        ('Step 5 \u2014 Decision Gate: ',
         'Based on preliminary results, make a go/no-go decision on the full experiment. '
         'If the compliance rate is below 70% or systematic prompt issues are detected, '
         'revise prompts before proceeding. If the models show no meaningful engagement '
         'with the images (random-level accuracy), document this finding and discuss '
         'whether alternative approaches are warranted.'),
    ]

    for bold_part, normal_part in steps:
        p = doc.add_paragraph()
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(normal_part)

    doc.add_heading('7.3 Success Criteria for Preliminary Experiment', level=2)

    p = doc.add_paragraph(
        'The following criteria determine whether the full experiment should proceed '
        'without modification:'
    )

    headers5 = ['Criterion', 'Threshold', 'Action if Not Met']
    rows5 = [
        ['Instruction compliance\n(parseable responses)',
         '\u2265 80% per model',
         'Revise prompt templates;\nconsider adding format\nenforcement instructions'],
        ['Consistency\n(Fleiss\' kappa at temp=0)',
         '\u03ba = 1.0\n(or \u2265 0.95)',
         'Investigate non-deterministic\nqueries; consider alternative\nsampling strategy'],
        ['Meaningful engagement\n(not random-level)',
         'SDR@2 > 15%\n(above chance)',
         'Evaluate whether models\nengage with grid overlay;\nconsider visual prompt\nadjustments'],
        ['No systematic refusals',
         'Refusal rate < 10%',
         'Adjust system prompt to\nreduce safety-triggered\nrefusals on medical images'],
        ['Strategy differentiation',
         'Observable difference\nbetween zero-shot\nand guided',
         'If no difference, consider\nwhether guided strategy\nadds value to study design'],
    ]

    create_table(doc, headers5, rows5)
    doc.add_paragraph('')

    doc.add_heading('7.4 Expected Outcomes and Contingencies', level=2)

    doc.add_paragraph(
        'Based on the literature, we anticipate the following outcomes from the '
        'preliminary experiment:'
    )

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Moderate accuracy on point landmarks: ')
    run.bold = True
    p.add_run(
        'Given Jeong et al.\'s finding of <30% image-based accuracy for MLLMs [4], '
        'we expect SDR@0 (exact match) to be low (5\u201320%). However, SDR@2 '
        '(\u00b12 cells) may be substantially higher if the grid overlay provides '
        'effective spatial anchoring. An SDR@2 of 30\u201350% for point landmarks '
        'would represent a meaningful improvement over unstructured spatial reasoning.'
    )

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Lower accuracy on area landmarks: ')
    run.bold = True
    p.add_run(
        'Area landmarks (mandibular canal, maxillary sinus, external oblique ridge) '
        'require identifying multiple cells, which is inherently more difficult. '
        'We expect Jaccard indices of 0.1\u20130.4, with higher Dice coefficients '
        'due to partial overlap credit.'
    )

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Guided strategy improvement: ')
    run.bold = True
    p.add_run(
        'We hypothesise that the guided prompting strategy will improve performance, '
        'as explicit grid explanation provides the spatial anchoring that Jeong et al. '
        'identified as lacking in general MLLM visual reasoning. The magnitude of '
        'improvement is uncertain but even a 5\u201310 percentage-point SDR@1 gain '
        'would be a publishable finding.'
    )

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('High compliance rate: ')
    run.bold = True
    p.add_run(
        'Both GPT-5.4 and Gemini 3.1 Pro are frontier models with strong instruction '
        'following. With temperature = 0 and explicit format instructions, we expect '
        '\u226590% compliance (parseable coordinate responses).'
    )

    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('Deterministic outputs: ')
    run.bold = True
    p.add_run(
        'At temperature = 0, we expect Fleiss\' kappa = 1.0 or very close to it, '
        'confirming that single-run evaluation is valid for the full experiment.'
    )

    # ================================================================
    # 8. CONCLUSION
    # ================================================================
    doc.add_heading('8. Conclusion', level=1)

    doc.add_paragraph(
        'Our study is well-positioned to make a novel and timely contribution to '
        'the DMFR literature. The grid-based assessment framework, the first MLLM '
        'evaluation on structured radiographic landmark identification, and the '
        'student comparison collectively address a genuine gap in the field. The '
        'core methodology is sound, with comprehensive metrics and a rigorous '
        'statistical plan.'
    )

    doc.add_paragraph(
        'However, this analysis has identified several areas where the experimental '
        'design can be strengthened before commencing the full experiment. The most '
        'critical additions are: (1) consistency verification through repeated runs, '
        '(2) grid-cell-to-millimetre conversion factors for clinical interpretability, '
        '(3) formal power calculation, and (4) intra-rater reliability data. '
        'Additionally, response categorisation, Bland-Altman analysis, and post-hoc '
        'test alignment would bring our methodology into closer conformity with DMFR '
        'publication standards.'
    )

    doc.add_paragraph(
        'The proposed preliminary experiment (60 queries, 600 total API calls, '
        '~$2\u20134) provides a low-cost, low-risk validation step that will '
        'calibrate expectations, verify pipeline correctness, and identify any '
        'prompt or methodology adjustments needed before committing to the full '
        '3,600-call experiment. We recommend completing all revisions and the '
        'preliminary experiment before proceeding to the full-scale evaluation.'
    )

    # ================================================================
    # REFERENCES
    # ================================================================
    doc.add_heading('References', level=1)

    references = [
        '[1] Bornstein MM, Staedler YM, Cattin PC. Artificial intelligence will change '
        'the research environment in dental medicine dramatically: will algorithms '
        'replace literature reviews in the near future? Dentomaxillofac Radiol. '
        '2025;54(5):329\u2013331. doi:10.1093/dmfr/twaf025',

        '[2] Liu Z, Nalley A, Hao J, Ai QYH, Yeung AWK, Tanaka R, Hung KF. '
        'The performance of large language models in dentomaxillofacial radiology: '
        'a systematic review. Dentomaxillofac Radiol. 2025;54(8):613+. '
        'doi:10.1093/dmfr/twaf060',

        '[3] Jeong H, Han SS, Yu Y, Kim S, Jeon KJ. How well do large language '
        'model-based chatbots perform in oral and maxillofacial radiology? '
        'Dentomaxillofac Radiol. 2024;53(6):390+. doi:10.1093/dmfr/twae025',

        '[4] Jeong H, Jeon KJ, Lee C, Choi YJ, Jo GD, Han SS. Leveraging multimodal '
        'large language model chatbots in oral radiology: a comprehensive evaluation '
        'using questions from a Korean dental university. Dentomaxillofac Radiol. '
        '2026;55(3):276\u2013286. doi:10.1093/dmfr/twaf083',

        '[5] Najmuddin M. Assessing the accuracy of multiple-choice questions using '
        'different artificial intelligence-driven tools\u2014an observational study. '
        'Dentomaxillofac Radiol. 2025;55(3):287+. doi:10.1093/dmfr/twaf081',

        '[6] Jiang F, Guo Y, Yang C, Zhou Y, Lin Y, Cheng F, Quan S, Feng Q, Li J. '
        'Artificial intelligence system for automated landmark localization and '
        'analysis of cephalometry. Dentomaxillofac Radiol. 2023;52(1):20220081. '
        'doi:10.1259/dmfr.20220081',

        '[7] Weingart JV, Schlager S, Metzger MC, Brandenburg LS, Hein A, '
        'Schmelzeisen R, Bamberg F, Kim S, Kellner E, Reisert M, Russe MF. '
        'Automated detection of cephalometric landmarks using deep neural patchworks. '
        'Dentomaxillofac Radiol. 2023;52(6):20230059. doi:10.1259/dmfr.20230059',

        '[8] Indermun S, Shaik S, Nyirenda C, Johannes K, Mulder R. Human examination '
        'and artificial intelligence in cephalometric landmark detection\u2014is AI '
        'ready to take over? Dentomaxillofac Radiol. 2023;52(6):20220362. '
        'doi:10.1259/dmfr.20220362',

        '[9] Menezes LS, Silva TP, Santos MAL, Hughes MM, Souza SRM, Ribeiro PML, '
        'de Freitas PHL, Takeshita WM. Assessment of landmark detection in '
        'cephalometric radiographs with different conditions of brightness and '
        'contrast using artificial intelligence software. Dentomaxillofac Radiol. '
        '2023;52(8):20230065. doi:10.1259/dmfr.20230065',

        '[10] Kazimierczak N, Kazimierczak W, Serafin Z, Nowicki P, Jankowski T, '
        'Jankowska A, Janiszewska-Olszowska J. Skeletal facial asymmetry: reliability '
        'of manual and artificial intelligence-driven analysis. Dentomaxillofac Radiol. '
        '2024;53(1):52\u201359. doi:10.1093/dmfr/twad006',

        '[11] Semper-Hogg W, Rau A, Fuessinger MA, Zimmermann S, Bamberg F, '
        'Metzger MC, Schmelzeisen R, Rau S, Reisert M, Russe MF. Deep learning-based '
        'segmentation of the mandibular canals in cone-beam CT reaches human-level '
        'performance. Dentomaxillofac Radiol. 2025;54(4):279\u2013285. '
        'doi:10.1093/dmfr/twaf006',

        '[12] Hamdan MH, Uribe SE, Tuzova L, Tuzoff D, Badr Z, Mol A, Tyndall DA. '
        'The influence of a deep learning tool on the performance of oral and '
        'maxillofacial radiologists in the detection of apical radiolucencies. '
        'Dentomaxillofac Radiol. 2024;54(2):118\u2013124. doi:10.1093/dmfr/twae058',

        '[13] Celik B, Savasteer EF, Kaya HI, Celik ME. The role of deep learning '
        'for periapical lesion detection on panoramic radiographs. Dentomaxillofac '
        'Radiol. 2023;52(8):20230118. doi:10.1259/dmfr.20230118',

        '[14] Fernandes FA, Ge M, Chaltikyan G, Gerdes MW, Omlin CW. Preparing for '
        'downstream tasks in artificial intelligence for dental radiology: a baseline '
        'performance comparison of deep learning models. Dentomaxillofac Radiol. '
        '2024;54(2):149+. doi:10.1093/dmfr/twae056',

        '[15] Collins GS, Moons KGM, Dhiman P, Riley RD, Beam AL, Van Calster B, '
        'et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction '
        'models that use regression or machine learning methods. BMJ. '
        '2024;385:e078378. doi:10.1136/bmj-2023-078378',

        '[16] Kottlors J, Iuga AI, Bluethgen C, Bressem K, Kather JN, Moy L, et al. '
        'Guidelines for reporting studies on large language models in radiology: an '
        'international Delphi expert survey (FLAIR). Radiology. 2026;318(2):e250913. '
        'doi:10.1148/radiol.250913',
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(10)

    return doc


if __name__ == '__main__':
    doc = build_document()
    output_path = 'docs/sota_analysis_and_positioning_of_our_work.docx'
    doc.save(output_path)
    print(f'Report saved: {output_path}')
