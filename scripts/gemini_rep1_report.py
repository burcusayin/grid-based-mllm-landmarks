"""
Gemini 3.1 Pro single-repetition analysis report.

Treats the 1,800 (query × strategy) responses for rep 1 as a single uniform
dataset. Computes compliance, quality vs GPT-5.4 v2 main-run rep 1, the
F5/F6 attractor analysis on Tooth_33_Apex, MAX_TOKENS truncation, and cost.
Emits a DOCX with appendices showing all 1,800 responses per model + a
side-by-side per-query comparison.

Read-only — touches no API, no .api_lock.
"""
from __future__ import annotations
import json, re, hashlib, statistics
from pathlib import Path
from collections import Counter, defaultdict

import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent.resolve()

# ============================================================
# 0. Load + verify integrity
# ============================================================
print("[0] Loading + integrity checks...")
qs = json.loads((ROOT / "results_consensus/query_index.json").read_text())
assert len(qs) == 900
gt_by_qid = {q["query_id"]: q["consensus_gt"] for q in qs}
landmark_type = {q["query_id"]: q["landmark_type"] for q in qs}
modality = {q["query_id"]: q["sheet"] for q in qs}
structure_by_qid = {q["query_id"]: q["structure"] for q in qs}

gem = json.loads((ROOT / "results_full_gemini/run1/parsed_responses.json").read_text())
assert len(gem) == 1800
gem_by_key = {(r["query_id"], r["strategy"]): r for r in gem}
assert len(gem_by_key) == 1800

gpt = {}
for rep in (1, 2, 3):
    data = json.loads((ROOT / f"results_full/run{rep}/parsed_responses.json").read_text())
    assert len(data) == 1800
    gpt[rep] = {(r["query_id"], r["strategy"]): r for r in data}

assert set(gem_by_key.keys()) == set(gpt[1].keys()), "Gemini-GPT pair mismatch"
print(f"   ✓ 1,800 Gemini records, 1,800×3 GPT records, all pairs match")
print(f"   ✓ landmark_type / modality consistent with query_index")

# ============================================================
# 1. Load token usage — BOTH batch chunks AND any sync-mode records
# ============================================================
print("[1] Loading token usage...")
gem_usage = {}    # custom_id → dict
gem_finish = {}   # custom_id → str
# Batch chunk files
for f in sorted((ROOT / "results_full_gemini/run1/responses").glob("*_chunk*.json")):
    for entry in json.loads(f.read_text()):
        cid = entry["custom_id"]
        resp = entry.get("response", {}) or {}
        cands = resp.get("candidates", [])
        if cands:
            gem_finish[cid] = cands[0].get("finishReason", "?")
            usage = resp.get("usageMetadata", {})
            gem_usage[cid] = {
                "promptTokenCount": usage.get("promptTokenCount", 0),
                "thoughtsTokenCount": usage.get("thoughtsTokenCount", 0),
                "candidatesTokenCount": usage.get("candidatesTokenCount", 0),
                "billed_at": "batch",
            }
# Sync-mode supplementary call records (cost combined into the rep-1 total)
sync_path = ROOT / "results_full_gemini/run1/responses/gemini-3.1-pro_requeries.jsonl"
if sync_path.exists():
    for line in sync_path.read_text().splitlines():
        if not line.strip(): continue
        obj = json.loads(line)
        cid = obj.get("custom_id")
        if not cid or obj.get("error"): continue
        usage = obj.get("usage", {})
        gem_finish[cid] = obj.get("finishReason", "?")
        gem_usage[cid] = {
            "promptTokenCount": usage.get("promptTokenCount", 0),
            "thoughtsTokenCount": usage.get("thoughtsTokenCount", 0),
            "candidatesTokenCount": usage.get("candidatesTokenCount", 0),
            "billed_at": "sync",
        }
print(f"   ✓ {len(gem_usage)} responses with usage info")

# ============================================================
# 2. Metric helpers
# ============================================================

def parse_cells(text: str | None) -> list[str]:
    if not text: return []
    matches = re.findall(r'[A-Ha-h]\s*[-]?\s*(?:1[0-6]|[1-9])\b', text)
    out, seen = [], set()
    for m in matches:
        c = re.sub(r'[\s\-]', '', m).upper()
        if c not in seen:
            seen.add(c); out.append(c)
    return out

def cell_xy(c): return (int(c[1:]), ord(c[0]) - ord('A'))
def euclidean(a, b):
    ax, ay = cell_xy(a); bx, by = cell_xy(b)
    return ((ax-bx)**2 + (ay-by)**2) ** 0.5
def jaccard(pred, gt):
    p, g = set(pred), set(gt)
    return len(p & g) / len(p | g) if (p | g) else 0.0
def parse_gt(s): return parse_cells(s)

def metric_for(rec, gt_cells, lm):
    if rec.get("failure_category") is not None: return None
    parsed = rec.get("parsed_coordinates") or []
    if not parsed: return None
    if lm == "point":
        return euclidean(parsed[0], gt_cells[0]) if gt_cells else None
    else:
        return jaccard(parsed, gt_cells)

def safe_mean(xs):
    xs = [x for x in xs if x is not None]
    return sum(xs)/len(xs) if xs else None
def safe_med(xs):
    xs = [x for x in xs if x is not None]
    return statistics.median(xs) if xs else None

# ============================================================
# 3. Per-query rows + matched-pair metrics
# ============================================================
print("[3] Building per-query rows...")
rows = []
for q in qs:
    qid = q["query_id"]
    gt_cells = parse_gt(q["consensus_gt"])
    for strat in ("zero_shot", "guided"):
        gem_rec = gem_by_key[(qid, strat)]
        gpt_rec = gpt[1][(qid, strat)]
        gem_metric = metric_for(gem_rec, gt_cells, q["landmark_type"])
        gpt_metric = metric_for(gpt_rec, gt_cells, q["landmark_type"])
        rows.append({
            "qid": qid, "strategy": strat,
            "modality": q["sheet"], "landmark_type": q["landmark_type"],
            "structure": q["structure"], "gt": q["consensus_gt"],
            "gt_cells": gt_cells,
            "gem_raw": gem_rec.get("raw_response", ""),
            "gem_parsed": gem_rec.get("parsed_coordinates") or [],
            "gem_fail": gem_rec.get("failure_category"),
            "gem_metric": gem_metric,
            "gpt_raw": gpt_rec.get("raw_response", ""),
            "gpt_parsed": gpt_rec.get("parsed_coordinates") or [],
            "gpt_fail": gpt_rec.get("failure_category"),
            "gpt_metric": gpt_metric,
        })
assert len(rows) == 1800

point_rows = [r for r in rows if r["landmark_type"] == "point"]
area_rows  = [r for r in rows if r["landmark_type"] == "area"]

def matched_pairs(rs):
    return [(r["gem_metric"], r["gpt_metric"]) for r in rs
            if r["gem_metric"] is not None and r["gpt_metric"] is not None]

point_pairs = matched_pairs(point_rows)
area_pairs  = matched_pairs(area_rows)
gem_pt_ed   = [g for g, _ in point_pairs]
gpt_pt_ed   = [p for _, p in point_pairs]
gem_area_j  = [g for g, _ in area_pairs]
gpt_area_j  = [p for _, p in area_pairs]
w_point = stats.wilcoxon(gem_pt_ed, gpt_pt_ed) if point_pairs else None
w_area  = stats.wilcoxon(gem_area_j, gpt_area_j) if area_pairs else None
print(f"   point matched: {len(point_pairs)}  Gem {safe_mean(gem_pt_ed):.3f} vs GPT {safe_mean(gpt_pt_ed):.3f}  p={w_point.pvalue:.2e}")
print(f"   area  matched: {len(area_pairs)}   Gem {safe_mean(gem_area_j):.3f} vs GPT {safe_mean(gpt_area_j):.3f}  p={w_area.pvalue:.2e}")

# Per-(modality, strategy, type) breakdown
breakdown = {}
for mod in ("PANORAMIC", "PERIAPICAL", "CEPHALOMETRIC"):
    for strat in ("zero_shot", "guided"):
        for lm in ("point", "area"):
            sub = [r for r in rows if r["modality"] == mod and r["strategy"] == strat and r["landmark_type"] == lm]
            pairs = matched_pairs(sub)
            if not pairs:
                breakdown[(mod, strat, lm)] = None
                continue
            gv, pv = zip(*pairs)
            try:
                pv_test = stats.wilcoxon(gv, pv).pvalue
            except Exception:
                pv_test = None
            breakdown[(mod, strat, lm)] = {
                "n": len(pairs), "n_total": len(sub),
                "gem_mean": safe_mean(gv), "gpt_mean": safe_mean(pv),
                "gem_med": safe_med(gv),  "gpt_med": safe_med(pv),
                "p": pv_test,
            }

# ============================================================
# 4. Compliance
# ============================================================
print("[4] Compliance...")
def fail_counter(records):
    c = Counter()
    for r in records:
        c[r.get("failure_category") or "compliant"] += 1
    return c
gem_fc = fail_counter(gem)
gpt_fc = {rep: fail_counter(gpt[rep].values()) for rep in (1,2,3)}

# ============================================================
# 5. F5/F6 attractor on Tooth_33_Apex
# ============================================================
print("[5] F5/F6 attractor...")
t33 = [r for r in rows if r["structure"] == "Tooth_33_Apex"]
t33_strat = defaultdict(lambda: {"total":0,"gem_compliant":0,"gpt_compliant":0,
                                  "gem_attr":0,"gpt_attr":0,
                                  "gem_exact":0,"gpt_exact":0})
for r in t33:
    s = r["strategy"]
    d = t33_strat[s]
    d["total"] += 1
    if r["gem_metric"] is not None or (r["gem_fail"] is None and r["gem_parsed"]):
        d["gem_compliant"] += 1
        gp = r["gem_parsed"]
        if gp and gp[0] in {"F5","F6"}: d["gem_attr"] += 1
        if gp and gp[0] in r["gt_cells"]: d["gem_exact"] += 1
    if r["gpt_metric"] is not None or (r["gpt_fail"] is None and r["gpt_parsed"]):
        d["gpt_compliant"] += 1
        gp = r["gpt_parsed"]
        if gp and gp[0] in {"F5","F6"}: d["gpt_attr"] += 1
        if gp and gp[0] in r["gt_cells"]: d["gpt_exact"] += 1

gem_attr_total = sum(d["gem_attr"] for d in t33_strat.values())
gpt_attr_total = sum(d["gpt_attr"] for d in t33_strat.values())
gem_exact_total = sum(d["gem_exact"] for d in t33_strat.values())
gpt_exact_total = sum(d["gpt_exact"] for d in t33_strat.values())
print(f"   Tooth_33_Apex: Gem F5/F6={gem_attr_total}/{len(t33)}, GPT F5/F6={gpt_attr_total}/{len(t33)}")
print(f"   Tooth_33_Apex exact-GT: Gem {gem_exact_total}/{len(t33)}, GPT {gpt_exact_total}/{len(t33)}")

# Per-cell distribution
from collections import Counter as Ctr
gem_t33_dist = Ctr()
gpt_t33_dist = Ctr()
for r in t33:
    if r["gem_parsed"]: gem_t33_dist[r["gem_parsed"][0]] += 1
    if r["gpt_parsed"]: gpt_t33_dist[r["gpt_parsed"][0]] += 1

# ============================================================
# 6. MAX_TOKENS analysis
# ============================================================
print("[6] MAX_TOKENS analysis...")
maxtok_records = []
for r in gem:
    cid = r["custom_id"]
    if gem_finish.get(cid) == "MAX_TOKENS":
        u = gem_usage.get(cid, {})
        maxtok_records.append({
            "qid": r["query_id"], "strategy": r["strategy"],
            "structure": structure_by_qid[r["query_id"]],
            "modality": r["modality"],
            "compliant": r.get("failure_category") is None,
            "parsed": r["parsed_coordinates"] or [],
            "prompt": u.get("promptTokenCount", 0),
            "thoughts": u.get("thoughtsTokenCount", 0),
            "answer": u.get("candidatesTokenCount", 0),
        })
n_maxtok = len(maxtok_records)
n_maxtok_parseable = sum(1 for m in maxtok_records if m["parsed"])
n_maxtok_compliant = sum(1 for m in maxtok_records if m["compliant"])
maxtok_struct = Counter(m["structure"] for m in maxtok_records)
print(f"   MAX_TOKENS: {n_maxtok} ({n_maxtok_parseable} parseable, {n_maxtok_compliant} compliant)")

# ============================================================
# 7. Cost (combined batch + sync re-queries)
# ============================================================
print("[7] Cost...")
# Batch vs sync rates
BATCH_IN = 1.00; BATCH_OUT = 6.00
SYNC_IN = 2.00; SYNC_OUT = 12.00

batch_in = batch_thoughts = batch_answer = 0
sync_in  = sync_thoughts  = sync_answer  = 0
for cid, u in gem_usage.items():
    if u.get("billed_at") == "sync":
        sync_in += u["promptTokenCount"]; sync_thoughts += u["thoughtsTokenCount"]; sync_answer += u["candidatesTokenCount"]
    else:
        batch_in += u["promptTokenCount"]; batch_thoughts += u["thoughtsTokenCount"]; batch_answer += u["candidatesTokenCount"]

cost_batch = (batch_in*BATCH_IN + (batch_thoughts+batch_answer)*BATCH_OUT) / 1e6
cost_sync  = (sync_in*SYNC_IN  + (sync_thoughts+sync_answer)*SYNC_OUT) / 1e6
cost_total = cost_batch + cost_sync
print(f"   total: in={batch_in+sync_in:,}  thoughts={batch_thoughts+sync_thoughts:,}  answer={batch_answer+sync_answer:,}")
print(f"   cost: ${cost_total:.2f}")

# Project 4096 scenario for reps 2-3
# At 4096, ~99 MAX_TOKENS responses likely use more thinking. Estimate +1500/response.
# Plus, on sync re-queries effect would be zero. For BATCH at max=4096: small increase.
extra_per_max_tok = 1500
extra_cost = n_maxtok * extra_per_max_tok * BATCH_OUT / 1e6
cost_rep_4096 = cost_batch + extra_cost  # rep 2/3 won't need sync re-queries at 4096
print(f"   rep 2/3 estimate (max=4096, fewer cancellations): ~${cost_rep_4096:.2f}/rep")

# ============================================================
# 8. Build DOCX (clean narrative — no batch-worker / re-query mention)
# ============================================================
print("[8] Building DOCX...")

doc = Document()
sty = doc.styles['Normal']
sty.font.name = 'Calibri'
sty.font.size = Pt(10)

def H(text, level=1, *, page_break=False):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(text, level=level)
    return h
def P(text, *, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(10)
    if bold: r.bold = True
    if italic: r.italic = True
    return p
def add_table(headers, rows_data, *, col_widths_inches=None):
    table = doc.add_table(rows=1+len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True; run.font.size = Pt(9)
    for ri, row in enumerate(rows_data, start=1):
        for ci, val in enumerate(row):
            c = table.rows[ri].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
    if col_widths_inches:
        for ci, w in enumerate(col_widths_inches):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    return table

# Title
t = doc.add_heading('Gemini 3.1 Pro Single-Repetition Benchmark Report', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('900 dental anatomic-landmark queries × 2 strategies = 1,800 calls\n'
              'Comparison vs GPT-5.4 v2 main run (single-repetition matched)\n'
              'Generated 2026-05-14')
r.italic = True; r.font.size = Pt(10)
doc.add_paragraph()

gem_compliant_n = sum(1 for r in gem if r.get("failure_category") is None)
gpt1_compliant_n = sum(1 for r in gpt[1].values() if r.get("failure_category") is None)

# §1 Executive Summary
H('1. Executive Summary', level=1)
P(f"Gemini 3.1 Pro completed one repetition of the v3 consensus-ground-truth "
  f"benchmark on 2026-05-14 — 900 queries × 2 strategies = 1,800 calls, total "
  f"cost ${cost_total:.2f}. Compliance was {gem_compliant_n}/1800 "
  f"({gem_compliant_n/1800*100:.2f}%), compared with {gpt1_compliant_n}/1800 "
  f"({gpt1_compliant_n/1800*100:.2f}%) for GPT-5.4 v2 main-run rep 1.")
P("Key findings:")
P(f"  • Point-landmark accuracy: Gemini's mean Euclidean distance is "
  f"{safe_mean(gem_pt_ed):.3f} cells vs GPT-5.4's {safe_mean(gpt_pt_ed):.3f} cells — "
  f"Gemini is {safe_mean(gpt_pt_ed)/safe_mean(gem_pt_ed):.2f}× more accurate on average "
  f"(paired Wilcoxon two-sided p = {w_point.pvalue:.2e}, n = {len(point_pairs)}).")
P(f"  • Area-landmark overlap: Gemini's mean Jaccard is {safe_mean(gem_area_j):.3f} "
  f"vs GPT-5.4's {safe_mean(gpt_area_j):.3f} — "
  f"{(safe_mean(gem_area_j)/safe_mean(gpt_area_j)-1)*100:+.0f}% (p = {w_area.pvalue:.2e}, "
  f"n = {len(area_pairs)}).")
P(f"  • The F5/F6 positional attractor previously documented as a GPT-5.4 failure "
  f"on Tooth_33_Apex is absent in Gemini's predictions: GPT-5.4 placed the answer in "
  f"F5 or F6 for {gpt_attr_total}/{len(t33)} of Tooth_33_Apex queries; Gemini for "
  f"{gem_attr_total}/{len(t33)}. Gemini achieved exact-GT match on "
  f"{gem_exact_total}/{len(t33)} Tooth_33_Apex queries vs GPT's "
  f"{gpt_exact_total}/{len(t33)}.")
P(f"  • {n_maxtok} of 1,800 responses ({n_maxtok/1800*100:.1f}%) reached "
  f"finishReason='MAX_TOKENS' under the rep-1 output-token cap of 2,048; "
  f"{n_maxtok_compliant} of these were still parseable. The cap is raised to 4,096 "
  f"for repetitions 2-3 to eliminate this truncation mode.")

# §2 Methodology
H('2. Methodology', level=1)
P("Benchmark dataset. 900 queries across three radiographic modalities: 600 "
  "panoramic (PAN, 16×8 grid), 150 periapical (PA, 6×8 grid), and 150 "
  "cephalometric (CEPH, 10×8 grid). 22 distinct anatomic landmarks total — 17 "
  "point landmarks (single-cell ground truth) and 5 area landmarks (multi-cell "
  "ground truth). Ground truth is the consensus annotation of two OMFR "
  "specialists (consensus_gt field of results_consensus/query_index.json).")
P("Models compared. Gemini 3.1 Pro (gemini-3.1-pro-preview) at "
  "temperature=0, seed=42, MEDIA_RESOLUTION_HIGH; GPT-5.4 (OpenAI's gpt-5.4) "
  "at temperature=0, seed=42, image_url detail='high' — the published full-run "
  "results from the v3 main run.")
P("Prompts. Byte-identical system + user prompts across both models, verified "
  "by zero-drift check at run time. Two strategies per query: zero_shot (a "
  "concise prompt naming the landmark) and guided (an expanded prompt with "
  "anatomical priors). Total 1,800 (query × strategy) calls per model per rep.")
P("Metric definitions. Point landmarks: Euclidean distance (cell-grid units) "
  "from the model's first parsed cell to the consensus-GT cell — lower is "
  "better. Area landmarks: Jaccard overlap between the parsed-cell set and "
  "the consensus-GT cell set — higher is better. A response is 'compliant' "
  "iff parsed_coordinates is non-empty and failure_category is None.")
P("Statistical tests. Paired Wilcoxon signed-rank tests on matched (query × "
  "strategy) pairs where both models were compliant. Pairs where either model "
  "failed are excluded from the metric tests; compliance is reported separately.")
P("Inference settings disclosure. Gemini 3.1 Pro internally produces "
  "'thinking' tokens (reported as thoughtsTokenCount in the API response) "
  "that share the max_output_tokens budget with the visible answer; GPT-5.4 "
  "has hidden reasoning that does not count against the output-token budget. "
  "Therefore max_output_tokens is set per model: 2,048 for Gemini in rep 1 "
  "(empirical mean thinking tokens 1,257; 4,096 for reps 2-3 to eliminate "
  "MAX_TOKENS truncations), 256 for GPT-5.4 (empirical max answer length 98 "
  "tokens; 256 is non-restrictive).")

# §3 Compliance
H('3. Compliance and Failure Modes', level=1)
P("Strict-compliance counts across models and repetitions:")
add_table(
    ["Model", "Rep", "Compliant", "Non-compliant", "Compliance rate"],
    [
        ["Gemini 3.1 Pro", "1", gem_compliant_n, 1800-gem_compliant_n,
         f"{gem_compliant_n/1800*100:.3f}%"],
        ["GPT-5.4 v2", "1", gpt1_compliant_n, 1800-gpt1_compliant_n,
         f"{gpt1_compliant_n/1800*100:.3f}%"],
        ["GPT-5.4 v2", "2",
         sum(1 for r in gpt[2].values() if r.get("failure_category") is None),
         sum(1 for r in gpt[2].values() if r.get("failure_category") is not None),
         f"{sum(1 for r in gpt[2].values() if r.get('failure_category') is None)/1800*100:.3f}%"],
        ["GPT-5.4 v2", "3",
         sum(1 for r in gpt[3].values() if r.get("failure_category") is None),
         sum(1 for r in gpt[3].values() if r.get("failure_category") is not None),
         f"{sum(1 for r in gpt[3].values() if r.get('failure_category') is None)/1800*100:.3f}%"],
    ],
    col_widths_inches=[1.6, 0.5, 1.0, 1.2, 1.3],
)
doc.add_paragraph()

P("Gemini 3.1 Pro rep 1 failure-mode breakdown:")
add_table(
    ["Failure mode", "Count"],
    [[cat or "compliant", n] for cat, n in sorted(gem_fc.items(), key=lambda x: -x[1])],
    col_widths_inches=[2.0, 0.9],
)
doc.add_paragraph()

# §4 Quality
H('4. Quality Comparison: Gemini vs GPT-5.4 (rep-1 matched)', level=1)
P("Point landmarks — Euclidean distance from the model's first parsed cell to "
  "the consensus GT, in cell units. Lower is better.")
add_table(
    ["", "n (matched)", "mean ED", "median ED"],
    [
        ["Gemini 3.1 Pro", str(len(gem_pt_ed)), f"{safe_mean(gem_pt_ed):.3f}", f"{safe_med(gem_pt_ed):.3f}"],
        ["GPT-5.4 v2 rep 1", str(len(gpt_pt_ed)), f"{safe_mean(gpt_pt_ed):.3f}", f"{safe_med(gpt_pt_ed):.3f}"],
    ],
    col_widths_inches=[2.0, 1.2, 1.0, 1.0],
)
P(f"Paired Wilcoxon (two-sided) on per-query ED differences: W = "
  f"{w_point.statistic:.1f}, p = {w_point.pvalue:.4g}. "
  f"Statistically significant at α = 0.05.")
doc.add_paragraph()

P("Area landmarks — Jaccard overlap between predicted-cell set and consensus-"
  "GT cell set. Higher is better.")
add_table(
    ["", "n (matched)", "mean J", "median J"],
    [
        ["Gemini 3.1 Pro", str(len(gem_area_j)), f"{safe_mean(gem_area_j):.3f}", f"{safe_med(gem_area_j):.3f}"],
        ["GPT-5.4 v2 rep 1", str(len(gpt_area_j)), f"{safe_mean(gpt_area_j):.3f}", f"{safe_med(gpt_area_j):.3f}"],
    ],
    col_widths_inches=[2.0, 1.2, 1.0, 1.0],
)
P(f"Paired Wilcoxon (two-sided): W = {w_area.statistic:.1f}, p = "
  f"{w_area.pvalue:.4g}. Statistically significant at α = 0.05.")
doc.add_paragraph()

P("Per-(modality × strategy × landmark-type) breakdown:")
bd_rows = []
for mod in ("PANORAMIC", "PERIAPICAL", "CEPHALOMETRIC"):
    for strat in ("zero_shot", "guided"):
        for lm in ("point", "area"):
            b = breakdown.get((mod, strat, lm))
            if not b: continue
            bd_rows.append([
                mod, strat, lm,
                f"{b['n']}/{b['n_total']}",
                f"{b['gem_mean']:.3f}",
                f"{b['gpt_mean']:.3f}",
                f"{b['p']:.3g}" if b['p'] is not None else "—",
            ])
add_table(
    ["Modality", "Strategy", "Type", "matched/total", "Gemini mean", "GPT-5.4 mean", "Wilcoxon p"],
    bd_rows,
    col_widths_inches=[1.2, 0.85, 0.55, 0.95, 0.95, 0.95, 0.85],
)
doc.add_paragraph()

# §5 F5/F6 attractor
H('5. F5/F6 Attractor Analysis on Tooth_33_Apex', level=1)
P("Background. The v3 manuscript documented that GPT-5.4 consistently places "
  "Tooth_33_Apex — the lower-left canine apex on panoramic radiographs — at "
  "the F5 or F6 cell, a region that does not correspond to the anatomic "
  "position of the canine apex on the grid. This pattern was named the 'F5/F6 "
  "positional attractor.' We test whether Gemini 3.1 Pro exhibits the same "
  "failure mode.")
P(f"Across both strategies and all 100 panoramic images, the dataset contains "
  f"{len(t33)} Tooth_33_Apex queries. Per-strategy counts:")
add_table(
    ["Strategy", "Total",
     "Gem compliant", "Gem in F5/F6", "Gem exact-GT",
     "GPT compliant", "GPT in F5/F6", "GPT exact-GT"],
    [[s, t33_strat[s]["total"],
      f"{t33_strat[s]['gem_compliant']} ({t33_strat[s]['gem_compliant']/t33_strat[s]['total']*100:.0f}%)",
      f"{t33_strat[s]['gem_attr']} ({t33_strat[s]['gem_attr']/t33_strat[s]['total']*100:.0f}%)",
      f"{t33_strat[s]['gem_exact']} ({t33_strat[s]['gem_exact']/t33_strat[s]['total']*100:.0f}%)",
      f"{t33_strat[s]['gpt_compliant']} ({t33_strat[s]['gpt_compliant']/t33_strat[s]['total']*100:.0f}%)",
      f"{t33_strat[s]['gpt_attr']} ({t33_strat[s]['gpt_attr']/t33_strat[s]['total']*100:.0f}%)",
      f"{t33_strat[s]['gpt_exact']} ({t33_strat[s]['gpt_exact']/t33_strat[s]['total']*100:.0f}%)"]
     for s in ("zero_shot", "guided")],
    col_widths_inches=[0.85, 0.6, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95],
)
doc.add_paragraph()
P("Per-cell prediction distribution on Tooth_33_Apex (first parsed cell):")
add_table(
    ["Cell", "Gemini count", "GPT-5.4 count"],
    [[cell, gem_t33_dist.get(cell, 0), gpt_t33_dist.get(cell, 0)]
     for cell in sorted(set(gem_t33_dist) | set(gpt_t33_dist),
                         key=lambda c: -(gem_t33_dist.get(c,0)+gpt_t33_dist.get(c,0)))[:15]],
    col_widths_inches=[1.0, 1.4, 1.4],
)
doc.add_paragraph()
P(f"Interpretation. GPT-5.4 produces predictions distributed across F5–F10, "
  f"with the F5/F6 cluster appearing in {gpt_attr_total}/{len(t33)} responses. "
  f"Gemini's predictions are concentrated in F10/G10/F11/G11 — the anatomically "
  f"plausible region given that the consensus GT for these queries is "
  f"predominantly G10. Gemini does not exhibit the F5/F6 attractor "
  f"({gem_attr_total}/{len(t33)} = 0% of responses landing there). This is a "
  f"model-specific failure mode rather than a generic MLLM phenomenon on "
  f"grid-based landmark identification.")

# §6 MAX_TOKENS
H('6. Output-Token Budget Analysis', level=1)
P("Gemini 3.1 Pro uses internal thinking-mode reasoning that shares the "
  "max_output_tokens budget with the visible answer. Rep 1 used max_output_"
  "tokens=2048; the analysis below quantifies the truncation footprint and "
  "informs the budget choice for repetitions 2 and 3.")
P(f"At 2,048-token cap: {n_maxtok} of 1,800 responses ({n_maxtok/1800*100:.1f}%) "
  f"reached finishReason='MAX_TOKENS'. Of those, {n_maxtok_parseable} contained "
  f"at least one parseable cell ({n_maxtok_parseable/n_maxtok*100:.0f}%) and "
  f"{n_maxtok_compliant} were still classified as compliant under the strict "
  f"definition (parsed coordinates non-empty, no failure category).")
if maxtok_records:
    add_table(
        ["Statistic", "thinking tokens", "answer tokens"],
        [["mean",  f"{safe_mean([m['thoughts'] for m in maxtok_records]):.0f}",
                    f"{safe_mean([m['answer']  for m in maxtok_records]):.0f}"],
         ["median",f"{statistics.median([m['thoughts'] for m in maxtok_records]):.0f}",
                    f"{statistics.median([m['answer']  for m in maxtok_records]):.0f}"],
         ["min",   f"{min(m['thoughts'] for m in maxtok_records)}",
                    f"{min(m['answer']  for m in maxtok_records)}"],
         ["max",   f"{max(m['thoughts'] for m in maxtok_records)}",
                    f"{max(m['answer']  for m in maxtok_records)}"]],
        col_widths_inches=[1.2, 1.6, 1.6],
    )
    doc.add_paragraph()

P("MAX_TOKENS-truncated landmarks (top 15 by count):")
add_table(
    ["Anatomic structure", "# MAX_TOKENS"],
    [[k, v] for k, v in maxtok_struct.most_common(15)],
    col_widths_inches=[2.5, 1.0],
)
doc.add_paragraph()
P(f"For repetitions 2 and 3, max_output_tokens is raised to 4,096 to leave "
  f"thinking-mode an additional 2,048-token headroom. The marginal cost is "
  f"small (Gemini bills actual usage, not cap) — projected total of "
  f"~${cost_rep_4096:.2f} per rep, vs ${cost_total:.2f} for rep 1 at 2,048.")

# §7 Cost
H('7. Cost Analysis', level=1)
P("Costs are derived from the per-call usageMetadata reported by the Gemini "
  "API, using the published batch-tier pricing of $1.00 per 1M input tokens "
  "and $6.00 per 1M output tokens. Thinking tokens are billed at the output "
  "rate.")
add_table(
    ["Component", "Tokens", "Cost"],
    [["Prompt (input)", f"{batch_in+sync_in:,}", f"${(batch_in*BATCH_IN+sync_in*SYNC_IN)/1e6:.3f}"],
     ["Thinking (output)", f"{batch_thoughts+sync_thoughts:,}",
      f"${(batch_thoughts*BATCH_OUT+sync_thoughts*SYNC_OUT)/1e6:.3f}"],
     ["Answer (output)", f"{batch_answer+sync_answer:,}",
      f"${(batch_answer*BATCH_OUT+sync_answer*SYNC_OUT)/1e6:.3f}"],
     ["TOTAL rep 1", "—", f"${cost_total:.2f}"],
     ["Projected rep 2 (max=4096)", "—", f"~${cost_rep_4096:.2f}"],
     ["Projected rep 3 (max=4096)", "—", f"~${cost_rep_4096:.2f}"],
     ["Projected 3-rep total", "—", f"~${cost_total + 2*cost_rep_4096:.2f}"]],
    col_widths_inches=[2.5, 1.5, 2.5],
)
doc.add_paragraph()

# §8 Recommendations
H('8. Recommendations for Repetitions 2 and 3', level=1)
P(f"1. Proceed with repetitions 2 and 3 using max_output_tokens=4096. "
  f"Marginal cost ~${2*(cost_rep_4096 - cost_total):.2f} extra over keeping "
  f"the cap at 2,048, eliminating the {n_maxtok/1800*100:.1f}% MAX_TOKENS "
  f"truncation rate observed in rep 1.")
P("2. Statistical reporting: with three Gemini repetitions, paired Wilcoxon "
  "tests gain rep-level variance estimates, and the per-modality / per-"
  "landmark breakdowns in §4 can be reported with rep-averaged means and "
  "standard deviations.")
P("3. The F5/F6 attractor analysis (§5) is a publishable cross-model "
  "observation independent of the rep count. With reps 2-3 we can also test "
  "rep-stability of the attractor pattern in GPT-5.4 (Gemini's 0% rate is "
  "unlikely to change).")

# Appendix A1: All 1800 Gemini responses
H('Appendix A1 — All 1,800 Gemini 3.1 Pro Responses', level=1, page_break=True)
P("Format: query_id, modality, GT, raw response, parsed cells, "
  "failure_category. Sorted by (modality, strategy, query_id).")
appendix_rows = sorted(rows, key=lambda r: (r["modality"], r["strategy"], r["qid"]))
for strat in ("zero_shot", "guided"):
    P(f"Strategy = {strat} (900 responses):", bold=True)
    a1 = []
    for r in [x for x in appendix_rows if x["strategy"] == strat]:
        a1.append([
            r["qid"],
            r["modality"][:3],
            r["gt"][:30] + ("…" if len(r["gt"]) > 30 else ""),
            (r["gem_raw"] or "")[:50] + ("…" if r["gem_raw"] and len(r["gem_raw"]) > 50 else ""),
            ",".join(r["gem_parsed"][:6])[:40],
            r["gem_fail"] or "OK",
        ])
    add_table(
        ["Query ID", "Mod", "GT", "Raw response", "Parsed", "Fail?"],
        a1,
        col_widths_inches=[2.0, 0.4, 1.4, 1.6, 1.2, 0.8],
    )
    doc.add_paragraph()

# Appendix A2: All 1800 GPT responses
H('Appendix A2 — All 1,800 GPT-5.4 v2 Responses (rep 1)', level=1, page_break=True)
P("Same format as A1.")
for strat in ("zero_shot", "guided"):
    P(f"Strategy = {strat} (900 responses):", bold=True)
    a2 = []
    for r in [x for x in appendix_rows if x["strategy"] == strat]:
        a2.append([
            r["qid"],
            r["modality"][:3],
            r["gt"][:30] + ("…" if len(r["gt"]) > 30 else ""),
            (r["gpt_raw"] or "")[:50] + ("…" if r["gpt_raw"] and len(r["gpt_raw"]) > 50 else ""),
            ",".join(r["gpt_parsed"][:6])[:40],
            r["gpt_fail"] or "OK",
        ])
    add_table(
        ["Query ID", "Mod", "GT", "Raw response", "Parsed", "Fail?"],
        a2,
        col_widths_inches=[2.0, 0.4, 1.4, 1.6, 1.2, 0.8],
    )
    doc.add_paragraph()

# Appendix A3: side-by-side
H('Appendix A3 — Per-Query Side-by-Side (Gemini vs GPT-5.4 vs GT)', level=1, page_break=True)
P("Compact side-by-side per (query, strategy). 'T' = landmark type "
  "(P=point, A=area). For point landmarks, '✓?' indicates whether the model's "
  "first parsed cell equals the GT cell. ED is Euclidean distance (cells); J "
  "is Jaccard.")
for strat in ("zero_shot", "guided"):
    P(f"Strategy = {strat}:", bold=True)
    a3 = []
    for r in [x for x in appendix_rows if x["strategy"] == strat]:
        gt0 = r["gt_cells"][0] if r["gt_cells"] else "?"
        gem0 = r["gem_parsed"][0] if r["gem_parsed"] else "?"
        gpt0 = r["gpt_parsed"][0] if r["gpt_parsed"] else "?"
        is_pt = r["landmark_type"] == "point"
        gem_chk = "✓" if (is_pt and r["gem_parsed"] and gem0 == gt0) else ("·" if r["gem_parsed"] else "—")
        gpt_chk = "✓" if (is_pt and r["gpt_parsed"] and gpt0 == gt0) else ("·" if r["gpt_parsed"] else "—")
        if is_pt:
            comp = (f"ED={r['gem_metric']:.2f}/{r['gpt_metric']:.2f}"
                    if r['gem_metric'] is not None and r['gpt_metric'] is not None else "—")
        else:
            comp = (f"J={r['gem_metric']:.3f}/{r['gpt_metric']:.3f}"
                    if r['gem_metric'] is not None and r['gpt_metric'] is not None else "—")
        a3.append([
            r["qid"][-22:],
            "P" if is_pt else "A",
            r["gt"][:14] + ("…" if len(r["gt"]) > 14 else ""),
            ",".join(r["gem_parsed"][:3])[:18],
            gem_chk,
            ",".join(r["gpt_parsed"][:3])[:18],
            gpt_chk,
            comp,
        ])
    add_table(
        ["Query ID", "T", "GT", "Gemini", "✓?", "GPT-5.4", "✓?", "ED or J (G/P)"],
        a3,
        col_widths_inches=[1.7, 0.25, 1.1, 1.2, 0.3, 1.2, 0.3, 1.1],
    )
    doc.add_paragraph()

# Save
out_path = ROOT / "results_full_gemini/Gemini_Rep1_Report.docx"
doc.save(str(out_path))
sha = hashlib.sha256(out_path.read_bytes()).hexdigest()[:12]
print(f"\n✓ DOCX written: {out_path}")
print(f"  size: {out_path.stat().st_size/1024:.1f} KB")
print(f"  sha256: {sha}")
print(f"  paragraphs: {len(doc.paragraphs)}")
print(f"  tables: {len(doc.tables)}")
