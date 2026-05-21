"""
Generate the v2 docx listing every GPT-5.4 response with consensus_gt as the GT column.

Same six-column layout as the v1 appendix
(Image | Landmark | GT | Run 1 | Run 2 | Run 3) so it can be diffed against the
v1 appendix to spot the GT changes between OMFR_1 and consensus. Uses fixed
table layout (matching v1) so long no-space cell lists wrap inside the column.

Reads:    results_consensus/full_run_records.pkl
Writes:   results_consensus/Full_Run_Model_Responses_Appendix_v2.docx
"""
from __future__ import annotations
import pickle
from pathlib import Path
from datetime import datetime, UTC
import subprocess

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SANDBOX = ROOT / 'results_consensus'
OUT  = SANDBOX / 'Full_Run_Model_Responses_Appendix_v2.docx'

records = pickle.load(open(SANDBOX / 'full_run_records.pkl', 'rb'))
git_sha = subprocess.check_output(['git','-C',str(ROOT),'rev-parse','HEAD']).decode().strip()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(11)
sec = doc.sections[0]
sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.85)

def H(text, level=1): return doc.add_heading(text, level=level)
def P(text, *, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size=Pt(size)
    if italic: r.italic=True
    return p
def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic=True; r.font.size=Pt(10)
    return p
def shade_cell(cell, fill='D9E1F2'):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), fill); tcPr.append(sh)
def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'808080')
        tcBorders.append(b)
    tcPr.append(tcBorders)
def set_table_fixed_layout(t):
    tblPr = t._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); t._element.insert(0, tblPr)
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
def add_response_table(headers, rows, col_widths, *, header_size=12, body_size=11):
    n_cols = len(headers)
    t = doc.add_table(rows=len(rows)+1, cols=n_cols)
    t.style = 'Light Grid'
    set_table_fixed_layout(t)
    for j,h in enumerate(headers):
        c = t.cell(0,j); c.text=''
        run = c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(header_size)
        shade_cell(c,'D9E1F2'); set_cell_borders(c)
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.cell(i+1,j); c.text=''
            run = c.paragraphs[0].add_run(str(val)); run.font.size=Pt(body_size)
            set_cell_borders(c)
    for j,w in enumerate(col_widths):
        for r in t.rows:
            r.cells[j].width = Inches(w)
    return t
def fmt_response(raw):
    if raw is None or not str(raw).strip(): return '<empty>'
    return str(raw).strip()

# Title
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('GPT-5.4 Model Responses — Full Benchmark (v2 Consensus)')
r.bold=True; r.font.size=Pt(18)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Companion data appendix to the v2 Consensus Results Report')
r.italic=True; r.font.size=Pt(13)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f'Generated {datetime.now(UTC).strftime("%Y-%m-%d")} from raw API outputs '
                 f'(commit {git_sha[:12]}, sandbox: results_consensus)')
r.font.size=Pt(10); r.italic=True
doc.add_paragraph()

# Overview
H('Overview', level=1)
n_records = len(records)
P(f'This appendix lists every GPT-5.4 response for all {n_records:,} unique (query, strategy) '
  f'pairs across the full benchmark, including all three repetitions. The "GT" column shows '
  f'the consensus ground truth (CONSENSUS_Ground_Truth from the Final benchmark Excel). '
  f'For point landmarks where consensus differs from v1\'s OMFR_1 (PAN_074 Mental_Foramen_L: '
  f'OMFR_1=G11 → consensus=G10; PAN_079 Condylar_Head_R: OMFR_1=B2 → consensus=B1), the v2 '
  f'GT column reflects the consensus value.')

groups = [
    ('PANORAMIC',     'zero_shot', 'point', 'B1'),
    ('PANORAMIC',     'guided',    'point', 'B2'),
    ('PANORAMIC',     'zero_shot', 'area',  'B3'),
    ('PANORAMIC',     'guided',    'area',  'B4'),
    ('PERIAPICAL',    'zero_shot', 'point', 'B5'),
    ('PERIAPICAL',    'guided',    'point', 'B6'),
    ('CEPHALOMETRIC', 'zero_shot', 'point', 'B7'),
    ('CEPHALOMETRIC', 'guided',    'point', 'B8'),
]
for mod, strat, ltype, label in groups:
    n = sum(1 for r in records if r['modality']==mod and r['strategy']==strat
            and r['landmark_type']==ltype)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'Table {label}: {mod}, {strat}, {ltype} — {n} queries × 3 reps = {n*3} responses.')

P('Compliance footnote: the four strict-parser failures appear verbatim — "12F", "12E", "6F", '
  '"4E" — corresponding to reversed-coordinate strings that encode the correct cell in '
  'inverted form. See the v2 report\'s Appendix B for the full failure list.', italic=True)

PT_WIDTHS = [0.95, 1.65, 1.30, 0.93, 0.93, 0.93]   # = 6.69" — same as v1
landmark_order = {
    'Mental_Foramen_L':0,'Condylar_Head_R':1,'Tooth_33_Apex':2,
    'Mandibular_Canal_L':3,'Maxillary_Sinus_R':4,'External_Oblique_Ridge_R':5,
    'Tooth_36_Distal_Apex':0,'Tooth_36_Distal_CEJ':1,'Tooth_36_Mesial_CEJ':2,
    'Sella_S':0,'Nasion_N':1,'Menton_Me':2,
}
for mod, strat, ltype, label in groups:
    items = [r for r in records
             if r['modality']==mod and r['strategy']==strat and r['landmark_type']==ltype]
    items.sort(key=lambda r: (r['image_id'], landmark_order.get(r['structure'], 99)))

    doc.add_page_break()
    H(f'Table {label} — {mod}, {strat}, {ltype} landmarks (vs consensus_gt)', level=2)
    caption(f'Table {label}: All GPT-5.4 raw responses for the {mod} {strat} {ltype} group '
            f'({len(items)} queries × 3 reps = {len(items)*3} responses). GT column = '
            f'consensus_gt.')
    rows = []
    for r in items:
        runs = [fmt_response(x) for x in r['rep_raw']]
        while len(runs) < 3:
            runs.append('<empty>')
        rows.append([r['image_id'], r['structure'], r['consensus_gt'] or '',
                     runs[0], runs[1], runs[2]])
    body_size = 10 if ltype == 'area' else 11
    add_response_table(['Image','Landmark','GT','Run 1','Run 2','Run 3'],
                       rows, PT_WIDTHS, header_size=12, body_size=body_size)

doc.save(str(OUT))
import os
print(f'Wrote {OUT}')
print(f'  Tables: {len(doc.tables)}')
print(f'  Total data rows: {sum(len(t.rows)-1 for t in doc.tables)}')
print(f'  Size: {os.path.getsize(OUT):,} bytes')
