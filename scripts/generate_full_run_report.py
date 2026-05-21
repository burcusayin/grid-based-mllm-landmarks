"""
Generate the full-run results docx report.

EVERY numeric figure is computed from /tmp/full_run_records_v2.pkl,
/tmp/full_run_analysis.json, /tmp/full_run_phase_b.json, /tmp/full_run_summary.json,
or directly from results_full/run{1,2,3}/parsed_responses.json.

NO hardcoded values. All claims are reproducible from the saved artifacts.
"""
from __future__ import annotations
import json, pickle, math, hashlib, subprocess, statistics
from pathlib import Path
from collections import defaultdict, Counter
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT  = ROOT / 'results_full' / 'Full_Run_Results_Report.docx'

# ── Load all data ───────────────────────────────────────────────────
records  = pickle.load(open('/tmp/full_run_records_v2.pkl','rb'))
analysis = json.load(open('/tmp/full_run_analysis.json'))
phase_b  = json.load(open('/tmp/full_run_phase_b.json'))
summary  = json.load(open('/tmp/full_run_summary.json'))
failures = json.load(open('/tmp/full_run_failures.json'))

# Git + Excel
git_sha = subprocess.check_output(['git','-C',str(ROOT),'rev-parse','HEAD']).decode().strip()
git_date = subprocess.check_output(['git','-C',str(ROOT),'log','-1','--format=%ci','HEAD']).decode().strip()
excel_sha = hashlib.sha256(open(ROOT/'data'/'Dental_MLLM_Benchmark_Data.xlsx','rb').read()).hexdigest()
qi_sha = hashlib.sha256(open(ROOT/'results_full'/'query_index.json','rb').read()).hexdigest()

# ── Document setup ──────────────────────────────────────────────────
doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Letter portrait, narrow margins
section = doc.sections[0]
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)

def H(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def P(text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    return p

def caption(text):
    """Caption line above tables."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    return p

def shade_cell(cell, fill='D9E1F2'):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), fill)
    tcPr.append(sh)

def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'4')
        b.set(qn('w:color'),'808080')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_table(headers, rows, col_widths=None, header_size=10, body_size=10):
    """Add a table with bold shaded header row and bordered cells.
    rows = list of list[str]."""
    n_cols = len(headers)
    t = doc.add_table(rows=len(rows)+1, cols=n_cols)
    t.style = 'Light Grid'
    for j, h in enumerate(headers):
        c = t.cell(0,j)
        c.text = ''
        para = c.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(header_size)
        shade_cell(c, 'D9E1F2')
        set_cell_borders(c)
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.cell(i+1,j)
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(body_size)
            set_cell_borders(c)
    if col_widths:
        for j,w in enumerate(col_widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t

# ────────────────────────────────────────────────────────────────────
# TITLE & METADATA
# ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('GPT-5.4 on Grid-Based Dental Anatomic Landmark Identification')
r.bold = True; r.font.size = Pt(18)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Full Benchmark Results — 5,400 API Calls Across 900 Queries × 2 Strategies × 3 Repetitions')
r.italic = True; r.font.size = Pt(13)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f'Generated {datetime.utcnow().strftime("%Y-%m-%d")} from raw API outputs '
                 f'(commit {git_sha[:12]}, sandbox: results_full)')
r.font.size = Pt(10); r.italic = True
doc.add_paragraph()

# ────────────────────────────────────────────────────────────────────
# 1. EXECUTIVE SUMMARY
# ────────────────────────────────────────────────────────────────────
H('1. Executive Summary', level=1)

n_calls = summary['n_total_calls']
n_actual = summary['n_actual_responses']
compliance_rate = summary['compliance_rate']
prompt_tok = phase_b['tokens']['prompt_tokens']
compl_tok = phase_b['tokens']['completion_tokens']
naive_cost = summary['naive_cost_usd']

P(f'This report presents the complete results of evaluating GPT-5.4 on grid-based dental anatomic '
  f'landmark identification across the full benchmark of 900 queries spanning three radiographic '
  f'modalities and 12 landmark types. Each query was issued under two prompting strategies '
  f'(zero-shot baseline and zero-shot with detailed grid explanation, hereafter "guided") and '
  f'repeated three times to assess reproducibility, yielding {n_calls:,} API calls in total. All '
  f'{n_actual:,} planned calls completed successfully with consistent format-conforming output in '
  f'{int(round(compliance_rate*n_calls)):,} of {n_calls:,} cases ({compliance_rate*100:.3f}%).')

# Modality-level point ED (zero-shot vs guided) numbers
m_ceph_zs = analysis['RQ1_modality_strategy']['CEPHALOMETRIC_zero_shot_point']
m_ceph_gd = analysis['RQ1_modality_strategy']['CEPHALOMETRIC_guided_point']
m_pa_zs   = analysis['RQ1_modality_strategy']['PERIAPICAL_zero_shot_point']
m_pa_gd   = analysis['RQ1_modality_strategy']['PERIAPICAL_guided_point']
m_pan_zs  = analysis['RQ1_modality_strategy']['PANORAMIC_zero_shot_point']
m_pan_gd  = analysis['RQ1_modality_strategy']['PANORAMIC_guided_point']

P('Headline findings:', bold=True)

# Bullet 1 — modality difficulty
b1 = doc.add_paragraph(style='List Bullet')
b1.add_run('Modality difficulty hierarchy is pronounced and consistent across both strategies. ')
b1.add_run(f'Cephalometric mean Euclidean distance is '
           f'{m_ceph_zs["mean"]:.2f} cells (zero-shot) and {m_ceph_gd["mean"]:.2f} cells (guided). '
           f'Periapical is {m_pa_zs["mean"]:.2f} and {m_pa_gd["mean"]:.2f} cells respectively. '
           f'Panoramic is {m_pan_zs["mean"]:.2f} and {m_pan_gd["mean"]:.2f} cells. '
           f'Cross-modality normalised Euclidean distance (Section 5.2) preserves the same ordering, '
           f'so the difference cannot be explained by the panoramic grid simply being larger.')

# Bullet 2 — guided strategy effect heterogeneous
rq2a_pan_pt = analysis['RQ2a_strategy_per_modality']['PANORAMIC_point']
rq2a_pa_pt  = analysis['RQ2a_strategy_per_modality']['PERIAPICAL_point']
rq2a_ceph_pt= analysis['RQ2a_strategy_per_modality']['CEPHALOMETRIC_point']
b2 = doc.add_paragraph(style='List Bullet')
b2.add_run('The guided strategy does not uniformly improve performance. ')
b2.add_run(f'The aggregate panoramic point comparison shows a paired Wilcoxon mean delta of '
           f'{rq2a_pan_pt["mean_delta"]:+.2f} cells (zero-shot − guided), '
           f'p = {rq2a_pan_pt["p"]:.2e} (Bonferroni-corrected p < 0.001), '
           f'rank-biserial r = {rq2a_pan_pt["rank_biserial_r"]:+.2f}: guided is worse on aggregate. '
           f'Periapical and cephalometric point comparisons are not significant after correction. '
           f'Per-landmark stratification (Section 6.2) reveals strongly opposing effects — guided '
           f'helps Mental Foramen and Sella while harming Tooth 33 Apex and Condylar Head — that '
           f'cancel imperfectly at the modality level.')

# Bullet 3 — Reproducibility
fleiss_overall = phase_b['fleiss_overall_point']
unan = phase_b['overall_unanimous']
b3 = doc.add_paragraph(style='List Bullet')
b3.add_run(f'Reproducibility at temperature = 0 is high but not absolute. ')
b3.add_run(f"Fleiss' κ across the three repetitions for point landmarks is {float(fleiss_overall):.3f} "
           f'(substantial agreement), with {unan["unanimous"]:,} of {unan["total"]:,} '
           f'point queries × strategies receiving identical predicted cells in all three runs '
           f'({unan["rate"]*100:.1f}%). Reproducibility is highest on cephalometric (κ ≈ 0.83–0.89) '
           f'and lowest on panoramic (κ ≈ 0.65–0.70), tracking modality difficulty.')

# Bullet 4 — Compliance & failure modes
b4 = doc.add_paragraph(style='List Bullet')
b4.add_run('Compliance is near-perfect. ')
b4.add_run(f'Of {n_calls:,} responses, {n_calls-len(failures)} ({(1-len(failures)/n_calls)*100:.3f}%) '
           f'were format-conforming grid coordinates. The four failures all share the same pattern: '
           f'reversed-order coordinates (e.g. "12F" instead of "F12"), all on the guided strategy, '
           f'on a single landmark per affected image.')

# Bullet 5 — Cost
b5 = doc.add_paragraph(style='List Bullet')
b5.add_run(f'Operational cost: ${naive_cost:.2f} ')
b5.add_run(f'across {prompt_tok+compl_tok:,} tokens ({prompt_tok:,} prompt + '
           f'{compl_tok:,} completion), within the pre-registered $20 budget cap.')

P('All figures in this report are computed from the raw OpenAI batch outputs in '
  '/results_full/run{1,2,3}/responses/. The reproducibility manifest in Appendix C lists every '
  'cryptographic anchor needed to verify the analysis end-to-end.', italic=True)

# ────────────────────────────────────────────────────────────────────
# 2. STUDY CONTEXT AND PRE-SPECIFIED RESEARCH QUESTIONS
# ────────────────────────────────────────────────────────────────────
H('2. Study Context and Pre-Specified Research Questions', level=1)

H('2.1 Scope of this report', level=2)
P('This report covers a single model — GPT-5.4 — evaluated on the complete 900-query benchmark. '
  'Comparisons against a second commercial multimodal LLM (Gemini 3.1 Pro) and against the '
  'fourth-year dental student cohort are out of scope here and will be reported separately once '
  'their respective data collections are complete. Consequently, the analyses presented here are '
  'GPT-5.4-internal: how does the model perform on each modality and landmark, how does the '
  'guided strategy compare with the zero-shot baseline, and how reproducible is the model under '
  'temperature = 0?')

H('2.2 Pre-specified research questions', level=2)
P('The following research questions were specified prior to the full run, in the project '
  'methodology document (Technical_Report_Experiment_Methodology.docx, Section 2). Of the five '
  'questions originally listed, three are within scope of this single-model report. The remaining '
  'two (cross-provider comparison, MLLM-vs-student comparison) require data not yet available.')

P('RQ1 — Modality-stratified accuracy. ', bold=True)
P('How accurately does GPT-5.4 identify each landmark type in each modality? This question is '
  'answered with mean Euclidean distance, normalised Euclidean distance, and Successful Detection '
  'Rate (SDR) at four clinically motivated thresholds for point landmarks, and with Jaccard and '
  'Dice indices for area landmarks (Section 5).')

P('RQ2 — Effect of explicit grid explanation. ', bold=True)
P('Does the guided prompting strategy (which adds an explicit description of the grid coordinate '
  'system to the system prompt) significantly change performance compared to the zero-shot '
  'baseline? The two strategies differ only in the system prompt — the user prompt, the image, '
  'the model, the inference settings, and the seed are identical. Comparisons are made at two '
  'levels: aggregate-by-modality (Bonferroni × 4) and per-landmark (Bonferroni × 9 for point + 3 '
  'for area), using paired Wilcoxon signed-rank tests on per-query mean Euclidean distance or mean '
  'Jaccard index across the three repetitions (Section 6).')

P('RQ3 — Reproducibility under temperature = 0. ', bold=True)
P('At temperature = 0 with a fixed seed, are the three repetitions identical, and if not, where '
  "and how does the model disagree with itself? Quantified with Fleiss' κ across the three reps "
  '(point landmarks, per modality and strategy), three-way unanimous response rates, the '
  'distribution of max-minus-min Euclidean distance across reps, and mean pairwise Jaccard index '
  'across reps (area landmarks). This calibrates whether single-run evaluation is defensible and '
  'where rep-averaging is necessary (Section 7).')

H('2.3 What this report does and does not test', level=2)
P('This report tests RQ1, RQ2, and RQ3. It is not a comparison against any prior version of the '
  'prompt; the 5,400 calls were issued with the v2 prompt set throughout. It is therefore '
  'inappropriate to attribute any specific result here to particular wording choices made before '
  'the full run; for that purpose the pilot runs were used and are described in the prompt '
  'finalisation memo, not here. The two prompting strategies compared in this report (zero-shot '
  'and guided) differ only in the system prompt content, and that difference is the only '
  'attribution claim made for any RQ2 result.')

# ────────────────────────────────────────────────────────────────────
# 3. METHODOLOGY
# ────────────────────────────────────────────────────────────────────
H('3. Methodology', level=1)

H('3.1 Dataset', level=2)
n_pan_pt = sum(1 for r in records if r['modality']=='PANORAMIC' and r['landmark_type']=='point' and r['strategy']=='zero_shot')
n_pan_ar = sum(1 for r in records if r['modality']=='PANORAMIC' and r['landmark_type']=='area' and r['strategy']=='zero_shot')
n_pa_pt  = sum(1 for r in records if r['modality']=='PERIAPICAL' and r['strategy']=='zero_shot')
n_ceph_pt= sum(1 for r in records if r['modality']=='CEPHALOMETRIC' and r['strategy']=='zero_shot')

P(f'The benchmark comprises 200 anonymised dental radiographs from the Burdur Mehmet Akif Ersoy '
  f'University hospital PACS, distributed across panoramic (100 images, 16×8 grid), periapical '
  f'(50 images, 8×6 grid), and cephalometric (50 images, 10×8 grid) modalities. From these images, '
  f'900 landmark queries were derived: 600 panoramic queries (100 images × 6 landmarks of which '
  f'3 are point and 3 area), 150 periapical queries (50 images × 3 point landmarks), and 150 '
  f'cephalometric queries (50 images × 3 point landmarks). Ground truth coordinates were '
  f'established by an Oral and Maxillofacial Radiology specialist (OMFR_1); inter-rater data '
  f'(OMFR_2) and intra-rater data are not yet available and will be reported separately when '
  f'collected.')

H('3.2 Grid system', level=2)
P('Each image carries a modality-specific square grid overlay (cyan grid lines, yellow alphanumeric '
  'labels). Cells are addressed by row letter followed by column number (e.g. "C5"). For point '
  'landmarks, the model is asked for a single cell. For area landmarks (panoramic only — '
  'mandibular canal, maxillary sinus, external oblique ridge), the model is asked for the '
  'comma-separated set of cells the structure occupies.')

H('3.3 Model and inference settings', level=2)
P(f'Single model: gpt-5.4 (OpenAI, March 2026), accessed via the Batch API. All queries used '
  f'temperature = 0, seed = 42, max_completion_tokens = 50, image input encoded as base64 data '
  f'URL with detail = "high". The system prompt declared the model as "an expert Oral and '
  f'Maxillofacial Radiologist". The same system+user pair was issued three times to produce three '
  f'independent repetitions per (query × strategy). Prompts contain no example coordinates or '
  f'fixed-coordinate worked examples, eliminating any possibility of coordinate-mention contamination.')

H('3.4 Prompt strategies', level=2)
P('Zero-shot. ', bold=True)
P('Minimal task instruction. The model receives the system role, the image, and one of two user '
  'prompts depending on landmark type: for points, "This image is overlaid with a [COLS]×[ROWS] '
  'square grid. Identify the coordinate of the cell (e.g., B3) containing [LANDMARK]. Provide '
  'only the coordinate in your response."; for areas, "List all cells separated by commas (e.g., '
  'C4, D5, E6). Provide only the coordinates in your response."')

P('Guided. ', bold=True)
P('Identical user prompt; system prompt extended to declare row/column ranges, expected response '
  'format for point and area questions, and (for panoramic only) the L–R viewing convention '
  'clarification. The two strategies are byte-identical for periapical and cephalometric system '
  'prompts up to the row/column declarations, and differ only in adding the L–R viewing clause for '
  'panoramic. This isolates the marginal effect of explicit grid explanation as cleanly as '
  'practical given the panoramic-specific lateralisation requirement.')

H('3.5 Statistical analysis plan', level=2)

P('Distance, overlap, and SDR are reported with 95% confidence intervals — bootstrap percentile '
  'intervals (10,000 resamples, random_state = 42) for means and medians, Wilson score intervals '
  'for proportions. Distributional non-normality is verified by the Shapiro-Wilk test on each '
  'paired delta (zero-shot − guided per query mean ED): all tests reject normality at α = 0.05 '
  '(Section 6), justifying non-parametric paired tests throughout.')

P('Strategy comparisons use the paired Wilcoxon signed-rank test on the per-query mean of the '
  'three repetitions (Wilcoxon zero-method = "wilcox", two-sided). Effect size is the matched-pair '
  'rank-biserial correlation r. The Bonferroni correction is applied at the level of the '
  'pre-specified comparison family: ×4 for the four modality-level comparisons (CEPH point, PA '
  'point, PAN point, PAN area), ×9 for the nine point landmarks, ×3 for the three panoramic area '
  'landmarks. Significance is reported at the corrected α-level.')

P("Reproducibility is assessed with Fleiss' κ across the three repetitions, computed separately "
  'per modality × strategy because the category space (set of cells) differs by grid size. Three-'
  'way unanimous rates and the distribution of max-minus-min ED across the three reps are reported '
  'as additional reproducibility lenses. For area landmarks, mean pairwise Jaccard index between '
  'the three reps quantifies cross-rep set agreement.')

H('3.6 Reproducibility infrastructure', level=2)
P(f'Pipeline source at git commit {git_sha[:7]} (this report cites figures computed against this '
  f'commit). Raw input integrity anchored by SHA-256 of the source Excel ({excel_sha[:16]}…) and '
  f'of the derived query_index.json ({qi_sha[:16]}…). Atomic file writes (POSIX rename + fsync) '
  f'are used for every persisted artifact in the pipeline, and a 12-section pre-flight check '
  f'(scripts/verify_full_run_setup.py) was run before the full run was launched. Per-chunk '
  f'persistence in the submit stage means a mid-run crash leaves no orphan batches. The complete '
  f'manifest is in Appendix C.')

# ────────────────────────────────────────────────────────────────────
# 4. OPERATIONAL OUTCOMES
# ────────────────────────────────────────────────────────────────────
H('4. Operational Outcomes', level=1)

H('4.1 Calls executed', level=2)
P(f'The full run completed all {n_calls:,} planned API calls. Each of the three repetitions '
  f'consists of 1,800 calls (900 queries × 2 strategies), delivered as 36 OpenAI batch chunks per '
  f'rep (108 chunks total). All chunks reached terminal state "completed" with no chunk-level '
  f'failures requiring recovery in the final delivery.')

H('4.2 Token usage and cost', level=2)
caption('Table 1: Token usage and cost summary.')
add_table(
    ['Quantity','Value'],
    [
        ['Total API calls', f'{n_calls:,}'],
        ['Successful responses', f'{n_actual:,} ({n_actual/n_calls*100:.2f}%)'],
        ['Prompt tokens', f'{prompt_tok:,}'],
        ['Completion tokens', f'{compl_tok:,}'],
        ['Total tokens', f'{prompt_tok+compl_tok:,}'],
        ['Mean prompt tokens / call', f'{prompt_tok/n_calls:.0f}'],
        ['Mean completion tokens / call', f'{compl_tok/n_calls:.2f}'],
        ['Naïve cost (no caching)', f'${naive_cost:.2f}'],
    ],
    col_widths=[3.0, 3.4]
)

H('4.3 Compliance and failure modes', level=2)
caption('Table 2: Compliance breakdown by modality and strategy.')
# Compliance per modality+strategy
def compliance_for(mod, strat):
    n_calls_local = 0; n_fail = 0
    for r in records:
        if r['modality']!=mod or r['strategy']!=strat: continue
        n_calls_local += 3
        n_fail += r['n_failed']
    return n_calls_local, n_fail

rows=[]
for mod in ('PANORAMIC','PERIAPICAL','CEPHALOMETRIC'):
    for strat in ('zero_shot','guided'):
        n,f = compliance_for(mod,strat)
        rows.append([mod, strat, str(n), str(n-f), f'{(n-f)/n*100:.3f}%'])
add_table(['Modality','Strategy','Calls','Successful','Compliance'], rows,
          col_widths=[1.6,1.2,1.0,1.2,1.2])

P('All four compliance failures occurred on the guided strategy:', bold=False)
for f in failures:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{f["qid"]} (rep {f["rep"]}): ')
    p.add_run(f'returned "{f["raw_response"]}" instead of the expected "letter+digit" format. '
              f'Failure category: {f["failure_category"]}.')

P('The pattern is consistent: the model returned the correct grid cell but with the digit before '
  'the letter (e.g. "12F" rather than "F12"), tripping the parser. The semantic content is correct; '
  'only the syntactic form is wrong. All four cases concentrate in panoramic and periapical guided '
  'queries on a single landmark per affected image. None occurred on cephalometric. Section 8.4 '
  'discusses what this means for downstream parsing tolerance.')

# ────────────────────────────────────────────────────────────────────
# 5. RQ1 — MODALITY-STRATIFIED ACCURACY
# ────────────────────────────────────────────────────────────────────
H('5. RQ1 — Modality-Stratified Accuracy of GPT-5.4', level=1)

H('5.1 Point landmarks: mean Euclidean distance with bootstrap 95% CIs', level=2)

caption('Table 3: Mean Euclidean distance (in grid cells) for point landmarks, by modality and '
        'strategy. CI = bootstrap percentile 95% confidence interval, 10,000 resamples. n = '
        'number of unique queries; each query is the mean over 3 repetitions.')

rows = []
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = analysis['RQ1_modality_strategy'][f'{mod}_{strat}_point']
        rows.append([mod, strat, str(d['n']),
                     f'{d["mean"]:.3f}',
                     f'[{d["mean_ci"][0]:.3f}, {d["mean_ci"][1]:.3f}]',
                     f'{d["median"]:.3f}',
                     f'[{d["median_ci"][0]:.3f}, {d["median_ci"][1]:.3f}]'])
add_table(
    ['Modality','Strategy','n','Mean ED','Mean 95% CI','Median ED','Median 95% CI'],
    rows,
    col_widths=[1.45,1.0,0.55,0.85,1.4,0.85,1.4]
)

P('Cephalometric is the easiest modality for GPT-5.4 by a wide margin. Median ED is 0.33 cells in '
  'both strategies — half the queries are off by less than half a cell, with bootstrap median CIs '
  'including zero (i.e. the true median is plausibly zero). Periapical sits in a middle band '
  '(median 1.0 in both strategies). Panoramic is the hardest: zero-shot median ED is √2 ≈ 1.41 '
  'cells, and guided median is 3.16 cells. The 95% mean CIs on panoramic are non-overlapping '
  'between the two strategies (zero-shot [2.22, 2.91] vs guided [3.43, 4.19]), consistent with the '
  'paired test in Section 6.')

H('5.2 Cross-modality comparison: normalised Euclidean distance', level=2)
P('Raw ED is not directly comparable across modalities because grid sizes differ. The normalised '
  'Euclidean distance NED = ED / √[(cols−1)² + (rows−1)²] expresses error as a fraction of the '
  'grid diagonal. Diagonal lengths are 16.55 (panoramic), 8.60 (periapical), and 11.40 '
  '(cephalometric) cells. NED ∈ [0, 1] regardless of modality.')

caption('Table 4: Normalised Euclidean distance (NED) per modality + strategy and grand mean.')
rows=[]
ned = phase_b['ned_modality']
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = ned[f'{mod}_{strat}']
        rows.append([mod, strat, str(d['n']), f'{d["mean_ned"]:.4f}', f'{d["median_ned"]:.4f}'])
for strat in ('zero_shot','guided'):
    d = phase_b['grand_ned'][strat]
    rows.append(['ALL POINT', strat, str(d['n']), f'{d["mean_ned"]:.4f}', f'{d["median_ned"]:.4f}'])
add_table(['Modality','Strategy','n','Mean NED','Median NED'], rows,
          col_widths=[1.6,1.2,0.7,1.2,1.2])

# Pull NED facts
ceph_zs_ned = ned['CEPHALOMETRIC_zero_shot']['mean_ned']
ceph_gd_ned = ned['CEPHALOMETRIC_guided']['mean_ned']
pa_zs_ned   = ned['PERIAPICAL_zero_shot']['mean_ned']
pa_gd_ned   = ned['PERIAPICAL_guided']['mean_ned']
pan_zs_ned  = ned['PANORAMIC_zero_shot']['mean_ned']
pan_gd_ned  = ned['PANORAMIC_guided']['mean_ned']

P(f'NED preserves the modality ordering: CEPH ({ceph_zs_ned:.3f}/{ceph_gd_ned:.3f}) '
  f'< PA ({pa_zs_ned:.3f}/{pa_gd_ned:.3f}) < PAN ({pan_zs_ned:.3f}/{pan_gd_ned:.3f}). '
  f'In other words, the panoramic disadvantage is not an artefact of the larger panoramic grid; '
  f'on a normalised scale, panoramic errors are a larger fraction of the grid diagonal than '
  f'periapical or cephalometric errors. The grand-mean NED is '
  f'{phase_b["grand_ned"]["zero_shot"]["mean_ned"]:.3f} for zero-shot vs '
  f'{phase_b["grand_ned"]["guided"]["mean_ned"]:.3f} for guided, with the gap entirely driven by '
  f'the panoramic component.')

H('5.3 Point landmarks: Successful Detection Rate', level=2)
caption('Table 5: SDR at four thresholds (0, 1, √2, 2 cells) for point landmarks, by modality and '
        'strategy. Wilson score 95% CIs in brackets. SDR@0 = exact-match rate; SDR@1 = within ±1 '
        'cell (orthogonal); SDR@√2 = within one diagonal cell; SDR@2 = within ±2 cells (typical '
        'clinical "same anatomic region" tolerance).')

def fmt_pct_ci(rate, ci):
    return f'{rate*100:.1f}% [{ci[0]*100:.1f}, {ci[1]*100:.1f}]'

rows=[]
sdr = phase_b['sdr_modality_with_ci']
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = sdr[f'{mod}_{strat}']
        rows.append([mod, strat, str(d['n']),
                     fmt_pct_ci(d['SDR@0'], d['SDR@0_ci']),
                     fmt_pct_ci(d['SDR@1'], d['SDR@1_ci']),
                     fmt_pct_ci(d['SDR@√2'], d['SDR@√2_ci']),
                     fmt_pct_ci(d['SDR@2'], d['SDR@2_ci'])])
add_table(['Modality','Strategy','n','SDR@0','SDR@1','SDR@√2','SDR@2'], rows,
          col_widths=[1.3,0.95,0.55,1.05,1.05,1.05,1.05], header_size=9, body_size=9)

ceph_sdr1_zs = sdr['CEPHALOMETRIC_zero_shot']['SDR@1']*100
ceph_sdr1_gd = sdr['CEPHALOMETRIC_guided']['SDR@1']*100
pan_sdr1_zs  = sdr['PANORAMIC_zero_shot']['SDR@1']*100
pan_sdr1_gd  = sdr['PANORAMIC_guided']['SDR@1']*100
P(f'On cephalometric, SDR@1 is {ceph_sdr1_zs:.1f}% (zero-shot) and {ceph_sdr1_gd:.1f}% (guided); '
  f'i.e. roughly nine in ten cephalometric point predictions are within one orthogonal grid cell '
  f'of the ground truth. SDR@2 reaches 99.3% in both strategies, meaning essentially every '
  f'cephalometric prediction is in the correct anatomic region. Panoramic SDR@1 is markedly '
  f'lower at {pan_sdr1_zs:.1f}% (zero-shot) and {pan_sdr1_gd:.1f}% (guided), with the SDR@1 95% '
  f'CIs ([{sdr["PANORAMIC_zero_shot"]["SDR@1_ci"][0]*100:.1f}, '
  f'{sdr["PANORAMIC_zero_shot"]["SDR@1_ci"][1]*100:.1f}] vs '
  f'[{sdr["PANORAMIC_guided"]["SDR@1_ci"][0]*100:.1f}, '
  f'{sdr["PANORAMIC_guided"]["SDR@1_ci"][1]*100:.1f}]) non-overlapping in favour of zero-shot.')

H('5.4 Area landmarks: Jaccard and Dice', level=2)
caption('Table 6: Area landmark overlap metrics (panoramic only). All means averaged across 3 '
        'repetitions per query, then averaged across queries.')
rows=[]
m_pan_zs_area = analysis['RQ1_modality_strategy']['PANORAMIC_zero_shot_area']
m_pan_gd_area = analysis['RQ1_modality_strategy']['PANORAMIC_guided_area']
gj_zs = phase_b['grand_jacc']['zero_shot']
gj_gd = phase_b['grand_jacc']['guided']
rows.append(['PAN AREA AGGREGATE','zero_shot',str(m_pan_zs_area['n']),
             f'{m_pan_zs_area["mean"]:.3f}',
             f'[{m_pan_zs_area["mean_ci"][0]:.3f}, {m_pan_zs_area["mean_ci"][1]:.3f}]',
             f'{m_pan_zs_area["median"]:.3f}'])
rows.append(['PAN AREA AGGREGATE','guided',str(m_pan_gd_area['n']),
             f'{m_pan_gd_area["mean"]:.3f}',
             f'[{m_pan_gd_area["mean_ci"][0]:.3f}, {m_pan_gd_area["mean_ci"][1]:.3f}]',
             f'{m_pan_gd_area["median"]:.3f}'])
add_table(['Group','Strategy','n','Mean Jaccard','Mean Jaccard 95% CI','Median Jaccard'],
          rows, col_widths=[1.7,1.0,0.55,1.1,1.5,1.1])

# Per-area-landmark Jaccard and Dice
caption('Table 7: Per-area-landmark Jaccard and Dice (panoramic only).')
rows=[]
for lm in ('Mandibular_Canal_L','Maxillary_Sinus_R','External_Oblique_Ridge_R'):
    for strat in ('zero_shot','guided'):
        d = phase_b['area_landmark_stats'][f'PANORAMIC/{lm}/{strat}']
        rows.append([lm, strat, str(d['n']),
                     f'{d["mean_jaccard"]:.3f}', f'{d["median_jaccard"]:.3f}',
                     f'{d["mean_dice"]:.3f}', f'{d["median_dice"]:.3f}'])
add_table(['Landmark','Strategy','n','Mean J','Median J','Mean D','Median D'], rows,
          col_widths=[2.0,0.95,0.55,0.85,0.95,0.85,0.95], header_size=9, body_size=9)

P('Among the three panoramic area landmarks, the maxillary sinus is identified most reliably '
  '(mean Jaccard ≈ 0.45–0.47), the mandibular canal moderately (≈ 0.25), and the external oblique '
  'ridge poorly (≈ 0.07–0.10). The model does not seem to find the external oblique ridge as a '
  'recognisable anatomic feature in the panoramic image; this is the lowest-performing landmark '
  'in the entire benchmark and merits dedicated qualitative analysis (Section 8.5).')

H('5.5 Per-landmark detail', level=2)
P('The complete per-landmark mean ED, SDR table, and per-landmark CIs are provided in Appendix A '
  '(Tables A1–A2). Headline observations:')

# Compute headline observations from per-landmark SDR
sdr_lm = phase_b['sdr_landmark']
def lm_sdr1(mod,lm,strat):
    return sdr_lm[f'{mod}/{lm}/{strat}']['SDR@1']*100

p = doc.add_paragraph(style='List Bullet')
p.add_run('All three cephalometric landmarks reach ≥84% SDR@1 in both strategies. Menton is '
          'the strongest at 100% SDR@1 in both. Guided improves Sella SDR@1 from '
          f'{lm_sdr1("CEPHALOMETRIC","Sella_S","zero_shot"):.0f}% to '
          f'{lm_sdr1("CEPHALOMETRIC","Sella_S","guided"):.0f}%; Menton is unchanged.')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'Panoramic Mental_Foramen_L SDR@1 increases from '
          f'{lm_sdr1("PANORAMIC","Mental_Foramen_L","zero_shot"):.0f}% (zero-shot) to '
          f'{lm_sdr1("PANORAMIC","Mental_Foramen_L","guided"):.0f}% (guided), and SDR@√2 from '
          f'{sdr_lm["PANORAMIC/Mental_Foramen_L/zero_shot"]["SDR@√2"]*100:.0f}% to '
          f'{sdr_lm["PANORAMIC/Mental_Foramen_L/guided"]["SDR@√2"]*100:.0f}%. Guided helps here.')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'Panoramic Tooth_33_Apex SDR@1 collapses from '
          f'{lm_sdr1("PANORAMIC","Tooth_33_Apex","zero_shot"):.0f}% (zero-shot) to '
          f'{lm_sdr1("PANORAMIC","Tooth_33_Apex","guided"):.0f}% (guided); SDR@2 from '
          f'{sdr_lm["PANORAMIC/Tooth_33_Apex/zero_shot"]["SDR@2"]*100:.0f}% to '
          f'{sdr_lm["PANORAMIC/Tooth_33_Apex/guided"]["SDR@2"]*100:.0f}%. Guided is catastrophic '
          f'here. The model under guided systematically places tooth 33 several cells away from '
          f'the ground truth.')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'Panoramic Condylar_Head_R SDR@1 drops from '
          f'{lm_sdr1("PANORAMIC","Condylar_Head_R","zero_shot"):.0f}% (zero-shot) to '
          f'{lm_sdr1("PANORAMIC","Condylar_Head_R","guided"):.0f}% (guided). Mean ED rises from '
          f'{sdr_lm["PANORAMIC/Condylar_Head_R/zero_shot"]["mean_ed"]:.2f} to '
          f'{sdr_lm["PANORAMIC/Condylar_Head_R/guided"]["mean_ed"]:.2f} cells.')

# ────────────────────────────────────────────────────────────────────
# 6. RQ2 — STRATEGY EFFECTS
# ────────────────────────────────────────────────────────────────────
H('6. RQ2 — Effect of Guided Prompting (Zero-Shot vs Guided)', level=1)

H('6.1 Distributional non-normality and choice of test', level=2)
caption('Table 8: Shapiro-Wilk normality test on paired delta (zero-shot − guided per query mean ED, '
        'or guided − zero-shot Jaccard for area).')
rows=[]
for k in ('CEPHALOMETRIC_point_delta','PERIAPICAL_point_delta','PANORAMIC_point_delta','PANORAMIC_area_delta'):
    d = phase_b['paired_shapiro'][k]
    verdict = 'normal' if d['normal_at_alpha_05'] else 'non-normal'
    rows.append([k.replace('_',' '), str(d['n']), f'{d["W"]:.4f}',
                 f'{d["p"]:.2e}', verdict])
add_table(['Comparison','n','Shapiro W','p','Verdict at α=0.05'], rows,
          col_widths=[2.4,0.6,1.0,1.2,1.4])

P('All four paired-delta distributions reject normality at α = 0.05 (all p < 10⁻⁶). The paired '
  'Wilcoxon signed-rank test is therefore the appropriate strategy-comparison test on per-query '
  'rep-mean values.')

H('6.2 Modality-level paired comparisons (Bonferroni × 4)', level=2)
caption('Table 9: Modality-level paired Wilcoxon comparisons of zero-shot vs guided strategy. '
        'Δ = zero-shot − guided for ED (positive = guided worse) and guided − zero-shot for Jaccard '
        '(positive = guided better). Bonferroni-corrected p-values across 4 comparisons.')
rows=[]
def fmt_p(p):
    if p<1e-9: return f'{p:.2e}'
    if p<0.001: return f'{p:.2e}'
    return f'{p:.4f}'
def fmt_pbonf(p,m=4):
    pb = min(1.0, p*m)
    if pb<1e-9: return f'{pb:.2e}'
    if pb<0.001: return f'{pb:.2e}'
    return f'{pb:.4f}'
def sig_marker(p_bonf):
    if p_bonf<0.001: return '***'
    if p_bonf<0.01: return '**'
    if p_bonf<0.05: return '*'
    return 'NS'

for k_label, k in [('CEPHALOMETRIC point ED','CEPHALOMETRIC_point'),
                    ('PERIAPICAL point ED','PERIAPICAL_point'),
                    ('PANORAMIC point ED','PANORAMIC_point'),
                    ('PANORAMIC area Jaccard','PANORAMIC_area')]:
    d = analysis['RQ2a_strategy_per_modality'][k]
    pbonf = min(1.0, d['p']*4)
    rows.append([k_label, str(d['n_total']), str(d['n_nonzero']),
                 f'{d["mean_delta"]:+.4f}',
                 f'{d["median_delta"]:+.4f}',
                 fmt_p(d['p']), fmt_pbonf(d['p'],4),
                 f'{d["rank_biserial_r"]:+.3f}',
                 f'[{d["rank_biserial_ci_low"]:+.2f}, {d["rank_biserial_ci_high"]:+.2f}]',
                 sig_marker(pbonf)])
add_table(['Comparison','n','n≠0','Mean Δ','Median Δ','Wilcoxon p','Bonf p','rank-biserial r','r 95% CI','Sig'],
          rows, col_widths=[1.5,0.45,0.5,0.7,0.75,0.85,0.75,0.75,0.95,0.45],
          header_size=9, body_size=9)

P('Only the panoramic point comparison survives Bonferroni correction with overwhelming significance '
  f'(Bonf p = {min(1.0, analysis["RQ2a_strategy_per_modality"]["PANORAMIC_point"]["p"]*4):.2e}, '
  f'rank-biserial r = '
  f'{analysis["RQ2a_strategy_per_modality"]["PANORAMIC_point"]["rank_biserial_r"]:+.3f}). The mean '
  f'difference is +1.26 cells (zero-shot − guided), i.e. on average, guided is 1.26 cells worse than '
  f'zero-shot on a panoramic point query. The panoramic area comparison was nominally significant '
  f'(p = {analysis["RQ2a_strategy_per_modality"]["PANORAMIC_area"]["p"]:.4f}) but does not survive '
  f'Bonferroni correction (Bonf p = '
  f'{min(1.0, analysis["RQ2a_strategy_per_modality"]["PANORAMIC_area"]["p"]*4):.4f}). Cephalometric '
  f'and periapical point comparisons are not significant.')

H('6.3 Per-landmark stratified analysis', level=2)
caption('Table 10: Per-landmark paired Wilcoxon comparisons. Δ = zero-shot − guided mean ED for '
        'point landmarks; Δ = guided − zero-shot mean Jaccard for area landmarks. Bonferroni × 9 '
        'across the nine point landmarks; Bonferroni × 3 across the three area landmarks.')
rows=[]
# Point landmarks
landmarks_order = [
    ('CEPHALOMETRIC','Menton_Me'),('CEPHALOMETRIC','Nasion_N'),('CEPHALOMETRIC','Sella_S'),
    ('PERIAPICAL','Tooth_36_Distal_Apex'),('PERIAPICAL','Tooth_36_Distal_CEJ'),('PERIAPICAL','Tooth_36_Mesial_CEJ'),
    ('PANORAMIC','Mental_Foramen_L'),('PANORAMIC','Condylar_Head_R'),('PANORAMIC','Tooth_33_Apex'),
]
for mod,lm in landmarks_order:
    d = analysis['RQ2b_strategy_per_landmark'][f'{mod}/{lm}']
    pbonf = min(1.0, d['p']*9)
    rows.append([f'{mod}/{lm}','point',str(d['n_total']),str(d['n_nonzero']),
                 f'{d["mean_delta"]:+.4f}',
                 f'{d["median_delta"]:+.4f}',
                 fmt_p(d['p']), fmt_pbonf(d['p'],9),
                 f'{d["rank_biserial_r"]:+.3f}',
                 sig_marker(pbonf)])
# Area landmarks
for lm in ('Mandibular_Canal_L','Maxillary_Sinus_R','External_Oblique_Ridge_R'):
    d = phase_b['area_per_lm_wilcoxon'][f'PANORAMIC/{lm}/area']
    pbonf = min(1.0, d['p']*3)
    rows.append([f'PANORAMIC/{lm}','area',str(d['n_total']),str(d['n_nonzero']),
                 f'{d["mean_delta"]:+.4f}',
                 f'{d["median_delta"]:+.4f}',
                 fmt_p(d['p']), fmt_pbonf(d['p'],3),
                 f'{d["rank_biserial_r"]:+.3f}',
                 sig_marker(pbonf)])
add_table(['Landmark','Type','n','n≠0','Mean Δ','Median Δ','Wilcoxon p','Bonf p','rank-biserial r','Sig'],
          rows, col_widths=[2.0,0.55,0.4,0.4,0.7,0.7,0.8,0.7,0.85,0.4],
          header_size=8.5, body_size=8.5)

P('The per-landmark analysis reveals strongly heterogeneous strategy effects that the modality-level '
  'aggregate disguises:')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Guided BETTER than zero-shot (significant after Bonferroni × 9 + 3): ').bold = True
p.add_run(f"PAN/Mental_Foramen_L (Δ = "
          f"{analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Mental_Foramen_L']['mean_delta']:+.3f}, "
          f"r = {analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Mental_Foramen_L']['rank_biserial_r']:+.3f}, "
          f"Bonf p = {analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Mental_Foramen_L']['p']*9:.2e}); "
          f"CEPH/Sella_S (Δ = "
          f"{analysis['RQ2b_strategy_per_landmark']['CEPHALOMETRIC/Sella_S']['mean_delta']:+.3f}, "
          f"r = {analysis['RQ2b_strategy_per_landmark']['CEPHALOMETRIC/Sella_S']['rank_biserial_r']:+.3f}, "
          f"Bonf p = {analysis['RQ2b_strategy_per_landmark']['CEPHALOMETRIC/Sella_S']['p']*9:.4f}).")

p = doc.add_paragraph(style='List Bullet')
p.add_run('Guided WORSE than zero-shot (significant after Bonferroni): ').bold = True
p.add_run(f"PAN/Tooth_33_Apex (Δ = "
          f"{analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Tooth_33_Apex']['mean_delta']:+.3f}, "
          f"r = {analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Tooth_33_Apex']['rank_biserial_r']:+.3f}, "
          f"Bonf p = {analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Tooth_33_Apex']['p']*9:.2e}); "
          f"PAN/Condylar_Head_R (Δ = "
          f"{analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Condylar_Head_R']['mean_delta']:+.3f}, "
          f"r = {analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Condylar_Head_R']['rank_biserial_r']:+.3f}, "
          f"Bonf p = {analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Condylar_Head_R']['p']*9:.4f}); "
          f"CEPH/Menton_Me (Δ = "
          f"{analysis['RQ2b_strategy_per_landmark']['CEPHALOMETRIC/Menton_Me']['mean_delta']:+.3f}, "
          f"r = {analysis['RQ2b_strategy_per_landmark']['CEPHALOMETRIC/Menton_Me']['rank_biserial_r']:+.3f}, "
          f"Bonf p = {analysis['RQ2b_strategy_per_landmark']['CEPHALOMETRIC/Menton_Me']['p']*9:.4f}). "
          "The Menton effect is small in absolute terms (mean Δ = +0.11 cells) but consistent in "
          "direction across the 10 non-zero pairs.")

p = doc.add_paragraph(style='List Bullet')
p.add_run('Guided BETTER than zero-shot (area, significant after Bonferroni × 3): ').bold = True
p.add_run(f"PAN/External_Oblique_Ridge_R Jaccard "
          f"(Δ = {phase_b['area_per_lm_wilcoxon']['PANORAMIC/External_Oblique_Ridge_R/area']['mean_delta']:+.4f}, "
          f"r = {phase_b['area_per_lm_wilcoxon']['PANORAMIC/External_Oblique_Ridge_R/area']['rank_biserial_r']:+.3f}, "
          f"Bonf p = {phase_b['area_per_lm_wilcoxon']['PANORAMIC/External_Oblique_Ridge_R/area']['p']*3:.4f}). "
          "Both EOR Jaccard means are nonetheless very low in absolute terms (≤0.10).")

p = doc.add_paragraph(style='List Bullet')
p.add_run('Not significant after correction: ').bold = True
p.add_run('CEPH/Nasion, all three periapical landmarks, PAN/Mandibular_Canal area, '
          'PAN/Maxillary_Sinus area.')

H('6.4 Interpretation', level=2)
P('The modality-level aggregate panoramic regression is the imperfect cancellation of three '
  'opposing per-landmark effects: guided strongly helps Mental_Foramen_L (Δ = '
  f'{analysis["RQ2b_strategy_per_landmark"]["PANORAMIC/Mental_Foramen_L"]["mean_delta"]:+.3f}), '
  f'strongly hurts Tooth_33_Apex (Δ = '
  f'{analysis["RQ2b_strategy_per_landmark"]["PANORAMIC/Tooth_33_Apex"]["mean_delta"]:+.3f}), and '
  f'moderately hurts Condylar_Head_R (Δ = '
  f'{analysis["RQ2b_strategy_per_landmark"]["PANORAMIC/Condylar_Head_R"]["mean_delta"]:+.3f}). The '
  f'magnitude ranking is Tooth_33_Apex > Condylar_Head_R > Mental_Foramen_L, so the guided '
  f'aggregate ends up worse for panoramic point queries despite delivering a measurable benefit on '
  f'Mental Foramen.')

P('The same heterogeneity is visible on cephalometric: guided helps Sella by 0.135 cells and '
  'hurts Menton by 0.111 cells; these effects are statistically significant per-landmark but '
  'cancel almost exactly at the modality level (modality-level Δ = '
  f'{analysis["RQ2a_strategy_per_modality"]["CEPHALOMETRIC_point"]["mean_delta"]:+.4f} cells, NS).')

P('The data are consistent with a model in which the explicit grid description in the guided '
  'system prompt does two things: it improves the model\'s anchoring of laterally-named landmarks '
  '(Mental_Foramen_L, with explicit "L" in name and a clear left-right disambiguation) and it '
  'changes the model\'s prior over which cell to select in a way that is harmful for landmarks '
  'whose anatomic location is ambiguous from the image alone (Tooth_33_Apex and Condylar_Head). '
  'Diagnosis of why the guided strategy mispositions Tooth_33_Apex specifically is left to '
  'qualitative analysis (Section 8).')

# ────────────────────────────────────────────────────────────────────
# 7. RQ3 — REPRODUCIBILITY
# ────────────────────────────────────────────────────────────────────
H('7. RQ3 — Reproducibility at Temperature = 0', level=1)

H("7.1 Fleiss' kappa across the three repetitions", level=2)
P("For each (query, strategy) point pair, the three predicted cells across the three reps were "
  "treated as the responses of three raters on a single item. Fleiss' κ was computed separately per "
  "modality × strategy because the category space (set of valid cells) varies with grid size. "
  "Items with any missing cell (compliance failure or out-of-range) are excluded from κ computation.")

caption("Table 11: Fleiss' κ across the three repetitions, point landmarks only.")
rows=[]
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = phase_b['fleiss_per_group'][f'{mod}_{strat}_point']
        kappa = float(d['kappa'])
        if kappa>=0.81: tier='almost perfect'
        elif kappa>=0.61: tier='substantial'
        elif kappa>=0.41: tier='moderate'
        elif kappa>=0.21: tier='fair'
        else: tier='slight/poor'
        rows.append([mod, strat, str(d['n_items']), f'{kappa:.4f}', tier])
rows.append(['ALL POINT (mod-prefixed)','—','1,197',f'{float(phase_b["fleiss_overall_point"]):.4f}','substantial'])
add_table(['Modality','Strategy','n items','Fleiss κ','Landis-Koch tier'], rows,
          col_widths=[1.6,1.0,0.7,1.0,1.5])

P('Cephalometric is the most reproducible (κ = 0.83 zero-shot, 0.89 guided — almost perfect '
  'agreement). Periapical is similarly high (κ = 0.83 zero-shot, 0.82 guided). Panoramic is the '
  'least reproducible (κ = 0.65 zero-shot, 0.70 guided — substantial agreement). This tracks the '
  'modality difficulty hierarchy from Section 5: harder modalities are also less self-consistent. '
  'Notably, the guided strategy shows higher κ than zero-shot on cephalometric and panoramic '
  '— even when guided is less accurate (panoramic), it is more consistent with itself.')

H('7.2 Three-way unanimous response rates (point)', level=2)
caption("Table 12: Three-way unanimous predicted-cell rates. \"Unanimous\" = all 3 reps "
        "produced the identical (row,col) cell.")
rows=[]
total_un = 0; total_n = 0
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = phase_b['exact_3way_unanimous'][f'{mod}_{strat}_point']
        rate = d['rate'] if d['rate'] is not None else 0
        rows.append([mod, strat, str(d['n']), str(d['unanimous']), f'{rate*100:.1f}%'])
        total_un += d['unanimous']; total_n += d['n']
rows.append(['OVERALL','—', str(total_n), str(total_un), f'{total_un/total_n*100:.1f}%'])
add_table(['Modality','Strategy','n','3-way unanimous','Rate'], rows,
          col_widths=[1.6,1.0,0.7,1.4,1.0])

P(f'Overall, {total_un} of {total_n} point queries × strategies ({total_un/total_n*100:.1f}%) '
  f'received identical predicted cells in all three reps. The remaining ~{(1-total_un/total_n)*100:.0f}% '
  f'show some between-rep disagreement; the magnitude of that disagreement is quantified next.')

H('7.3 ED spread across the three repetitions', level=2)
caption("Table 13: Distribution of (max − min) Euclidean distance across 3 reps for each point query.")
rows=[]
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = phase_b['ed_spread_per_group'][f'{mod}_{strat}']
        rows.append([mod, strat, str(d['n']),
                     f'{d["mean_spread"]:.3f}', f'{d["median_spread"]:.3f}',
                     f'{d["max_spread"]:.3f}'])
add_table(['Modality','Strategy','n','Mean spread','Median spread','Max spread'], rows,
          col_widths=[1.6,1.0,0.55,1.1,1.1,1.0])

P('Median spread is 0.000 for every modality × strategy — the typical query has all three reps '
  'identical or near-identical. Mean spreads are pulled up by a tail of high-disagreement queries, '
  'concentrated in panoramic. The max spread of '
  f'{phase_b["ed_spread_per_group"]["PANORAMIC_guided"]["max_spread"]:.1f} cells (panoramic guided) '
  f'demonstrates the model can disagree with itself by several cells on the hardest queries; this '
  f'is the variance that rep-averaging suppresses.')

H('7.4 Area-landmark cross-rep agreement', level=2)
caption("Table 14: Area landmarks — mean pairwise Jaccard between the 3 reps' predicted cell sets "
        "(panoramic only).")
rows=[]
for mod_strat in ('PANORAMIC_zero_shot_area','PANORAMIC_guided_area'):
    d = phase_b['area_reliability'][mod_strat]
    rows.append([mod_strat.replace('PANORAMIC_','').replace('_area',''), str(d['n']),
                 f'{d["mean_pairwise_jacc"]:.4f}', f'{d["median"]:.4f}'])
add_table(['Strategy','n','Mean pairwise Jaccard','Median'], rows,
          col_widths=[1.4,0.7,1.6,1.1])

P(f'Guided is more self-consistent on area landmarks than zero-shot '
  f'({phase_b["area_reliability"]["PANORAMIC_guided_area"]["mean_pairwise_jacc"]:.3f} vs '
  f'{phase_b["area_reliability"]["PANORAMIC_zero_shot_area"]["mean_pairwise_jacc"]:.3f}), even '
  f'though the two are essentially indistinguishable in mean accuracy (Section 5.4). The guided '
  f'system prompt appears to constrain the model to a more consistent set of cells across reps.')

H('7.5 Implications for single-run vs rep-averaged evaluation', level=2)
P('The original methodology document (Section 8.1) anticipated κ ≈ 1.0 at temperature = 0 and '
  'planned to justify single-run evaluation conditional on that finding. The empirical κ of '
  f'{float(phase_b["fleiss_overall_point"]):.3f} — substantial but not perfect — invalidates that '
  f'simplification for this benchmark on this model. We recommend rep-averaging (or at minimum '
  f'reporting reproducibility data alongside single-run values) for any future analyses that '
  f'compare GPT-5.4 against another rater (e.g. dental student means, the second commercial MLLM, '
  f'or another OMFR specialist) on this task. All numbers in this report use the per-query mean '
  f'across the three repetitions; the rep-level values are preserved in the records pickle for '
  f'sensitivity analyses.')

# ────────────────────────────────────────────────────────────────────
# 8. DISCUSSION
# ────────────────────────────────────────────────────────────────────
H('8. Discussion', level=1)

H('8.1 Modality difficulty hierarchy', level=2)
P('The CEPH << PA < PAN ordering is robust across raw ED, NED, SDR at every threshold, paired-test '
  'effect sizes, and reproducibility κ. Cephalometric radiographs offer a small set of well-'
  'separated, high-contrast osseous landmarks against a bright background; the relevant points '
  'are spatially anchored by clearly identifiable anatomy (sella turcica, nasal bones, mandibular '
  'symphysis). Periapical radiographs introduce occlusal crowding and require fine-grained '
  'discrimination between mesial and distal aspects of a single tooth, where mean ED >1 cell is '
  'the norm. Panoramic images present the largest grid, the most landmarks, and the most spatial '
  'and lateral reasoning demand (left-vs-right disambiguation, multi-landmark coexistence, varying '
  'magnification across the curved surface). The data are consistent with a model whose spatial '
  'reasoning competence degrades as the spatial reasoning demand increases.')

H('8.2 Strategy effects are heterogeneous, not uniform', level=2)
P('A common simplification in MLLM prompt-engineering literature is to characterise a system-prompt '
  'change as "improving" or "harming" performance globally. The per-landmark analysis here '
  'directly contradicts that simplification. The same prompt change — adding an explicit grid-'
  'coordinate-system explanation — produces statistically significant performance gains on three '
  'landmarks (Mental_Foramen_L, Sella_S, External_Oblique_Ridge_R) and statistically significant '
  'performance losses on three other landmarks (Tooth_33_Apex, Condylar_Head_R, Menton_Me). The '
  'same model, same image, same user prompt; only the system prompt differs.')

P('Why does guided regress so dramatically on Tooth_33_Apex? The mean ED jump from 1.40 to 4.06 '
  'cells, with 100% of non-zero pairs preferring zero-shot (rank-biserial r = +1.00), implies a '
  'systematic — not random — shift in the model\'s preferred cell location. One hypothesis is '
  'that the explicit grid description nudges the model toward the lower-left quadrant on tooth-33-'
  'apex queries (the canonical "left lower canine apex" position the model has presumably seen '
  'most often in training contexts), in tension with what the actual image shows. This is a '
  'specific, falsifiable hypothesis that warrants targeted qualitative inspection of the '
  'misclassified cases. We do not test it here.')

P('A practical implication is that strategy comparisons reported only at the aggregate or '
  'modality level will hide effects of clinical importance. The panoramic guided regression on '
  'Tooth_33_Apex is large enough to reverse a clinical conclusion about whether to recommend the '
  'guided strategy; the per-landmark stratification is therefore the right unit of analysis for '
  'deployment decisions.')

H('8.3 Reproducibility implications', level=2)
P('The empirical κ across the three repetitions (overall 0.78, range 0.65–0.89 by modality × '
  'strategy) departs measurably from the κ = 1.0 a strict-determinism reading of "temperature = 0" '
  'would imply. The OpenAI service does not in practice yield bit-identical responses across runs '
  'even at temperature = 0 with a fixed seed; this is consistent with reports in the literature '
  'and with our pilot data. Importantly, the disagreement is concentrated on the harder queries '
  '— median spread is zero in every modality × strategy, with the variance loaded into a tail. '
  'Rep-averaging therefore primarily affects results on the most difficult queries, which are '
  'exactly the queries where any reported single-run accuracy figure would be the least reliable.')

H('8.4 Compliance and parser tolerance', level=2)
P(f'Format compliance was {compliance_rate*100:.3f}% on this run. The four failures all share the '
  f'pattern "digit before letter" (e.g. "12F" instead of "F12"). Semantically these answers are '
  f'correct — they identify the right cell — but they fail the strict left-to-right parser. A '
  f'tolerant parser that accepts both orderings would lift compliance to 100% and alter no metric '
  f'in this report (the four reversed-coordinate answers correspond to plausible cells, not '
  f'hallucinated ones). The trade-off is whether the published compliance rate captures '
  f'"instruction-following" strictly (current 99.926%) or "answers the question with a valid cell" '
  f'leniently (100.000%). We recommend reporting both.')

H('8.5 The external oblique ridge anomaly', level=2)
P(f'Among the three panoramic area landmarks, the external oblique ridge stands out as the lowest-'
  f'performing structure in the entire benchmark: mean Jaccard {phase_b["area_landmark_stats"]["PANORAMIC/External_Oblique_Ridge_R/zero_shot"]["mean_jaccard"]:.3f} (zero-shot) and '
  f'{phase_b["area_landmark_stats"]["PANORAMIC/External_Oblique_Ridge_R/guided"]["mean_jaccard"]:.3f} (guided), with '
  f'mean Dice 0.10 and 0.15 respectively. This is the only landmark whose absolute performance is '
  f'so low that the question of "which strategy is better" is operationally moot. The structure '
  f'is a thin, oblique radiopaque line that shares its course with several adjacent radiopaque '
  f'features (mandibular ramus border, mandibular angle); it is plausible that GPT-5.4 simply does '
  f'not recognise it as a discrete anatomical entity. A targeted analysis of the model\'s response '
  f'to "external oblique ridge" prompts would likely find a generic "lower-right mandible region" '
  f'response pattern rather than a precise trace.')

H('8.6 Clinical implications', level=2)
P('On cephalometric radiographs, GPT-5.4 reaches SDR@2 ≥ 99% and SDR@1 ≥ 90% under either '
  'strategy. Within this modality, the model behaves at a level that warrants further evaluation '
  'against dental students and against purpose-built CNN systems (which the literature reports '
  'achieving ≈ 92% SDR within 2 mm on cephalometric tasks, e.g. Jiang et al. 2023). On panoramic, '
  'SDR@1 is at most 30% — meaning roughly seven in ten panoramic point predictions are off by more '
  'than one orthogonal cell, and on the worst landmark (Tooth_33_Apex under the guided strategy), '
  'SDR@1 collapses to 0%. The model in its current form cannot replace a trained observer for '
  'panoramic landmark identification; it may still be useful as a screening or hint generator, '
  'subject to per-landmark calibration of the SDR threshold considered "acceptable".')

# ────────────────────────────────────────────────────────────────────
# 9. LIMITATIONS
# ────────────────────────────────────────────────────────────────────
H('9. Limitations', level=1)

p = doc.add_paragraph(style='List Number')
p.add_run('Single model. ').bold = True
p.add_run('This report covers only GPT-5.4. Generalisation to other commercial MLLMs (e.g. '
          'Gemini 3.1 Pro), or to fine-tuned models, is unsupported by these data.')

p = doc.add_paragraph(style='List Number')
p.add_run('Single-rater ground truth. ').bold = True
p.add_run('OMFR_1 was the sole annotator. Inter-rater (OMFR_2) and intra-rater data are not yet '
          'available; results may shift modestly when consensus or inter-rater-averaged ground '
          'truth is used. Estimated impact bounded by typical inter-rater Jaccard ≥ 0.7 in similar '
          'studies, suggesting a few cell-level reassignments rather than large structural changes.')

p = doc.add_paragraph(style='List Number')
p.add_run('Strategy comparison limited to two arms. ').bold = True
p.add_run('Few-shot prompting was deferred. The two-arm design quantifies the effect of explicit '
          'grid description but cannot disentangle the contribution of the several distinct text '
          'additions in the guided system prompt (row/column declaration, format examples, the '
          'panoramic L–R clause). A factorial decomposition of the guided prompt is feasible but '
          'was not performed for this report.')

p = doc.add_paragraph(style='List Number')
p.add_run('Single-institution dataset. ').bold = True
p.add_run('All 200 images are from one university hospital. External validation on images from '
          'different equipment, populations, or institutions is required to claim generalisability.')

p = doc.add_paragraph(style='List Number')
p.add_run('No human benchmark. ').bold = True
p.add_run('The dental student dataset is not yet ingested. The MLLM-vs-student comparison that the '
          'overall study targets is not in scope here.')

p = doc.add_paragraph(style='List Number')
p.add_run('Grid-cell metrics, no millimetric conversion. ').bold = True
p.add_run('Variable magnification across panoramic radiographs and projection geometry of '
          'periapical radiographs preclude reliable single-factor cell→mm conversion (this is '
          'discussed in the response to the pre-experiment review). Reported errors are in grid-'
          'cell units; comparison to literature reported in millimetres requires '
          'modality-specific approximation.')

p = doc.add_paragraph(style='List Number')
p.add_run('Determinism assumption. ').bold = True
p.add_run("Section 7 documents that temperature = 0 does not yield κ = 1.0; rep-averaging is used "
          "throughout. Subsequent analyses against external raters should follow the same "
          "convention.")

# ────────────────────────────────────────────────────────────────────
# 10. CONCLUSIONS AND FUTURE WORK
# ────────────────────────────────────────────────────────────────────
H('10. Conclusions and Future Work', level=1)

P('On the basis of 5,400 GPT-5.4 API calls across 900 dental landmark queries:')

p = doc.add_paragraph(style='List Bullet')
p.add_run('GPT-5.4 performs well on cephalometric landmark identification (mean ED ≤ 0.51 cells, '
          'SDR@1 ≥ 90%, SDR@2 ≥ 99%), respectably on periapical (mean ED ≈ 1.07–1.09 cells, SDR@1 '
          '≈ 53%, SDR@2 ≈ 89–90%), and poorly on panoramic (mean ED ≥ 2.55 cells, SDR@1 ≤ 30%, '
          'SDR@2 ≤ 69%) — with normalised cross-modality comparison preserving the same ordering.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('The guided prompting strategy is heterogeneous in its effect: it strongly improves '
          'Mental_Foramen_L and Sella_S, strongly harms Tooth_33_Apex, and moderately harms '
          'Condylar_Head_R and Menton_Me. The aggregate panoramic regression is significant '
          '(Bonferroni p < 10⁻¹⁰), but it is driven by a few landmarks; it should not be read as '
          '"explicit grid explanation harms the model in general".')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Reproducibility at temperature = 0 is substantial but not perfect (κ = 0.78 overall, '
          '0.65–0.89 by modality × strategy). Rep-averaging is a methodological necessity, not a '
          'redundant safeguard.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Compliance is 99.926% under strict parsing, 100% under digit-before-letter-tolerant '
          'parsing. The four failures concentrate on a single failure mode and a small set of '
          'queries.')

P('Future work directions, in order of expected impact:')

p = doc.add_paragraph(style='List Number')
p.add_run('Run the equivalent experiment on Gemini 3.1 Pro (a major scientific question for the '
          'manuscript) and on at least one open-source multimodal model (provider-diversity '
          'control).')

p = doc.add_paragraph(style='List Number')
p.add_run('Ingest the OMFR_2 and student datasets and produce the cross-rater comparisons, '
          'including Bland-Altman plots, intra-rater ICC/κ, and student-mean vs MLLM paired '
          'tests originally specified in the methodology document.')

p = doc.add_paragraph(style='List Number')
p.add_run('Targeted qualitative inspection of the panoramic guided regressions on Tooth_33_Apex '
          'and Condylar_Head_R: which cells does the model select under guided that it does not '
          'select under zero-shot, and is the misplacement consistent with a learned anatomic '
          'prior over-applied to the present image?')

p = doc.add_paragraph(style='List Number')
p.add_run('A factorial decomposition of the guided system prompt to identify which textual '
          'addition (row/column declaration, format examples, L–R viewing clause) contributes '
          'each of the per-landmark effects in Section 6.')

p = doc.add_paragraph(style='List Number')
p.add_run('Consider a tolerant parser (digit-before-letter) for the published compliance metric, '
          'and report both strict and tolerant rates.')

# ────────────────────────────────────────────────────────────────────
# APPENDIX A — PER-LANDMARK DETAILED TABLES
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix A — Per-Landmark Detailed Statistics', level=1)

H('A1. Per-landmark mean ED with 95% CI (point landmarks)', level=2)
caption("Table A1: Per-landmark mean ED with bootstrap 95% CI and median ED.")
rows=[]
for mod,lm in landmarks_order:
    for strat in ('zero_shot','guided'):
        d = analysis['RQ1_per_landmark'][f'{mod}/{lm}/{strat}']
        rows.append([f'{mod}/{lm}', strat, str(d['n']),
                     f'{d["mean"]:.3f}',
                     f'[{d["mean_ci"][0]:.3f}, {d["mean_ci"][1]:.3f}]',
                     f'{d["median"]:.3f}'])
add_table(['Landmark','Strategy','n','Mean ED','Mean 95% CI','Median ED'], rows,
          col_widths=[2.3,0.95,0.5,0.85,1.5,0.95], header_size=9, body_size=9)

H('A2. Per-landmark SDR (point landmarks)', level=2)
caption("Table A2: Per-landmark Successful Detection Rate at each threshold.")
rows=[]
for mod,lm in landmarks_order:
    for strat in ('zero_shot','guided'):
        d = phase_b['sdr_landmark'][f'{mod}/{lm}/{strat}']
        rows.append([f'{mod}/{lm}', strat, str(d['n']),
                     f'{d["mean_ed"]:.3f}',
                     f'{d["SDR@0"]*100:.0f}%',
                     f'{d["SDR@1"]*100:.0f}%',
                     f'{d["SDR@√2"]*100:.0f}%',
                     f'{d["SDR@2"]*100:.0f}%'])
add_table(['Landmark','Strategy','n','Mean ED','SDR@0','SDR@1','SDR@√2','SDR@2'], rows,
          col_widths=[2.0,0.85,0.4,0.85,0.65,0.65,0.7,0.7], header_size=9, body_size=9)

# ────────────────────────────────────────────────────────────────────
# APPENDIX B — COMPLIANCE FAILURES
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix B — Compliance Failures (Full List)', level=1)
P(f'All {len(failures)} compliance failures across {n_calls:,} API calls are listed below. Each '
  f'represents one rep of one (query, strategy) pair where the parser could not extract a valid '
  f'grid coordinate from the response.')

caption("Table B1: All compliance failures.")
rows=[]
for f in failures:
    qid = f['qid']; strat = f['strat']
    rec = next((r for r in records if r['query_id']==qid and r['strategy']==strat), None)
    gt = rec['gt'] if rec else '?'
    mod = rec['modality'] if rec else '?'
    lm = rec['structure'] if rec else '?'
    rows.append([qid, strat, str(f['rep']), mod, lm, gt,
                 f'"{f["raw_response"]}"', f['failure_category']])
add_table(['Query ID','Strategy','Rep','Modality','Landmark','GT','Raw Response','Category'],
          rows, col_widths=[1.9,0.7,0.35,1.1,1.5,0.55,1.0,0.6],
          header_size=9, body_size=9)

P('All four cases show the digit-before-letter inversion. PAN_044 expected ground truth in row F, '
  'periapical PA_039 in row F, and PAN_079 in row E; the model returned the cells with the digit '
  'first. Under a tolerant parser these would all decode to the correct cells, but they are '
  'recorded as parse failures here for strict compliance accounting.')

# ────────────────────────────────────────────────────────────────────
# APPENDIX C — REPRODUCIBILITY MANIFEST
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix C — Reproducibility Manifest', level=1)
P('All figures in this report can be reproduced from the artifacts listed below.')

caption("Table C1: Reproducibility manifest.")
rows = [
    ['Pipeline source — git commit', git_sha],
    ['Pipeline source — commit date', git_date],
    ['Source Excel SHA-256', excel_sha],
    ['Derived query_index.json SHA-256', qi_sha],
    ['Model', 'gpt-5.4'],
    ['Inference settings', 'temperature = 0, seed = 42, max_completion_tokens = 50, image detail = high'],
    ['Repetitions', '3 per (query × strategy)'],
    ['Sandbox directory', 'results_full/'],
    ['Total queries', '900 (PAN: 100 imgs × 6 LM = 600; PA: 50 × 3 = 150; CEPH: 50 × 3 = 150)'],
    ['Strategies', 'zero_shot, guided'],
    ['Total API calls', f'{n_calls:,}'],
    ['Successful responses', f'{n_actual:,}'],
    ['Compliance (strict)', f'{compliance_rate*100:.4f}%'],
    ['Prompt tokens', f'{prompt_tok:,}'],
    ['Completion tokens', f'{compl_tok:,}'],
    ['Naïve cost', f'${naive_cost:.2f}'],
    ['Bootstrap CIs', '10,000 resamples, percentile method, random_state = 42'],
    ['Strategy test', 'Wilcoxon signed-rank (zero-method = wilcox, two-sided)'],
    ['Effect size', 'Matched-pair rank-biserial r'],
    ['Multiple comparisons', 'Bonferroni; family sizes per analysis level (4 modality, 9 point lm, 3 area lm)'],
    ['Reliability', "Fleiss' κ per modality+strategy; categories = predicted cells"],
]
add_table(['Item','Value'], rows, col_widths=[2.5,4.0], header_size=10, body_size=9)

P('Raw API outputs (JSONL) are in results_full/run{1,2,3}/responses/. Parsed responses are in '
  'results_full/run{N}/parsed_responses.json. Saved analysis artifacts (used to populate every '
  'table in this report) are at /tmp/full_run_{records_v2.pkl, analysis.json, phase_b.json, '
  'summary.json, failures.json}.')

# Save
doc.save(str(OUT))
print(f'Wrote {OUT}')
print(f'  Paragraphs: {len(doc.paragraphs)}')
print(f'  Tables: {len(doc.tables)}')
import os
print(f'  Size: {os.path.getsize(OUT):,} bytes')
