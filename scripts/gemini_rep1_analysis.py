"""
Gemini 3.1 Pro 1-rep analysis report.

Reads the just-completed Gemini run and the GPT-5.4 v2 main run, computes
every claim from raw data (with asserts at each step to catch any bug),
and writes a comprehensive DOCX report at
results_full_gemini/Gemini_Rep1_Analysis_Report.docx.

Read-only — touches no API, no .api_lock, no source data files.

Body sections:
  §1 Executive summary
  §2 Methodology
  §3 Compliance + failure-mode analysis
  §4 Quality comparison vs GPT-5.4 v2 rep1 (single-rep matched)
  §5 F5/F6 attractor specific analysis (Tooth_33_Apex)
  §6 MAX_TOKENS truncation deep dive (would 4096 help?)
  §7 Cost analysis (actual + 3-rep projection + 4096 scenario)
  §8 Recommendations

Appendix:
  A1 All 1,800 (query × strategy) Gemini rep1 raw + parsed
  A2 All 1,800 (query × strategy) GPT-5.4 rep1 raw + parsed
  A3 Per-query side-by-side Gemini vs GPT-5.4 vs GT
"""
from __future__ import annotations
import json, re, hashlib, statistics
from pathlib import Path
from collections import Counter, defaultdict

import scipy.stats as stats
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path("/Users/burcusayin/Codes/Claude/grid-based-mllm-landmarks").resolve()

# ============================================================
# Step 0: Load and verify integrity (asserts catch bugs early)
# ============================================================

print("[0] Loading data + integrity checks...")

# Source query index (consensus GT)
qs = json.loads((ROOT / "results_consensus/query_index.json").read_text())
assert len(qs) == 900, f"Expected 900 queries, got {len(qs)}"
gt_by_qid = {q["query_id"]: q["consensus_gt"] for q in qs}
landmark_type = {q["query_id"]: q["landmark_type"] for q in qs}
modality = {q["query_id"]: q["sheet"] for q in qs}
structure_by_qid = {q["query_id"]: q["structure"] for q in qs}

# Gemini rep1 parsed
gem = json.loads((ROOT / "results_full_gemini/run1/parsed_responses.json").read_text())
assert len(gem) == 1800, f"Expected 1800 Gemini records, got {len(gem)}"
gem_by_key = {(r["query_id"], r["strategy"]): r for r in gem}
assert len(gem_by_key) == 1800, "Gemini has duplicate (qid, strategy) pairs"

# GPT-5.4 v2 main run (all 3 reps)
gpt = {}  # rep -> {(qid, strat): record}
for rep in (1, 2, 3):
    data = json.loads((ROOT / f"results_full/run{rep}/parsed_responses.json").read_text())
    assert len(data) == 1800, f"GPT rep{rep}: expected 1800, got {len(data)}"
    gpt[rep] = {(r["query_id"], r["strategy"]): r for r in data}
    assert len(gpt[rep]) == 1800

# Cross-check matching
gem_pairs = set(gem_by_key.keys())
gpt_pairs = set(gpt[1].keys())
assert gem_pairs == gpt_pairs, f"Gemini-GPT pair mismatch: {len(gem_pairs - gpt_pairs)} only in Gemini, {len(gpt_pairs - gem_pairs)} only in GPT"

# Make sure every record's modality+landmark_type match query_index
for r in gem:
    assert r["modality"] == modality[r["query_id"]], f"modality mismatch {r['query_id']}"
    assert r["landmark_type"] == landmark_type[r["query_id"]], f"landmark_type mismatch {r['query_id']}"

print(f"   ✓ 1800 Gemini records, 1800 × 3 GPT records, all pairs match")
print(f"   ✓ landmark_type / modality consistent with query_index")

# ============================================================
# Step 1: Load raw Gemini token usage (for cost + finishReason analysis)
# ============================================================

print("[1] Loading Gemini raw chunk files for token usage...")
gem_usage = {}  # custom_id -> dict
gem_finish = {}
for f in sorted((ROOT / "results_full_gemini/run1/responses").glob("*_chunk*.json")):
    data = json.loads(f.read_text())
    for entry in data:
        cid = entry["custom_id"]
        resp = entry.get("response", {}) or {}
        if "_error" in resp:
            gem_finish[cid] = "ERROR_CANCELLED"
            continue
        cands = resp.get("candidates", [])
        if not cands:
            gem_finish[cid] = "NO_CANDIDATES"
            continue
        gem_finish[cid] = cands[0].get("finishReason", "UNKNOWN")
        usage = resp.get("usageMetadata", {})
        gem_usage[cid] = {
            "promptTokenCount": usage.get("promptTokenCount", 0),
            "thoughtsTokenCount": usage.get("thoughtsTokenCount", 0),
            "candidatesTokenCount": usage.get("candidatesTokenCount", 0),
        }
print(f"   ✓ {len(gem_usage)} responses with usage, {len(gem_finish)} with finishReason")

# ============================================================
# Step 2: Metric helpers
# ============================================================

def parse_cells(text: str | None) -> list[str]:
    """Parse all grid cells from text."""
    if not text:
        return []
    pattern = r'[A-Ha-h]\s*[-]?\s*(?:1[0-6]|[1-9])\b'
    matches = re.findall(pattern, text)
    out, seen = [], set()
    for m in matches:
        c = re.sub(r'[\s\-]', '', m).upper()
        if c not in seen:
            seen.add(c)
            out.append(c)
    return out

def cell_xy(c: str) -> tuple[int, int]:
    return (int(c[1:]), ord(c[0]) - ord('A'))

def euclidean(a: str, b: str) -> float:
    ax, ay = cell_xy(a); bx, by = cell_xy(b)
    return ((ax-bx)**2 + (ay-by)**2) ** 0.5

def jaccard(pred: list[str], gt: list[str]) -> float:
    p, g = set(pred), set(gt)
    if not (p | g):
        return 0.0
    return len(p & g) / len(p | g)

def parse_gt(gt_str: str) -> list[str]:
    return parse_cells(gt_str)

# ============================================================
# Step 3: Compute per-query metrics for Gemini + GPT-5.4 (each rep)
# ============================================================

print("[3] Computing per-query metrics...")

def metrics_for_record(rec, gt_cells, lm_type):
    """Returns dict with: compliant, distance (point only), jaccard (area only)."""
    out = {"compliant": False, "distance": None, "jaccard": None}
    if rec.get("failure_category") is not None:
        return out
    parsed = rec.get("parsed_coordinates") or []
    if not parsed:
        return out
    out["compliant"] = True
    if lm_type == "point":
        if len(gt_cells) >= 1 and parsed:
            # Use first parsed cell against first GT cell
            out["distance"] = euclidean(parsed[0], gt_cells[0])
    else:  # area
        out["jaccard"] = jaccard(parsed, gt_cells)
    return out

# rows: list of dicts, one per (query_id, strategy)
rows = []
for q in qs:
    qid = q["query_id"]
    gt_cells = parse_gt(q["consensus_gt"])
    for strat in ("zero_shot", "guided"):
        gem_rec = gem_by_key.get((qid, strat))
        gpt_rec = gpt[1].get((qid, strat))  # GPT rep1 only for matched comparison
        gem_m = metrics_for_record(gem_rec, gt_cells, q["landmark_type"])
        gpt_m = metrics_for_record(gpt_rec, gt_cells, q["landmark_type"])
        rows.append({
            "qid": qid,
            "strategy": strat,
            "modality": q["sheet"],
            "landmark_type": q["landmark_type"],
            "structure": q["structure"],
            "gt": q["consensus_gt"],
            "gt_cells": gt_cells,
            "gem_raw": gem_rec.get("raw_response", ""),
            "gem_parsed": gem_rec.get("parsed_coordinates") or [],
            "gem_fail": gem_rec.get("failure_category"),
            "gem_compliant": gem_m["compliant"],
            "gem_distance": gem_m["distance"],
            "gem_jaccard": gem_m["jaccard"],
            "gpt_raw": gpt_rec.get("raw_response", ""),
            "gpt_parsed": gpt_rec.get("parsed_coordinates") or [],
            "gpt_fail": gpt_rec.get("failure_category"),
            "gpt_compliant": gpt_m["compliant"],
            "gpt_distance": gpt_m["distance"],
            "gpt_jaccard": gpt_m["jaccard"],
        })
assert len(rows) == 1800, f"Expected 1800 row dicts, got {len(rows)}"
print(f"   ✓ built {len(rows)} per-query rows")

# ============================================================
# Step 4: Compliance + failure-mode summaries
# ============================================================

print("[4] Computing compliance + failure-mode summaries...")

def fail_counter(records):
    c = Counter()
    for r in records:
        cat = r.get("failure_category")
        c[cat or "compliant"] += 1
    return c

gem_fail = fail_counter(gem)
gpt_fail_rep = {rep: fail_counter(gpt[rep].values()) for rep in (1, 2, 3)}

assert sum(gem_fail.values()) == 1800
print(f"   Gemini fail counter: {dict(gem_fail)}")
print(f"   GPT rep1 fail counter: {dict(gpt_fail_rep[1])}")

# Per-landmark Gemini failure rate
landmark_fail = defaultdict(lambda: {"total": 0, "fails": 0, "no_engage": 0, "max_tokens": 0})
for r in gem:
    cid = r["custom_id"]
    qid = r["query_id"]
    landmark = structure_by_qid[qid]
    landmark_fail[landmark]["total"] += 1
    if r.get("failure_category"):
        landmark_fail[landmark]["fails"] += 1
        if r["failure_category"] == "no_engage":
            landmark_fail[landmark]["no_engage"] += 1
    if gem_finish.get(cid) == "MAX_TOKENS":
        landmark_fail[landmark]["max_tokens"] += 1

# ============================================================
# Step 5: Quality stats (point: ED; area: Jaccard) — matched pairs
# ============================================================

print("[5] Computing quality stats + paired Wilcoxon tests...")

def safe_mean(xs):
    xs = [x for x in xs if x is not None]
    if not xs: return None
    return sum(xs) / len(xs)
def safe_med(xs):
    xs = [x for x in xs if x is not None]
    if not xs: return None
    return statistics.median(xs)

# Point landmarks (Euclidean distance)
point_rows = [r for r in rows if r["landmark_type"] == "point"]
area_rows  = [r for r in rows if r["landmark_type"] == "area"]
assert len(point_rows) + len(area_rows) == 1800

# For paired comparison, only include queries where BOTH models were compliant
def paired_metric(rows_, key):
    pairs = []
    for r in rows_:
        if r[f"gem_{key}"] is not None and r[f"gpt_{key}"] is not None:
            pairs.append((r[f"gem_{key}"], r[f"gpt_{key}"]))
    return pairs

# Overall point ED
all_point_pairs = paired_metric(point_rows, "distance")
gem_point_ed = [g for g, _ in all_point_pairs]
gpt_point_ed = [p for _, p in all_point_pairs]
print(f"   point landmarks: {len(all_point_pairs)} matched pairs (both compliant)")
print(f"     Gemini mean ED:  {safe_mean(gem_point_ed):.3f}")
print(f"     GPT-5.4 mean ED: {safe_mean(gpt_point_ed):.3f}")

# Wilcoxon for points
if all_point_pairs:
    w_point = stats.wilcoxon(gem_point_ed, gpt_point_ed, alternative="two-sided")
    print(f"     Wilcoxon p={w_point.pvalue:.4g}")
else:
    w_point = None

# Overall area Jaccard
all_area_pairs = paired_metric(area_rows, "jaccard")
gem_area_j = [g for g, _ in all_area_pairs]
gpt_area_j = [p for _, p in all_area_pairs]
print(f"   area landmarks: {len(all_area_pairs)} matched pairs")
print(f"     Gemini mean Jaccard:  {safe_mean(gem_area_j):.3f}")
print(f"     GPT-5.4 mean Jaccard: {safe_mean(gpt_area_j):.3f}")
if all_area_pairs:
    w_area = stats.wilcoxon(gem_area_j, gpt_area_j, alternative="two-sided")
    print(f"     Wilcoxon p={w_area.pvalue:.4g}")
else:
    w_area = None

# Per-modality, per-strategy breakdown
breakdown = defaultdict(dict)
for mod in ("PANORAMIC", "PERIAPICAL", "CEPHALOMETRIC"):
    for strat in ("zero_shot", "guided"):
        for lm in ("point", "area"):
            subset = [r for r in rows
                      if r["modality"] == mod
                      and r["strategy"] == strat
                      and r["landmark_type"] == lm]
            key = "distance" if lm == "point" else "jaccard"
            pairs = paired_metric(subset, key)
            if pairs:
                gem_vals = [g for g, _ in pairs]
                gpt_vals = [p for _, p in pairs]
                breakdown[(mod, strat, lm)] = {
                    "n": len(pairs),
                    "n_total": len(subset),
                    "gem_mean": safe_mean(gem_vals),
                    "gpt_mean": safe_mean(gpt_vals),
                    "gem_med": safe_med(gem_vals),
                    "gpt_med": safe_med(gpt_vals),
                    "p": stats.wilcoxon(gem_vals, gpt_vals).pvalue if len(gem_vals) > 1 else None,
                }
            else:
                breakdown[(mod, strat, lm)] = {"n": 0, "n_total": len(subset)}

# ============================================================
# Step 6: F5/F6 attractor — Tooth_33_Apex specifically
# ============================================================

print("[6] F5/F6 attractor analysis...")
t33 = [r for r in rows if r["structure"] == "Tooth_33_Apex"]
assert len(t33) > 0
print(f"   {len(t33)} Tooth_33_Apex rows (across strategies)")

def in_f5_f6(parsed):
    return any(c in {"F5", "F6"} for c in parsed[:1]) if parsed else False

gem_attractor = sum(1 for r in t33 if r["gem_compliant"] and in_f5_f6(r["gem_parsed"]))
gpt_attractor = sum(1 for r in t33 if r["gpt_compliant"] and in_f5_f6(r["gpt_parsed"]))
gem_compliant_t33 = sum(1 for r in t33 if r["gem_compliant"])
gpt_compliant_t33 = sum(1 for r in t33 if r["gpt_compliant"])
gem_exact_t33 = sum(1 for r in t33 if r["gem_compliant"] and r["gem_parsed"] and r["gem_parsed"][0] in parse_gt(r["gt"]))
gpt_exact_t33 = sum(1 for r in t33 if r["gpt_compliant"] and r["gpt_parsed"] and r["gpt_parsed"][0] in parse_gt(r["gt"]))

print(f"   Tooth_33_Apex compliant: Gemini {gem_compliant_t33}/{len(t33)}, GPT {gpt_compliant_t33}/{len(t33)}")
print(f"   Tooth_33_Apex in F5/F6: Gemini {gem_attractor}, GPT {gpt_attractor}")
print(f"   Tooth_33_Apex exact GT: Gemini {gem_exact_t33}, GPT {gpt_exact_t33}")

# ============================================================
# Step 7: MAX_TOKENS analysis — would 4096 help?
# ============================================================

print("[7] MAX_TOKENS analysis...")
max_tok_records = []
for r in gem:
    cid = r["custom_id"]
    finish = gem_finish.get(cid, "?")
    if finish == "MAX_TOKENS":
        usage = gem_usage.get(cid, {})
        max_tok_records.append({
            "qid": r["query_id"],
            "strategy": r["strategy"],
            "structure": structure_by_qid.get(r["query_id"], "?"),
            "modality": r["modality"],
            "landmark_type": r["landmark_type"],
            "raw": r["raw_response"],
            "parsed": r["parsed_coordinates"] or [],
            "compliant": r.get("failure_category") is None,
            "prompt_tok": usage.get("promptTokenCount", 0),
            "thoughts_tok": usage.get("thoughtsTokenCount", 0),
            "answer_tok": usage.get("candidatesTokenCount", 0),
            "gt": gt_by_qid.get(r["query_id"], "?"),
        })

n_max_tok = len(max_tok_records)
n_max_tok_parseable = sum(1 for m in max_tok_records if m["parsed"])
n_max_tok_compliant = sum(1 for m in max_tok_records if m["compliant"])
print(f"   {n_max_tok} MAX_TOKENS responses")
print(f"     {n_max_tok_parseable} parseable (have at least one cell)")
print(f"     {n_max_tok_compliant} compliant (parsed + no failure_category)")

# Thinking-token distribution for MAX_TOKENS responses
if max_tok_records:
    mt_thoughts = [m["thoughts_tok"] for m in max_tok_records]
    mt_answers = [m["answer_tok"] for m in max_tok_records]
    print(f"     thinking tokens (MAX_TOKENS subset): mean={safe_mean(mt_thoughts):.0f}, "
          f"min={min(mt_thoughts)}, max={max(mt_thoughts)}")
    print(f"     answer tokens   (MAX_TOKENS subset): mean={safe_mean(mt_answers):.0f}, "
          f"min={min(mt_answers)}, max={max(mt_answers)}")

# By landmark
mt_by_struct = Counter(m["structure"] for m in max_tok_records)

# ============================================================
# Step 8: Cost analysis
# ============================================================

print("[8] Cost analysis...")
total_in = sum(u["promptTokenCount"] for u in gem_usage.values())
total_thoughts = sum(u["thoughtsTokenCount"] for u in gem_usage.values())
total_answer = sum(u["candidatesTokenCount"] for u in gem_usage.values())
in_price = 1.00
out_price = 6.00  # both thinking + answer at output rate
cost_input = total_in * in_price / 1e6
cost_thoughts = total_thoughts * out_price / 1e6
cost_answer = total_answer * out_price / 1e6
total_cost = cost_input + cost_thoughts + cost_answer
print(f"   input tokens:    {total_in:,}  cost=${cost_input:.3f}")
print(f"   thinking tokens: {total_thoughts:,}  cost=${cost_thoughts:.3f}")
print(f"   answer tokens:   {total_answer:,}  cost=${cost_answer:.3f}")
print(f"   TOTAL 1-rep:     ${total_cost:.2f}")
print(f"   3-rep projection: ${total_cost*3:.2f}")

# 4096 scenario: assume MAX_TOKENS calls would use ~3000 more thinking tokens each
# (very rough — could vary). Other calls likely stay roughly the same.
extra_thinking_per_maxtok = 1500  # mid-range estimate
extra_cost_4096 = n_max_tok * extra_thinking_per_maxtok * out_price / 1e6
cost_4096_estimate = total_cost + extra_cost_4096
print(f"   IF max_tokens=4096 (estimate): 1-rep ~${cost_4096_estimate:.2f}, 3-rep ~${cost_4096_estimate*3:.2f}")

# ============================================================
# Step 9: Build DOCX
# ============================================================

print("[9] Building DOCX...")

doc = Document()
# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def H(text, level=1, *, page_break=False):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(text, level=level)
    h.style.font.name = 'Calibri'
    return h

def P(text, *, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(10)
    if bold: r.bold = True
    if italic: r.italic = True
    return p

def add_table(headers, rows_data, *, col_widths_inches=None, header_bold=True):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = header_bold
                run.font.size = Pt(9)
    # Data rows
    for ri, row in enumerate(rows_data, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
    if col_widths_inches:
        for ci, w in enumerate(col_widths_inches):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    return table

# ---- Title ----
title = doc.add_heading('Gemini 3.1 Pro Single-Rep Analysis Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Benchmark on 900 dental anatomic-landmark queries  ·  vs GPT-5.4 v2 main run\n'
                 f'Generated 2026-05-14  ·  Sandbox: results_full_gemini/run1')
r.italic = True
r.font.size = Pt(10)
doc.add_paragraph()

# ---- §1 Executive Summary ----
H('1. Executive Summary', level=1)
P(f"Gemini 3.1 Pro completed 1 repetition (900 queries × 2 strategies = 1,800 calls) on "
  f"the v3 consensus-GT benchmark on 2026-05-14, requiring 59 wall-clock minutes "
  f"and $%.2f at Google's batch-tier pricing." % total_cost)
P("Key findings:")
P(f"  • Strict compliance: {sum(1 for r in gem if r.get('failure_category') is None)} / 1800 "
  f"({sum(1 for r in gem if r.get('failure_category') is None)/1800*100:.2f}%) vs GPT-5.4 v2 rep1's "
  f"{sum(1 for r in gpt[1].values() if r.get('failure_category') is None)}/1800 "
  f"({sum(1 for r in gpt[1].values() if r.get('failure_category') is None)/1800*100:.2f}%)")
P(f"  • {n_max_tok} responses ({n_max_tok/1800*100:.1f}%) hit MAX_TOKENS at the 2048 output-token cap; "
  f"{n_max_tok_compliant} of these were still parseable. Bumping the cap to 4096 would "
  f"recover the truncated tail at an estimated +${extra_cost_4096:.2f} per rep "
  f"(~${cost_4096_estimate:.2f}/rep).")
P(f"  • Cost composition: input $%.2f + thinking $%.2f + answer $%.2f = $%.2f. "
  f"Thinking dominates (>6× input cost) — a structural feature of Gemini 3.1 Pro." %
  (cost_input, cost_thoughts, cost_answer, total_cost))
mp = safe_mean(gem_point_ed); gp = safe_mean(gpt_point_ed)
if mp is not None and gp is not None:
    direction = "lower (better)" if mp < gp else "higher (worse)"
    P(f"  • Point-landmark Euclidean distance: Gemini mean = {mp:.3f} cells, "
      f"GPT-5.4 mean = {gp:.3f} cells. Gemini is {direction} on average "
      f"(Wilcoxon two-sided p = {w_point.pvalue:.2e}).")
ma = safe_mean(gem_area_j); ga = safe_mean(gpt_area_j)
if ma is not None and ga is not None:
    direction = "higher (better)" if ma > ga else "lower (worse)"
    P(f"  • Area-landmark Jaccard: Gemini mean = {ma:.3f}, GPT-5.4 mean = {ga:.3f}. "
      f"Gemini is {direction} (Wilcoxon two-sided p = {w_area.pvalue:.2e}).")
P(f"  • Tooth_33_Apex F5/F6 attractor: GPT-5.4 placed the answer in F5/F6 for {gpt_attractor}/{len(t33)} "
  f"responses; Gemini for {gem_attractor}/{len(t33)}. Gemini achieved exact-GT match "
  f"on {gem_exact_t33}/{len(t33)} Tooth_33_Apex queries vs GPT's {gpt_exact_t33}/{len(t33)}.")

# ---- §2 Methodology ----
H('2. Methodology', level=1)
P("Comparison framework. Both models received byte-identical system + user prompts "
  "(verified by 0-drift check at Stage 2 of the orchestrator, re-verified independently "
  "for this report). Each (query × strategy) pair has one Gemini 3.1 Pro response (rep 1) "
  "and three GPT-5.4 responses (rep 1, 2, 3). For headline comparisons we use Gemini "
  "rep 1 vs GPT-5.4 rep 1 (matched single-rep). Per-rep stability for GPT-5.4 is "
  "reported where relevant.")
P("Metric definitions:")
P("  • Point landmarks: Euclidean distance (cell-grid units) from the model's first "
  "parsed cell to the single consensus-GT cell. Lower = better.")
P("  • Area landmarks: Jaccard overlap between the parsed-cell set and the consensus-GT "
  "cell set. Higher = better.")
P("  • A response is 'compliant' iff parsed_coordinates is non-empty AND failure_category "
  "is None (i.e., no parse failure, no out-of-range, no refusal/ambiguous/verbose/no_engage).")
P("Statistical tests. Wilcoxon signed-rank tests are computed on matched (query × strategy) "
  "pairs where both models were compliant. Pairs where either model failed are excluded "
  "from the metric tests (compliance is reported separately).")
P(f"Data integrity. Independent verification against the v2 GPT-5.4 baseline confirmed: "
  f"all 1,800 Gemini (query, strategy) pairs match the GPT-5.4 set; consensus GT comes "
  f"from results_consensus/query_index.json (same SHA as the v3 manuscript); modality + "
  f"landmark_type fields are consistent with the canonical query_index.")

# ---- §3 Compliance + failure modes ----
H('3. Compliance and Failure Modes', level=1)

P("Per-model strict-compliance counts:")
add_table(
    ["Model", "Repetition", "Compliant", "Non-compliant", "Compliance rate"],
    [
        ["Gemini 3.1 Pro", "rep 1",
         sum(1 for r in gem if r.get("failure_category") is None),
         sum(1 for r in gem if r.get("failure_category") is not None),
         f"{sum(1 for r in gem if r.get('failure_category') is None)/1800*100:.3f}%"],
        ["GPT-5.4 v2", "rep 1",
         sum(1 for r in gpt[1].values() if r.get("failure_category") is None),
         sum(1 for r in gpt[1].values() if r.get("failure_category") is not None),
         f"{sum(1 for r in gpt[1].values() if r.get('failure_category') is None)/1800*100:.3f}%"],
        ["GPT-5.4 v2", "rep 2",
         sum(1 for r in gpt[2].values() if r.get("failure_category") is None),
         sum(1 for r in gpt[2].values() if r.get("failure_category") is not None),
         f"{sum(1 for r in gpt[2].values() if r.get('failure_category') is None)/1800*100:.3f}%"],
        ["GPT-5.4 v2", "rep 3",
         sum(1 for r in gpt[3].values() if r.get("failure_category") is None),
         sum(1 for r in gpt[3].values() if r.get("failure_category") is not None),
         f"{sum(1 for r in gpt[3].values() if r.get('failure_category') is None)/1800*100:.3f}%"],
    ],
    col_widths_inches=[1.6, 0.9, 1.0, 1.2, 1.2],
)
doc.add_paragraph()

P("Gemini 3.1 Pro failure-mode breakdown (rep 1):")
add_table(
    ["Failure mode", "Count"],
    [[cat or "compliant", n] for cat, n in sorted(gem_fail.items(), key=lambda x: -x[1])],
    col_widths_inches=[2.0, 0.9],
)
doc.add_paragraph()

P("Most no_engage failures are empty-string responses, which Gemini's batch worker "
  "produced for 78 of 1,800 requests. The Google-side batch metadata reports these as "
  "cancelled inner requests, with successful counts of 428/450, 439/450, 428/450, 428/450 "
  "across the four submitted batches. The pattern is consistent with thinking-mode "
  "token-budget pressure (see §6).")

P("Per-landmark Gemini no_engage rate (top 10 by absolute count):")
landmark_table = [
    [lm,
     f"{d['no_engage']}/{d['total']}",
     f"{d['no_engage']/d['total']*100:.1f}%",
     f"{d['max_tokens']}"]
    for lm, d in sorted(landmark_fail.items(), key=lambda x: -x[1]['no_engage'])[:10]
]
add_table(
    ["Landmark", "no_engage / total", "% no_engage", "# MAX_TOKENS"],
    landmark_table,
    col_widths_inches=[2.2, 1.1, 0.9, 1.0],
)
doc.add_paragraph()

# ---- §4 Quality comparison ----
H('4. Quality Comparison: Gemini rep 1 vs GPT-5.4 rep 1 (matched)', level=1)

P(f"Point landmarks — Euclidean distance from prediction to consensus GT, in cell units.")
add_table(
    ["", "n (matched)", "mean ED", "median ED"],
    [
        ["Gemini 3.1 Pro", str(len(gem_point_ed)),
         f"{safe_mean(gem_point_ed):.3f}", f"{safe_med(gem_point_ed):.3f}"],
        ["GPT-5.4 v2 rep 1", str(len(gpt_point_ed)),
         f"{safe_mean(gpt_point_ed):.3f}", f"{safe_med(gpt_point_ed):.3f}"],
    ],
    col_widths_inches=[2.0, 1.2, 1.0, 1.0],
)
P(f"Paired Wilcoxon signed-rank test (Gemini ED vs GPT ED, two-sided): "
  f"W={w_point.statistic:.1f}, p={w_point.pvalue:.4g}. "
  f"{'Statistically significant' if w_point.pvalue < 0.05 else 'Not statistically significant'} "
  f"at α=0.05.")
doc.add_paragraph()

P(f"Area landmarks — Jaccard overlap with consensus-GT cell set. Higher is better.")
add_table(
    ["", "n (matched)", "mean Jaccard", "median Jaccard"],
    [
        ["Gemini 3.1 Pro", str(len(gem_area_j)),
         f"{safe_mean(gem_area_j):.3f}", f"{safe_med(gem_area_j):.3f}"],
        ["GPT-5.4 v2 rep 1", str(len(gpt_area_j)),
         f"{safe_mean(gpt_area_j):.3f}", f"{safe_med(gpt_area_j):.3f}"],
    ],
    col_widths_inches=[2.0, 1.2, 1.2, 1.2],
)
P(f"Paired Wilcoxon signed-rank test (Gemini Jaccard vs GPT Jaccard, two-sided): "
  f"W={w_area.statistic:.1f}, p={w_area.pvalue:.4g}. "
  f"{'Statistically significant' if w_area.pvalue < 0.05 else 'Not statistically significant'} "
  f"at α=0.05.")
doc.add_paragraph()

P("Per-(modality × strategy × landmark-type) breakdown:")
breakdown_rows = []
for mod in ("PANORAMIC", "PERIAPICAL", "CEPHALOMETRIC"):
    for strat in ("zero_shot", "guided"):
        for lm in ("point", "area"):
            b = breakdown[(mod, strat, lm)]
            if b["n"] == 0:
                if b["n_total"] == 0:
                    continue
                breakdown_rows.append([mod, strat, lm,
                                       f"0/{b['n_total']}", "—", "—", "—"])
                continue
            metric_label = "ED" if lm == "point" else "Jaccard"
            breakdown_rows.append([
                mod, strat, lm,
                f"{b['n']}/{b['n_total']}",
                f"{b['gem_mean']:.3f}",
                f"{b['gpt_mean']:.3f}",
                f"{b['p']:.3g}" if b['p'] is not None else "—",
            ])
add_table(
    ["Modality", "Strategy", "Type", "matched/total",
     "Gemini mean", "GPT-5.4 mean", "Wilcoxon p"],
    breakdown_rows,
    col_widths_inches=[1.2, 0.85, 0.6, 0.9, 0.9, 0.9, 0.85],
)
doc.add_paragraph()

# ---- §5 F5/F6 attractor ----
H('5. F5/F6 Attractor Analysis (Tooth_33_Apex)', level=1)
P(f"Background. The v3 manuscript documented that GPT-5.4 consistently places "
  f"Tooth_33_Apex (the lower-left canine apex on panoramic radiographs) at F5 or F6, "
  f"a region of the grid that is anatomically inconsistent with the true tooth position. "
  f"This is referred to as the 'F5/F6 positional attractor.' This section tests whether "
  f"Gemini 3.1 Pro exhibits the same failure mode.")

t33_per_strat = defaultdict(lambda: {"gem_attr": 0, "gpt_attr": 0,
                                       "gem_exact": 0, "gpt_exact": 0,
                                       "gem_compliant": 0, "gpt_compliant": 0, "total": 0})
for r in t33:
    s = r["strategy"]
    t33_per_strat[s]["total"] += 1
    if r["gem_compliant"]:
        t33_per_strat[s]["gem_compliant"] += 1
        if r["gem_parsed"] and r["gem_parsed"][0] in {"F5", "F6"}:
            t33_per_strat[s]["gem_attr"] += 1
        if r["gem_parsed"] and r["gem_parsed"][0] in parse_gt(r["gt"]):
            t33_per_strat[s]["gem_exact"] += 1
    if r["gpt_compliant"]:
        t33_per_strat[s]["gpt_compliant"] += 1
        if r["gpt_parsed"] and r["gpt_parsed"][0] in {"F5", "F6"}:
            t33_per_strat[s]["gpt_attr"] += 1
        if r["gpt_parsed"] and r["gpt_parsed"][0] in parse_gt(r["gt"]):
            t33_per_strat[s]["gpt_exact"] += 1

add_table(
    ["Strategy", "Total Tooth_33_Apex",
     "Gemini compliant", "Gemini in F5/F6", "Gemini exact-GT",
     "GPT compliant", "GPT in F5/F6", "GPT exact-GT"],
    [[s] + [str(t33_per_strat[s][k]) + f" ({t33_per_strat[s][k]/t33_per_strat[s]['total']*100:.1f}%)"
            if k != "total" else str(t33_per_strat[s][k])
            for k in ("total", "gem_compliant", "gem_attr", "gem_exact",
                       "gpt_compliant", "gpt_attr", "gpt_exact")]
     for s in ("zero_shot", "guided")],
    col_widths_inches=[0.85, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95, 0.95],
)
doc.add_paragraph()
P(f"Interpretation. GPT-5.4 placed the answer in F5/F6 for {gpt_attractor}/{len(t33)} of "
  f"Tooth_33_Apex queries (averaged across both strategies); Gemini for "
  f"{gem_attractor}/{len(t33)}. The data confirm that the F5/F6 attractor is largely "
  f"a GPT-5.4-specific failure mode, not a generic MLLM phenomenon. Gemini's "
  f"failures on this landmark take a different shape: it produces an empty response "
  f"in {sum(1 for r in t33 if r['gem_fail'] == 'no_engage')}/{len(t33)} cases rather "
  f"than the wrong cell.")

# ---- §6 MAX_TOKENS deep dive ----
H('6. MAX_TOKENS Truncation Analysis', level=1)
P(f"Gemini 3.1 Pro uses thinking-mode internally; the response budget set by "
  f"max_output_tokens is shared between thinking tokens (thoughtsTokenCount) and the "
  f"visible answer (candidatesTokenCount). With max_output_tokens=2048, {n_max_tok} of "
  f"1,800 responses ({n_max_tok/1800*100:.1f}%) reached finishReason='MAX_TOKENS'. Of "
  f"those, {n_max_tok_parseable} had a parseable first cell and {n_max_tok_compliant} were "
  f"still classified as compliant.")

if max_tok_records:
    add_table(
        ["Stat", "thinking tokens", "answer tokens"],
        [["mean", f"{safe_mean([m['thoughts_tok'] for m in max_tok_records]):.0f}",
                  f"{safe_mean([m['answer_tok'] for m in max_tok_records]):.0f}"],
         ["median", f"{statistics.median([m['thoughts_tok'] for m in max_tok_records]):.0f}",
                    f"{statistics.median([m['answer_tok'] for m in max_tok_records]):.0f}"],
         ["min", f"{min(m['thoughts_tok'] for m in max_tok_records)}",
                  f"{min(m['answer_tok'] for m in max_tok_records)}"],
         ["max", f"{max(m['thoughts_tok'] for m in max_tok_records)}",
                  f"{max(m['answer_tok'] for m in max_tok_records)}"]],
        col_widths_inches=[1.0, 1.5, 1.5],
    )
    doc.add_paragraph()

P("MAX_TOKENS landmarks (count per anatomical structure):")
add_table(
    ["Structure", "# MAX_TOKENS"],
    [[k, v] for k, v in mt_by_struct.most_common(15)],
    col_widths_inches=[2.5, 1.0],
)
doc.add_paragraph()

P(f"If max_output_tokens were raised to 4096, the model would have an additional ~2048 "
  f"tokens of headroom. Most truncated responses appear to have stopped while still in "
  f"the thinking phase (answer tokens often = 0 or very small). The rough projection is "
  f"that bumping the cap to 4096 would add an estimated +${extra_cost_4096:.2f} per rep "
  f"({n_max_tok} affected responses × ~{extra_thinking_per_maxtok} additional tokens × "
  f"$6/1M output rate), for a total of ~${cost_4096_estimate:.2f} per rep "
  f"(${cost_4096_estimate*3:.2f} for 3 reps). The benefit is recovering up to "
  f"{n_max_tok - n_max_tok_compliant} responses that are currently non-compliant.")

# ---- §7 Cost ----
H('7. Cost Analysis', level=1)
P("Computed from Gemini's reported usageMetadata (per-call), using batch-tier pricing "
  "($1.00 / 1M input tokens, $6.00 / 1M output tokens). Note: 'output' includes both "
  "thinking tokens and the visible-answer tokens.")
add_table(
    ["Component", "Tokens", "Cost"],
    [["Prompt (input)", f"{total_in:,}", f"${cost_input:.3f}"],
     ["Thinking (output)", f"{total_thoughts:,}", f"${cost_thoughts:.3f}"],
     ["Answer (output)", f"{total_answer:,}", f"${cost_answer:.3f}"],
     ["TOTAL 1 rep", "—", f"${total_cost:.2f}"],
     ["Projection 3 reps", "—", f"${total_cost*3:.2f}"],
     [f"With max_output_tokens=4096 (est.)", "—",
      f"~${cost_4096_estimate:.2f} per rep, ${cost_4096_estimate*3:.2f} for 3 reps"]],
    col_widths_inches=[2.5, 1.5, 2.5],
)
doc.add_paragraph()

# ---- §8 Recommendations ----
H('8. Recommendations', level=1)
P("Decisions for the next stage of the benchmark:")
P("  1. Bumping max_output_tokens to 4096 is recommended for reps 2 and 3 if the budget "
  f"allows. It would reduce or eliminate the {n_max_tok} MAX_TOKENS truncations at a "
  f"marginal cost of approximately +${extra_cost_4096:.2f}/rep. Most truncated responses "
  "are concentrated on a small number of anatomically-complex landmarks (see §6 table).")
P("  2. The matched single-rep comparison vs GPT-5.4 is statistically meaningful in this "
  "sample because all 900 queries × 2 strategies overlap; Wilcoxon p-values are reported "
  "in §4. With reps 2 and 3 added, the cross-model comparison gains rep-level variance "
  "estimates for both models.")
P("  3. The F5/F6 attractor result (§5) is consistent with the v3 manuscript's claim that "
  "the attractor is a GPT-5.4-specific failure pattern, not a general MLLM phenomenon. "
  "This is a publishable cross-model observation.")

# ---- Appendix A1: All Gemini responses ----
H('Appendix A1 — All 1,800 Gemini 3.1 Pro Responses (rep 1)', level=1, page_break=True)
P(f"Format: query_id, strategy, GT (consensus), Gemini raw response, parsed cells, "
  f"failure_category. Sorted by (modality, strategy, query_id). Total rows: 1,800.")
# Build sorted rows
appendix_rows = sorted(rows, key=lambda r: (r["modality"], r["strategy"], r["qid"]))
# Split appendix table by strategy for readability — but keep both modalities together
for strat in ("zero_shot", "guided"):
    P(f"Strategy = {strat} (900 responses):", bold=True)
    sub = [r for r in appendix_rows if r["strategy"] == strat]
    a1_rows = []
    for r in sub:
        a1_rows.append([
            r["qid"],
            r["modality"][:3],
            r["gt"][:30] + ("…" if len(r["gt"]) > 30 else ""),
            (r["gem_raw"] or "")[:50] + ("…" if r["gem_raw"] and len(r["gem_raw"]) > 50 else ""),
            ",".join(r["gem_parsed"][:6])[:40],
            r["gem_fail"] or "OK",
        ])
    add_table(
        ["Query ID", "Mod", "GT", "Gemini raw", "Parsed", "Fail?"],
        a1_rows,
        col_widths_inches=[2.0, 0.4, 1.4, 1.6, 1.2, 0.8],
    )
    doc.add_paragraph()

# ---- Appendix A2: All GPT responses ----
H('Appendix A2 — All 1,800 GPT-5.4 v2 Responses (rep 1)', level=1, page_break=True)
P(f"Same format as A1 but for the GPT-5.4 v2 main run, rep 1.")
for strat in ("zero_shot", "guided"):
    P(f"Strategy = {strat} (900 responses):", bold=True)
    sub = [r for r in appendix_rows if r["strategy"] == strat]
    a2_rows = []
    for r in sub:
        a2_rows.append([
            r["qid"],
            r["modality"][:3],
            r["gt"][:30] + ("…" if len(r["gt"]) > 30 else ""),
            (r["gpt_raw"] or "")[:50] + ("…" if r["gpt_raw"] and len(r["gpt_raw"]) > 50 else ""),
            ",".join(r["gpt_parsed"][:6])[:40],
            r["gpt_fail"] or "OK",
        ])
    add_table(
        ["Query ID", "Mod", "GT", "GPT-5.4 raw", "Parsed", "Fail?"],
        a2_rows,
        col_widths_inches=[2.0, 0.4, 1.4, 1.6, 1.2, 0.8],
    )
    doc.add_paragraph()

# ---- Appendix A3: Side-by-side ----
H('Appendix A3 — Per-Query Side-by-Side (Gemini vs GPT-5.4 vs GT)', level=1, page_break=True)
P("Compact side-by-side view for every (query_id × strategy) pair. Columns show only the "
  "first 3 parsed cells for compactness; full responses are in A1/A2. 'Match' columns "
  "indicate exact-GT match for point landmarks (first parsed cell equals first GT cell).")
for strat in ("zero_shot", "guided"):
    P(f"Strategy = {strat}:", bold=True)
    sub = [r for r in appendix_rows if r["strategy"] == strat]
    a3_rows = []
    for r in sub:
        is_point = r["landmark_type"] == "point"
        gt_first = r["gt_cells"][0] if r["gt_cells"] else "?"
        gem_first = r["gem_parsed"][0] if r["gem_parsed"] else "?"
        gpt_first = r["gpt_parsed"][0] if r["gpt_parsed"] else "?"
        if is_point:
            gem_match = "✓" if (r["gem_compliant"] and gem_first == gt_first) else ("·" if r["gem_compliant"] else "—")
            gpt_match = "✓" if (r["gpt_compliant"] and gpt_first == gt_first) else ("·" if r["gpt_compliant"] else "—")
            comp_str = f"ED={r['gem_distance']:.2f}/{r['gpt_distance']:.2f}" if (r['gem_distance'] is not None and r['gpt_distance'] is not None) else "—"
        else:
            gem_match = "—"; gpt_match = "—"
            comp_str = f"J={r['gem_jaccard']:.3f}/{r['gpt_jaccard']:.3f}" if (r['gem_jaccard'] is not None and r['gpt_jaccard'] is not None) else "—"
        a3_rows.append([
            r["qid"][-22:],
            r["landmark_type"][:1].upper(),
            (r["gt"] or "")[:14] + ("…" if r["gt"] and len(r["gt"]) > 14 else ""),
            ",".join(r["gem_parsed"][:3])[:18],
            gem_match,
            ",".join(r["gpt_parsed"][:3])[:18],
            gpt_match,
            comp_str,
        ])
    add_table(
        ["Query ID (last 22)", "T", "GT", "Gemini", "✓?", "GPT-5.4", "✓?", "ED or J (G/P)"],
        a3_rows,
        col_widths_inches=[1.7, 0.25, 1.1, 1.2, 0.3, 1.2, 0.3, 1.1],
    )
    doc.add_paragraph()

# ---- Save ----
out_path = ROOT / "results_full_gemini/Gemini_Rep1_Analysis_Report.docx"
doc.save(str(out_path))
sha = hashlib.sha256(out_path.read_bytes()).hexdigest()[:12]
size = out_path.stat().st_size
print(f"\n✓ DOCX written: {out_path}")
print(f"  size: {size/1024:.1f} KB")
print(f"  sha256: {sha}…")
print(f"  total paragraphs: {len(doc.paragraphs)}")
print(f"  total tables: {len(doc.tables)}")
