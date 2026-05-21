"""
Generate a standalone docx listing every GPT-5.4 response from the full run.

Format follows v2 pilot Appendix B:
  - One table per (modality × strategy × landmark_type) combination
  - Columns: Image | Landmark | GT | Run 1 | Run 2 | Run 3
  - Headers 12pt bold; body 11pt for point sections, 10pt for area sections
    (area cells can hold long comma-separated cell lists)
  - Captions above tables
  - Letter portrait, totals fit ~6.69" content width

All values come from /tmp/full_run_records_v2.pkl. Each rep entry stores the
exact raw_response string the parser saw; that is what is shown here, with
"<empty>" for compliance failures (no parseable cells).
"""
from __future__ import annotations
import pickle
from pathlib import Path
from datetime import datetime, UTC
from collections import defaultdict
import subprocess

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUT  = ROOT / 'results_full' / 'Full_Run_Model_Responses_Appendix.docx'

records = pickle.load(open('/tmp/full_run_records_v2.pkl','rb'))
git_sha = subprocess.check_output(['git','-C',str(ROOT),'rev-parse','HEAD']).decode().strip()

# ── Document setup ──────────────────────────────────────────────────
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)

def H(text, level=1):
    return doc.add_heading(text, level=level)

def P(text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    return p

def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    return p

def shade_cell(cell, fill='D9E1F2'):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), fill)
    tcPr.append(sh)

def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'4')
        b.set(qn('w:color'),'808080')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_table_fixed_layout(t):
    """Force fixed column widths so Word respects them and wraps long
    no-space strings (e.g. 'H15,H16,G13,G14,...') instead of squeezing
    the neighbouring columns. Without this, autofit overrides widths."""
    tblPr = t._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        t._element.insert(0, tblPr)
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

def add_response_table(headers, rows, col_widths, *, header_size=12, body_size=11):
    n_cols = len(headers)
    t = doc.add_table(rows=len(rows)+1, cols=n_cols)
    t.style = 'Light Grid'
    set_table_fixed_layout(t)
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(header_size)
        shade_cell(c, 'D9E1F2')
        set_cell_borders(c)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i+1, j)
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(body_size)
            set_cell_borders(c)
    for j, w in enumerate(col_widths):
        for r in t.rows:
            r.cells[j].width = Inches(w)
    return t

def fmt_response(raw):
    """Return the raw response string, or "<empty>" for compliance failures."""
    if raw is None or not str(raw).strip():
        return '<empty>'
    return str(raw).strip()

# ── Title page ──────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('GPT-5.4 Model Responses — Full Benchmark')
r.bold = True; r.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Companion data appendix to the Full Run Results Report')
r.italic = True; r.font.size = Pt(13)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(
    f'Generated {datetime.now(UTC).strftime("%Y-%m-%d")} from raw API outputs '
    f'(commit {git_sha[:12]}, sandbox: results_full)'
)
r.font.size = Pt(10); r.italic = True
doc.add_paragraph()

# ── Introduction ────────────────────────────────────────────────────
H('Overview', level=1)
n_records = len(records)
total_rep_responses = sum(len([x for x in r['raw_responses'] if x is not None]) for r in records)

P(f'This appendix lists every GPT-5.4 response for all {n_records:,} unique '
  f'(query, strategy) pairs across the full benchmark, including each of the three repetitions. '
  f'Total cells of model output: {n_records*3:,}. The data are partitioned into '
  f'eight tables, one per (modality × strategy × landmark type) combination:')

groups = [
    ('PANORAMIC',     'zero_shot', 'point', 'B1'),
    ('PANORAMIC',     'guided',    'point', 'B2'),
    ('PANORAMIC',     'zero_shot', 'area',  'B3'),
    ('PANORAMIC',     'guided',    'area',  'B4'),
    ('PERIAPICAL',    'zero_shot', 'point', 'B5'),
    ('PERIAPICAL',    'guided',    'point', 'B6'),
    ('CEPHALOMETRIC', 'zero_shot', 'point', 'B7'),
    ('CEPHALOMETRIC', 'guided',    'point', 'B8'),
]

for mod, strat, ltype, label in groups:
    n = sum(1 for r in records if r['modality']==mod and r['strategy']==strat and r['landmark_type']==ltype)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'Table {label}: {mod}, {strat}, {ltype} — {n} queries × 3 reps = {n*3} responses.')

P('Within each table, rows are ordered first by image ID and then by landmark, so all queries on '
  'the same image appear adjacently. The "GT" column is the ground truth from OMFR_1. The "Run N" '
  'columns hold the verbatim raw response string the parser received from the model. Empty model '
  'responses (the four strict-compliance failures documented in the main report\'s Appendix B) are '
  'shown as "<empty>" — note that all four were in fact reversed-coordinate strings such as "12F" '
  'and the parser logged them with the original raw text; that exact original text appears in '
  'this column.', italic=False)

P('Compliance footnote: cells are shown verbatim, so a few responses display reversed-coordinate '
  'strings like "12F" or "6F". These are the strict-parser failures: they encode the correct cell '
  'in inverted form. See the main report\'s Appendix B for the full failure list and Section 8.4 '
  'for the parser-tolerance discussion.', italic=True)

# ── One table per group ─────────────────────────────────────────────
PT_SIZES   = {'header_size': 12, 'body_size': 11}
AREA_SIZES = {'header_size': 12, 'body_size': 10}

# Letter portrait usable width = 6.70" (8.5" page, 0.9" margins each side).
# Same six-column layout for both point and area tables, matching v2 pilot
# Appendix B convention. With fixed table layout (set_table_fixed_layout
# above), Word respects these widths and wraps long no-space cell lists
# inside their column instead of squeezing neighbours.
RESP_WIDTHS = [0.95, 1.65, 1.30, 0.93, 0.93, 0.93]   # = 6.69"
PT_WIDTHS   = RESP_WIDTHS
AREA_WIDTHS = RESP_WIDTHS

# Sort key: image_id, then a stable landmark order
landmark_order = {
    # PAN
    'Mental_Foramen_L': 0, 'Condylar_Head_R': 1, 'Tooth_33_Apex': 2,
    'Mandibular_Canal_L': 3, 'Maxillary_Sinus_R': 4, 'External_Oblique_Ridge_R': 5,
    # PA
    'Tooth_36_Distal_Apex': 0, 'Tooth_36_Distal_CEJ': 1, 'Tooth_36_Mesial_CEJ': 2,
    # CEPH
    'Sella_S': 0, 'Nasion_N': 1, 'Menton_Me': 2,
}

for mod, strat, ltype, label in groups:
    items = [r for r in records
             if r['modality']==mod and r['strategy']==strat and r['landmark_type']==ltype]
    items.sort(key=lambda r: (r['image_id'], landmark_order.get(r['structure'], 99)))

    doc.add_page_break()
    H(f'Table {label} — {mod}, {strat}, {ltype} landmarks', level=2)
    caption(f'Table {label}: All GPT-5.4 raw responses for the {mod} {strat} {ltype} group '
            f'({len(items)} queries × 3 reps = {len(items)*3} responses).')

    rows = []
    for r in items:
        runs = [fmt_response(x) for x in r['raw_responses']]
        # pad to 3 if any reps missing (shouldn't happen for v2)
        while len(runs) < 3:
            runs.append('<empty>')
        rows.append([r['image_id'], r['structure'], r['gt'], runs[0], runs[1], runs[2]])

    if ltype == 'area':
        add_response_table(
            ['Image', 'Landmark', 'GT', 'Run 1', 'Run 2', 'Run 3'],
            rows, AREA_WIDTHS, **AREA_SIZES,
        )
    else:
        add_response_table(
            ['Image', 'Landmark', 'GT', 'Run 1', 'Run 2', 'Run 3'],
            rows, PT_WIDTHS, **PT_SIZES,
        )

# ── Save ────────────────────────────────────────────────────────────
doc.save(str(OUT))

import os
print(f'Wrote {OUT}')
print(f'  Paragraphs: {len(doc.paragraphs)}')
print(f'  Tables: {len(doc.tables)}')
total_data_rows = sum(len(t.rows)-1 for t in doc.tables)
print(f'  Total data rows across all tables: {total_data_rows}')
print(f'  Size: {os.path.getsize(OUT):,} bytes')
