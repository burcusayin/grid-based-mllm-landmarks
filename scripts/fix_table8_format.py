"""Re-format Table 8 in DMFR_Grid_Study_v2.docx to exactly match the style
of Table 6 (GPT vs students) — same column layout, group subheaders, font,
alignment, value formatting, and Bonferroni p convention.

DOES NOT touch any other content. Preserves all user revisions to v2.

Format matched against Table 6:
  - Times New Roman 10pt
  - Header row: bold, center-aligned (except column 0)
  - Group subheader rows: bold italic, only column 0 populated
    ("Zero-shot strategy" / "Guided strategy")
  - Data rows: column 0 left-aligned (modality label like
    "Cephalometric (point, ED)"); columns 1-6 center-aligned
  - Δ and r: 2 decimal places with en-dash for negative ("−0.05" not "-0.05")
  - p column convention:
     * If Bonferroni-corrected p (raw × 4) < 1.0 AND significant → show raw × 4 with ***/**/* marker
     * If NS after Bonferroni (raw × 4 > 0.05) → show raw with (NS) marker
       (this also covers cases where raw × 4 ≥ 1.0)
"""
from __future__ import annotations
import json
import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "docs" / "submission" / "DMFR_Grid_Study_v2.docx"

# ── Load source data ──────────────────────────────────────────────
gem_vs_stu = json.load(open(ROOT / "results_full_gemini" / "gemini_vs_student.json"))

# ── Number formatting helpers ─────────────────────────────────────
SUP_TRANSLATE = str.maketrans("0123456789-+", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻⁺")

def fmt_neg(x: float, digits: int = 2) -> str:
    """Format x with en-dash for negative (Table 6 convention)."""
    if x >= 0:
        return f"+{x:.{digits}f}"
    return f"−{abs(x):.{digits}f}"

def fmt_p_table6(raw_p: float, k: int = 4) -> tuple[str, str]:
    """Return (display_p, sig_marker) following Table 6 convention.

    Rule (verified against Table 6 numbers):
      - Bonferroni-corrected p = min(raw × k, 1.0).
      - If corrected p < 0.05 → display the corrected p (Bonferroni × k)
                                 with appropriate stars.
      - If corrected p ≥ 0.05 → display the raw p with "(NS)" marker.
    """
    corr = min(1.0, raw_p * k)
    if corr < 0.05:
        # Significant after Bonferroni — show corrected p
        if corr < 0.001:
            marker = "***"
        elif corr < 0.01:
            marker = "**"
        else:
            marker = "*"
        # Choose display format
        if corr >= 0.01:
            disp = f"{corr:.3f}"
        else:
            # Scientific: "X.XX×10⁻Y"
            s = f"{corr:.2e}"
            mant, exp = s.split('e')
            exp_int = int(exp)
            disp = f"{mant}×10{str(exp_int).translate(SUP_TRANSLATE)}"
        return f"{disp} {marker}", marker
    else:
        # NS after Bonferroni — show raw p with (NS) marker
        if raw_p >= 0.01:
            disp = f"{raw_p:.3f}"
        else:
            s = f"{raw_p:.2e}"
            mant, exp = s.split('e')
            exp_int = int(exp)
            disp = f"{mant}×10{str(exp_int).translate(SUP_TRANSLATE)}"
        return f"{disp} (NS)", "NS"


# ── Build Table 8 data exactly matching Table 6's row structure ───
def gvs(strat: str, key: str) -> dict:
    return gem_vs_stu["per_strategy"][strat]["per_modality"][key]

ROW_ORDER = [
    ("Cephalometric (point, ED)", "CEPHALOMETRIC_point", "ED"),
    ("Periapical (point, ED)",   "PERIAPICAL_point",   "ED"),
    ("Panoramic (point, ED)",    "PANORAMIC_point",    "ED"),
    ("Panoramic (area, Jaccard)", "PANORAMIC_area",     "Jaccard"),
]

def row_for(strat: str, label: str, key: str, metric: str) -> list[str]:
    d = gvs(strat, key)
    n = d["n_total"]
    gem_v = d.get("mean_gemini_ed", d.get("mean_gemini_jaccard"))
    stu_v = d.get("mean_student_ed", d.get("mean_student_jaccard"))
    delta = d["mean_delta"]
    r_val = d["rank_biserial_r"]
    p_disp, _ = fmt_p_table6(d["p"], k=4)
    return [
        label,
        str(n),
        f"{gem_v:.2f}",
        f"{stu_v:.2f}",
        fmt_neg(delta, 2),
        p_disp,
        fmt_neg(r_val, 2),
    ]

# 11 rows total: header + 2 group headers + 4 zero-shot + 4 guided
header = ["Group", "n", "Gemini 3.1 Pro Mean", "Student Mean", "Mean Δᵃ", "p (Bonf. ×4)ᵇ", "rᶜ"]
group1 = ["Zero-shot strategy", "", "", "", "", "", ""]
group2 = ["Guided strategy",    "", "", "", "", "", ""]

zs_rows = [row_for("zero_shot", lbl, k, m) for lbl, k, m in ROW_ORDER]
g_rows  = [row_for("guided",    lbl, k, m) for lbl, k, m in ROW_ORDER]

all_rows = [header, group1] + zs_rows + [group2] + g_rows

# Print preview before applying
print("=" * 100)
print("New Table 8 content (Table 6 format):")
print("=" * 100)
for i, row in enumerate(all_rows):
    print(f"  Row {i:2d}: " + " │ ".join(f"{c:<26}" for c in row))
print()


# ── Edit the docx ─────────────────────────────────────────────────
doc = Document(DOCX)
old_table = doc.tables[7]   # Table 8 (zero-indexed)
template = doc.tables[5]    # Table 6 (we'll clone its STYLE + layout)

# Strategy: clone Table 6's XML wholesale, replace cell contents.
# This guarantees identical formatting (borders, widths, font, alignment).

# Get position of old Table 8 in body
old_elem = old_table._element
parent = old_elem.getparent()
old_index = list(parent).index(old_elem)

# Deep-copy Table 6's XML to use as the new Table 8 base
new_tbl = copy.deepcopy(template._element)

# Now we need to:
#  - Make sure the new table has the right number of rows (11)
#  - Replace cell contents with new_rows
#  - Update Bonferroni multiplier in header if needed (Table 6 has ×4, same)

# Count rows in cloned table (should be 11 — same as Table 6)
rows = new_tbl.findall(qn("w:tr"))
print(f"Cloned table has {len(rows)} rows (expecting 11)")

# Iterate rows, replace cell content with new_rows
for i, tr in enumerate(rows):
    cells = tr.findall(qn("w:tc"))
    if i >= len(all_rows):
        break
    desired = all_rows[i]
    for j, tc in enumerate(cells):
        if j >= len(desired):
            break
        # Find the paragraph(s) in this cell
        paras = tc.findall(qn("w:p"))
        if not paras:
            continue
        # Clear all runs in the first paragraph (and remove subsequent paragraphs)
        first_p = paras[0]
        # Remove extra paragraphs (keep first)
        for extra_p in paras[1:]:
            tc.remove(extra_p)
        # Clear all runs and child elements other than pPr
        for child in list(first_p):
            if child.tag != qn("w:pPr"):
                first_p.remove(child)
        # Add a new run with the desired text, preserving formatting
        new_text = desired[j]
        if not new_text:
            continue
        new_r = OxmlElement("w:r")
        new_rPr = OxmlElement("w:rPr")
        # Apply Times New Roman 10pt
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        new_rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "20")  # 20 half-points = 10pt
        new_rPr.append(sz)
        # Bold for header row (i=0) and group rows (i=1, 6)
        if i == 0:
            b = OxmlElement("w:b")
            new_rPr.append(b)
        elif i in (1, 6):
            # Group header rows: bold italic (only first column populated)
            b = OxmlElement("w:b")
            new_rPr.append(b)
            it = OxmlElement("w:i")
            new_rPr.append(it)
        new_r.append(new_rPr)
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        first_p.append(new_r)

# Replace old Table 8 with new
parent.remove(old_elem)
parent.insert(old_index, new_tbl)

# Save
doc.save(DOCX)
print(f"\nSaved: {DOCX}")
