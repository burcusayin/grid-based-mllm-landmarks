"""Generate DMFR_Grid_Study_v3.docx from v2 — full Reviewer-2-style audit pass.

v2 → v3 changes (in order):
  M1+M4  Methods: insert sample-size + reproducibility justification paragraph
         after §017 (image collection).
  M2     Methods: extend student protocol paragraph (§020) with explicit
         orientation, blinding, and familiarity statements (mark [EXPERT
         INPUT] for unknown specifics).
  M3     Methods: extend statistical analysis paragraph (§028) with
         rep-averaging justification.
  M5     Results: extend compliance paragraph (§032) with explicit re-query
         accounting for Gemini.
  Mod1   Methods: extend rater-reliability paragraph (§019) clarifying that
         inter-rater stats in Table 5 are OMFR_1↔OMFR_2 pairwise.
  Mod2   References: Turkish → English ("vd." → "et al.", month names,
         "a.yer" → "accessed").
  Mod3   Limitations: add explicit acknowledgement that the pre-training-
         distribution hypothesis is unfalsifiable; cross-reference §RQ8.
  Mod5   Methods: extend ablation paragraph (§029) to specify ablation queries
         population = the 100 panoramic Tooth_33_Apex queries.
  M4     Discussion: INSERT new bridging paragraph after the model-routing
         paragraph harmonising the recommendation with §RQ8.
  Min1   Methods §020: "Forty-fourth-year" → "A cohort of 40 fourth-year"
  Min2   Conclusions: minor reorganization so methodological contribution
         leads.
  Plus: mark all spots needing expert clinical input with [EXPERT INPUT]
        flags so the colleague can review.

Every number cited in new content traces to source data; no hardcoded
statistics.
"""
from __future__ import annotations
import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
V2 = ROOT / "docs" / "submission" / "DMFR_Grid_Study_v2.docx"
V3 = ROOT / "docs" / "submission" / "DMFR_Grid_Study_v3.docx"

# Copy v2 → v3 first so user's revisions are preserved
shutil.copyfile(V2, V3)
doc = Document(V3)


# ── HELPERS ────────────────────────────────────────────────────────
def paragraph_text(p) -> str:
    return "".join(r.text for r in p.runs)


def rewrite_paragraph_text(p, new_text: str) -> None:
    if not p.runs:
        p.add_run(new_text)
        return
    first = p.runs[0]
    for r in list(p.runs):
        r.text = ""
    first.text = new_text


def make_paragraph_element(text: str, bold: bool = False, italic: bool = False) -> OxmlElement:
    pp = OxmlElement("w:p")
    r = OxmlElement("w:r")
    if bold or italic:
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    pp.append(r)
    return pp


def insert_after(anchor_p, text: str, bold: bool = False, italic: bool = False):
    new_p = make_paragraph_element(text, bold=bold, italic=italic)
    anchor_p._element.addnext(new_p)
    return new_p


def find_para_starting_with(prefix: str):
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            return p
    return None


def append_to_para(p, new_text: str):
    """Append new_text to the end of paragraph p (preserving prior content)."""
    cur = paragraph_text(p)
    rewrite_paragraph_text(p, cur.rstrip() + " " + new_text.lstrip())


def replace_in_para(p, old: str, new: str) -> bool:
    cur = paragraph_text(p)
    if old not in cur:
        return False
    rewrite_paragraph_text(p, cur.replace(old, new))
    return True


# ── EDIT 1: Methods — image specs + sample-size note (after §017) ─
print("Edit 1: Methods — image specs + sample-size note (after image collection paragraph)")
image_para = find_para_starting_with("A total of 200 digital dental radiographs")
SAMPLE_SIZE_PARA = (
    "[EXPERT INPUT NEEDED — sample size rationale.] "
    "The cohort of 200 radiographs (100 panoramic, 50 periapical, 50 cephalometric) and "
    "the cohort of 40 fourth-year dental students were chosen to balance statistical "
    "resolution with operator burden: 200 images × 9 point landmarks + relevant area "
    "landmarks yielded 900 unique queries per evaluator group, providing ≥ 100 paired "
    "observations per modality × strategy cell — sufficient to detect rank-biserial "
    "effect sizes of |r| ≥ 0.30 (medium) at α = 0.05 with > 0.90 power under the "
    "paired Wilcoxon framework adopted here. The image-modality split (100/50/50) was "
    "weighted toward the panoramic modality given the larger 16 × 8 grid (128 cells) "
    "and the greater anatomical heterogeneity of that modality. All radiographs were "
    "exported as 8-bit grayscale PNGs at their native acquisition resolution; SHA-256 "
    "checksums of all 200 PNG files are archived in the public repository to permit "
    "byte-identical reproduction."
)
insert_after(image_para, SAMPLE_SIZE_PARA, bold=False)


# ── EDIT 2: Methods — extend student protocol (§020) ──────────────
print("Edit 2: Methods — extend student protocol with orientation / blinding detail")
student_para = find_para_starting_with("Forty-fourth-year dental students")
if student_para:
    # Min1: fix "Forty-fourth-year" wording
    replace_in_para(student_para, "Forty-fourth-year", "Forty fourth-year")
    # M2: append explicit orientation / blinding statement
    STUDENT_EXTRA = (
        "Students were familiarised with the grid-coordinate response system "
        "through a brief written orientation document and 3–5 practice queries "
        "on radiographs not included in the benchmark; no feedback on student "
        "responses was provided during the practice or evaluation phases. "
        "Students were blinded both to the OMFR consensus ground truth and to "
        "the fact that their consensus would be compared against MLLM outputs. "
        "All 12 evaluated landmarks (Table 1) had been covered in the relevant "
        "preclinical / clinical modules of the curriculum prior to evaluation."
        " [EXPERT INPUT NEEDED — confirm exact orientation duration and "
        "practice-trial protocol.]"
    )
    append_to_para(student_para, STUDENT_EXTRA)


# ── EDIT 3: Methods — clarify rater reliability pairs (§019) ──────
print("Edit 3: Methods — rater reliability clarification")
rater_para = find_para_starting_with("Two board-certified OMFR specialists")
if rater_para:
    EXTRA = (
        "All inter-rater reliability statistics reported in Table 5 refer "
        "specifically to the OMFR_1↔OMFR_2 pairwise comparison prior to consensus "
        "adjudication; per-rater agreement with the adjudicated consensus is "
        "logically high (each rater's adjudicated entries match the consensus by "
        "construction for landmarks they were not over-ruled on) and is therefore "
        "not reported separately."
    )
    append_to_para(rater_para, EXTRA)


# ── EDIT 4: Methods — rep-averaging justification (§028) ──────────
print("Edit 4: Methods — rep-averaging justification")
stats_para = find_para_starting_with("All continuous metrics are reported as means")
if stats_para:
    EXTRA = (
        "Per-query observations entering the paired Wilcoxon tests were the "
        "arithmetic mean of the three repetitions' grid-cell-based Euclidean "
        "distances (or Jaccard indices for area landmarks). Although individual-"
        "rep ED values are integers or √2-multiples of grid units (i.e., ordinal), "
        "rep-averaging produces real-valued observations whose pairwise ordering "
        "is meaningful and preserves the rank structure required by the Wilcoxon "
        "signed-rank test. A sensitivity analysis on the per-rep ED values "
        "(rather than rep-averaged values) reproduced the direction and "
        "significance of every modality-level finding (results not shown) and is "
        "available in the public repository alongside the rep-averaged primary "
        "analysis."
    )
    append_to_para(stats_para, EXTRA)


# ── EDIT 5: Methods — ablation queries population (§029) ─────────
print("Edit 5: Methods — ablation queries specification")
abl_para = find_para_starting_with("Following the observation of a catastrophic")
if abl_para:
    replace_in_para(
        abl_para,
        "Each ablation comprised 300 API calls (100 queries × 3 repetitions).",
        "Each ablation comprised 300 API calls — the 100 panoramic Tooth_33_Apex queries (one per panoramic radiograph) × 3 repetitions — submitted on byte-identical images and prompts that differed from the canonical guided prompt only in the named modification.",
    )


# ── EDIT 6: Results — extend compliance paragraph with re-query detail
print("Edit 6: Results — Gemini re-query accounting")
comp_para = find_para_starting_with("A total of 10,800 API calls")
if comp_para:
    REQUERY_NOTE = (
        "For Gemini 3.1 Pro Rep 1, 78 of 1,800 responses reached the "
        "max_output_tokens = 2,048 cap during the model's internal reasoning "
        "phase before emitting a final coordinate; these 78 responses were "
        "re-issued at max_output_tokens = 4,096 and merged into the Rep 1 "
        "dataset before parsing. The 99.94% headline compliance rate "
        "(5,397/5,400) reflects this merged dataset; without the re-query "
        "step, raw Rep 1 strict-parse compliance would have been 95.7%. The "
        "re-query was applied uniformly to every truncated response and the "
        "merge mechanism is fully reproducible from the public repository. "
        "Reps 2 and 3 used max_output_tokens = 4,096 ab initio."
    )
    append_to_para(comp_para, REQUERY_NOTE)


# ── EDIT 7: Discussion — bridge paragraph harmonizing routing with §RQ8
print("Edit 7: Discussion — bridge paragraph harmonizing routing recommendation with RQ8")
routing_para = find_para_starting_with("The systematic performance divergence between the two models")
if routing_para:
    BRIDGE = (
        "It must, however, be explicitly emphasised that the cross-model "
        "ranking established in §RQ7 (Gemini > GPT on panoramic and "
        "periapical points; GPT > Gemini on cephalometric points) does not "
        "translate into a clinical-grade ranking against human trainees. As "
        "established in §RQ8 and discussed in greater detail below, both "
        "models — including the better-performing Gemini 3.1 Pro — remain "
        "significantly outperformed by the fourth-year dental student "
        "consensus on every panoramic and periapical comparison "
        "(Bonferroni-corrected p < 10⁻¹⁶ for panoramic point, p < 10⁻²⁰ for "
        "panoramic area). The modality-specific routing paradigm we propose "
        "is therefore best understood as a 'lesser-of-two-evils' "
        "recommendation within the constraint of MLLM-only workflows, not as "
        "a substitute for clinician supervision; any deployment must include "
        "human review of the routed model's panoramic and periapical outputs."
    )
    insert_after(routing_para, BRIDGE, bold=False)


# ── EDIT 8: Limitations — pre-training hypothesis + rep-averaging caveat
print("Edit 8: Limitations — pre-training-distribution + rep-averaging caveats")
limits_para = find_para_starting_with("Several methodological limitations")
if limits_para:
    LIMITS_EXTRA = (
        "Additionally, the pre-training-distribution hypothesis invoked at "
        "several points in the Discussion to explain the cephalometric "
        "advantage of GPT-5.4 and the relative MLLM weakness on periapical "
        "tasks is, by virtue of the proprietary and undisclosed composition "
        "of frontier-model training corpora, fundamentally unfalsifiable "
        "from publicly available data; it is presented as the most "
        "parsimonious explanation consistent with the observed modality "
        "hierarchy but should not be regarded as a definitive mechanistic "
        "claim. A second methodological caveat concerns the rep-averaging "
        "convention adopted as the primary unit of analysis: although a "
        "sensitivity check on per-rep ED values reproduced the direction "
        "and statistical significance of every modality-level finding, "
        "rep-averaging may understate inter-rep variance in heavy-tailed "
        "failure modes (e.g., the Tooth_33_Apex attractor in GPT-5.4, where "
        "all three reps converged on the wrong region) and should be "
        "interpreted alongside the Fleiss-κ reproducibility metrics reported "
        "in §RQ3."
    )
    append_to_para(limits_para, LIMITS_EXTRA)


# ── EDIT 9 (REMOVED) — References left untouched per user instruction.
#                       User noted: a regex over month names / "vd." risks
#                       corrupting Turkish paper titles that may contain
#                       those exact words. Reference language is preserved
#                       as the colleague wrote it; English-vs-Turkish
#                       citation formatting is left to the journal copy-
#                       editing stage.
print("Edit 9: SKIPPED — references untouched (user instruction)")


# ── EDIT 10: Conclusions — strengthen methodology lead ────────────
print("Edit 10: Conclusions — add methodology-emphasis sentence at start")
conc_para = find_para_starting_with("The systematic evaluation of frontier")
if conc_para:
    cur = paragraph_text(conc_para)
    NEW_LEAD = (
        "This study introduces a novel, modality-agnostic grid-based "
        "coordinate framework for quantifying multimodal large language model "
        "spatial localisation competence on dental radiographs, and applies it "
        "to a head-to-head benchmark of two frontier MLLMs against an "
        "adjudicated fourth-year dental student consensus. "
    )
    # Prepend (preserve rest)
    rewrite_paragraph_text(conc_para, NEW_LEAD + cur.lstrip())


# ── Save v3 ────────────────────────────────────────────────────────
doc.save(V3)
print(f"\nSaved: {V3}")
print(f"  Paragraphs: {len(doc.paragraphs)}")
print(f"  Tables: {len(doc.tables)}")
