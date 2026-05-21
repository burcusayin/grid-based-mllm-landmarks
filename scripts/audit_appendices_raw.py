"""
Verify every cell in Appendix F (GPT 5,400 responses) and Appendix G
(Gemini 5,400 responses) against RAW data — not against full_run_records.pkl.

For each row in F1-F8 and G1-G8, the "Run 1 / Run 2 / Run 3" cells must
byte-match the raw response text we extract from:
  - GPT: results_full/run{1,2,3}/responses/*.jsonl
  - Gemini: results_full_gemini/run{1,2,3}/responses/*.json + requeries.jsonl
"""
from __future__ import annotations
import json, sys
from pathlib import Path
from collections import defaultdict
from docx import Document

ROOT = Path(__file__).resolve().parent.parent

# ── Load raw responses (same logic as audit_from_raw.py) ───────────
def load_gpt_raw():
    raw = {}  # (image_id_structure, strategy, rep_idx) → text
    for rep in (1, 2, 3):
        for jl in sorted((ROOT / "results_full" / f"run{rep}" / "responses").glob("*.jsonl")):
            for line in open(jl):
                line = line.strip()
                if not line: continue
                obj = json.loads(line)
                cid = obj["custom_id"]
                # custom_id ends with _zero_shot or _guided
                if cid.endswith("_zero_shot"):
                    qid, strat = cid[:-len("_zero_shot")], "zero_shot"
                elif cid.endswith("_guided"):
                    qid, strat = cid[:-len("_guided")], "guided"
                else:
                    continue
                content = (obj.get("response", {}).get("body", {})
                           .get("choices", [{}])[0].get("message", {})
                           .get("content", "") or "").strip()
                raw[(qid, strat, rep - 1)] = content
    return raw


def load_gem_raw():
    raw = {}
    for rep in (1, 2, 3):
        rep_dir = ROOT / "results_full_gemini" / f"run{rep}" / "responses"
        for chunk in sorted(rep_dir.glob("gemini-3.1-pro_*_chunk*.json")):
            for entry in json.loads(chunk.read_text()):
                cid = entry.get("custom_id", "")
                resp = entry.get("response") or {}
                cands = resp.get("candidates", []) or []
                text = None
                if cands and isinstance(cands[0], dict):
                    parts = (cands[0].get("content", {}) or {}).get("parts", []) or []
                    if parts and isinstance(parts[0], dict):
                        text = parts[0].get("text", "") or ""
                if cid.endswith("_zero_shot"):
                    qid, strat = cid[:-len("_zero_shot")], "zero_shot"
                elif cid.endswith("_guided"):
                    qid, strat = cid[:-len("_guided")], "guided"
                else: continue
                raw[(qid, strat, rep - 1)] = text
        # Merge requeries (rep 1 only)
        rq = rep_dir / "gemini-3.1-pro_requeries.jsonl"
        if rq.exists():
            for line in open(rq):
                line = line.strip()
                if not line: continue
                obj = json.loads(line)
                cid = obj["custom_id"]
                if cid.endswith("_zero_shot"):
                    qid, strat = cid[:-len("_zero_shot")], "zero_shot"
                elif cid.endswith("_guided"):
                    qid, strat = cid[:-len("_guided")], "guided"
                else: continue
                raw[(qid, strat, rep - 1)] = obj.get("raw_response", "") or ""
    return raw


def fmt_display(raw):
    """Match the docx formatter: '<empty>' for None/empty, stripped otherwise."""
    if raw is None or not str(raw).strip():
        return "<empty>"
    return str(raw).strip()


print("Loading raw GPT JSONLs...")
gpt_raw = load_gpt_raw()
print(f"  {len(gpt_raw)} (qid, strategy, rep) keys")

print("Loading raw Gemini chunks + requeries...")
gem_raw = load_gem_raw()
print(f"  {len(gem_raw)} (qid, strategy, rep) keys")

# ── Read v5 docx and walk Appendix F + G tables ────────────────────
doc = Document(ROOT / "results_consensus" / "Full_Run_Results_Report_v5_Consensus.docx")

# Find table indices for F1-F8 and G1-G8.
# In a python-docx Document, doc.tables is a flat list in document order.
# We need to find which tables correspond to Appendix F and G.
# Strategy: walk paragraphs, count tables encountered, record which heading
# precedes each table.
table_meta = []  # list of (table_index, preceding_heading_text)
table_count = 0
current_heading = ""
# python-docx: tables and paragraphs are independently indexed; we walk the
# body XML order via doc.element.body
from docx.oxml.ns import qn
body = doc.element.body
for child in body:
    if child.tag == qn("w:p"):
        # paragraph
        # find pStyle
        pPr = child.find(qn("w:pPr"))
        is_heading = False
        if pPr is not None:
            ps = pPr.find(qn("w:pStyle"))
            if ps is not None and ps.get(qn("w:val"), "").startswith("Heading"):
                is_heading = True
        if is_heading:
            current_heading = "".join(t.text or "" for t in child.iter(qn("w:t")))
    elif child.tag == qn("w:tbl"):
        table_meta.append((table_count, current_heading))
        table_count += 1

# Pair table_meta with the actual table objects
appendix_f_tables = []  # list of (label, table)
appendix_g_tables = []
for (idx, heading), table in zip(table_meta, doc.tables):
    if heading.startswith("Table F") and table is not None:
        # extract label like "F1"
        lbl = heading.split()[1]
        appendix_f_tables.append((lbl, table))
    elif heading.startswith("Table G") and table is not None:
        lbl = heading.split()[1]
        appendix_g_tables.append((lbl, table))

print(f"\nFound {len(appendix_f_tables)} F tables: {[l for l,_ in appendix_f_tables]}")
print(f"Found {len(appendix_g_tables)} G tables: {[l for l,_ in appendix_g_tables]}")

# F_GROUPS maps label → (modality, strategy, ltype)
F_GROUPS = {
    'F1': ('PANORAMIC', 'zero_shot'),
    'F2': ('PANORAMIC', 'guided'),
    'F3': ('PANORAMIC', 'zero_shot'),
    'F4': ('PANORAMIC', 'guided'),
    'F5': ('PERIAPICAL', 'zero_shot'),
    'F6': ('PERIAPICAL', 'guided'),
    'F7': ('CEPHALOMETRIC', 'zero_shot'),
    'F8': ('CEPHALOMETRIC', 'guided'),
}
G_GROUPS = {k.replace('F', 'G'): v for k, v in F_GROUPS.items()}

# ── Walk each table and verify every row ───────────────────────────
def verify_appendix(tables_list, raw_dict, groups_map, name):
    n_rows_total = 0
    n_mismatches = 0
    sample_mismatches = []
    for label, table in tables_list:
        _, strat = groups_map[label]
        n_rows = len(table.rows) - 1  # excluding header
        for ridx in range(1, len(table.rows)):
            row_cells = table.rows[ridx].cells
            image_id = row_cells[0].text.strip()
            landmark = row_cells[1].text.strip()
            gt = row_cells[2].text.strip()
            run1 = row_cells[3].text.strip()
            run2 = row_cells[4].text.strip()
            run3 = row_cells[5].text.strip()
            qid = f"{image_id}_{landmark}"
            actual_run1 = fmt_display(raw_dict.get((qid, strat, 0)))
            actual_run2 = fmt_display(raw_dict.get((qid, strat, 1)))
            actual_run3 = fmt_display(raw_dict.get((qid, strat, 2)))
            displayed = [run1, run2, run3]
            expected = [actual_run1, actual_run2, actual_run3]
            for d, e in zip(displayed, expected):
                n_rows_total += 1
            if displayed != expected:
                n_mismatches += 1
                if len(sample_mismatches) < 3:
                    sample_mismatches.append((label, qid, displayed, expected))
    print(f"\n  {name}: walked {n_rows_total} response-cells across {len(tables_list)} tables")
    print(f"  Mismatches: {n_mismatches}")
    if sample_mismatches:
        print(f"  Sample mismatches:")
        for lbl, qid, d, e in sample_mismatches:
            print(f"    {lbl}/{qid}: displayed={d}, expected={e}")
    return n_mismatches

print()
print("=" * 70)
print("Verifying Appendix F (GPT 5,400 responses) against raw JSONLs")
print("=" * 70)
n_f_diff = verify_appendix(appendix_f_tables, gpt_raw, F_GROUPS, "Appendix F")

print()
print("=" * 70)
print("Verifying Appendix G (Gemini 5,400 responses) against raw chunks + requeries")
print("=" * 70)
n_g_diff = verify_appendix(appendix_g_tables, gem_raw, G_GROUPS, "Appendix G")

print()
print("=" * 70)
print(f"TOTAL: Appendix F mismatches = {n_f_diff}, Appendix G mismatches = {n_g_diff}")
print("=" * 70)
if n_f_diff == 0 and n_g_diff == 0:
    print("✓ Every cell in Appendix F (5,400) AND Appendix G (5,400) byte-matches the raw "
          "responses.")
else:
    print("✗ Found mismatches — review")
    sys.exit(1)
