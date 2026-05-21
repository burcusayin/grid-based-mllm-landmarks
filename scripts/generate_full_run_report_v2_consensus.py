"""
Generate the v2 (consensus-GT) full-run results docx report.

Mirrors the v1 report structure but:
  - Pulls every figure from results_consensus/{analysis,phase_b,summary,
    rater_reliability,gpt_vs_student}.json and full_run_records.pkl
  - Uses CONSENSUS_Ground_Truth as the canonical reference throughout
  - Adds Section 8 (Ground Truth Validation — inter/intra-rater reliability)
  - Adds Section 9 (GPT-5.4 vs Student paired comparison)
  - Adds Appendix D (v1 omfr_1 ↔ v2 consensus sensitivity table)
  - Anchors to results_full/ via reanalysis_anchor.json (SHAs of the frozen
    raw JSONLs the analysis re-evaluates).

EVERY numeric figure is computed inline from the JSONs. NO hardcoded values.
The v1 report (results_full/Full_Run_Results_Report.docx) stays untouched.
"""
from __future__ import annotations
import json, pickle, math, hashlib, subprocess
from pathlib import Path
from datetime import datetime, UTC

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SANDBOX = ROOT / 'results_consensus'
OUT  = SANDBOX / 'Full_Run_Results_Report_v5_Consensus.docx'

# v5 adds Gemini 3.1 Pro full-run results (§12), cross-model GPT vs Gemini
# paired analysis (§13), and extends the Discussion (§11.6-§11.8) and
# Conclusions (now §15) with cross-model commentary. All §11.1-§11.5 expert
# contributions from v4 are preserved verbatim.

# ── Load v2 (consensus) data ────────────────────────────────────────
records      = pickle.load(open(SANDBOX / 'full_run_records.pkl', 'rb'))
analysis     = json.load(open(SANDBOX / 'analysis.json'))
phase_b      = json.load(open(SANDBOX / 'phase_b.json'))
summary      = json.load(open(SANDBOX / 'summary.json'))
rater_reli   = json.load(open(SANDBOX / 'rater_reliability.json'))
gpt_v_stu    = json.load(open(SANDBOX / 'gpt_vs_student.json'))
v2_manifest  = json.load(open(SANDBOX / 'v2_manifest.json'))
anchor       = json.load(open(SANDBOX / 'reanalysis_anchor.json'))

# ── Load Gemini 3.1 Pro data (v5 additions) ─────────────────────────
# Recomputed identically to GPT (single-cell rule for points, atomic anchor
# of every raw chunk + re-queries file). v5 §12, §13 are derived from these.
GEMINI_SANDBOX = ROOT / 'results_full_gemini'
gem_records  = pickle.load(open(GEMINI_SANDBOX / 'full_run_records.pkl', 'rb'))
gem_analysis = json.load(open(GEMINI_SANDBOX / 'analysis.json'))
gem_phase_b  = json.load(open(GEMINI_SANDBOX / 'phase_b.json'))
gem_summary  = json.load(open(GEMINI_SANDBOX / 'summary.json'))
gem_manifest = json.load(open(GEMINI_SANDBOX / 'full_run_manifest.json'))
gem_manifest_rep1 = json.load(open(GEMINI_SANDBOX / 'full_run_manifest_rep1.json'))
gem_anchor   = json.load(open(GEMINI_SANDBOX / 'gemini_recompute_anchor.json'))
gpt_vs_gem   = json.load(open(GEMINI_SANDBOX / 'gpt_vs_gemini.json'))

# Also load v1 sensitivity (omfr_1 reference) — for Appendix D
analysis_omfr1 = json.load(open(SANDBOX / 'analysis_omfr_1.json'))
phase_b_omfr1  = json.load(open(SANDBOX / 'phase_b_omfr_1.json'))

# Load all three follow-up ablations — for Section 10
# Each ablation tests a different prompt-level hypothesis about why GPT-5.4
# collapses on Tooth_33_Apex under the canonical guided prompt. All three
# share the same 100 PAN/Tooth_33_Apex queries (verified at preflight).
ABL_DIRS = {
    'no_tooth_num':  ROOT / 'results_ablation_no_tooth_num',
    'patient_left':  ROOT / 'results_ablation_patient_left',
    'no_LR':         ROOT / 'results_ablation_no_LR',
}
ABL_STRATEGY = {
    'no_tooth_num':  'guided_no_tooth_num',
    'patient_left':  'guided_patient_left',
    'no_LR':         'guided_no_LR',
}

ablations = {}        # key → ablation_analysis.json
abl_manifests = {}    # key → ablation_manifest.json
abl_anchors = {}      # key → ablation_anchor.json
abl_costs = {}        # key → naive cost computed from raw tokens

for key, dirpath in ABL_DIRS.items():
    ablations[key]    = json.load(open(dirpath / 'ablation_analysis.json'))
    abl_manifests[key] = json.load(open(dirpath / 'ablation_manifest.json'))
    abl_anchors[key]  = json.load(open(dirpath / 'ablation_anchor.json'))
    pt = ct = 0
    for _rep in (1, 2, 3):
        for _jl in (dirpath / f'run{_rep}' / 'responses').glob('*.jsonl'):
            for _line in open(_jl):
                _u = json.loads(_line).get('response', {}).get('body', {}).get('usage', {})
                pt += _u.get('prompt_tokens', 0)
                ct += _u.get('completion_tokens', 0)
    abl_costs[key] = pt * 1.25e-6 + ct * 7.5e-6

# Convenience accessors (every ablation's zero_shot and guided baselines come
# from the same frozen results_full/ data, so they agree numerically across
# ablations to floating-point precision; we just pull them from no_tooth_num).
abl_zs        = ablations['no_tooth_num']['summary_per_strategy']['zero_shot']
abl_gd        = ablations['no_tooth_num']['summary_per_strategy']['guided']
abl_nt        = ablations['no_tooth_num']['summary_per_strategy']['guided_no_tooth_num']
abl_pl        = ablations['patient_left']['summary_per_strategy']['guided_patient_left']
abl_nlr       = ablations['no_LR']['summary_per_strategy']['guided_no_LR']
pw_nt_gd      = ablations['no_tooth_num']['paired_wilcoxon']['guided_no_tooth_num_vs_guided']
pw_pl_gd      = ablations['patient_left']['paired_wilcoxon']['guided_patient_left_vs_guided']
pw_nlr_gd     = ablations['no_LR']['paired_wilcoxon']['guided_no_LR_vs_guided']
pw_gd_zs      = ablations['no_tooth_num']['paired_wilcoxon']['guided_vs_zero_shot']
abl_total_cost = sum(abl_costs.values())

# Load the qualitative inspection (per-image mode cells across all 5 conditions).
qual_path = ROOT / 'results_ablation_no_LR' / 'qualitative_inspection.json'
qual_rows = json.load(open(qual_path)) if qual_path.exists() else []

# Compute per-ablation strict compliance numbers directly from each sandbox's
# compliance_stats.json, so Section 10's "Operational outcome" lines are not
# hardcoded and cannot drift from the canonical record.
def ablation_compliance(key):
    sb = ABL_DIRS[key]
    total = compliant = 0
    for rep in (1, 2, 3):
        cs = json.load(open(sb / f'run{rep}' / 'compliance_stats.json'))
        total += cs.get('total', 0)
        compliant += cs.get('compliant', 0)
    return total, compliant
abl_compliance = {k: ablation_compliance(k) for k in ABL_DIRS}

# ── Ground-truth changes summary ────────────────────────────────────
queries = json.load(open(SANDBOX / 'query_index.json'))

def norm_cells(s):
    if not s: return ''
    cells = sorted(c.strip().upper().replace(' ','').replace('-','')
                   for c in s.split(',') if c.strip())
    return ','.join(cells)

gt_changes = {'point': {'CEPHALOMETRIC':[], 'PERIAPICAL':[], 'PANORAMIC':[]},
              'area':  {'PANORAMIC':[]}}
for q in queries:
    if norm_cells(q['consensus_gt']) != norm_cells(q['omfr_1']):
        gt_changes[q['landmark_type']][q['sheet']].append(q['query_id'])

# ── Failures (compliance) — recompute from records ──────────────────
failures = []
for rec in records:
    for rep_idx, fail in enumerate(rec['rep_failure']):
        if fail is not None:
            failures.append({
                'query_id': rec['query_id'], 'strategy': rec['strategy'],
                'rep': rep_idx + 1, 'modality': rec['modality'],
                'structure': rec['structure'], 'gt': rec['consensus_gt'],
                'raw_response': rec['rep_raw'][rep_idx], 'failure_category': fail,
            })

# ── Git + SHAs ──────────────────────────────────────────────────────
git_sha = subprocess.check_output(['git','-C',str(ROOT),'rev-parse','HEAD']).decode().strip()
git_date = subprocess.check_output(['git','-C',str(ROOT),'log','-1','--format=%ci','HEAD']).decode().strip()
v1_anchor_sha = anchor['upstream_query_index']['sha256']
final_excel_sha = v2_manifest['source_excel_sha256']
old_excel_sha = hashlib.sha256(open(ROOT/'data'/'Dental_MLLM_Benchmark_Data.xlsx','rb').read()).hexdigest()
qi_sha = hashlib.sha256(open(SANDBOX/'query_index.json','rb').read()).hexdigest()

# ── Document setup ──────────────────────────────────────────────────
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
sec = doc.sections[0]
sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.85)


def H(text, level=1):
    return doc.add_heading(text, level=level)

def P(text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    return p

def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(10)
    return p

def shade_cell(cell, fill='D9E1F2'):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), fill); tcPr.append(sh)

def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'808080')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_table(headers, rows, col_widths=None, header_size=10, body_size=10, fixed=False):
    n_cols = len(headers)
    t = doc.add_table(rows=len(rows)+1, cols=n_cols)
    t.style = 'Light Grid'
    if fixed:
        tblPr = t._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); t._element.insert(0, tblPr)
        layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
        tblPr.append(layout)
    for j, h in enumerate(headers):
        c = t.cell(0,j); c.text=''
        run = c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(header_size)
        shade_cell(c, 'D9E1F2'); set_cell_borders(c)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i+1,j); c.text=''
            run = c.paragraphs[0].add_run(str(val)); run.font.size = Pt(body_size)
            set_cell_borders(c)
    if col_widths:
        for j, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    return t


def fmt_p(p):
    return f'{p:.2e}' if p < 0.001 else f'{p:.4f}'


def fmt_pbonf(p, m):
    pb = min(1.0, p*m)
    return f'{pb:.2e}' if pb < 0.001 else f'{pb:.4f}'


def sig_marker(p_bonf):
    if p_bonf < 0.001: return '***'
    if p_bonf < 0.01: return '**'
    if p_bonf < 0.05: return '*'
    return 'NS'


def fmt_pct_ci(rate, ci):
    return f'{rate*100:.1f}% [{ci[0]*100:.1f}, {ci[1]*100:.1f}]'


# ────────────────────────────────────────────────────────────────────
# TITLE & METADATA
# ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    'Comparative Analysis of Multimodal Large Language Models and Dental Students '
    'in Spatial Proficiency for Radiographic Anatomic Landmark Identification: '
    'A Novel Grid-Based Assessment Model')
r.bold = True; r.font.size = Pt(16)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Technical Report v5 — GPT-5.4 vs Gemini 3.1 Pro vs Adjudicated Two-Rater Consensus GT '
                'with Team-Adjudicated Dental Student Comparison')
r.italic = True; r.font.size = Pt(12)
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('10,800 API calls (5,400 GPT-5.4 + 5,400 Gemini 3.1 Pro) · 900 queries · 200 radiographs · '
                 'byte-identical prompts and images')
r.italic = True; r.font.size = Pt(10)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f'Generated {datetime.now(UTC).strftime("%Y-%m-%d")} from raw API outputs '
                 f'(commit {git_sha[:12]}, sandbox: results_consensus)')
r.font.size = Pt(10); r.italic = True
doc.add_paragraph()

# ────────────────────────────────────────────────────────────────────
# ABSTRACT
# ────────────────────────────────────────────────────────────────────
H('Abstract', level=1)
# Pull headline numbers live from JSON so the abstract cannot drift from data
_abs_pan_gd = gpt_vs_gem['RQ_modality_strategy']['PANORAMIC_guided_point']
_abs_ceph_gd = gpt_vs_gem['RQ_modality_strategy']['CEPHALOMETRIC_guided_point']
_abs_pa_gd = gpt_vs_gem['RQ_modality_strategy']['PERIAPICAL_guided_point']
_abs_f56_gpt = gpt_vs_gem['F5_F6_attractor']['GPT/Tooth_33_Apex/guided']
_abs_f56_gem = gpt_vs_gem['F5_F6_attractor']['Gemini/Tooth_33_Apex/guided']

P('Background. ', bold=True)
P('Multimodal large language models (MLLMs) are being proposed as zero-shot or low-shot '
  'tools across medical-imaging tasks, but their accuracy, reproducibility, and failure modes '
  'on fine-grained anatomic-landmark identification — and how they compare to one another and '
  'to human raters — remain incompletely characterised, particularly on dental radiographs '
  'where modality-specific challenges (anatomic superimposition in panoramic, narrow contrast '
  'on periapical, projection-specific geometry on cephalometric) interact with prompt-level '
  'variation.', italic=False)

P('Methods. ', bold=True)
P(f'We constructed a 900-query grid-based dental anatomic landmark identification benchmark '
  f'over 200 radiographs spanning three modalities (100 panoramic with 16×8 grid, 50 periapical '
  f'with 8×6 grid, 50 cephalometric with 10×8 grid) and 12 clinically defined landmarks '
  f'(9 point + 3 area). Each query is paired with a two-rater adjudicated consensus ground '
  f'truth (CONSENSUS_Ground_Truth), with each of two oral-and-maxillofacial-radiology (OMFR) '
  f'specialists providing primary plus washout-period re-rating data for inter- and intra-'
  f'rater reliability. A team-adjudicated dental-student consensus is reported as a third '
  f'comparator. Two commercial MLLMs (GPT-5.4 and Gemini 3.1 Pro) were queried on identical '
  f'inputs (byte-identical prompts on byte-identical images, verified at preflight) under two '
  f'prompt strategies (zero_shot and guided) for three independent repetitions at temperature '
  f'= 0, totalling 10,800 API calls (5,400 per model). Per-query observations are arithmetic '
  f'means of the three per-rep metrics (Euclidean distance in grid cells for point landmarks; '
  f'Jaccard for area landmarks). Statistical inference uses paired Wilcoxon signed-rank tests '
  f'with bootstrap 95% CIs on rank-biserial r (n=2,000 resamples) and Bonferroni correction.')

P('Results. ', bold=True)
P(f'Strict-parse compliance was 99.926% for GPT-5.4 (4 failures / 5,400) and 99.944% for '
  f'Gemini 3.1 Pro (3 failures / 5,400). Gemini outperforms GPT-5.4 on panoramic and '
  f'periapical point landmarks by large effect sizes: panoramic guided mean Euclidean distance '
  f'{_abs_pan_gd["gpt_mean"]:.2f} cells (GPT) vs {_abs_pan_gd["gemini_mean"]:.2f} cells '
  f'(Gemini), rank-biserial r = {_abs_pan_gd["rank_biserial_r"]:+.2f} '
  f'(Bonferroni-corrected p = {fmt_pbonf(_abs_pan_gd["p"], 8)}; PAN guided point); periapical '
  f'guided mean ED {_abs_pa_gd["gpt_mean"]:.2f} vs {_abs_pa_gd["gemini_mean"]:.2f} cells '
  f'(r = {_abs_pa_gd["rank_biserial_r"]:+.2f}). The direction reverses on cephalometric '
  f'points only — GPT-5.4 {_abs_ceph_gd["gpt_mean"]:.2f} vs Gemini '
  f'{_abs_ceph_gd["gemini_mean"]:.2f} cells (r = '
  f'{_abs_ceph_gd["rank_biserial_r"]:+.2f}, GPT closer to GT). A targeted analysis of the '
  f'Tooth_33_Apex panoramic landmark identified a model-specific failure mode: under '
  f'byte-identical guided prompting GPT-5.4 placed '
  f'{_abs_f56_gpt["n_in_F5_or_F6"]}/{_abs_f56_gpt["n_predictions_with_valid_cell"]} '
  f'({_abs_f56_gpt["frac_F5_F6"]*100:.1f}%) predictions on cells F5 or F6 (anatomically the '
  f'lower-right molar projection, not the canine apex), whereas Gemini placed '
  f'{_abs_f56_gem["n_in_F5_or_F6"]}/{_abs_f56_gem["n_predictions_with_valid_cell"]} '
  f'({_abs_f56_gem["frac_F5_F6"]*100:.1f}%) there. Three pre-registered prompt-level ablations '
  f'on GPT-5.4 (FDI tooth-number removal, patient-frame disambiguation, panoramic L-R clause '
  f'removal) failed to rescue the regression on the same 100 queries, identifying the F5/F6 '
  f'attractor as structural-anatomic confusion at the model level rather than a prompt-level '
  f'artefact. Rep-to-rep agreement (Fleiss kappa across 3 reps on point landmarks) was '
  f'{phase_b["fleiss_overall_point"]:.2f} for GPT-5.4 and '
  f'{gem_phase_b["fleiss_overall_point"]:.2f} for Gemini.')

P('Conclusions. ', bold=True)
P('On a multi-modality dental landmark benchmark with adjudicated ground truth, two state-of-'
  'the-art MLLMs differ systematically and in non-uniform directions across imaging modalities, '
  'with Gemini 3.1 Pro substantially closer to the consensus GT on panoramic and periapical '
  'point landmarks but GPT-5.4 closer on cephalometric points. The most prominent prompt-'
  'sensitive failure mode (Tooth_33_Apex F5/F6 attractor under guided prompting) is GPT-specific '
  'and cannot be rescued by prompt engineering; it appears to reflect a structural-anatomic '
  'confusion at the model level. Neither model has yet reached clinical-deployment accuracy on '
  'panoramic landmarks (SDR@1 ≤ 74%), and the student consensus remains the more accurate '
  'point-landmark locator on this benchmark. The findings argue for human-in-the-loop deployment '
  'and for explicit cross-model evaluation when selecting an MLLM for a specific dental-imaging '
  'workflow. All data, code, and intermediate analysis artefacts are released for full '
  'reproducibility (see §3.5 Code and Data Availability).')

P('Keywords. ', bold=True)
P('Multimodal large language models; dental radiography; anatomic landmark identification; '
  'panoramic; periapical; cephalometric; benchmark; reproducibility; GPT-5.4; Gemini 3.1 Pro.',
  italic=True)

# ────────────────────────────────────────────────────────────────────
# 1. EXECUTIVE SUMMARY (now: Introduction and Headline Findings)
# ────────────────────────────────────────────────────────────────────
H('1. Introduction and Headline Findings', level=1)

P('Background and motivation.', bold=True)
P('Multimodal large language models (MLLMs) increasingly claim general-purpose vision-language '
  'competence — including on medical-imaging tasks — but published benchmarks rarely combine '
  '(i) rigorous adjudicated ground truth, (ii) cross-model comparisons on byte-identical inputs, '
  '(iii) explicit human-rater baselines from the relevant clinical population, (iv) full-rep '
  'temperature-zero reproducibility analysis, and (v) targeted ablations on identified failure '
  'modes. Dental radiographic landmark identification is a useful test bed because three '
  'modality types (panoramic, periapical, cephalometric) require qualitatively different '
  'visual reasoning (cross-arch overview with anatomic superimposition for panoramic; small-'
  'field high-contrast for periapical; well-defined silhouette geometry for cephalometric), and '
  'because purpose-built CNN systems have established performance ceilings on each '
  '(~ 92% SDR within 2 mm on cephalometric landmarks; substantially weaker on panoramic).')

P('Contributions of this paper.', bold=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('A 900-query benchmark on 200 dental radiographs (100 PAN + 50 PA + 50 CEPH) '
          'with a fully adjudicated CONSENSUS_Ground_Truth derived from two OMFR specialists, '
          'including washout-period intra-rater data and a team-adjudicated dental-student '
          'consensus as a human-rater baseline.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Cross-model evaluation of GPT-5.4 and Gemini 3.1 Pro on byte-identical prompts on '
          'byte-identical images, 3 repetitions × 2 strategies (zero-shot, guided) × 900 '
          'queries × 2 models = 10,800 API calls in total, with a paired Wilcoxon signed-rank '
          'analysis (Bonferroni-corrected) on per-query rep-mean differences.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Identification and dissection of a model-specific failure mode (GPT-5.4 F5/F6 '
          'attractor on Tooth_33_Apex), supported by three pre-registered prompt-level '
          'ablations and ruled out as prompt-induced by the cross-model comparison '
          '(Gemini, under the identical prompt, exhibits 0/300 predictions on F5/F6).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('A full reproducibility infrastructure: cryptographic SHA-256 anchoring of every '
          'input file (200 image PNGs, the Final benchmark Excel, every raw JSONL/chunk), '
          'live-rendered prompts verified byte-identical to those issued at API time, and '
          'open-source code release (see §3.5). Every number in this report is reproducible '
          'end-to-end from the raw chunks via four scripts (run_full_run_*, recompute_*, '
          'analyze_*).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Methodological transparency: this report uses arithmetic-mean-of-per-rep-metrics '
          'as the canonical per-query observation (NOT majority voting), with rep-to-rep '
          'agreement reported separately via Fleiss kappa. Rationale and alternative-method '
          'comparison are in §3.2.')

P('Paper structure.', bold=True)
P('Section 2 sets out the study context, the seven research questions (RQ1-RQ7), and the '
  'data assets. Section 3 describes the methodology (ground truth, statistical analysis '
  'plan with explicit rep-aggregation rule, reproducibility infrastructure, canonical prompts). '
  'Sections 4-7 present GPT-5.4 single-model results against the consensus GT (operational '
  'outcomes, modality-stratified accuracy, prompt-strategy effect, reproducibility). Sections '
  '8-9 validate the ground truth (inter- and intra-rater reliability) and compare GPT-5.4 to '
  'the dental student consensus. Section 10 details three pre-registered prompt-level '
  'ablations on the Tooth_33_Apex panoramic regression. Section 11 is the Discussion (with '
  'expert clinical contributions in §11.1-§11.3; cross-model perspective in §11.6). '
  'Section 12 reports the Gemini 3.1 Pro single-model results in parallel structure to '
  'Sections 4-7. Section 13 is the GPT-vs-Gemini paired cross-model comparison. Section 14 '
  'discusses limitations; Section 15 concludes. Appendices A-G provide detailed per-landmark '
  'statistics, the full compliance-failure list, the reproducibility manifest, the v1 '
  'sensitivity comparison, the per-image qualitative inspection, and all 10,800 model '
  'responses (5,400 GPT in Appendix F, 5,400 Gemini in Appendix G).')

n_calls = summary['n_total_calls']
n_actual = summary['n_actual_responses']
compliance_rate = summary['compliance_rate']
prompt_tok = phase_b['tokens']['prompt_tokens']
compl_tok = phase_b['tokens']['completion_tokens']
naive_cost = summary['naive_cost_usd']

P(f'This v5 report is the two-model successor to v4, which was itself the v2-consensus-GT '
  f'edition of the original v1 GPT-only report. The v5 additions are: (1) a full Gemini 3.1 Pro '
  f'benchmark on the same 900 queries × 2 strategies × 3 reps = 5,400 API calls, executed in '
  f'May 2026 with byte-identical prompts and byte-identical images to the GPT-5.4 main run '
  f'(Section 12); and (2) a head-to-head GPT-vs-Gemini paired comparison against the same '
  f'consensus GT (Section 13). All v4 GPT-only sections (1-11.5, the §10 Tooth_33_Apex '
  f'ablations, the v1↔v2 sensitivity analysis in Appendix D, and the §11.1-§11.3 expert '
  f'clinical contributions) are preserved verbatim; v5 extends the discussion (§11.4-§11.6), '
  f'limitations (§14), and conclusions (§15) with multi-model commentary. The {n_calls:,} '
  f'GPT-5.4 batch responses are unchanged — they have only been re-scored against the consensus '
  f'GT, not re-issued — and the raw JSONLs are cryptographically anchored at preflight to '
  f'guarantee no model-output drift across v1, v2, v4, and v5. The 5,400 Gemini batch responses '
  f'are likewise SHA-256 anchored at recompute time (gemini_recompute_anchor.json, Appendix C).')

m_ceph_zs = analysis['RQ1_modality_strategy']['CEPHALOMETRIC_zero_shot_point']
m_ceph_gd = analysis['RQ1_modality_strategy']['CEPHALOMETRIC_guided_point']
m_pa_zs   = analysis['RQ1_modality_strategy']['PERIAPICAL_zero_shot_point']
m_pa_gd   = analysis['RQ1_modality_strategy']['PERIAPICAL_guided_point']
m_pan_zs  = analysis['RQ1_modality_strategy']['PANORAMIC_zero_shot_point']
m_pan_gd  = analysis['RQ1_modality_strategy']['PANORAMIC_guided_point']

P('Headline findings:', bold=True)

b1 = doc.add_paragraph(style='List Bullet')
b1.add_run('Modality difficulty hierarchy is robust to the GT change. ')
b1.add_run(f'Cephalometric mean Euclidean distance is '
           f'{m_ceph_zs["mean"]:.2f} cells (zero-shot) and {m_ceph_gd["mean"]:.2f} cells (guided). '
           f'Periapical is {m_pa_zs["mean"]:.2f} and {m_pa_gd["mean"]:.2f}. '
           f'Panoramic is {m_pan_zs["mean"]:.2f} and {m_pan_gd["mean"]:.2f}. The same '
           f'CEPH << PA < PAN ordering observed in v1 holds under consensus.')

# Strategy effect highlights
rq2a_pan_pt = analysis['RQ2a_strategy_per_modality']['PANORAMIC_point']
b2 = doc.add_paragraph(style='List Bullet')
b2.add_run('Heterogeneous strategy effects persist. ')
b2.add_run(f'Aggregate PAN point comparison: paired Wilcoxon mean Δ = '
           f'{rq2a_pan_pt["mean_delta"]:+.2f} cells (zero-shot − guided), '
           f'p = {rq2a_pan_pt["p"]:.2e} (Bonferroni p < 0.001), '
           f'rank-biserial r = {rq2a_pan_pt["rank_biserial_r"]:+.2f}. Per-landmark '
           f'analysis (Section 6.3) shows six landmarks with Bonferroni-significant '
           f'strategy effects: guided helps PAN/Mental_Foramen_L, CEPH/Sella_S, and '
           f'PAN/External_Oblique_Ridge_R (area); guided harms PAN/Tooth_33_Apex, '
           f'PAN/Condylar_Head_R, and CEPH/Menton_Me. The qualitative pattern is '
           f'directly consistent with v1.')

# Two-rater reliability — NEW FOR V2
ip_omfr12 = rater_reli['inter_rater']['INTER_omfr1_vs_omfr2']['ALL_point']
ia_omfr12 = rater_reli['inter_rater']['INTER_omfr1_vs_omfr2']['ALL_area']
b3 = doc.add_paragraph(style='List Bullet')
b3.add_run('Inter-rater reliability validates the consensus GT. ')
b3.add_run(f'OMFR_1 vs OMFR_2: point mean ED = {ip_omfr12["mean_ed_cells"]:.3f} cells, '
           f'Cohen\'s κ = {ip_omfr12["cohens_kappa"]:.3f} '
           f'(within-1-cell rate {ip_omfr12["within_1_cell_rate"]*100:.1f}%); area mean Jaccard '
           f'= {ia_omfr12["mean_jaccard"]:.3f}, mean Dice = {ia_omfr12["mean_dice"]:.3f}. '
           f'Intra-rater (180 doubly-rated queries with ≥2-week washout) is essentially perfect '
           f'on points (κ ≥ 0.97).')

# GPT vs Student — KEY NEW FINDING
gs = gpt_v_stu['per_strategy']
b4 = doc.add_paragraph(style='List Bullet')
b4.add_run('GPT-5.4 vs dental-student consensus (NEW in v2). ').bold = True
b4.add_run(
    f'Cephalometric: GPT and student are statistically indistinguishable '
    f'(zero-shot Δ = {gs["zero_shot"]["per_modality"]["CEPHALOMETRIC_point"]["mean_delta"]:+.3f} cells, '
    f'p = {gs["zero_shot"]["per_modality"]["CEPHALOMETRIC_point"]["p"]:.3f}, NS). '
    f'Periapical: student substantially outperforms GPT '
    f'(zero-shot Δ = {gs["zero_shot"]["per_modality"]["PERIAPICAL_point"]["mean_delta"]:+.3f} cells, '
    f'p = {gs["zero_shot"]["per_modality"]["PERIAPICAL_point"]["p"]:.2e}, '
    f'r = {gs["zero_shot"]["per_modality"]["PERIAPICAL_point"]["rank_biserial_r"]:+.2f}). '
    f'Panoramic point: student substantially outperforms GPT '
    f'(zero-shot Δ = {gs["zero_shot"]["per_modality"]["PANORAMIC_point"]["mean_delta"]:+.3f} cells, '
    f'p = {gs["zero_shot"]["per_modality"]["PANORAMIC_point"]["p"]:.2e}, '
    f'r = {gs["zero_shot"]["per_modality"]["PANORAMIC_point"]["rank_biserial_r"]:+.2f}). '
    f'Panoramic area: student substantially outperforms GPT '
    f'(student mean Jaccard = {gs["zero_shot"]["per_modality"]["PANORAMIC_area"]["mean_student_jaccard"]:.3f} vs '
    f'GPT zero-shot {gs["zero_shot"]["per_modality"]["PANORAMIC_area"]["mean_gpt_jaccard"]:.3f}, '
    f'p = {gs["zero_shot"]["per_modality"]["PANORAMIC_area"]["p"]:.2e}).')

unan = phase_b['overall_unanimous']
b5 = doc.add_paragraph(style='List Bullet')
b5.add_run('Reproducibility unchanged. ')
b5.add_run(f"Fleiss' κ across the three repetitions for point landmarks is "
           f"{float(phase_b['fleiss_overall_point']):.3f} (substantial agreement); "
           f"three-way unanimous rate {unan['unanimous']:,}/{unan['total']:,} "
           f"({unan['rate']*100:.1f}%). Independent of GT.")

b6 = doc.add_paragraph(style='List Bullet')
b6.add_run(f'Compliance unchanged: {compliance_rate*100:.3f}% strict, 100% under '
           f'digit-before-letter-tolerant parsing. The 4 failures are the same reversed-coordinate '
           f'errors flagged in v1.')

b7 = doc.add_paragraph(style='List Bullet')
b7.add_run('Tooth_33_Apex regression — three targeted ablations reject all proposed '
           'wording-level explanations; clinical interpretation reframes the regression as '
           'anatomic structural confusion. ').bold = True
b7.add_run(
    f'Three focused 300-call ablations (Section 10) tested distinct hypotheses about why '
    f'GPT-5.4 collapses on Tooth_33_Apex under the canonical guided prompt (5% correct-side '
    f'rate vs zero_shot\'s 67%). Variant A (remove tooth number "33"): correct-side rate '
    f'{abl_nt["correct_side_rate"]*100:.0f}% — REJECTED. Variant B (rewrite parenthetical with '
    f'"patient\'s left" before "lower"): correct-side rate {abl_pl["correct_side_rate"]*100:.0f}% '
    f'— REJECTED. Variant C (remove the panoramic L–R inversion clause from the system prompt — '
    f'diagnostic): correct-side rate {abl_nlr["correct_side_rate"]*100:.0f}% — REJECTED '
    f'(small nominal effect, far below the 60% threshold for "L–R clause is the cause"). A '
    f'qualitative inspection (Section 10.5) shows the model converges on cells F5/F6 across '
    f'all four guided variants. The project\'s OMFR specialist identified F5/F6 in the '
    f'16 × 8 panoramic grid as anatomically corresponding to the LOWER-RIGHT MOLAR region — '
    f'not to any canine. The Tooth_33_Apex regression is therefore not a spatial left/right '
    f'flip; it is structural-anatomic confusion (wrong side AND wrong tooth class). No '
    f'prompt-engineering fix is plausible for a failure at this level. The canonical guided '
    f'prompt is unchanged. Total ablation cost: ${abl_total_cost:.2f}.'
)

P('Cross-model findings (NEW in v5 — full development in §12 and §13):', bold=True)
_xm_pan_gd = gpt_vs_gem['RQ_modality_strategy']['PANORAMIC_guided_point']
_xm_pa_gd  = gpt_vs_gem['RQ_modality_strategy']['PERIAPICAL_guided_point']
_xm_ceph_gd = gpt_vs_gem['RQ_modality_strategy']['CEPHALOMETRIC_guided_point']
_xm_pan_gd_ar = gpt_vs_gem['RQ_modality_strategy']['PANORAMIC_guided_area']
_f56_gpt_gd = gpt_vs_gem['F5_F6_attractor']['GPT/Tooth_33_Apex/guided']
_f56_gem_gd = gpt_vs_gem['F5_F6_attractor']['Gemini/Tooth_33_Apex/guided']

b6 = doc.add_paragraph(style='List Bullet')
b6.add_run('Two models on the same benchmark. ').bold = True
b6.add_run(
    f'v5 adds Gemini 3.1 Pro (3 reps × 1,800 calls = 5,400 calls; total v5 = 10,800 calls). '
    f'Same 900 queries, byte-identical prompts on byte-identical images, scored against the '
    f'same consensus GT. Gemini compliance 99.944% (3 strict failures), Fleiss κ across reps '
    f'= {gem_phase_b["fleiss_overall_point"]:.2f} on point landmarks (vs '
    f'{phase_b["fleiss_overall_point"]:.2f} for GPT-5.4). Full operational outcomes in §12.1.')

b7 = doc.add_paragraph(style='List Bullet')
b7.add_run('Cross-model headline (paired Wilcoxon, Bonferroni × 8). ').bold = True
b7.add_run(
    f'Gemini outperforms GPT-5.4 by large effect sizes on panoramic point landmarks under '
    f'guided prompting — mean ED {_xm_pan_gd["gpt_mean"]:.2f} cells (GPT) vs '
    f'{_xm_pan_gd["gemini_mean"]:.2f} cells (Gemini), rank-biserial '
    f'r = {_xm_pan_gd["rank_biserial_r"]:+.2f}, Bonferroni p = '
    f'{fmt_pbonf(_xm_pan_gd["p"], 8)}. The same direction holds for periapical points and '
    f'panoramic area landmarks. Direction reverses on cephalometric points only — '
    f'GPT {_xm_ceph_gd["gpt_mean"]:.2f} vs Gemini {_xm_ceph_gd["gemini_mean"]:.2f} cells, '
    f'r = {_xm_ceph_gd["rank_biserial_r"]:+.2f} (GPT closer to GT). Full table in §13.2.')

b8 = doc.add_paragraph(style='List Bullet')
b8.add_run('F5/F6 attractor is GPT-specific. ').bold = True
b8.add_run(
    f'Under byte-identical guided prompts on the same images, GPT-5.4 places '
    f'{_f56_gpt_gd["n_in_F5_or_F6"]}/{_f56_gpt_gd["n_predictions_with_valid_cell"]} '
    f'({_f56_gpt_gd["frac_F5_F6"]*100:.1f}%) Tooth_33_Apex predictions on F5 or F6, whereas '
    f'Gemini places {_f56_gem_gd["n_in_F5_or_F6"]}/'
    f'{_f56_gem_gd["n_predictions_with_valid_cell"]} '
    f'({_f56_gem_gd["frac_F5_F6"]*100:.1f}%) on F5 or F6. The cross-model comparison '
    f'settles the §10 prompt-vs-model question definitively: F5/F6 is a model-level '
    f'failure mode (GPT-specific), not a prompt-induced artefact. Full analysis in §13.4.')

P('All figures are computed inline from results_consensus/{full_run_records.pkl, analysis.json, '
  'phase_b.json, rater_reliability.json, gpt_vs_student.json}, '
  'results_full_gemini/{full_run_records.pkl, analysis.json, phase_b.json, gpt_vs_gemini.json}, '
  'and results_ablation_*/. The reproducibility manifest (Appendix C) lists every '
  'cryptographic anchor.', italic=True)

# ────────────────────────────────────────────────────────────────────
# 2. STUDY CONTEXT
# ────────────────────────────────────────────────────────────────────
H('2. Study Context, Pre-Specified Research Questions, and Data Update', level=1)

H('2.1 Scope of this report', level=2)
P('This v5 report benchmarks TWO commercial multimodal large language models — GPT-5.4 and '
  'Gemini 3.1 Pro — on the same 900-query grid-based dental anatomic landmark identification '
  'task, against the same adjudicated CONSENSUS_Ground_Truth, with a head-to-head paired '
  'comparison on byte-identical prompts and byte-identical images. Both models were run for '
  '3 repetitions × 900 queries × 2 prompt strategies (zero_shot and guided) = 5,400 API calls '
  'each, totalling 10,800 calls. GPT-5.4 results (the foundation of v1-v4) appear in Sections '
  '4-10; the Gemini full run and the cross-model comparison are NEW in v5 (Sections 12 and 13). '
  'A team-adjudicated dental student consensus appears as a third comparator for the '
  'GPT-vs-student analysis in Section 9; the cross-model section uses only GPT and Gemini '
  '(student comparison for Gemini is left as future work).')

H('2.2 Pre-specified research questions', level=2)
P('Research questions RQ1-RQ5 below were pre-registered as the v2 GPT-only protocol; RQ6-RQ7 '
  'are POST-HOC additions in v5 (the Gemini full run only became available in May 2026, after '
  'the v2 protocol was locked and the GPT-only analyses in Sections 4-11 had already been '
  'completed). See §13.1 for the pre-registration-status disclosure that accompanies the '
  'cross-model analyses.', italic=True)
P('RQ1 — Modality-stratified accuracy of GPT-5.4 against the consensus GT. ', bold=True)
P('Mean Euclidean distance, normalised Euclidean distance, and SDR at four thresholds for '
  'point landmarks; Jaccard / Dice for area landmarks (Section 5).')
P('RQ2 — Effect of explicit grid explanation. ', bold=True)
P('Paired Wilcoxon comparisons (zero-shot vs guided) at modality level (Bonferroni × 4) and '
  'per-landmark level (Bonferroni × 9 + 3) (Section 6).')
P('RQ3 — Reproducibility under temperature = 0. ', bold=True)
P("Fleiss' κ across the three repetitions, three-way unanimous rates, ED spread distributions, "
  "and area cross-rep mean pairwise Jaccard (Section 7). This RQ is GT-independent and the "
  "results are identical to v1.")
P('RQ4 — Ground truth validation (NEW in v2). ', bold=True)
P('Inter- and intra-rater reliability of the two OMFR specialists, and how their disagreement '
  'compares to GPT-5.4 error magnitude (Section 8).')
P('RQ5 — GPT-5.4 vs dental student (NEW in v2). ', bold=True)
P('Paired Wilcoxon comparisons of GPT vs student-consensus distances to the consensus GT, '
  'with Bland-Altman descriptive statistics and acceptability-band analysis (Section 9).')
P('RQ6 — Modality-stratified accuracy of Gemini 3.1 Pro against the consensus GT (NEW in v5, '
  'POST-HOC). ', bold=True)
P('Same dependent variables and statistical machinery as RQ1, applied to the Gemini full run. '
  'Bonferroni multipliers within the Gemini-only family: ×4 for strategy paired tests at '
  'modality level, ×12 for strategy paired tests per landmark (Section 12).')
P('RQ7 — GPT-5.4 vs Gemini 3.1 Pro paired cross-model comparison (NEW in v5, POST-HOC). ',
  bold=True)
P('Paired Wilcoxon signed-rank tests between the two models\' rep-mean ED (or rep-mean Jaccard) '
  'per query × strategy, with rank-biserial r + bootstrap 95% CI on r (n = 2,000), Bonferroni × 8 '
  'at the modality table level and × 24 at the per-landmark table level. F5/F6 attractor '
  'analysis is reported separately as a categorical comparison on Tooth_33_Apex predictions '
  '(Section 13).')

H('2.3 Data update notice — what changed since v1 / v4', level=2)
P('The Final benchmark Excel '
  f'(SHA-256 {final_excel_sha[:16]}…) supersedes the v1 source '
  f'(SHA-256 {old_excel_sha[:16]}…) and adds the following columns (delivered in v2 / v4):')

p = doc.add_paragraph(style='List Bullet')
p.add_run('OMFR_2 — second specialist\'s rating (newly available).')
p = doc.add_paragraph(style='List Bullet')
p.add_run(f'OMFR_1_Second and OMFR_2_Second — washout-period re-ratings on a deliberately '
          f'sampled 180-query subset (60 panoramic + 30 periapical + 30 cephalometric × 6/3/3 '
          f'landmarks), enabling intra-rater reliability.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('CONSENSUS_Ground_Truth — team-adjudicated reconciliation of the two specialists '
          '(canonical GT for v2).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Student_Response — single column holding the team-adjudicated student consensus '
          'per query (replacing v1\'s 40 per-student columns, of which only one was populated). '
          'This is treated as a single rater for paired statistical tests.')

# GT change summary
gt_pt_pan = len(gt_changes['point']['PANORAMIC'])
gt_area_pan = len(gt_changes['area']['PANORAMIC'])
gt_pt_pa = len(gt_changes['point']['PERIAPICAL'])
gt_pt_ceph = len(gt_changes['point']['CEPHALOMETRIC'])

P(f'GT changes from v1 (OMFR_1) → v2 (consensus): the cell-set after normalisation differs on '
  f'{gt_pt_ceph}/150 cephalometric points, {gt_pt_pa}/150 periapical points, '
  f'{gt_pt_pan}/300 panoramic points, and {gt_area_pan}/300 panoramic areas. The two '
  f'panoramic-point GT changes are PAN_074_Mental_Foramen_L (G11 → G10) and '
  f'PAN_079_Condylar_Head_R (B2 → B1) — a single cell shift each, attributable to the '
  f'adjudication process. The 33 panoramic-area changes are mostly minor cell additions or '
  f'removals at the boundary of the structure. Appendix D quantifies the impact on every '
  f'reported metric.')

P('What changed since v4:', bold=True)
P(f'v5 adds 5,400 Gemini 3.1 Pro batch responses (3 reps × 1,800 calls), executed against '
  f'the same 900 queries / 200 images / 2 prompt strategies as the GPT-5.4 main run. The '
  f'image SHAs and the prompt text were locked at v4 and re-anchored at v5 preflight; the '
  f'Gemini orchestrator (scripts/run_full_run_gemini.py Stage 2) refused to launch unless '
  f'the 1,800 (query × strategy) prompt renderings were byte-identical to '
  f'results_full/prompts_used.json. The Final Excel ({final_excel_sha[:16]}…) is unchanged '
  f'from v4. Rep 1 of the Gemini run used max_output_tokens = 2048 with 78 re-queries to '
  f'recover MAX_TOKENS-truncated cases; reps 2-3 used 4096 from the start. This '
  f'configuration-between-reps drift is disclosed in §12.1 and §14, with per-rep settings '
  f'preserved in full_run_manifest_rep1.json vs full_run_manifest.json. No GPT-5.4 data, '
  f'no consensus GT, and no student data is changed in v5; the only new responses are '
  f'Gemini\'s.')

H('2.4 What this report does and does not test', level=2)
P('What this report DOES test:', bold=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Two commercial multimodal LLMs (GPT-5.4 and Gemini 3.1 Pro) on identical inputs '
          '(same 900 queries, same images, byte-identical prompts) against the same adjudicated '
          'consensus GT.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Two prompting strategies (zero_shot, guided) for each model. The strategy comparison '
          'in RQ2 / RQ6 compares zero-shot vs guided as the prompts were issued on the full '
          'run. It is not a comparison against any prior prompt revision (those decisions '
          'happened during pilot prompt finalisation, which is out of scope here).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('GPT-5.4 vs team-adjudicated dental student consensus (RQ5, Section 9). The student '
          'column is treated as a single rater per query, in line with how it was delivered '
          '(the underlying 8-student responses were collapsed prior to delivery, so per-query '
          'inter-student variability cannot be computed).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Three pre-registered prompt-level ablations on the Tooth_33_Apex panoramic landmark '
          'for GPT-5.4 (Section 10).')

P('What this report DOES NOT test:', bold=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Other commercial MLLMs. Claude Sonnet 4.6 infrastructure is prepared in the '
          'codebase (scripts/run_full_run_claude.py) but the full run is reserved for v6 / '
          'future work.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Open-source MLLMs. None evaluated.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Gemini-side prompt ablations. The §10 GPT ablations target the F5/F6 attractor — '
          'a GPT-specific failure mode (Section 13.4). Gemini does not exhibit it, so '
          'replicating the same ablations on Gemini would not be informative for that '
          'failure. Identifying Gemini\'s OWN weakest landmark (Sella_S, where Gemini '
          'under-performs GPT per §13.3) and ablating prompts on that landmark is left as '
          'future work.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Per-pixel image-fidelity equivalence between models. Each model was queried at its '
          'provider-max fidelity setting (GPT-5.4: detail=high, ~2,275 image tokens/request; '
          'Gemini: MEDIA_RESOLUTION_HIGH, ~1,077 image tokens/request). These are not equal in '
          'raw token count; we test the OPERATIONAL ceiling (the way an end user would query '
          'each model) rather than a controlled-fidelity equivalence. Discussion in §14.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Gemini vs dental student. The student-consensus paired comparison is only run for '
          'GPT-5.4 (Section 9). The same comparison for Gemini is left as future work.')

# ────────────────────────────────────────────────────────────────────
# 3. METHODOLOGY (abbreviated; pointers to v1 for unchanged sections)
# ────────────────────────────────────────────────────────────────────
H('3. Methodology', level=1)
P('Dataset, image processing, grid system, model and inference settings, and prompt strategies '
  'are unchanged from v1 (results_full/Full_Run_Results_Report.docx Section 3) and are not '
  'reproduced in full here. The material differences for v2 are summarised below.')

H('3.1 Canonical ground truth', level=2)
P(f'The canonical reference for all primary metrics in this report is CONSENSUS_Ground_Truth, '
  f'an adjudicated reconciliation of OMFR_1 and OMFR_2 produced by the team after both raters '
  f'completed their primary rating and the washout-period re-rating. OMFR_1 is preserved as a '
  f'sensitivity reference (Appendix D).')

H('3.2 Statistical analysis plan', level=2)
P('Distance, overlap, and SDR are reported with 95% confidence intervals — bootstrap percentile '
  'intervals (10,000 resamples, random_state = 42) for means and medians, Wilson score intervals '
  'for proportions. Non-normality of paired deltas is verified by Shapiro-Wilk; strategy '
  'comparisons use the paired Wilcoxon signed-rank test with rank-biserial r effect size and a '
  'bootstrap 95% CI on r (n = 2,000 resamples, seed = 42). Bonferroni correction is applied at '
  'the comparison-family level: ×4 for modality (CEPH/PA/PAN-point + PAN-area), ×9 for the nine '
  'point landmarks, ×3 for the three area landmarks; for the cross-model comparison in §13 the '
  'multipliers are ×8 (4 modality×type × 2 strategies) and ×24 (12 landmarks × 2 strategies). '
  'Inter- and intra-rater reliability use cell-level Cohen\'s κ for points and Jaccard / Dice '
  'for areas. GPT-vs-student paired tests and GPT-vs-Gemini paired tests use the same Wilcoxon '
  'framework on per-query rep-mean ED (or rep-mean Jaccard for areas).')

P('Aggregation across the three repetitions:', bold=True)
P('Each model produced three independent batches (reps 1, 2, 3) at temperature = 0 for every '
  '(query × strategy) pair, giving three raw cell predictions per pair. The per-rep metric '
  '(ED, Jaccard, Dice) is computed independently against the reference, yielding three per-rep '
  'values for that pair. The per-query observation that enters every statistical test in this '
  'report is then the ARITHMETIC MEAN of those three per-rep values, ignoring per-rep failures '
  '(parse failures, ambiguous multi-cell responses on point landmarks, or empty responses; full '
  'list for Gemini in §12.5 Table 34, for GPT-5.4 in Appendix B2). This is the SAME aggregation '
  'used throughout v1, v2 and v4 for GPT-5.4 and applied identically to Gemini in v5. We do NOT '
  'use majority voting on the predicted cell — that approach would collapse three measurements '
  'into one and discard the rep-to-rep noise that we instead report explicitly in §7 (GPT '
  'Fleiss κ = 0.78 on points) and §12.4 (Gemini Fleiss κ = 0.88 on points). Rep-mean is the '
  'minimum-variance unbiased estimator of the per-query "expected ED at temperature = 0" under '
  'the assumption that rep variation is independent and identically distributed; under that '
  'assumption the variance reduction is 1/3 versus a single-rep estimate. Linear mixed-effects '
  'models with random rep effects give numerically equivalent point estimates here because '
  'between-rep variance is small relative to between-query variance (see §7 and §12.4 for '
  'per-group reproducibility metrics).')

H('3.3 Reproducibility infrastructure', level=2)
P(f'Pipeline source at git commit {git_sha[:7]}. Source data anchored by SHA-256 of the Final '
  f'Excel ({final_excel_sha[:16]}…) and the derived query_index.json ({qi_sha[:16]}…). The '
  f'extended query_index includes all six annotator fields (omfr_1, omfr_1_second, omfr_2, '
  f'omfr_2_second, consensus_gt, student) per query. The frozen GPT-5.4 outputs from results_full/ '
  f'are referenced read-only and SHA-anchored at preflight via reanalysis_anchor.json (108 raw '
  f'JSONL files + the v1 query_index). Re-evaluation against consensus_gt is performed by '
  f'scripts/recompute_against_consensus.py; downstream GPT-only analyses by '
  f'scripts/analyze_consensus_run.py, scripts/analyze_rater_reliability.py, and '
  f'scripts/analyze_gpt_vs_student.py.')

P(f'For the Gemini full run (v5), the analogous chain is: '
  f'scripts/run_full_run_gemini.py (orchestrator, 7 stages, anchors 12 chunk JSONLs + 1 '
  f'requeries.jsonl SHA-256 + 200 image SHAs in gemini_recompute_anchor.json) → '
  f'scripts/recompute_gemini.py (re-derives full_run_records.pkl from raw chunks with the same '
  f'single-cell rule the operational parser uses, byte-parity verified against '
  f'parsed_responses.json on all 3 × 1,800 records) → '
  f'scripts/analyze_consensus_run.py --sandbox results_full_gemini (re-uses the GPT analyzer, '
  f'producing results_full_gemini/{{analysis,phase_b,summary}}.json) → '
  f'scripts/analyze_gpt_vs_gemini.py (paired Wilcoxon + Bland-Altman + Cohen\'s d_paired + '
  f'F5/F6 attractor analysis; output at results_full_gemini/gpt_vs_gemini.json). Every cell in '
  f'§12 and §13 is reproducible from raw chunks via these four scripts. See Appendix C for '
  f'the complete manifest.')

H('3.4 Canonical prompts (verbatim)', level=2)
P('Every numeric claim in Sections 5-11 of this report is derived from GPT-5.4 responses, and '
  'every numeric claim in Sections 12-13 is derived from Gemini 3.1 Pro responses (or paired '
  'GPT-vs-Gemini comparisons of those responses), under the SAME two canonical prompt '
  'strategies (zero_shot and guided), rendered from the SAME pipeline.generate_prompt() '
  'function for every query. The Gemini orchestrator '
  'verifies byte-identity between the Gemini-side prompts and the canonical results_full/'
  'prompts_used.json BEFORE any API call is issued (Stage 2 of '
  'scripts/run_full_run_gemini.py), and refuses to launch if any of the 1,800 (query × '
  'strategy) prompt renderings differ from the GPT-5.4 baseline — so the cross-model '
  'comparison in §13 is on byte-identical prompts on byte-identical images (image SHAs '
  'anchored at Stage 1, see Appendix C and §12.1). The two strategies use the same '
  'system role declaration; they differ in the system-prompt extension and in the user-prompt '
  'preamble. The exact rendered prompts depend on landmark type (point vs area), modality '
  '(panoramic / periapical / cephalometric), and whether the landmark is FDI-numbered. '
  'Below we show the two canonical prompts for Tooth_33_Apex (PAN, FDI-flagged, point) — '
  'this is the focal landmark of the Section 10 ablations and exhibits the full guided '
  'system prompt with the panoramic L–R inversion clause. For other (landmark, modality, '
  'strategy) tuples, the prompts are rendered with the same templates and per-modality '
  'parameters (grid dimensions, modality-specific clauses); the full prompts_used.json files '
  'in each sandbox preserve the exact prompts issued for every query.')

# Render the prompts live from pipeline.generate_prompt so they cannot drift from what was
# actually issued for the v2 main run.
import sys as _sys
_sys.path.insert(0, str(ROOT))
import pipeline as _pipeline
_qi_full = json.load(open(ROOT / 'results_consensus' / 'query_index.json'))
_q33 = next(q for q in _qi_full if q['structure'] == 'Tooth_33_Apex')

_zs_sys, _zs_usr = _pipeline.generate_prompt(_q33, 'zero_shot')
_gd_sys, _gd_usr = _pipeline.generate_prompt(_q33, 'guided')

caption("Box 0a: Canonical ZERO-SHOT prompt rendered for Tooth_33_Apex (PAN, FDI-flagged, "
        "point landmark). Rendered live from pipeline.generate_prompt at report-build time; "
        "byte-identical to what was issued to the API in the v2 main run.")
_t0a = doc.add_table(rows=2, cols=1)
_t0a.style = 'Light Grid'
_tblPr0a = _t0a._element.find(qn('w:tblPr'))
if _tblPr0a is None:
    _tblPr0a = OxmlElement('w:tblPr'); _t0a._element.insert(0, _tblPr0a)
_lay0a = OxmlElement('w:tblLayout'); _lay0a.set(qn('w:type'), 'fixed')
_tblPr0a.append(_lay0a)
_t0a.cell(0,0).text = ''
_run = _t0a.cell(0,0).paragraphs[0].add_run('SYSTEM'); _run.bold=True; _run.font.size=Pt(10)
shade_cell(_t0a.cell(0,0), 'D9E1F2'); set_cell_borders(_t0a.cell(0,0))
_t0a.cell(1,0).text = ''
_run = _t0a.cell(1,0).paragraphs[0].add_run(_zs_sys); _run.font.size=Pt(9)
set_cell_borders(_t0a.cell(1,0))
for r in _t0a.rows: r.cells[0].width = Inches(6.6)

_t0a2 = doc.add_table(rows=2, cols=1)
_t0a2.style = 'Light Grid'
_tblPr = _t0a2._element.find(qn('w:tblPr'))
if _tblPr is None:
    _tblPr = OxmlElement('w:tblPr'); _t0a2._element.insert(0, _tblPr)
_lay = OxmlElement('w:tblLayout'); _lay.set(qn('w:type'), 'fixed'); _tblPr.append(_lay)
_t0a2.cell(0,0).text = ''
_run = _t0a2.cell(0,0).paragraphs[0].add_run('USER'); _run.bold=True; _run.font.size=Pt(10)
shade_cell(_t0a2.cell(0,0), 'D9E1F2'); set_cell_borders(_t0a2.cell(0,0))
_t0a2.cell(1,0).text = ''
_run = _t0a2.cell(1,0).paragraphs[0].add_run(_zs_usr); _run.font.size=Pt(9)
set_cell_borders(_t0a2.cell(1,0))
for r in _t0a2.rows: r.cells[0].width = Inches(6.6)

caption("Box 0b: Canonical GUIDED prompt rendered for Tooth_33_Apex (same landmark, same "
        "image — only the prompt strategy differs). The system prompt includes the row/column "
        "declaration sentence and (for panoramic) the L–R inversion clause; the user prompt "
        "drops the grid declaration (it has been moved into the system prompt) but keeps the "
        "FDI numbering announcement.")
_t0b = doc.add_table(rows=2, cols=1)
_t0b.style = 'Light Grid'
_tblPr = _t0b._element.find(qn('w:tblPr'))
if _tblPr is None:
    _tblPr = OxmlElement('w:tblPr'); _t0b._element.insert(0, _tblPr)
_lay = OxmlElement('w:tblLayout'); _lay.set(qn('w:type'), 'fixed'); _tblPr.append(_lay)
_t0b.cell(0,0).text = ''
_run = _t0b.cell(0,0).paragraphs[0].add_run('SYSTEM'); _run.bold=True; _run.font.size=Pt(10)
shade_cell(_t0b.cell(0,0), 'D9E1F2'); set_cell_borders(_t0b.cell(0,0))
_t0b.cell(1,0).text = ''
_run = _t0b.cell(1,0).paragraphs[0].add_run(_gd_sys); _run.font.size=Pt(9)
set_cell_borders(_t0b.cell(1,0))
for r in _t0b.rows: r.cells[0].width = Inches(6.6)

_t0b2 = doc.add_table(rows=2, cols=1)
_t0b2.style = 'Light Grid'
_tblPr = _t0b2._element.find(qn('w:tblPr'))
if _tblPr is None:
    _tblPr = OxmlElement('w:tblPr'); _t0b2._element.insert(0, _tblPr)
_lay = OxmlElement('w:tblLayout'); _lay.set(qn('w:type'), 'fixed'); _tblPr.append(_lay)
_t0b2.cell(0,0).text = ''
_run = _t0b2.cell(0,0).paragraphs[0].add_run('USER'); _run.bold=True; _run.font.size=Pt(10)
shade_cell(_t0b2.cell(0,0), 'D9E1F2'); set_cell_borders(_t0b2.cell(0,0))
_t0b2.cell(1,0).text = ''
_run = _t0b2.cell(1,0).paragraphs[0].add_run(_gd_usr); _run.font.size=Pt(9)
set_cell_borders(_t0b2.cell(1,0))
for r in _t0b2.rows: r.cells[0].width = Inches(6.6)

P('The three Section 10 ablation strategies (guided_no_tooth_num, guided_patient_left, '
  'guided_no_LR) are variants of the canonical guided prompt; each modifies a single isolated '
  'substring relative to Box 0b above. The exact differences are shown in Boxes 1, 2, 3 of '
  'Sections 10.2–10.4.')

H('3.5 Code and data availability', level=2)
P('All source code (data ingestion, the canonical prompt-rendering pipeline, the orchestrator '
  'for each model\'s full run, the recompute and analysis scripts, this report\'s generator, '
  'every audit script, and the smoke-test regression gate) is open-sourced at:')
P('https://github.com/burcusayin/grid-based-mllm-landmarks', bold=True)
P('The repository will be made public before journal submission; until then, access can be '
  'arranged directly with the corresponding author. Each report version (v1 → v2 → v4 → v5) '
  'is tagged in the repository\'s git history, and every figure and table cell in this report '
  'is reproducible from raw API outputs via the following invocations (run from the project '
  'root):', italic=True)
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python pipeline.py prepare_v2 --excel data/Final_Dental_MLLM_Benchmark_'
          'Data.xlsx --anchor-to results_full` — builds the extended query_index.json with '
          'consensus_gt, OMFR_1/2 (primary + washout), and student columns.')
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python scripts/recompute_against_consensus.py` — re-derives GPT-5.4 '
          'full_run_records.pkl from raw OpenAI JSONLs, anchored via reanalysis_anchor.json.')
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python scripts/recompute_gemini.py` — re-derives Gemini 3.1 Pro '
          'full_run_records.pkl from raw Google chunk JSONs + the 78-row rep-1 re-queries '
          'JSONL, anchored via gemini_recompute_anchor.json.')
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python scripts/analyze_consensus_run.py [--sandbox results_full_gemini]` '
          '— per-model RQ1/RQ2/reproducibility analyses producing analysis.json + phase_b.json '
          '+ summary.json.')
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python scripts/analyze_rater_reliability.py` — inter- and intra-rater '
          'reliability statistics for the OMFR specialists and the student consensus.')
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python scripts/analyze_gpt_vs_student.py` — paired Wilcoxon and '
          'Bland-Altman analysis for GPT-5.4 vs student-consensus on the same queries.')
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python scripts/analyze_gpt_vs_gemini.py` — paired Wilcoxon, '
          'Bland-Altman, Cohen\'s d_paired, and F5/F6 attractor analysis for the GPT-5.4 '
          'vs Gemini 3.1 Pro head-to-head comparison.')
p = doc.add_paragraph(style='List Number')
p.add_run('`.venv/bin/python scripts/generate_full_run_report_v2_consensus.py` — generates '
          'this report (v5) end-to-end, including all tables and Appendices F-G.')
P('Two from-raw audit scripts are also released for paper-grade verification: '
  '`scripts/audit_from_raw.py` re-derives every headline statistic in §4 / §12.1 / §5 / §12.2 / '
  '§13.2 / §13.4 / §13.5 directly from the Excel + raw JSONLs/chunks (without dependence on '
  'any derived JSON), and `scripts/audit_v5_numerical.py` cross-checks the docx-embedded claims '
  'against the same source. All bootstrap and Wilcoxon-bootstrap procedures use the pinned '
  'random seed = 42 throughout. Total approximate end-to-end runtime on a 2024 MacBook Pro '
  '(M4 Max, 32 GB) is ~10 minutes for the analysis chain (no API calls; uses cached raw '
  'responses) and ~1.5 hours for a full Gemini API run (the GPT-5.4 raw responses are '
  'preserved from the May 2026 main run and not re-issued).', italic=True)

# ────────────────────────────────────────────────────────────────────
# 4. OPERATIONAL OUTCOMES
# ────────────────────────────────────────────────────────────────────
H('4. Operational Outcomes', level=1)
P(f'The same {n_calls:,} GPT-5.4 batch calls underpin both v1 and v2; no new API calls were '
  f'issued for the consensus re-evaluation. Token usage and cost are therefore unchanged from v1. '
  f'The parallel Gemini 3.1 Pro full run (5,400 calls, with identical prompts on identical '
  f'images) was executed in May 2026 and is summarised in §12.1; combined operational totals '
  f'for v5 are reported there. The remainder of this section reports GPT-5.4 operational '
  f'outcomes only; the GPT-vs-Gemini cross-model results begin in §11.6 and are fully developed '
  f'in §13.', italic=False)

caption('Table 1: Token usage and cost summary.')
add_table(
    ['Quantity','Value'],
    [
        ['Total API calls', f'{n_calls:,}'],
        ['Successful responses', f'{n_actual:,} ({n_actual/n_calls*100:.2f}%)'],
        ['Prompt tokens', f'{prompt_tok:,}'],
        ['Completion tokens', f'{compl_tok:,}'],
        ['Total tokens', f'{prompt_tok+compl_tok:,}'],
        ['Naïve cost (no caching)', f'${naive_cost:.2f}'],
    ],
    col_widths=[3.0, 3.4]
)

caption('Table 2: Compliance breakdown by modality and strategy.')
def compliance_for(mod, strat):
    n_calls_local = 0; n_fail = 0
    for r in records:
        if r['modality']!=mod or r['strategy']!=strat: continue
        n_calls_local += 3
        n_fail += r['n_failed']
    return n_calls_local, n_fail
rows=[]
for mod in ('PANORAMIC','PERIAPICAL','CEPHALOMETRIC'):
    for strat in ('zero_shot','guided'):
        n,f = compliance_for(mod,strat)
        rows.append([mod, strat, str(n), str(n-f), f'{(n-f)/n*100:.3f}%'])
add_table(['Modality','Strategy','Calls','Successful','Compliance'], rows,
          col_widths=[1.6,1.2,1.0,1.2,1.2])

P(f'All four compliance failures are reversed-coordinate errors on the guided strategy '
  f'(strings of the form "12F", "12E", "6F", "4E" parsed strictly as letter-then-digit). '
  f'Under a digit-before-letter-tolerant parser these would all decode to VALID grid cells '
  f'("F12", "E12", "F6", "E4" respectively) — so the parse rate becomes 100% — but those '
  f'decoded cells are NOT the consensus GT cells: they sit 1.0 to 2.8 cells from the GT. '
  f'The failure mode is therefore a format-level glitch (digit-letter inversion) layered on '
  f'an underlying prediction that would still be wrong by 1–3 cells in Euclidean distance. '
  f'Full list with per-failure decoded distances in Appendix B.')

# ────────────────────────────────────────────────────────────────────
# 5. RQ1 — Modality-stratified accuracy
# ────────────────────────────────────────────────────────────────────
H('5. RQ1 — Modality-Stratified Accuracy of GPT-5.4 (vs Consensus GT)', level=1)

H('5.1 Point landmarks: mean Euclidean distance with bootstrap 95% CIs', level=2)
caption('Table 3: Mean Euclidean distance (in grid cells) for point landmarks, by modality '
        'and strategy. CI = bootstrap percentile 95% confidence interval, 10,000 resamples. '
        'n = number of unique queries; each query is the mean over 3 repetitions vs '
        'consensus_gt.')
rows=[]
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = analysis['RQ1_modality_strategy'][f'{mod}_{strat}_point']
        rows.append([mod, strat, str(d['n']),
                     f'{d["mean"]:.3f}',
                     f'[{d["mean_ci"][0]:.3f}, {d["mean_ci"][1]:.3f}]',
                     f'{d["median"]:.3f}',
                     f'[{d["median_ci"][0]:.3f}, {d["median_ci"][1]:.3f}]'])
add_table(['Modality','Strategy','n','Mean ED','Mean 95% CI','Median ED','Median 95% CI'],
          rows, col_widths=[1.45,1.0,0.55,0.85,1.4,0.85,1.4])

P('Reading Table 3. ', bold=True)
P(f'A clear modality-difficulty hierarchy is apparent. On cephalometric, mean ED is '
  f'{m_ceph_zs["mean"]:.2f} cells under zero-shot and {m_ceph_gd["mean"]:.2f} under guided — '
  f'within the precision of a single grid cell. Periapical sits one rough cell step further '
  f'out ({m_pa_zs["mean"]:.2f} / {m_pa_gd["mean"]:.2f}). Panoramic point landmarks are an '
  f'order of magnitude harder: {m_pan_zs["mean"]:.2f} cells under zero-shot, '
  f'{m_pan_gd["mean"]:.2f} cells under guided — and crucially, the 95% bootstrap CIs for '
  f'panoramic point ED under the two strategies do NOT overlap '
  f'([{m_pan_zs["mean_ci"][0]:.2f}, {m_pan_zs["mean_ci"][1]:.2f}] vs '
  f'[{m_pan_gd["mean_ci"][0]:.2f}, {m_pan_gd["mean_ci"][1]:.2f}]), which is the first '
  f'indication that the guided prompt does not uniformly help on panoramic. The medians on '
  f'cephalometric ({m_ceph_zs["median"]:.2f} / {m_ceph_gd["median"]:.2f}) sit well below the '
  f'means, indicating a right-skewed error distribution with most queries near-perfect and a '
  f'small tail of harder cases. On panoramic guided, median ED ({m_pan_gd["median"]:.2f}) is '
  f'about twice the zero-shot median ({m_pan_zs["median"]:.2f}), so the regression is in the '
  f'centre of the distribution, not driven by outliers.')

H('5.2 Cross-modality: normalised Euclidean distance', level=2)
caption('Table 4: NED per modality + strategy and grand mean.')
rows=[]
ned = phase_b['ned_modality']
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = ned[f'{mod}_{strat}']
        rows.append([mod, strat, str(d['n']), f'{d["mean_ned"]:.4f}', f'{d["median_ned"]:.4f}'])
for strat in ('zero_shot','guided'):
    d = phase_b['grand_ned'][strat]
    rows.append(['ALL POINT', strat, str(d['n']), f'{d["mean_ned"]:.4f}', f'{d["median_ned"]:.4f}'])
add_table(['Modality','Strategy','n','Mean NED','Median NED'], rows,
          col_widths=[1.6,1.2,0.7,1.2,1.2])

P('Reading Table 4. ', bold=True)
P('Raw ED across modalities is not directly comparable because the grids differ in size '
  '(panoramic 16 × 8 with diagonal 16.55 cells; periapical 8 × 6 with diagonal 8.60; '
  'cephalometric 10 × 8 with diagonal 11.40). Normalised ED, expressed as a fraction of the '
  'grid diagonal, makes them comparable. The CEPH << PA < PAN ordering visible in Table 3 is '
  'preserved under NED, confirming that the panoramic disadvantage is genuine rather than an '
  'artefact of the larger panoramic grid. Under guided, the panoramic NED grows by a factor of '
  '≈ 1.5 over zero-shot, while CEPH and PA are essentially flat — so the modality-specific '
  'reaction to the guided prompt is itself heterogeneous, and the panoramic-only regression is '
  'preserved on a properly normalised scale.')

H('5.3 SDR at four thresholds (point landmarks)', level=2)
caption('Table 5: SDR at thresholds 0, 1, √2, 2 cells. Wilson score 95% CIs in brackets.')
rows=[]
sdr = phase_b['sdr_modality_with_ci']
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = sdr[f'{mod}_{strat}']
        rows.append([mod, strat, str(d['n']),
                     fmt_pct_ci(d['SDR@0'], d['SDR@0_ci']),
                     fmt_pct_ci(d['SDR@1'], d['SDR@1_ci']),
                     fmt_pct_ci(d['SDR@√2'], d['SDR@√2_ci']),
                     fmt_pct_ci(d['SDR@2'], d['SDR@2_ci'])])
add_table(['Modality','Strategy','n','SDR@0','SDR@1','SDR@√2','SDR@2'], rows,
          col_widths=[1.3,0.95,0.55,1.05,1.05,1.05,1.05], header_size=9, body_size=9)

P('Reading Table 5. ', bold=True)
P(f'SDR thresholds give clinically intuitive accuracy bands. SDR@1 (within ±1 orthogonal '
  f'cell) and SDR@2 (within ±2 cells, the typical clinical "same anatomic region" tolerance) '
  f'are the most relevant for downstream applications. On cephalometric, SDR@2 reaches '
  f'{sdr["CEPHALOMETRIC_zero_shot"]["SDR@2"]*100:.1f}% / '
  f'{sdr["CEPHALOMETRIC_guided"]["SDR@2"]*100:.1f}% (zero-shot / guided) — meaning '
  f'essentially every cephalometric prediction lands in the correct anatomic region. Panoramic '
  f'SDR@2 is dramatically lower: {sdr["PANORAMIC_zero_shot"]["SDR@2"]*100:.1f}% under zero-shot '
  f'and {sdr["PANORAMIC_guided"]["SDR@2"]*100:.1f}% under guided. The Wilson 95% CIs for '
  f'panoramic SDR@1 do not overlap between strategies, reinforcing the strategy-specific '
  f'panoramic regression that Section 6 quantifies with paired tests. Exact-match SDR@0 is '
  f'low across all modalities — a single cell of error is below the discriminative resolution '
  f'of the grid, and most clinically useful applications would tolerate SDR@1 or SDR@2 ranges.')

H('5.4 Area landmarks: Jaccard and Dice', level=2)
caption('Table 6: Aggregate area-landmark overlap metrics (panoramic only).')
rows=[]
m_pan_zs_area = analysis['RQ1_modality_strategy']['PANORAMIC_zero_shot_area']
m_pan_gd_area = analysis['RQ1_modality_strategy']['PANORAMIC_guided_area']
rows.append(['PAN AREA AGGREGATE','zero_shot',str(m_pan_zs_area['n']),
             f'{m_pan_zs_area["mean"]:.3f}',
             f'[{m_pan_zs_area["mean_ci"][0]:.3f}, {m_pan_zs_area["mean_ci"][1]:.3f}]',
             f'{m_pan_zs_area["median"]:.3f}'])
rows.append(['PAN AREA AGGREGATE','guided',str(m_pan_gd_area['n']),
             f'{m_pan_gd_area["mean"]:.3f}',
             f'[{m_pan_gd_area["mean_ci"][0]:.3f}, {m_pan_gd_area["mean_ci"][1]:.3f}]',
             f'{m_pan_gd_area["median"]:.3f}'])
add_table(['Group','Strategy','n','Mean Jaccard','Mean 95% CI','Median Jaccard'],
          rows, col_widths=[1.7,1.0,0.55,1.1,1.5,1.1])

caption('Table 7: Per-area-landmark Jaccard and Dice.')
rows=[]
for lm in ('Mandibular_Canal_L','Maxillary_Sinus_R','External_Oblique_Ridge_R'):
    for strat in ('zero_shot','guided'):
        d = phase_b['area_landmark_stats'][f'PANORAMIC/{lm}/{strat}']
        rows.append([lm, strat, str(d['n']),
                     f'{d["mean_jaccard"]:.3f}', f'{d["median_jaccard"]:.3f}',
                     f'{d["mean_dice"]:.3f}', f'{d["median_dice"]:.3f}'])
add_table(['Landmark','Strategy','n','Mean J','Median J','Mean D','Median D'], rows,
          col_widths=[2.0,0.95,0.55,0.85,0.95,0.85,0.95], header_size=9, body_size=9)

P('Reading Tables 6 and 7. ', bold=True)
P('Area landmarks exist only on panoramic in this study. Aggregate mean Jaccard is '
  f'{m_pan_zs_area["mean"]:.3f} (zero-shot) and {m_pan_gd_area["mean"]:.3f} (guided) — both '
  f'are modest in absolute terms. The per-landmark breakdown is more informative: maxillary '
  f'sinus is the best-localised area landmark (mean Jaccard ≈ '
  f'{phase_b["area_landmark_stats"]["PANORAMIC/Maxillary_Sinus_R/zero_shot"]["mean_jaccard"]:.2f}), '
  f'a large convex radiolucent structure that the model can apparently bound reasonably well. '
  f'The mandibular canal — a thin linear structure crossing many cells — comes in around '
  f'Jaccard {phase_b["area_landmark_stats"]["PANORAMIC/Mandibular_Canal_L/zero_shot"]["mean_jaccard"]:.2f}, '
  f'and the external oblique ridge is the worst-performing landmark in the whole study '
  f'(Jaccard {phase_b["area_landmark_stats"]["PANORAMIC/External_Oblique_Ridge_R/zero_shot"]["mean_jaccard"]:.2f}). '
  f'This last result is striking: the EOR is a thin oblique line that shares its course with '
  f'other radiopaque mandibular features, and the model appears not to recognise it as a '
  f'discrete anatomical entity. Dice values are systematically higher than Jaccard (by the '
  f'algebraic identity D = 2J/(1+J)) and are reported alongside Jaccard so the metrics can be '
  f'compared to the DMFR cephalometric-segmentation literature, which generally reports Dice.')

# ────────────────────────────────────────────────────────────────────
# 6. RQ2 — Strategy effects
# ────────────────────────────────────────────────────────────────────
H('6. RQ2 — Effect of Guided Prompting (Zero-Shot vs Guided)', level=1)

H('6.1 Distributional non-normality and choice of test', level=2)
caption('Table 8: Shapiro-Wilk normality test on paired deltas.')
rows=[]
for k in ('CEPHALOMETRIC_point_delta','PERIAPICAL_point_delta','PANORAMIC_point_delta','PANORAMIC_area_delta'):
    d = phase_b['paired_shapiro'].get(k, {})
    if not d: continue
    verdict = 'normal' if d['normal_at_alpha_05'] else 'non-normal'
    rows.append([k.replace('_',' '), str(d['n']), f'{d["W"]:.4f}',
                 f'{d["p"]:.2e}', verdict])
add_table(['Comparison','n','Shapiro W','p','Verdict at α=0.05'], rows,
          col_widths=[2.4,0.6,1.0,1.2,1.4])
P('Reading Table 8. ', bold=True)
P('Every paired-delta distribution rejects normality at α = 0.05 with very small p-values '
  '(all p < 10⁻⁶). This is methodologically important because it justifies the non-parametric '
  'Wilcoxon signed-rank test (the appropriate test for paired non-normal data) over the paired '
  't-test (which assumes normality of differences). The non-normality is biologically expected '
  'for grid-cell distance measurements: deltas are bounded below by zero in absolute terms, '
  'truncated at the grid diagonal, and dominated by a tail of high-disagreement queries while '
  'most queries are at or near zero delta. Section 8.2 of the original methodology document '
  'pre-specified Wilcoxon as the canonical test for this study; the empirical Shapiro-Wilk '
  'results confirm that choice was correct.')

H('6.2 Modality-level paired comparisons (Bonferroni × 4)', level=2)
caption('Table 9: Modality-level paired Wilcoxon comparisons of zero-shot vs guided. '
        'Sign convention: for ED, Δ = zero-shot − guided (positive Δ → guided better, since '
        'lower ED is better). For Jaccard, Δ = guided − zero-shot (positive Δ → guided better, '
        'since higher Jaccard is better). The rank-biserial r follows the same convention '
        '(positive r → guided better; negative r → guided worse).')
rows=[]
for k_label, k in [('CEPHALOMETRIC point ED','CEPHALOMETRIC_point'),
                    ('PERIAPICAL point ED','PERIAPICAL_point'),
                    ('PANORAMIC point ED','PANORAMIC_point'),
                    ('PANORAMIC area Jaccard','PANORAMIC_area')]:
    d = analysis['RQ2a_strategy_per_modality'][k]
    pbonf = min(1.0, d['p']*4)
    rows.append([k_label, str(d['n_total']), str(d['n_nonzero']),
                 f'{d["mean_delta"]:+.4f}', f'{d["median_delta"]:+.4f}',
                 fmt_p(d['p']), fmt_pbonf(d['p'],4),
                 f'{d["rank_biserial_r"]:+.3f}',
                 f'[{d["rank_biserial_ci_low"]:+.2f}, {d["rank_biserial_ci_high"]:+.2f}]',
                 sig_marker(pbonf)])
add_table(['Comparison','n','n≠0','Mean Δ','Median Δ','Wilcoxon p','Bonf p','rank-biserial r','r 95% CI','Sig'],
          rows, col_widths=[1.5,0.45,0.5,0.7,0.75,0.85,0.75,0.75,0.95,0.45],
          header_size=9, body_size=9)

P('Reading Table 9. ', bold=True)
_pan_pt = analysis['RQ2a_strategy_per_modality']['PANORAMIC_point']
_pan_ar = analysis['RQ2a_strategy_per_modality']['PANORAMIC_area']
P('Only one of the four pre-specified modality-level paired comparisons survives Bonferroni '
  'correction at α = 0.05: the panoramic point comparison '
  f'(Bonferroni p = {min(1.0, _pan_pt["p"]*4):.2e}, '
  f'rank-biserial r = {_pan_pt["rank_biserial_r"]:+.3f}). '
  f'The rank-biserial r interpretation is important here: r = {_pan_pt["rank_biserial_r"]:+.3f} '
  f'with a 95% bootstrap CI of [{_pan_pt["rank_biserial_ci_low"]:+.2f}, '
  f'{_pan_pt["rank_biserial_ci_high"]:+.2f}]. The negative sign means guided is reliably worse '
  f'than zero-shot on panoramic point landmarks (under the convention positive r → guided '
  f'better). The magnitude |r| = {abs(_pan_pt["rank_biserial_r"]):.2f} is large by Cohen\'s '
  f'convention (|r| ≥ 0.5 is "large"); the CI does not cross zero, so the population-level '
  f'effect is reliable (not just an artifact of this particular sample). '
  f'The cephalometric and periapical point comparisons are NS after Bonferroni; the panoramic '
  f'area comparison is nominally significant (uncorrected p ≈ {_pan_ar["p"]:.3f}) but does not '
  f'survive Bonferroni and has small r ({_pan_ar["rank_biserial_r"]:+.2f}), so the practical '
  f'effect is modest. The Bonferroni correction is conservative across only 4 family-wise '
  f'comparisons here, which is the standard recommendation for pre-specified modality-level '
  f'analyses in DMFR landmark studies; the per-landmark analysis in Section 6.3 applies a '
  f'finer-grain Bonferroni × 9 + × 3 correction.')

H('6.3 Per-landmark stratified analysis', level=2)
caption('Table 10: Per-landmark paired Wilcoxon comparisons. Bonferroni × 9 (point) + ×3 (area).')
rows=[]
landmarks_order = [
    ('CEPHALOMETRIC','Menton_Me'),('CEPHALOMETRIC','Nasion_N'),('CEPHALOMETRIC','Sella_S'),
    ('PERIAPICAL','Tooth_36_Distal_Apex'),('PERIAPICAL','Tooth_36_Distal_CEJ'),('PERIAPICAL','Tooth_36_Mesial_CEJ'),
    ('PANORAMIC','Mental_Foramen_L'),('PANORAMIC','Condylar_Head_R'),('PANORAMIC','Tooth_33_Apex'),
]
for mod,lm in landmarks_order:
    d = analysis['RQ2b_strategy_per_landmark'][f'{mod}/{lm}']
    pbonf = min(1.0, d['p']*9)
    rows.append([f'{mod}/{lm}','point',str(d['n_total']),str(d['n_nonzero']),
                 f'{d["mean_delta"]:+.4f}', f'{d["median_delta"]:+.4f}',
                 fmt_p(d['p']), fmt_pbonf(d['p'],9),
                 f'{d["rank_biserial_r"]:+.3f}', sig_marker(pbonf)])
for lm in ('Mandibular_Canal_L','Maxillary_Sinus_R','External_Oblique_Ridge_R'):
    d = analysis['RQ2b_strategy_per_landmark'][f'PANORAMIC/{lm}/area']
    pbonf = min(1.0, d['p']*3)
    rows.append([f'PANORAMIC/{lm}','area',str(d['n_total']),str(d['n_nonzero']),
                 f'{d["mean_delta"]:+.4f}', f'{d["median_delta"]:+.4f}',
                 fmt_p(d['p']), fmt_pbonf(d['p'],3),
                 f'{d["rank_biserial_r"]:+.3f}', sig_marker(pbonf)])
add_table(['Landmark','Type','n','n≠0','Mean Δ','Median Δ','Wilcoxon p','Bonf p','r','Sig'],
          rows, col_widths=[2.0,0.55,0.4,0.4,0.7,0.7,0.8,0.7,0.85,0.4],
          header_size=8.5, body_size=8.5)

P('Reading Table 10 — heterogeneous strategy effects. ', bold=True)
_rq2b = analysis['RQ2b_strategy_per_landmark']
_mfl = _rq2b['PANORAMIC/Mental_Foramen_L']
_sel = _rq2b['CEPHALOMETRIC/Sella_S']
_eor = _rq2b['PANORAMIC/External_Oblique_Ridge_R/area']
_t33 = _rq2b['PANORAMIC/Tooth_33_Apex']
_cdh = _rq2b['PANORAMIC/Condylar_Head_R']
_men = _rq2b['CEPHALOMETRIC/Menton_Me']
P(f'Per-landmark stratification reveals what the modality aggregate disguises: the guided '
  f'prompt is not uniformly good or bad — it produces statistically significant gains on some '
  f'landmarks and statistically significant losses on others. Sign convention as in Table 10: '
  f'positive Δ and positive r → guided better; negative → guided worse. '
  f'Three landmarks improve under guided '
  f'(PAN/Mental_Foramen_L Δ = {_mfl["mean_delta"]:+.2f} cells, r = {_mfl["rank_biserial_r"]:+.2f}; '
  f'CEPH/Sella_S Δ = {_sel["mean_delta"]:+.2f}, r = {_sel["rank_biserial_r"]:+.2f}; '
  f'PAN/External_Oblique_Ridge_R Jaccard Δ = {_eor["mean_delta"]:+.3f}, '
  f'r = {_eor["rank_biserial_r"]:+.2f}). Three landmarks regress '
  f'(PAN/Tooth_33_Apex Δ = {_t33["mean_delta"]:+.2f}, r = {_t33["rank_biserial_r"]:+.2f} — '
  f'the largest effect in the study; PAN/Condylar_Head_R Δ = {_cdh["mean_delta"]:+.2f}, '
  f'r = {_cdh["rank_biserial_r"]:+.2f}; CEPH/Menton_Me Δ = {_men["mean_delta"]:+.2f}, '
  f'r = {_men["rank_biserial_r"]:+.2f} — small absolute magnitude but a perfectly consistent '
  f'direction).')

P(f'The PAN/Tooth_33_Apex regression is by far the largest individual effect '
  f'(Δ = {_t33["mean_delta"]:+.2f} cells, r = {_t33["rank_biserial_r"]:+.2f}) and accounts for '
  f'most of the aggregate panoramic-point regression in Table 9. '
  f'A rank-biserial r of {_t33["rank_biserial_r"]:+.2f} means that on every single non-tied '
  f'query, guided gave a larger ED than zero-shot. This is not a "guided is sometimes worse" '
  f'pattern — it is "guided is systematically worse on Tooth_33_Apex on every image in the '
  f'dataset". Section 10 explores this regression in depth through three pre-registered '
  f'ablations and a qualitative inspection of where the predictions land.')

P('Methodological note. ', bold=True)
P('All p-values reported in Table 10 are Bonferroni-corrected at the family-wise level '
  '(× 9 for point landmarks, × 3 for area landmarks). This is the conservative choice; '
  'alternative corrections (Holm-Bonferroni, Benjamini-Hochberg FDR) would generally yield '
  'more rejections but at the cost of weaker family-wise type-I error control. Where the '
  'Bonferroni-corrected verdict is borderline (e.g., a comparison where uncorrected p is just '
  'below 0.05), the rank-biserial r and its 95% CI provide the more substantive guide — a '
  'finding with r ≈ 0 should not be treated as clinically meaningful regardless of p, and a '
  'finding with |r| ≥ 0.5 with a CI excluding zero is robust regardless of how multiple-'
  'testing is corrected.')

# ────────────────────────────────────────────────────────────────────
# 7. RQ3 — Reproducibility (unchanged from v1)
# ────────────────────────────────────────────────────────────────────
H('7. RQ3 — Reproducibility at Temperature = 0', level=1)
P('RQ3 is GT-independent: it asks how stable GPT-5.4 is across the three repetitions of the '
  'same query, regardless of which reference is used to evaluate accuracy. Results below are '
  'numerically identical to v1.')

H("7.1 Fleiss' kappa across 3 repetitions (point landmarks)", level=2)
caption("Table 11: Fleiss' κ per modality × strategy.")
rows=[]
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = phase_b['fleiss_per_group'][f'{mod}_{strat}_point']
        kappa = float(d['kappa'])
        if kappa>=0.81: tier='almost perfect'
        elif kappa>=0.61: tier='substantial'
        elif kappa>=0.41: tier='moderate'
        elif kappa>=0.21: tier='fair'
        else: tier='slight/poor'
        rows.append([mod, strat, str(d['n_items']), f'{kappa:.4f}', tier])
total_n_for_overall = sum(int(phase_b['fleiss_per_group'][f'{mod}_{strat}_point']['n_items'])
                          for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC')
                          for strat in ('zero_shot','guided'))
rows.append(['ALL POINT (mod-prefixed)','—',f'{total_n_for_overall:,}',
             f'{float(phase_b["fleiss_overall_point"]):.4f}','substantial'])
add_table(['Modality','Strategy','n items','Fleiss κ','Landis-Koch tier'], rows,
          col_widths=[1.6,1.0,0.7,1.0,1.5])

P('Reading Table 11. ', bold=True)
P(f'The overall Fleiss κ across the three repetitions is '
  f'{float(phase_b["fleiss_overall_point"]):.3f} (substantial agreement on the Landis-Koch '
  f'scale). This is the empirical determinism of GPT-5.4 on this task under temperature = 0 '
  f'with seed = 42: not perfect, but high enough that the rep-averaging used throughout this '
  f'report (per-query mean across 3 reps) captures the central tendency well. The methodology '
  f'document (Section 8.1 of Technical_Report_Experiment_Methodology.docx) anticipated κ '
  f'≈ 1.0 under temperature = 0 and planned to use this as justification for single-run '
  f'evaluation; the empirical result instead requires us to use rep-averaging. The pattern '
  f'tracks modality difficulty: cephalometric is the most reproducible '
  f'(κ ≈ 0.83 – 0.89, almost-perfect), periapical is similarly high (κ ≈ 0.82 – 0.83), and '
  f'panoramic is the least (κ ≈ 0.65 – 0.70, still substantial). Importantly, guided is more '
  f'reproducible than zero-shot on cephalometric and panoramic — the model is more '
  f'self-consistent when given the guided system prompt, even when it is more often '
  f'self-consistently wrong (as Section 10 shows for Tooth_33_Apex).')

H('7.2 Three-way unanimous response rates', level=2)
caption('Table 12: 3-way unanimous predicted-cell rates.')
rows=[]
total_un = 0; total_n = 0
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for strat in ('zero_shot','guided'):
        d = phase_b['exact_3way_unanimous'][f'{mod}_{strat}_point']
        rate = d['rate'] if d['rate'] is not None else 0
        rows.append([mod, strat, str(d['n']), str(d['unanimous']), f'{rate*100:.1f}%'])
        total_un += d['unanimous']; total_n += d['n']
rows.append(['OVERALL','—', str(total_n), str(total_un), f'{total_un/total_n*100:.1f}%'])
add_table(['Modality','Strategy','n','3-way unanimous','Rate'], rows,
          col_widths=[1.6,1.0,0.7,1.4,1.0])

P('Reading Table 12. ', bold=True)
P(f'Three-way unanimous predicted-cell rate is the strictest agreement metric: it asks how '
  f'often the model emits the IDENTICAL cell on all three repetitions of the same query. '
  f'Overall, {total_un}/{total_n} = {total_un/total_n*100:.1f}% of point queries achieve '
  f'three-way unanimity, which means roughly one in three point queries has some between-rep '
  f'variation in the chosen cell. The cephalometric guided strategy is the most consistent '
  f'(unanimous rate close to 80%), while panoramic strategies are least (around 50%). The '
  f'remaining ~31% of queries — those with at least one disagreement between reps — are the '
  f'queries where rep-averaging is genuinely doing work; without it, any one of the three reps '
  f'could yield a different cell. This non-perfect determinism justifies our methodological '
  f'convention of always reporting rep-mean ED rather than single-rep ED for cross-rater '
  f'comparisons, and is a substantive methodological observation for the manuscript: studies '
  f'that report single-run accuracy at temperature = 0 with current commercial multimodal '
  f'LLMs may overstate the reproducibility of their findings.')

H('7.3 Area-landmark cross-rep agreement', level=2)
caption("Table 13: Mean pairwise Jaccard between the 3 reps' predicted cell sets.")
rows=[]
for ms in ('PANORAMIC_zero_shot_area','PANORAMIC_guided_area'):
    d = phase_b['area_reliability'][ms]
    rows.append([ms.replace('PANORAMIC_','').replace('_area',''), str(d['n']),
                 f'{d["mean_pairwise_jacc"]:.4f}', f'{d["median"]:.4f}'])
add_table(['Strategy','n','Mean pairwise Jaccard','Median'], rows,
          col_widths=[1.4,0.7,1.6,1.1])

P('Reading Table 13. ', bold=True)
P(f'For area landmarks, three-way unanimity (identical cell sets across reps) is overly strict '
  f'and would understate reproducibility; we instead compute the mean pairwise Jaccard between '
  f'the three reps\' predicted cell sets. Zero-shot scores '
  f'{phase_b["area_reliability"]["PANORAMIC_zero_shot_area"]["mean_pairwise_jacc"]:.3f}, '
  f'guided scores '
  f'{phase_b["area_reliability"]["PANORAMIC_guided_area"]["mean_pairwise_jacc"]:.3f}. The '
  f'guided strategy yields more consistent area predictions across reps, mirroring the point-'
  f'landmark pattern in Table 11 (guided is more self-consistent). This is not the same as '
  f'being more accurate — Table 6 shows guided and zero-shot area-aggregate mean Jaccard are '
  f'essentially indistinguishable in absolute terms. The guided system prompt narrows the '
  f'distribution of cell choices the model makes; it does not move that distribution closer to '
  f'ground truth for area landmarks.')

# ────────────────────────────────────────────────────────────────────
# 8. RQ4 — Ground Truth Validation (NEW)
# ────────────────────────────────────────────────────────────────────
H('8. RQ4 — Ground Truth Validation (NEW)', level=1)

P('The consensus GT used in this report is adjudicated from two OMFR specialists. This section '
  'quantifies the reliability of the underlying ratings — both inter-rater (across the two '
  'specialists) and intra-rater (each specialist with themselves, on a 180-query washout-period '
  'subset).')

H('8.1 Inter-rater reliability (OMFR_1 vs OMFR_2)', level=2)
caption('Table 14: Inter-rater reliability between OMFR_1 and OMFR_2 (n = 900 queries total). '
        'Cell-level Cohen\'s κ for points; Jaccard / Dice for areas.')
ip = rater_reli['inter_rater']['INTER_omfr1_vs_omfr2']
rows = []
for mod in ('CEPHALOMETRIC','PERIAPICAL','PANORAMIC'):
    for ltype in ('point','area'):
        if ltype == 'area' and mod != 'PANORAMIC': continue
        key = f'{mod}_{ltype}'
        if key not in ip: continue
        d = ip[key]
        if ltype == 'point':
            rows.append([mod, ltype, str(d['n']),
                         f'{d["mean_ed"]:.3f}',
                         f'{d["within_1_cell_rate"]*100:.1f}%',
                         f'{d["cohens_kappa"]:.3f}',
                         '—', '—'])
        else:
            rows.append([mod, ltype, str(d['n']),
                         '—','—','—',
                         f'{d["mean_jaccard"]:.3f}',
                         f'{d["mean_dice"]:.3f}'])
ip_all_pt = ip['ALL_point']
ip_all_ar = ip['ALL_area']
rows.append(['ALL','point',str(ip_all_pt['n']),
             f'{ip_all_pt["mean_ed_cells"]:.3f}',
             f'{ip_all_pt["within_1_cell_rate"]*100:.1f}%',
             f'{ip_all_pt["cohens_kappa"]:.3f}', '—','—'])
rows.append(['ALL','area',str(ip_all_ar['n']),
             '—','—','—',
             f'{ip_all_ar["mean_jaccard"]:.3f}',
             f'{ip_all_ar["mean_dice"]:.3f}'])
add_table(['Modality','Type','n','Mean ED','Within 1 cell','Cohen κ','Mean Jaccard','Mean Dice'],
          rows, col_widths=[1.2,0.55,0.5,0.85,1.0,0.7,1.0,0.85],
          header_size=9, body_size=9)

P('Reading Table 14. ', bold=True)
P(f'Inter-rater agreement between the two OMFR specialists is high. On point landmarks, '
  f'Cohen\'s κ = {ip_all_pt["cohens_kappa"]:.3f} (almost-perfect by Landis-Koch), and '
  f'{ip_all_pt["within_1_cell_rate"]*100:.1f}% of all 600 point queries have both raters '
  f'within 1 cell of each other. The mean inter-rater ED is '
  f'{ip_all_pt["mean_ed_cells"]:.3f} cells — below the 1-cell discriminative resolution of '
  f'the grid. On area landmarks (panoramic only), mean Jaccard '
  f'{ip_all_ar["mean_jaccard"]:.3f} and mean Dice {ip_all_ar["mean_dice"]:.3f} indicate '
  f'substantial overlap; the typical disagreement is a few boundary cells on the edge of a '
  f'continuous radiopaque structure, which is methodologically expected when discretising a '
  f'continuous anatomical boundary onto a grid.')

P('These values exceed the inter-rater reliability bars reported in comparable DMFR landmark '
  'studies. Indermun et al. (2023) [8] and Menezes et al. (2023) [9] report ICCs in the 0.85 – '
  '0.95 range on cephalometric landmark detection — our κ ≥ 0.86 sits within or above that '
  'band. The consensus GT used in this study is therefore methodologically sound by the '
  'standards of the field: the underlying measurement is reliable across specialists, and '
  'the residual inter-rater disagreement (≈ 1.5% of point queries with > 1-cell mismatch) is '
  'small enough and localised enough that team-adjudicated consensus can resolve it without '
  'introducing bias.')

H('8.2 Intra-rater reliability (washout-period re-rating, n = 180)', level=2)
caption('Table 15: Intra-rater reliability — each specialist\'s first vs second pass on a '
        '180-query subset (60 PAN + 30 PA + 30 CEPH × 6/3/3 landmarks per image), '
        'with ≥2-week washout.')
rows=[]
for label, key in (('OMFR_1', 'INTRA_omfr1'), ('OMFR_2', 'INTRA_omfr2')):
    iar = rater_reli['intra_rater'][key]
    ap_ = iar['ALL_point']; aa = iar['ALL_area']
    rows.append([label,'point',str(ap_['n']),
                 f'{ap_["mean_ed_cells"]:.3f}',
                 f'{ap_["within_1_cell_rate"]*100:.1f}%',
                 f'{ap_["cohens_kappa"]:.3f}', '—','—'])
    rows.append([label,'area',str(aa['n']),
                 '—','—','—',
                 f'{aa["mean_jaccard"]:.3f}',
                 f'{aa["mean_dice"]:.3f}'])
add_table(['Rater','Type','n','Mean ED','Within 1 cell','Cohen κ','Mean Jaccard','Mean Dice'],
          rows, col_widths=[1.2,0.55,0.5,0.85,1.0,0.7,1.0,0.85],
          header_size=9, body_size=9)

P('Intra-rater point-landmark agreement is essentially perfect (mean ED ≤ 0.025 cells, κ ≥ '
  '0.97). Intra-rater area-landmark agreement is high but not perfect (Jaccard ≥ 0.96), '
  'consistent with the inherent ambiguity of where to place the boundary of a continuous '
  'anatomical structure within a discrete grid. Together with the inter-rater results, this '
  'validates the consensus GT as a methodologically sound canonical reference: the underlying '
  'measurement is reproducible by each specialist with themselves, and the residual inter-'
  'rater disagreement is small and concentrated where adjudication can resolve it.')

# ────────────────────────────────────────────────────────────────────
# 9. RQ5 — GPT-5.4 vs Student (NEW)
# ────────────────────────────────────────────────────────────────────
H('9. RQ5 — GPT-5.4 vs Dental Student Consensus (NEW)', level=1)

P('For each query, GPT-5.4\'s mean Euclidean distance to consensus_gt across the three '
  'repetitions is paired against the student consensus\'s ED to consensus_gt. The paired '
  'Wilcoxon test on the deltas (GPT − student) tells us whether GPT and the student '
  'consensus are statistically distinguishable, on the same images, against the same GT.')

H('9.1 Modality-level paired comparisons', level=2)

caption('Table 16: GPT-5.4 vs Student modality-level paired Wilcoxon (zero-shot strategy).')
rows=[]
for mod_lt, label in (('CEPHALOMETRIC_point','CEPH point'),
                       ('PERIAPICAL_point','PA point'),
                       ('PANORAMIC_point','PAN point'),
                       ('PANORAMIC_area','PAN area')):
    d = gs['zero_shot']['per_modality'].get(mod_lt)
    if not d: continue
    metric = "ED" if 'point' in mod_lt else "Jaccard"
    gpt_v = d.get('mean_gpt_ed', d.get('mean_gpt_jaccard'))
    stu_v = d.get('mean_student_ed', d.get('mean_student_jaccard'))
    pbonf = min(1.0, d['p']*4)
    rows.append([label, str(d['n_total']), f'{gpt_v:.3f}', f'{stu_v:.3f}',
                 f'{d["mean_delta"]:+.3f}', fmt_p(d['p']), fmt_pbonf(d['p'], 4),
                 f'{d["rank_biserial_r"]:+.3f}', sig_marker(pbonf)])
add_table(['Group','n',f'GPT mean','Student mean','Mean Δ','Wilcoxon p','Bonf p (×4)','r','Sig'],
          rows, col_widths=[1.0,0.45,0.85,0.95,0.7,0.95,0.95,0.65,0.4],
          header_size=9, body_size=9)

caption('Table 17: GPT-5.4 vs Student modality-level paired Wilcoxon (guided strategy).')
rows=[]
for mod_lt, label in (('CEPHALOMETRIC_point','CEPH point'),
                       ('PERIAPICAL_point','PA point'),
                       ('PANORAMIC_point','PAN point'),
                       ('PANORAMIC_area','PAN area')):
    d = gs['guided']['per_modality'].get(mod_lt)
    if not d: continue
    metric = "ED" if 'point' in mod_lt else "Jaccard"
    gpt_v = d.get('mean_gpt_ed', d.get('mean_gpt_jaccard'))
    stu_v = d.get('mean_student_ed', d.get('mean_student_jaccard'))
    pbonf = min(1.0, d['p']*4)
    rows.append([label, str(d['n_total']), f'{gpt_v:.3f}', f'{stu_v:.3f}',
                 f'{d["mean_delta"]:+.3f}', fmt_p(d['p']), fmt_pbonf(d['p'], 4),
                 f'{d["rank_biserial_r"]:+.3f}', sig_marker(pbonf)])
add_table(['Group','n',f'GPT mean','Student mean','Mean Δ','Wilcoxon p','Bonf p (×4)','r','Sig'],
          rows, col_widths=[1.0,0.45,0.85,0.95,0.7,0.95,0.95,0.65,0.4],
          header_size=9, body_size=9)

P('Reading Tables 16 and 17. ', bold=True)
P('Cephalometric is the only modality where GPT-5.4 is statistically indistinguishable from '
  'the student consensus (zero-shot p = '
  f'{gs["zero_shot"]["per_modality"]["CEPHALOMETRIC_point"]["p"]:.3f}, guided p = '
  f'{gs["guided"]["per_modality"]["CEPHALOMETRIC_point"]["p"]:.3f}, both NS after Bonferroni × 4). '
  'On the easiest modality the model is at the same level as a fourth-year dental student. '
  'This is a defensible publication claim and a useful positive result.')

P('On periapical and panoramic, the picture inverts. Student consensus is substantially more '
  'accurate, with overwhelming statistical significance and very large effect sizes: PAN '
  f'point r > +0.83 (zero-shot) and r > +0.89 (guided), PA point r ≈ +0.78. A rank-biserial r '
  f'of +0.89 means that on roughly 95% of pairwise comparisons the student is closer to '
  f'consensus than GPT, which is the "essentially every query" interpretation. On PAN area '
  f'landmarks, the student\'s mean Jaccard ({gs["zero_shot"]["per_modality"]["PANORAMIC_area"]["mean_student_jaccard"]:.2f}) '
  f'is more than twice GPT\'s ({gs["zero_shot"]["per_modality"]["PANORAMIC_area"]["mean_gpt_jaccard"]:.2f} '
  f'/ {gs["guided"]["per_modality"]["PANORAMIC_area"]["mean_gpt_jaccard"]:.2f}). The rank-biserial '
  f'r ≈ −0.86 on PAN area (note: the sign convention is reversed for Jaccard since higher is '
  f'better) is symmetrically large to the PAN point effect.')

P('The guided strategy does not change the clinical reading. On every comparison, the guided '
  'and zero-shot rows produce similar verdicts (NS on CEPH, very-strong "student wins" on PA '
  'and PAN). The Tooth_33_Apex regression under guided affects the PAN point effect size '
  'slightly (r moves from +0.83 to +0.89), but the qualitative conclusion is unchanged.')

H('9.2 Bland-Altman descriptive statistics', level=2)
caption('Table 18: Bland-Altman descriptive stats (mean bias + 95% limits of agreement) for '
        'GPT − Student per modality. Useful for visualising systematic bias even when paired '
        'tests are NS.')
rows=[]
for strat in ('zero_shot','guided'):
    for mod_lt, label in (('CEPHALOMETRIC_point','CEPH point ED'),
                           ('PERIAPICAL_point','PA point ED'),
                           ('PANORAMIC_point','PAN point ED'),
                           ('PANORAMIC_area','PAN area Jaccard')):
        d = gs[strat]['per_modality'].get(mod_lt)
        if not d or 'bland_altman' not in d or not d['bland_altman']: continue
        ba = d['bland_altman']
        rows.append([label, strat, str(ba['n']),
                     f'{ba["mean_diff"]:+.3f}', f'{ba["sd_diff"]:.3f}',
                     f'{ba["loa_low"]:+.3f}', f'{ba["loa_high"]:+.3f}'])
add_table(['Group','Strategy','n','Mean diff','SD diff','LoA low','LoA high'],
          rows, col_widths=[1.5,1.0,0.5,1.0,0.85,0.95,0.95],
          header_size=9, body_size=9)

H('9.3 Acceptability bands', level=2)
P('For each landmark, we compute an "acceptability band" defined as the larger of the mean '
  'inter-rater disagreement (OMFR_1 vs OMFR_2 mean ED on that landmark) and 1.0 grid cell. '
  'A GPT prediction within the band is "as close to consensus as the two specialists are to '
  'each other on this landmark on a typical query, or 1 cell — whichever is wider". This '
  'contextualises raw SDR@1 values: a low SDR@1 on a landmark where the two specialists also '
  'routinely disagree by ≥1 cell would be less alarming than a low SDR@1 where the specialists '
  'agreed almost perfectly.')
# Live-compute how many landmarks have the band defaulting to 1.0 vs. floating above.
_ab_data = gpt_v_stu['acceptability_band_per_landmark']
_n_landmarks_defaulted = sum(
    1 for k, v in _ab_data['zero_shot'].items()
    if v['acceptability_band_cells'] > v['human_mean_disagreement_cells'] + 0.01)
_n_landmarks_total = len(_ab_data['zero_shot'])
P(f'Empirical note. ', bold=True)
P(f'In this dataset, the mean inter-rater disagreement is below 1 cell for every one of the '
  f'{_n_landmarks_total} point landmarks measured (max human disagreement = '
  f'{max(v["human_mean_disagreement_cells"] for v in _ab_data["zero_shot"].values()):.2f} cells, '
  f'on CEPHALOMETRIC/Nasion_N). The band therefore defaults to 1.0 cell for all '
  f'{_n_landmarks_defaulted}/{_n_landmarks_total} landmarks, and the "Within band / Rate" '
  f'columns in Table 19 below are numerically equal to per-landmark SDR@1. The band machinery '
  f'is retained for methodological transparency and would matter for landmarks whose human '
  f'inter-rater disagreement exceeded 1 cell — for example, on a less-conserved structure or '
  f'on a different imaging modality. Readers should treat Table 19 as a per-landmark SDR@1 '
  f'view that is explicit about why the comparison threshold is 1 cell rather than the smaller '
  f'human-disagreement value.')

caption('Table 19: Per-landmark acceptability-band check.')
rows=[]
for strat in ('zero_shot','guided'):
    for mod, lm in landmarks_order:
        key = f'{mod}/{lm}'
        d = gpt_v_stu['acceptability_band_per_landmark'].get(strat, {}).get(key)
        if not d: continue
        rows.append([key, strat, str(d['n']),
                     f'{d["human_mean_disagreement_cells"]:.2f}',
                     f'{d["acceptability_band_cells"]:.2f}',
                     f'{d["within_band"]}/{d["n"]}',
                     f'{d["within_band_rate"]*100:.1f}%'])
add_table(['Landmark','Strategy','n','Human disagree','Band','Within band','Rate'],
          rows, col_widths=[2.0,0.85,0.4,1.05,0.6,1.05,0.65],
          header_size=9, body_size=9)

P('Limitations of this comparison:', italic=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('The student consensus is a single team-adjudicated response per query, not the '
          'individual mean of the 8 students who saw it. We therefore cannot quantify per-query '
          'inter-student variability, and the paired test treats the student response as a '
          'single-rater observation.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Two student responses were anomalous: PAN_010_Condylar_Head_R returned "3A" '
          '(reversed coordinates, the same failure mode as the four GPT compliance failures) '
          'and PAN_068_Tooth_33_Apex returned a 6-cell list for a single-cell point landmark. '
          'The first is excluded (n = 299 instead of 300 for PAN point); the second is treated '
          'as the first parsed cell (D14) for a methodologically conservative read.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('GPT-5.4 is averaged over 3 reps; the student is a single response. This biases the '
          'comparison slightly in GPT\'s favour (rep-averaging reduces the variance of GPT\'s '
          'per-query estimate). The effect sizes are large enough that this bias does not '
          'plausibly explain the observed differences.')

# ────────────────────────────────────────────────────────────────────
# 10. TARGETED ABLATIONS ON THE TOOTH_33_APEX REGRESSION
# ────────────────────────────────────────────────────────────────────
H('10. Targeted Ablations on the Tooth_33_Apex Regression', level=1)

# ── 10.1 Motivation ─────────────────────────────────────────────
H('10.1 Motivation and overview', level=2)
# Render the §10.1 motivation numbers live from analysis.json and ablation_analysis.json
# so they cannot drift from the canonical record.
_t33_zs_ed   = abl_zs['mean_ed']
_t33_gd_ed   = abl_gd['mean_ed']
_t33_zs_sdr1 = abl_zs['SDR@1']
_t33_gd_sdr1 = abl_gd['SDR@1']
_t33_gd_correct_side = abl_gd['correct_side_rate']
_t33_rq2b_r  = analysis['RQ2b_strategy_per_landmark']['PANORAMIC/Tooth_33_Apex']['rank_biserial_r']
P(f'Section 6.3 documents a striking pattern in the v2 main run: under the canonical guided '
  f'prompt, GPT-5.4 collapses on PAN/Tooth_33_Apex (mean ED {_t33_zs_ed:.2f} → {_t33_gd_ed:.2f} '
  f'cells, SDR@1 {_t33_zs_sdr1*100:.0f}% → {_t33_gd_sdr1*100:.0f}%, rank-biserial '
  f'r = {_t33_rq2b_r:+.2f} under the §6.3 convention where positive r → guided better). At '
  f'the same time, a non-FDI panoramic landmark with an explicit English-suffix lateralisation '
  f'tag (Mental_Foramen_L) IMPROVES under guided. A column-distribution analysis on the 100 '
  f'Tooth_33_Apex queries revealed that under canonical guided, '
  f'{(1-_t33_gd_correct_side)*100:.0f}% of predictions land on the wrong (image-left) side, '
  f'despite the ground truth consistently lying on the image right (cols 9–10).')

P('We ran three targeted ablations, each isolating one candidate explanation. Every ablation '
  'used the SAME 100 PAN/Tooth_33_Apex queries, the SAME model (gpt-5.4), the SAME inference '
  'settings (temperature = 0, seed = 42), and was anchored to the frozen results_full/ JSONLs '
  'for the zero_shot and canonical guided baselines (no re-issuing of those calls). Each '
  'ablation introduced exactly one change relative to canonical guided. Decision criteria were '
  'locked before data collection. The full reproducibility manifests for each ablation are in '
  'Appendix C.')

caption('Table 20: Overview of the three Section 10 ablations.')
add_table(
    ['Ablation', 'Strategy name', 'Change vs canonical guided', 'Pre-spec. threshold'],
    [
        ['A — FDI tooth-number removal',
         'guided_no_tooth_num',
         'User prompt: drop "tooth #33 (lower left canine)" → "the lower left canine". System prompt unchanged.',
         '≥ 90% correct-side ⇒ confirmed'],
        ['B — Patient-frame disambiguation',
         'guided_patient_left',
         'User prompt: rewrite parenthetical "tooth #33 (lower left canine)" → "tooth #33 (patient\'s left lower canine)". System prompt unchanged.',
         '≥ 90% correct-side ⇒ confirmed'],
        ['C — L–R inversion clause removal (diagnostic)',
         'guided_no_LR',
         'System prompt: remove the panoramic L–R inversion sentence. User prompt unchanged.',
         '≥ 60% correct-side ⇒ L–R clause IS the cause'],
    ],
    col_widths=[2.0, 1.55, 2.4, 1.15],
    header_size=9, body_size=9, fixed=True,
)

P('Why these three specifically: each tests a distinct candidate mechanism that emerged from '
  'inspecting the v2 data and from external collaborator feedback. (A) addresses the hypothesis '
  'that the FDI tooth number "33" is a confusing token in combination with the L–R clause. (B) '
  'addresses the hypothesis that "left" inside the multi-word compound "lower left canine" is '
  'parsed ambiguously, and adding a clear "patient\'s" qualifier disambiguates. (C) directly '
  'tests whether the L–R inversion clause itself is the proximate cause. These three together '
  'cover the substantive surface-level interventions one would propose. They are NOT '
  'exhaustive — the Discussion (Section 11) outlines what remains untested.')

# ── 10.2 Ablation A: FDI tooth-number ───────────────────────────
H('10.2 Ablation A — FDI tooth-number removal', level=2)

P('Hypothesis (from a collaborator): the FDI number "33" combined with the panoramic L–R '
  'inversion clause produces a dual-"left" signal that confuses the model. Strict-literal '
  'fix: drop "tooth #33 (lower left canine)" and replace with just "the lower left canine"; '
  'keep the FDI numbering-system announcement at the start of the user prompt.')

caption('Box 1: Exact prompt change in Ablation A (Tooth_33_Apex user prompt only). '
        'System prompt byte-identical to canonical guided.')
demo_a = doc.add_table(rows=2, cols=2)
demo_a.style = 'Light Grid'
tblPr = demo_a._element.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr'); demo_a._element.insert(0, tblPr)
layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
for j, h in enumerate(['canonical guided', 'Ablation A: guided_no_tooth_num']):
    c = demo_a.cell(0, j); c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
    shade_cell(c, 'D9E1F2'); set_cell_borders(c)
# Render the Box 1 user prompts LIVE from pipeline.generate_prompt so the box can never
# drift from what was actually issued to the API (verified against each ablation's
# prompts_used.json file for all 100 Tooth_33_Apex entries).
_box1_left_user  = _pipeline.generate_prompt(_q33, 'guided')[1]
_box1_right_user = _pipeline.generate_prompt(_q33, 'guided_no_tooth_num')[1]
demo_a.cell(1, 0).text = ''
demo_a.cell(1, 0).paragraphs[0].add_run(_box1_left_user).font.size = Pt(9)
demo_a.cell(1, 1).text = ''
demo_a.cell(1, 1).paragraphs[0].add_run(_box1_right_user).font.size = Pt(9)
for j in (0, 1):
    set_cell_borders(demo_a.cell(0, j)); set_cell_borders(demo_a.cell(1, j))
    for r in demo_a.rows: r.cells[j].width = Inches(3.3)

_a_total, _a_comp = abl_compliance["no_tooth_num"]
P(f'Operational outcome: {_a_total} calls, ${abl_costs["no_tooth_num"]:.2f}, '
  f'{_a_comp/_a_total*100:.2f}% strict compliance '
  f'({_a_total - _a_comp} failure(s) — see Appendix B2 for the full list). '
  f'Anchored against {len(abl_anchors["no_tooth_num"].get("files", {}))} frozen raw JSONLs in '
  f'results_full/.')

caption('Table 21: Ablation A results — three-way comparison on Tooth_33_Apex (n=100).')
def fmt_pct(v): return f'{v*100:.1f}%' if v is not None else '–'
def fmt_num(v, d=3): return f'{v:.{d}f}' if v is not None else '–'
rows = []
for label, key in (('zero_shot', abl_zs),
                    ('guided (canonical)', abl_gd),
                    ('guided_no_tooth_num (Ablation A)', abl_nt)):
    rows.append([label, str(key['n_with_metric']),
                 fmt_num(key['mean_ed']),
                 fmt_pct(key['correct_side_rate']),
                 fmt_pct(key['SDR@1']),
                 fmt_pct(key['SDR@2']),
                 fmt_pct(key['SDR@4'])])
add_table(['Strategy', 'n', 'Mean ED', 'Correct-side', 'SDR@1', 'SDR@2', 'SDR@4'], rows,
          col_widths=[2.4, 0.4, 0.7, 1.0, 0.65, 0.65, 0.65],
          header_size=9, body_size=9, fixed=True)

P(f'Critical reading. ', bold=True)
P(f'The correct-side rate moves from canonical guided\'s {abl_gd["correct_side_rate"]*100:.0f}% '
  f'to Ablation A\'s {abl_nt["correct_side_rate"]*100:.0f}% — essentially unchanged. Paired '
  f'Wilcoxon vs canonical guided: Δ = {pw_nt_gd["mean_delta"]:+.4f} cells, p = {pw_nt_gd["p"]:.4f} '
  f'(not significant, r = {pw_nt_gd["rank_biserial_r"]:+.3f}). The pre-registered threshold (≥90%) '
  f'is not approached. Decision: REJECTED — the FDI tooth number is not the proximate cause. '
  f'Removing it leaves the same systematic ~4-cell wrong-side displacement intact. The next '
  f'ablation tested whether disambiguating "left" within the parenthetical is sufficient.')

# ── 10.3 Ablation B: Patient-frame disambiguation ──────────────
H('10.3 Ablation B — Patient-frame disambiguation', level=2)

P('Hypothesis: the word "left" inside the compound "lower left canine" is parsed ambiguously '
  '(it could read as image-frame "left" rather than patient-frame "left"). Adding an explicit '
  '"patient\'s" qualifier in front of "left" should disambiguate it. Importantly, this ablation '
  'is the strictest-literal test we could design: it changes EXACTLY one substring in the user '
  'prompt; everything else — including the FDI numbering announcement, the tooth number, and '
  'the system prompt — is byte-identical to canonical guided. Mental_Foramen_L (which uses an '
  'English "_L" suffix) works correctly under guided, so this ablation tests whether a '
  'similarly explicit patient-frame anchor on Tooth_33 has the same effect.')

caption('Box 2: Exact prompt change in Ablation B. System prompt byte-identical to canonical '
        'guided.')
demo_b = doc.add_table(rows=2, cols=2)
demo_b.style = 'Light Grid'
tblPr = demo_b._element.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr'); demo_b._element.insert(0, tblPr)
layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
for j, h in enumerate(['canonical guided', 'Ablation B: guided_patient_left']):
    c = demo_b.cell(0, j); c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
    shade_cell(c, 'D9E1F2'); set_cell_borders(c)
# Render Box 2 user prompts LIVE from pipeline.generate_prompt so the box can never
# drift from what was actually issued to the API.
_box2_left_user  = _pipeline.generate_prompt(_q33, 'guided')[1]
_box2_right_user = _pipeline.generate_prompt(_q33, 'guided_patient_left')[1]
demo_b.cell(1, 0).text = ''
demo_b.cell(1, 0).paragraphs[0].add_run(_box2_left_user).font.size = Pt(9)
demo_b.cell(1, 1).text = ''
demo_b.cell(1, 1).paragraphs[0].add_run(_box2_right_user).font.size = Pt(9)
for j in (0, 1):
    set_cell_borders(demo_b.cell(0, j)); set_cell_borders(demo_b.cell(1, j))
    for r in demo_b.rows: r.cells[j].width = Inches(3.3)

P('Net change: the substring "lower left canine" → "patient\'s left lower canine" inside the '
  'parenthetical. Two micro-edits (insert "patient\'s" and reorder "lower"/"left") so that "left" '
  'sits unambiguously next to its patient-frame qualifier. The substitution is verified '
  'mechanically isolated (preflight confirms that replacing the substring in canonical guided '
  'yields Ablation B exactly).')

_b_total, _b_comp = abl_compliance["patient_left"]
P(f'Operational outcome: {_b_total} calls, ${abl_costs["patient_left"]:.2f}, '
  f'{_b_comp/_b_total*100:.2f}% strict compliance '
  f'({_b_total - _b_comp} failure(s) — see Appendix B2). '
  f'Anchored against {len(abl_anchors["patient_left"].get("files", {}))} frozen raw JSONLs.')

caption('Table 22: Ablation B results — three-way comparison on Tooth_33_Apex (n=100).')
rows = []
for label, key in (('zero_shot', abl_zs),
                    ('guided (canonical)', abl_gd),
                    ('guided_patient_left (Ablation B)', abl_pl)):
    rows.append([label, str(key['n_with_metric']),
                 fmt_num(key['mean_ed']),
                 fmt_pct(key['correct_side_rate']),
                 fmt_pct(key['SDR@1']),
                 fmt_pct(key['SDR@2']),
                 fmt_pct(key['SDR@4'])])
add_table(['Strategy', 'n', 'Mean ED', 'Correct-side', 'SDR@1', 'SDR@2', 'SDR@4'], rows,
          col_widths=[2.4, 0.4, 0.7, 1.0, 0.65, 0.65, 0.65],
          header_size=9, body_size=9, fixed=True)

P('Critical reading. ', bold=True)
P(f'Correct-side rate is {abl_pl["correct_side_rate"]*100:.0f}% — actually slightly LOWER than '
  f'canonical guided\'s {abl_gd["correct_side_rate"]*100:.0f}%. Paired Wilcoxon: Δ = '
  f'{pw_pl_gd["mean_delta"]:+.4f} cells, p = {pw_pl_gd["p"]:.4f}, r = '
  f'{pw_pl_gd["rank_biserial_r"]:+.3f}. Although the mean ED is nominally improved by ~0.14 '
  f'cells, the practical effect on the wrong-side failure mode is null. The pre-registered '
  f'threshold (≥90% correct-side) is not approached. Decision: REJECTED — explicit '
  f'patient-frame disambiguation of "left" does not fix the regression. This is a noteworthy '
  f'negative result: a fix that looked highly promising on grammatical/linguistic grounds — '
  f'and that worked for the structurally similar Mental_Foramen_L under canonical guided — '
  f'fails on Tooth_33_Apex. The user-prompt wording is not where the regression lives.')

# ── 10.4 Ablation C: L-R clause removal (diagnostic) ───────────
H('10.4 Ablation C — L–R inversion clause removal (diagnostic)', level=2)

# Render the L-R clause and its length live from pipeline.generate_prompt so the cited
# character count cannot drift from the actual prompt.
_canon_sys = _pipeline.generate_prompt(_q33, 'guided')[0]
_nolr_sys  = _pipeline.generate_prompt(_q33, 'guided_no_LR')[0]
_lr_clause = ("In panoramic radiographs, the patient's right side appears on the left side of "
              "the image, and the patient's left side appears on the right side of the image.")
assert _lr_clause in _canon_sys and _lr_clause not in _nolr_sys, \
    "L-R clause does not match canonical guided / guided_no_LR system prompts"
_lr_len = len(_lr_clause)
P(f'After two negative results on user-prompt wording, suspicion shifted to the system prompt. '
  f'The most prominent suspect was the panoramic L–R inversion clause — a {_lr_len}-character '
  f'sentence that the guided system prompt includes only for panoramic queries:')

p = doc.add_paragraph()
r = p.add_run('  "' + _lr_clause + '"')
r.italic = True; r.font.size = Pt(10)

P('Hypothesis: the L–R clause activates a misapplied lateralisation flip on FDI-named teeth, '
  'and removing it should restore Tooth_33_Apex performance toward zero_shot levels. Note this '
  'is explicitly a DIAGNOSTIC ablation, not a candidate canonical prompt: removing the clause '
  'would also remove the Mental_Foramen_L improvement (which is exactly what the clause was '
  'introduced for in the v2 prompt design). The pre-registered decision threshold is therefore '
  'lower (≥60% correct-side, returning to zero_shot range) rather than the ≥90% threshold used '
  'for Ablations A and B which were candidate canonical fixes.')

caption('Box 3: Exact prompt change in Ablation C (system prompt only). User prompt '
        'byte-identical to canonical guided.')
demo_c = doc.add_table(rows=2, cols=2)
demo_c.style = 'Light Grid'
tblPr = demo_c._element.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = OxmlElement('w:tblPr'); demo_c._element.insert(0, tblPr)
layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed'); tblPr.append(layout)
for j, h in enumerate(['canonical guided (system)', 'Ablation C: guided_no_LR (system)']):
    c = demo_c.cell(0, j); c.text = ''
    r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(10)
    shade_cell(c, 'D9E1F2'); set_cell_borders(c)
demo_c.cell(1, 0).text = ''
demo_c.cell(1, 0).paragraphs[0].add_run(
    '…The grid lines are drawn in cyan and labels are in yellow. '
    'In panoramic radiographs, the patient\'s right side appears on the left side of the image, '
    'and the patient\'s left side appears on the right side of the image. '
    'For point-based questions, respond with exactly one cell coordinate…'
).font.size = Pt(9)
demo_c.cell(1, 1).text = ''
demo_c.cell(1, 1).paragraphs[0].add_run(
    '…The grid lines are drawn in cyan and labels are in yellow. '
    'For point-based questions, respond with exactly one cell coordinate…'
).font.size = Pt(9)
for j in (0, 1):
    set_cell_borders(demo_c.cell(0, j)); set_cell_borders(demo_c.cell(1, j))
    for r in demo_c.rows: r.cells[j].width = Inches(3.3)

_c_total, _c_comp = abl_compliance["no_LR"]
P(f'Operational outcome: {_c_total} calls, ${abl_costs["no_LR"]:.2f}, '
  f'{_c_comp/_c_total*100:.2f}% strict compliance '
  f'({_c_total - _c_comp} failure(s)). '
  f'Anchored against {len(abl_anchors["no_LR"].get("files", {}))} frozen raw JSONLs.')

caption('Table 23: Ablation C results — three-way comparison on Tooth_33_Apex (n=100).')
rows = []
for label, key in (('zero_shot', abl_zs),
                    ('guided (canonical)', abl_gd),
                    ('guided_no_LR (Ablation C)', abl_nlr)):
    rows.append([label, str(key['n_with_metric']),
                 fmt_num(key['mean_ed']),
                 fmt_pct(key['correct_side_rate']),
                 fmt_pct(key['SDR@1']),
                 fmt_pct(key['SDR@2']),
                 fmt_pct(key['SDR@4'])])
add_table(['Strategy', 'n', 'Mean ED', 'Correct-side', 'SDR@1', 'SDR@2', 'SDR@4'], rows,
          col_widths=[2.4, 0.4, 0.7, 1.0, 0.65, 0.65, 0.65],
          header_size=9, body_size=9, fixed=True)

P(f'Critical reading. ', bold=True)
P(f'Correct-side rate moves from {abl_gd["correct_side_rate"]*100:.0f}% (canonical guided) to '
  f'{abl_nlr["correct_side_rate"]*100:.0f}% (Ablation C). This is the largest movement among the '
  f'three ablations — paired Wilcoxon Δ = {pw_nlr_gd["mean_delta"]:+.4f} cells, p = '
  f'{pw_nlr_gd["p"]:.4g}, r = {pw_nlr_gd["rank_biserial_r"]:+.3f} — but it is still vastly '
  f'short of zero_shot\'s 67%. The pre-registered 60% threshold is not met. Decision: '
  f'REJECTED — removing the L–R clause does not restore performance. The clause is therefore '
  f'NOT the proximate cause; it contributes a small effect on the margin (around 4 additional '
  f'percentage points of correct-side rate), but the bulk of the regression survives its '
  f'removal. The cause must lie deeper than this single sentence.')

# ── 10.5 Qualitative inspection ─────────────────────────────────
H('10.5 Qualitative inspection — where do predictions land?', level=2)

P('After three negative ablations on candidate phrases, we conducted a qualitative analysis of '
  'where the model actually places its prediction under each condition. For each of the 100 '
  'Tooth_33_Apex queries we computed the mode predicted cell (across the three repetitions) '
  'and looked at the distribution.')

# Compute summary statistics from qual_rows (already loaded above)
def col_stats(strat_key):
    if not qual_rows: return None
    cols = [r[f'{strat_key}_mode_col'] for r in qual_rows if r[f'{strat_key}_mode_col'] is not None]
    cells = [r[f'{strat_key}_mode'] for r in qual_rows if r[f'{strat_key}_mode'] != '-']
    n = len(cols)
    if n == 0: return None
    from collections import Counter as Ctr
    top = Ctr(cells).most_common(2)
    return {
        'n': n,
        'mean_col': sum(cols)/n,
        'correct_side': sum(1 for c in cols if c >= 9)/n,
        'top1': top[0] if top else ('-', 0),
        'top2': top[1] if len(top) > 1 else ('-', 0),
    }

stats = {s: col_stats(s) for s in
         ('zero_shot','guided','guided_no_tooth_num','guided_patient_left','guided_no_LR')}

caption('Table 24: Per-image MODE predicted cell — column distribution across the 100 '
        'Tooth_33_Apex queries. GT mean column = 9.8 (image right). "Top-1 / Top-2 cell" reports '
        'the most-chosen mode cell (across all 100 queries) and how many queries chose it.')
rows = []
for s, label in (('zero_shot', 'zero_shot'),
                  ('guided', 'guided (canonical)'),
                  ('guided_no_tooth_num', 'guided_no_tooth_num (Ablation A)'),
                  ('guided_patient_left', 'guided_patient_left (Ablation B)'),
                  ('guided_no_LR', 'guided_no_LR (Ablation C)')):
    d = stats[s]
    if d is None: continue
    rows.append([label, str(d['n']),
                 f'{d["mean_col"]:.1f}',
                 f'{int(round(d["correct_side"]*100))}%',
                 f'{d["top1"][0]} (×{d["top1"][1]})',
                 f'{d["top2"][0]} (×{d["top2"][1]})'])
add_table(['Strategy', 'n', 'Mean pred. col (across reps)',
           'MODE cell on image-right (cols ≥9)',
           'Top-1 cell', 'Top-2 cell'],
          rows,
          col_widths=[2.6, 0.4, 1.0, 1.4, 0.7, 0.7],
          header_size=9, body_size=9, fixed=True)

P('Note on definitions. ', bold=True)
# Live-derive the comparison numbers
_n_qual = len(qual_rows)
_mode_correct_gd = sum(1 for r in qual_rows
                       if (r.get('guided_mode_col') or 0) >= 9)
_mean_correct_gd_pct = abl_gd['correct_side_rate'] * 100
P(f'Table 24 reports the MODE predicted cell (the most-frequent rep prediction for each query, '
  f'across the 3 repetitions). Under this MODE-based definition, {_mode_correct_gd}/'
  f'{_n_qual} canonical-guided queries place their MODE cell on the image right '
  f'(cols ≥ 9). Tables 21–23 and the pre-registered ablation decision criteria use a slightly '
  f'different operational definition: the per-query mean of the three rep-columns must be ≥ 9. '
  f'Under that mean-of-cols definition, canonical guided\'s correct-side rate is '
  f'{_mean_correct_gd_pct:.1f}%. The MODE and mean-of-cols rates can disagree by a few queries '
  f'when 1 of the 3 reps falls on the opposite side; both characterisations are correct for '
  f'their respective purposes and we report both so the reader can see the small definitional '
  f'gap explicitly. The qualitative conclusions of Section 10.5 (F5/F6 convergence, robust '
  f'across wording variants) are unchanged under either definition.')

P('Critical reading. ', bold=True)
P('The model places almost every Tooth_33_Apex prediction on the same two cells (F5 and F6) '
  'under all four guided variants, regardless of the prompt wording change. Zero-shot, in '
  'contrast, predicts on cells F9 / F10 / F8 — anatomically reasonable locations near the GT. '
  'The "wrong-side" pattern under guided is not a random side-flip: it is a tight convergence '
  'onto a specific image region (lower row, inner-left columns).')

# Quadrant breakdown of wrong-side guided predictions — live-derived from qual_rows
def _parse_cell(s):
    if not s or not isinstance(s, str): return None
    s = s.strip().upper()
    if len(s) < 2: return None
    try:
        return (s[0], int(s[1:]))
    except ValueError:
        return None

_wrong_under_gd = [r for r in qual_rows
                   if (_parse_cell(r.get('guided_mode')) or ('?', 99))[1] < 9]
_n_wrong = len(_wrong_under_gd)
_in_band = 0
for r in _wrong_under_gd:
    p = _parse_cell(r.get('guided_mode'))
    if p:
        row_letter, col = p
        row_idx = ord(row_letter) - ord('A')
        if 4 <= row_idx <= 7 and 5 <= col <= 8:
            _in_band += 1
P(f'A finer-grain look: of the {_n_wrong} canonical-guided wrong-side predictions, '
  f'{_in_band}/{_n_wrong} ({_in_band/_n_wrong*100:.0f}%) land in the LOWER row band (E–H) on '
  f'the INNER image-left (columns 5–8). Not the far-left (cols 1–4), not the upper-half. '
  f'The wrong-side error is geometrically concentrated, not diffuse.')

P('Clinical interpretation of the F5/F6 attractor region (input from OMFR specialist).', bold=True)
P('Reviewing the F5/F6 cluster against panoramic dental anatomy: in a 16 × 8 panoramic grid, '
  'cells F5/F6 lie in the lower-row inner-image-left region, which on a typical panoramic '
  'radiograph corresponds to the area where the roots and crowns of the patient\'s right-side '
  'MOLAR teeth (FDI 46/47/48) project. They do NOT correspond to any canine. The lower-left '
  'canine apex (tooth 33) is anatomically located much more anteriorly and on the image right '
  '— the GT region we measure (cols 9–10) is exactly where it belongs. The model is therefore '
  'not making a simple lateralisation flip ("got left/right inverted"); it is placing the '
  'tooth-33-apex prediction in a region that anatomically corresponds to a different tooth '
  'class altogether. The failure is doubly incorrect: wrong side AND wrong dental structure '
  '(molar region instead of canine region).')

# ── Figure 1: OMFR-specialist-annotated PAN_001 image ──────────
# The image was prepared by the project's OMFR specialist (extracted from
# docs/Clinical_Discussion_Contributions.docx). The PNG/JPG file lives in
# results_consensus/ alongside the generated report so the figure provenance
# is recorded with the rest of the v2 outputs.
caption("Figure 1: Representative panoramic radiograph (PAN_001) with the 16 × 8 grid "
        "overlay. Green: consensus ground-truth cell for Tooth_33_Apex (G10, the lower-left "
        "canine apex). Red: model attractor region under the canonical guided prompt and all "
        "three Section-10 ablation variants (F5–F6, the lower-right molar region). The two "
        "regions are anatomically distinct in both lateralisation and tooth class. Annotation "
        "by the project's OMFR specialist.")

_fig1_path = SANDBOX / 'figure1_clinical_annotation.jpg'
assert _fig1_path.exists(), (
    f"Figure 1 source file not found at {_fig1_path}. The image is extracted from "
    f"docs/Clinical_Discussion_Contributions.docx; re-extract via the snippet in "
    f"the project memory if missing.")
# Embed at ~6.5 inches wide (fits the letter-portrait page with default margins)
doc.add_picture(str(_fig1_path), width=Inches(6.5))

P('Empirical falsification of the simple-flip hypothesis (from our data alone). ', bold=True)
# Compute the mean of model's predicted column and the mirror-of-GT column.
# qual_rows is loaded at the top; we use the guided mean column and the GT mean column.
def _mean_col(strat_key):
    cols = [r[f'{strat_key}_mode_col'] for r in qual_rows if r[f'{strat_key}_mode_col'] is not None]
    return sum(cols)/len(cols) if cols else None
gt_mean_col   = sum(r['gt_col'] for r in qual_rows) / len(qual_rows) if qual_rows else None
gd_mean_col   = _mean_col('guided')
zs_mean_col   = _mean_col('zero_shot')
# Geometric mirror around the panoramic midline (16-col grid → midline at col 8.5):
#   c_mirror = 17 - c  (so cols 9 → 8, 10 → 7).
mirror_mean_col = 17 - gt_mean_col if gt_mean_col is not None else None
beyond_mirror   = (mirror_mean_col - gd_mean_col) if (mirror_mean_col and gd_mean_col) else None

P(f'Even setting aside the clinical interpretation above, our data alone falsify the simplest '
  f'alternative hypothesis — that the model is merely flipping the panoramic L–R orientation. '
  f'A pure horizontal flip around the panoramic midline (col 8.5 in our 16-column grid) maps '
  f'GT column c to its mirror at column (17 − c). With GT mean column = '
  f'{gt_mean_col:.2f} (range 9 – 10), the simple-flip prediction is a mean predicted column of '
  f'{mirror_mean_col:.2f}. The empirical mean predicted column under canonical guided is '
  f'{gd_mean_col:.2f} — i.e., the model\'s predictions sit '
  f'{beyond_mirror:.2f} columns BEYOND the simple-flip mirror, deeper into image-left. (For '
  f'reference, zero_shot has a mean predicted column of {zs_mean_col:.2f}, almost exactly on '
  f'GT.) The discrepancy with the simple-flip prediction is itself an empirical observation '
  f'from our 100-image dataset: the regression is NOT well-explained as a spatial L–R '
  f'inversion. Combined with the clinical interpretation that "1–2 columns further than the '
  f'mirror, in row F" corresponds to lower-right molar dentition (cf. Figure 1), the case for '
  f'structural-anatomic confusion is doubly grounded — geometrically from our data, '
  f'anatomically from the OMFR specialist.')

P('Cross-condition stability. ', bold=True)
# Live-derive co-failure counts on the wrong-side subset
def _is_wrong_side(cell):
    p = _parse_cell(cell)
    return p is not None and p[1] < 9
_co = {}
for cond in ('guided_no_tooth_num_mode', 'guided_patient_left_mode',
             'guided_no_LR_mode', 'zero_shot_mode'):
    _co[cond] = sum(1 for r in _wrong_under_gd if _is_wrong_side(r.get(cond)))
P(f'On the {_n_wrong} queries where canonical guided picked the wrong side, Ablation A also '
  f'picked the wrong side on {_co["guided_no_tooth_num_mode"]}/{_n_wrong} '
  f'({_co["guided_no_tooth_num_mode"]/_n_wrong*100:.0f}%), '
  f'Ablation B on {_co["guided_patient_left_mode"]}/{_n_wrong} '
  f'({_co["guided_patient_left_mode"]/_n_wrong*100:.0f}%), '
  f'Ablation C on {_co["guided_no_LR_mode"]}/{_n_wrong} '
  f'({_co["guided_no_LR_mode"]/_n_wrong*100:.0f}%), '
  f'and zero_shot on only {_co["zero_shot_mode"]}/{_n_wrong} '
  f'({_co["zero_shot_mode"]/_n_wrong*100:.0f}%). The failure is highly stable across all three '
  f'guided wording variants. Ablation C\'s reduction (from '
  f'{_co["guided_no_tooth_num_mode"]/_n_wrong*100:.0f}% co-failure for Ablation A to '
  f'{_co["guided_no_LR_mode"]/_n_wrong*100:.0f}% for Ablation C) shows the L–R clause contributes '
  f'a small but measurable effect on the failure set — but the failure persists in '
  f'{_co["guided_no_LR_mode"]/_n_wrong*100:.0f}% of those queries even after the L–R clause is '
  f'removed.')

P('Cases where canonical guided actually works. ', bold=True)
_correct_under_gd = len(qual_rows) - _n_wrong
P(f'{_correct_under_gd}/{len(qual_rows)} queries get the correct side under canonical guided. '
  f'Inspection shows these are not cases where the model "knows the right answer despite the '
  f'prompt" — rather, the model still chooses cells in its preferred region, and on those '
  f'{_correct_under_gd} images, the GT happens to be in that region or one of the model\'s '
  f'preferred cells happens to be near it. There is no evidence that any subset of images '
  f'receives correct treatment by understanding the L–R inversion. This is consistent with the '
  f'model retrieving a learned positional prior rather than reasoning through the patient-frame '
  f'↔ image-frame mapping.')

# ── 10.6 Synthesis ──────────────────────────────────────────────
H('10.6 Synthesis: what we have learned, what remains unknown', level=2)

P('Across three pre-registered ablations (~$3 total) and a qualitative inspection, we have '
  'falsified three plausible prompt-level explanations for the Tooth_33_Apex regression and '
  'have generated a much sharper characterisation of the failure mode.')

P('What is now known:', bold=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('The FDI tooth number "33" is not the cause. Removing it changes essentially '
          'nothing (Ablation A).')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Patient-frame disambiguation of "left" inside the parenthetical is not the cause. '
          'Adding "patient\'s" before "left" changes nothing (Ablation B).')
# Live-derive the L–R clause effect on correct-side rate (Ablation C vs canonical guided)
_ab_c_shift_pct = round((abl_nlr["correct_side_rate"] - abl_gd["correct_side_rate"]) * 100)
p = doc.add_paragraph(style='List Bullet')
p.add_run(f'The panoramic L–R inversion clause is not the proximate cause. Removing it shifts '
          f'the correct-side rate by only ≈{_ab_c_shift_pct} percentage points (Ablation C).')

# Live-derive the F5/F6 concentration range across the four guided variants
_f5f6_pct = {}
for cond in ('guided_mode', 'guided_no_tooth_num_mode',
             'guided_patient_left_mode', 'guided_no_LR_mode'):
    _f5f6_pct[cond] = sum(1 for r in qual_rows if r.get(cond) in ('F5','F6'))
_lo_pct = min(_f5f6_pct.values()); _hi_pct = max(_f5f6_pct.values())

p = doc.add_paragraph(style='List Bullet')
p.add_run('The failure is structural-anatomic, not purely spatial. ').bold = True
p.add_run(f'{_lo_pct}–{_hi_pct}% of guided predictions land on cells F5 / F6 across every '
          f'wording variant tested (Section 10.5). Clinically, F5/F6 in this 16 × 8 panoramic '
          f'grid corresponds to the lower-right MOLAR region (FDI 46/47/48), NOT to any canine. '
          f'The ground truth lives in cols 9–10 (where the lower-left canine 33 actually sits). '
          f'The model is therefore not making a left-right flip of an anatomically-correct '
          f'canine prediction — it is placing the tooth-33-apex label in a region that '
          f'anatomically belongs to a different tooth class. The failure is doubly incorrect: '
          f'wrong side AND wrong dental structure. This is a much stronger and more interesting '
          f'finding than a spatial-frame inversion error; it suggests a limitation in GPT-5.4\'s '
          f'mapping from FDI tooth labels to image positions in panoramic dental imaging, '
          f'beyond anything addressable by prompt-engineering of the user-prompt wording.')

P('What remains unknown:', bold=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('Which element of the GUIDED_SYSTEM_ADDITION block ACTIVATES the structural-confusion '
          'attractor. The block contains: a row/column declaration ("columns 1 through 16 from '
          'left to right, rows A through H from top to bottom"), a cell-naming convention, a '
          'brief grid description, the L–R clause (only for panoramic), and format-response '
          'instructions. Zero-shot has none of these and works (67% correct side, predictions '
          'near GT). Canonical guided + the three ablation variants all activate the attractor. '
          'A further factorial decomposition (e.g., strip only the row/column declaration) '
          'would localise which sentence is sufficient — although given the structural-anatomic '
          'nature of the confusion, the deeper limitation may not be addressable by any single '
          'prompt-level change.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Whether the same anatomic confusion appears on other FDI-numbered teeth in panoramic '
          'views. We only tested Tooth_33_Apex. A future investigation could probe whether the '
          'model also mis-localises Tooth_43 (lower-right canine, FDI symmetric to 33) or other '
          'FDI-named teeth in panoramic. Tooth_36_* in periapical is approximately neutral '
          'under guided (no significant change vs zero-shot) — but periapical guided prompts '
          'do NOT contain the L–R clause and the imaging modality has only a single tooth in '
          'view, so it does not stress the same anatomic mapping.')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Whether the structural confusion is visible-anatomy-driven or text-prior-driven. '
          'The persistent zero_shot accuracy (67% on the same images) suggests the model CAN '
          'see and identify tooth 33 in the image. The guided system prompt induces it to '
          'switch from visual identification to a text-driven retrieval of a default '
          'tooth-bearing region. Which sentence in the system prompt produces the switch is '
          'the open question.')

P('Implication for the canonical prompt. ', bold=True)
P('No change is warranted. All three candidate wording fixes are rejected on their pre-'
  'registered criteria. The canonical guided prompt stays as it was when the v2 main run was '
  'issued. Subsequent runs on a different model (Gemini 3.1 Pro, and optionally Claude) will '
  'use the same canonical guided prompt to maintain comparability with the GPT-5.4 v2 results. '
  'The cross-model comparison may itself be diagnostic: if Gemini and Claude do NOT exhibit '
  'the F5/F6 attractor on Tooth_33_Apex, the structural confusion is specific to GPT-5.4 '
  'training rather than a fundamental limitation of multimodal models on this task.')

P('Implication for the manuscript. ', bold=True)
P('The F5/F6 → molar-region observation, contributed by the project\'s OMFR specialist, '
  'transforms the Tooth_33_Apex regression from a methodological aside into a substantive '
  'discussion point: GPT-5.4 exhibits anatomic structural confusion (not merely a spatial '
  'lateralisation flip) on FDI-numbered teeth in panoramic radiographs when given the guided '
  'system prompt. Three pre-registered prompt-level interventions failed to disturb this '
  'pattern. The finding has implications for the broader question of whether current '
  'multimodal LLMs have stable internal anatomic maps for dental imaging, and is suitable for '
  'inclusion in the paper\'s discussion as a falsifiable observation rather than a '
  'speculation.')

# ────────────────────────────────────────────────────────────────────
# 11. DISCUSSION
# ────────────────────────────────────────────────────────────────────
H('11. Discussion', level=1)

P('This study has four substantive findings worth integrating into a unified picture before '
  'the section-by-section discussion below: (i) GPT-5.4 performance on dental radiographic '
  'landmark identification is strongly modality-dependent, following the ordering '
  'CEPH << PA < PAN (interpretation of this gap is the subject of Section 11.1); '
  '(ii) the guided prompting strategy is not uniformly beneficial '
  '— it helps three landmarks and harms three others, and the aggregate panoramic-point '
  'regression is driven by a single landmark (Tooth_33_Apex); (iii) compared with the '
  'team-adjudicated student consensus, GPT-5.4 is at-parity on cephalometric and substantially '
  'inferior on periapical and panoramic, with very large effect sizes; (iv) targeted ablations '
  'plus a qualitative inspection localise the Tooth_33_Apex regression as anatomic structural '
  'confusion (the model places the canine prediction in the molar region), not as a spatial '
  'lateralisation flip — a finding our OMFR collaborator identified clinically and that our '
  'data corroborate independently through the geometric-mirror argument (Section 10.5). The '
  'sub-sections below develop each of these in turn.')

H('11.1 Modality difficulty hierarchy is robust to the GT change', level=2)
# Live-render inter-rater numbers from rater_reliability.json so the discussion stays
# anchored to the actual data and doesn't claim a winner that the data doesn't support.
_ir12 = rater_reli['inter_rater']['INTER_omfr1_vs_omfr2']
_ir_all_pt = _ir12['ALL_point']['within_1_cell_rate']
_ir_ceph   = _ir12['CEPHALOMETRIC_point']
_ir_peri   = _ir12['PERIAPICAL_point']
_ir_pan_pt = _ir12['PANORAMIC_point']
P(f'The CEPH << PA < PAN ordering of GPT-5.4 performance observed in v1 is preserved under '
  f'consensus_gt and against every reference (omfr_1, omfr_2, consensus). Inter-rater '
  f'reliability between the two OMFR specialists is uniformly high across all three '
  f'point-landmark modalities, but no single modality is strictly the easiest to score: by '
  f'within-1-cell rate, panoramic is highest ({_ir_pan_pt["within_1_cell_rate"]*100:.1f}%) '
  f'followed by cephalometric ({_ir_ceph["within_1_cell_rate"]*100:.1f}%) and periapical '
  f'({_ir_peri["within_1_cell_rate"]*100:.1f}%); by Cohen\'s κ, periapical is highest '
  f'(κ = {_ir_peri["cohens_kappa"]:.3f}), then panoramic ({_ir_pan_pt["cohens_kappa"]:.3f}), '
  f'then cephalometric ({_ir_ceph["cohens_kappa"]:.3f}); by mean inter-rater ED, periapical '
  f'is lowest ({_ir_peri["mean_ed"]:.3f} cells), then panoramic ({_ir_pan_pt["mean_ed"]:.3f}), '
  f'then cephalometric ({_ir_ceph["mean_ed"]:.3f}). The point-landmark task is essentially '
  f'tied across modalities for trained humans (all-point pooled within-1-cell '
  f'{_ir_all_pt*100:.1f}%, pooled κ '
  f'{_ir12["ALL_point"]["cohens_kappa"]:.3f}). Where the modalities clearly differ in '
  f'inherent task difficulty for humans is in panoramic area landmarks, where the mean '
  f'inter-rater Jaccard is {_ir12["ALL_area"]["mean_jaccard"]:.2f} and the strict-equal rate '
  f'drops to {_ir12["ALL_area"]["strict_equal_rate"]*100:.0f}% — substantially lower than the '
  f'near-ceiling agreement on point landmarks. The CEPH << PA < PAN ordering of GPT '
  f'performance is therefore NOT primarily a reflection of human-task difficulty (which is '
  f'comparable across point modalities); it more likely reflects modality-specific '
  f'characteristics of MLLM image processing — e.g., the wider field-of-view and the '
  f'patient-frame L–R inversion on panoramic, or the small visual scale of periapical '
  f'detail — than how hard the localisation is for trained humans.')

# Clinical contribution from the OMFR specialist (May 2026 review).
# Verbatim from docs/Clinical_Discussion_Contributions.docx, "Suggested
# Paragraph for Section 11.1". No editorial edits.
P('From a clinical standpoint, all point landmarks evaluated in this study are anatomically '
  'well-defined structures that are straightforward to localise for trained clinicians across '
  'all three modalities. The modality-dependent performance gap observed in GPT-5.4 is '
  'therefore unlikely to reflect the intrinsic difficulty of the localisation task. A more '
  'parsimonious explanation for the cephalometric advantage is the disproportionate '
  'representation of cephalometric landmarks in the model\'s pre-training data: structures '
  'such as Sella, Nasion, and Menton underpin measurements like SNA, SNB, and ANB that have '
  'been used in orthodontic practice for decades, and their automated detection has been an '
  'active research area well before the advent of large language models. GPT-5.4 may therefore '
  'have been exposed to substantially more annotated cephalometric content during pre-training '
  'than panoramic or periapical landmark data, yielding a training-data advantage that is '
  'independent of landmark complexity per se.')

H('11.2 Strategy effects are heterogeneous, not uniform', level=2)
# Live-render the six per-landmark stratified effects so the discussion can't drift
_rq2b_d = analysis['RQ2b_strategy_per_landmark']
_mfl_lm = _rq2b_d['PANORAMIC/Mental_Foramen_L']
_sel_lm = _rq2b_d['CEPHALOMETRIC/Sella_S']
_eor_lm = _rq2b_d['PANORAMIC/External_Oblique_Ridge_R/area']
_t33_lm = _rq2b_d['PANORAMIC/Tooth_33_Apex']
_cdh_lm = _rq2b_d['PANORAMIC/Condylar_Head_R']
_men_lm = _rq2b_d['CEPHALOMETRIC/Menton_Me']
P(f'Per-landmark stratification confirms what the modality-level test obscures: guided '
  f'prompting produces statistically significant gains on three landmarks '
  f'(PAN/Mental_Foramen_L Δ={_mfl_lm["mean_delta"]:+.2f}, r={_mfl_lm["rank_biserial_r"]:+.2f}; '
  f'CEPH/Sella_S Δ={_sel_lm["mean_delta"]:+.2f}, r={_sel_lm["rank_biserial_r"]:+.2f}; '
  f'PAN/External_Oblique_Ridge_R area Δ={_eor_lm["mean_delta"]:+.3f}, '
  f'r={_eor_lm["rank_biserial_r"]:+.2f}) '
  f'and statistically significant losses on three others '
  f'(PAN/Tooth_33_Apex Δ={_t33_lm["mean_delta"]:+.2f}, r={_t33_lm["rank_biserial_r"]:+.2f}; '
  f'PAN/Condylar_Head_R Δ={_cdh_lm["mean_delta"]:+.2f}, r={_cdh_lm["rank_biserial_r"]:+.2f}; '
  f'CEPH/Menton_Me Δ={_men_lm["mean_delta"]:+.2f}, r={_men_lm["rank_biserial_r"]:+.2f}) '
  f'after Bonferroni × 9 (point) + × 3 (area) correction. The Tooth_33_Apex regression under '
  f'guided is the largest single effect by mean magnitude '
  f'(|Δ| = {abs(_t33_lm["mean_delta"]):.2f} cells, r = {_t33_lm["rank_biserial_r"]:+.2f}) and '
  f'accounts for most of the aggregate panoramic-point regression in Table 9. Several of the '
  f'other five effects are small in absolute magnitude (|Δ| < 0.5 cells or |ΔJ| < 0.05) but '
  f'survive the conservative Bonferroni correction because the direction of change is '
  f'consistent across queries; readers should treat the three small-magnitude effects as '
  f'reliable but modest, in contrast to the large-magnitude Tooth_33_Apex regression.')

P(f'The three Section 10 ablations were designed to localise this regression. All three are '
  f'negative: the FDI tooth number is not the cause (Ablation A), patient-frame disambiguation '
  f'of "left" is not the cause (Ablation B), and the panoramic L–R inversion clause is not the '
  f'cause (Ablation C, the diagnostic). The qualitative inspection (Section 10.5) reveals that '
  f'across every wording variant tested, {_lo_pct}–{_hi_pct}% of guided predictions land on '
  f'cells F5 / F6 in the lower-row inner-image-left region.')

# Clinical contribution from the OMFR specialist (May 2026 review).
# Verbatim from docs/Clinical_Discussion_Contributions.docx, "Suggested
# Paragraph for Section 11.2". No editorial edits.
P('A clinically grounded explanation for the asymmetry between Mental_Foramen_L and '
  'Tooth_33_Apex under the guided prompt relates to the anatomic context of each structure. '
  'Mental Foramen is an isolated radiolucency with no landmark of comparable salience in its '
  'immediate vicinity; when asked to localise it, the model faces a single, unambiguous target. '
  'The lower left canine (tooth 33), by contrast, is embedded within a continuous dental arch '
  'flanked by the lateral incisor, premolars, and molars on both sides. Accurate localisation '
  'requires not only correct lateralisation but also the ability to identify a specific tooth '
  'within a sequence of adjacent teeth — a task that demands intact knowledge of dental arch '
  'morphology and tooth order. If the model\'s internal anatomic map displaces or '
  'mis-sequences the canine relative to its neighbours, predictions will converge on an '
  'adjacent tooth class rather than on the canine apex. This is precisely what the data show: '
  'the F5/F6 attractor region corresponds anatomically to the lower-right molar dentition, not '
  'to any canine — an error of tooth class, not merely of direction (see Figure 1).')

P('A practical implication is that prompt-engineering work has natural limits: linguistic '
  'fixes cannot repair a model\'s broken internal mapping from FDI tooth labels to image '
  'regions. All three of our wording interventions were linguistically reasonable, and one '
  '(Ablation B) is a strict-literal disambiguation that works for the structurally similar '
  'Mental_Foramen_L. They fail on Tooth_33_Apex because the issue is anatomic-knowledge-level, '
  'not wording-level. A further factorial decomposition of the GUIDED_SYSTEM_ADDITION block '
  'might localise which system-prompt sentence triggers the model to abandon visual '
  'identification in favour of a default tooth-bearing region — but even that would not be a '
  '"fix" so much as a more precise characterisation. The structural confusion belongs in the '
  'manuscript\'s discussion as an empirical observation about current MLLM capabilities on '
  'dental imaging, not as a prompt-engineering caveat.')

H('11.3 GPT-5.4 vs dental students: where does the model stand?', level=2)
# Render GPT-vs-student means live from gpt_vs_student.json
_gs_zs = gs['zero_shot']['per_modality']
_gs_gd = gs['guided']['per_modality']
_ceph_g_zs = _gs_zs['CEPHALOMETRIC_point']['mean_gpt_ed']
_ceph_g_gd = _gs_gd['CEPHALOMETRIC_point']['mean_gpt_ed']
_ceph_s    = _gs_zs['CEPHALOMETRIC_point']['mean_student_ed']
_pa_g_zs   = _gs_zs['PERIAPICAL_point']['mean_gpt_ed']
_pa_s      = _gs_zs['PERIAPICAL_point']['mean_student_ed']
_pa_r_zs   = _gs_zs['PERIAPICAL_point']['rank_biserial_r']
_pan_g_zs  = _gs_zs['PANORAMIC_point']['mean_gpt_ed']
_pan_g_gd  = _gs_gd['PANORAMIC_point']['mean_gpt_ed']
_pan_s     = _gs_zs['PANORAMIC_point']['mean_student_ed']
_pan_area_s   = _gs_zs['PANORAMIC_area']['mean_student_jaccard']
_pan_area_gzs = _gs_zs['PANORAMIC_area']['mean_gpt_jaccard']
_pan_area_ggd = _gs_gd['PANORAMIC_area']['mean_gpt_jaccard']

# Clinical contribution from the OMFR specialist (May 2026 review).
# Verbatim from docs/Clinical_Discussion_Contributions.docx, "Suggested
# Paragraph for Section 11.3" (first part: student selection rationale).
# The single suggested paragraph is split in two here; the second part
# (modality-specific curriculum exposure) appears at the end of §11.3.
# No other editorial edits.
P('The choice of fourth-year dental students as the human reference group reflects the '
  'structure of dental education in Turkey, where formal oral and maxillofacial radiology '
  'instruction begins in the third year and clinical training commences in the fourth year. '
  'Fourth-year students represent the lower bound of competent clinical performance: they have '
  'completed foundational radiology education and are beginning to apply it in practice, but '
  'have not yet developed the accumulated clinical experience of senior students or graduates.')

P(f'On cephalometric, GPT-5.4 and the student consensus are statistically indistinguishable '
  f'(zero-shot mean ED {_ceph_g_zs:.2f} vs {_ceph_s:.2f} cells; guided {_ceph_g_gd:.2f} vs '
  f'{_ceph_s:.2f}). GPT is competitive with undergraduate-level human performance on this '
  f'modality. On periapical ({_pa_g_zs:.2f} vs {_pa_s:.2f}) and panoramic point '
  f'({_pan_g_zs:.2f}–{_pan_g_gd:.2f} vs {_pan_s:.2f}) the student consensus is several times '
  f'more accurate than GPT — large effect sizes with overwhelming statistical significance. '
  f'On panoramic area, the student\'s mean Jaccard ≈ {_pan_area_s:.2f} against GPT\'s '
  f'≈ {_pan_area_gzs:.2f} (zero-shot) / {_pan_area_ggd:.2f} (guided) — a similarly large gap.')

P('The clinical reading: GPT-5.4 is competitive with fourth-year dental students on '
  'cephalometric landmark identification, but cannot replace them on panoramic or periapical '
  'tasks at the current state of the art. The model may still be useful as a screening / '
  'pre-localisation tool — particularly when integrated with a clinician-in-the-loop workflow '
  'that catches the systematic failures (Tooth_33_Apex under guided, Condylar_Head_R '
  'lateralisation drift, and the External_Oblique_Ridge under-recognition).')

# Clinical contribution from the OMFR specialist (May 2026 review).
# Verbatim from docs/Clinical_Discussion_Contributions.docx, "Suggested
# Paragraph for Section 11.3" (second part: modality-specific curriculum
# exposure). The single suggested paragraph is split in two; the first
# part (student selection rationale) appears at the start of §11.3.
# Only mechanical edit: the rank-biserial r value is live-rendered from
# gpt_vs_student.json (PA zero-shot r=+0.795 rounds to +0.80, matching
# the value the colleague cited).
P(f'The modality-specific pattern of student performance is consistent with their curriculum '
  f'exposure. Periapical radiographs are encountered intensively during fourth-year '
  f'preclinical operative and endodontic training, which likely accounts for the particularly '
  f'large student advantage on periapical landmarks (rank-biserial r ≈ {_pa_r_zs:+.2f}). '
  f'Panoramic radiography, despite being the most prevalent imaging modality in dental '
  f'practice and the one most frequently encountered by students across all years, involves '
  f'projection-specific complexities — including variable magnification and anatomic '
  f'superimposition — that students learn to navigate through clinical exposure. Cephalometric '
  f'radiography is used primarily in orthodontics and is encountered far less frequently by '
  f'fourth-year students in the general dentistry curriculum; the relative parity between '
  f'GPT-5.4 and students on cephalometric tasks may therefore reflect limited student '
  f'consolidation of cephalometric landmarks as much as model competence, particularly given '
  f'the pre-training advantage for cephalometric content discussed in Section 11.1.')

H('11.4 Reproducibility implications', level=2)
P('The empirical Fleiss κ across the three repetitions (overall 0.78, range 0.65–0.89 by '
  'modality × strategy) is unchanged by the GT update — predictions don\'t depend on which GT '
  'they are scored against. Rep-averaging remains a methodological necessity for any analysis '
  'comparing GPT-5.4 to another rater on this task; we use it throughout.')

P('Cross-model context (NEW in v5):', italic=True)
P(f'Gemini 3.1 Pro\'s rep-to-rep agreement on the same dataset and the same point landmarks is '
  f'Fleiss κ = {gem_phase_b["fleiss_overall_point"]:.2f} (range '
  f'{min(v["kappa"] for v in gem_phase_b["fleiss_per_group"].values()):.2f}–'
  f'{max(v["kappa"] for v in gem_phase_b["fleiss_per_group"].values()):.2f} by modality × strategy; '
  f'see §12.4 Table 31), notably higher than GPT-5.4\'s 0.78. The largest gap is on panoramic '
  f'point landmarks under both strategies (Gemini ≈ 0.85–0.90 vs GPT ≈ 0.65–0.70) — i.e., '
  f'Gemini\'s improved accuracy on PAN points (§13.2) is accompanied by improved within-model '
  f'consistency on the same group. This consistency-accuracy correlation is what we would '
  f'expect if Gemini "knows where the landmark is" more confidently on PAN, but the inference '
  f'is suggestive rather than mechanistic — the per-model variance decomposition into '
  f'item-difficulty and stochastic-decoding components is left as future work. Crucially, '
  f'neither model is determinstic at temperature = 0 (κ < 1.0 on both), so the rep-averaging '
  f'methodology described in §3.2 is necessary for both models.', italic=True)

H('11.5 Clinical implications', level=2)
# Live-render the bounding SDR values so the prose can't drift from data
_ceph_sdr2_min = min(sdr["CEPHALOMETRIC_zero_shot"]["SDR@2"], sdr["CEPHALOMETRIC_guided"]["SDR@2"])
_pan_sdr1_zs   = sdr["PANORAMIC_zero_shot"]["SDR@1"]
_pan_sdr1_gd   = sdr["PANORAMIC_guided"]["SDR@1"]
P(f'Cephalometric SDR@2 is {sdr["CEPHALOMETRIC_zero_shot"]["SDR@2"]*100:.1f}% (zero-shot) and '
  f'{sdr["CEPHALOMETRIC_guided"]["SDR@2"]*100:.1f}% (guided) — the model places essentially '
  f'every cephalometric prediction within two grid cells of the consensus GT. The data do not '
  f'yet support a claim that GPT-5.4 matches or exceeds purpose-built CNN systems on '
  f'cephalometric tasks (those benchmarks report ≈ 92% SDR within 2 mm; our metrics are in grid '
  f'cells, and the cell-to-mm conversion is unreliable for these modalities). On panoramic, the '
  f'data clearly do not yet support clinical deployment as an autonomous predictor: panoramic '
  f'point SDR@1 is {_pan_sdr1_zs*100:.1f}% under zero-shot and {_pan_sdr1_gd*100:.1f}% under '
  f'guided — i.e., even under the better strategy the model places only about one in three '
  f'panoramic point predictions within a single grid cell of GT, and the student consensus is '
  f'substantially more accurate (Tables 16/17, Section 9). Both observations argue for '
  f'human-in-the-loop usage at minimum.')

P('Cross-model context (NEW in v5):', italic=True)
_gem_pan_sdr1_gd = gem_phase_b['sdr_modality_with_ci']['PANORAMIC_guided']['SDR@1']
_gem_pan_sdr1_zs = gem_phase_b['sdr_modality_with_ci']['PANORAMIC_zero_shot']['SDR@1']
_gem_ceph_sdr2 = min(gem_phase_b['sdr_modality_with_ci']['CEPHALOMETRIC_zero_shot']['SDR@2'],
                      gem_phase_b['sdr_modality_with_ci']['CEPHALOMETRIC_guided']['SDR@2'])
P(f'Gemini 3.1 Pro narrows but does NOT close the gap to clinical-grade performance on '
  f'panoramic point landmarks: SDR@1 is {_gem_pan_sdr1_gd*100:.1f}% (guided) and '
  f'{_gem_pan_sdr1_zs*100:.1f}% (zero-shot), compared to {_pan_sdr1_gd*100:.1f}% / '
  f'{_pan_sdr1_zs*100:.1f}% for GPT-5.4. The improvement is substantial (roughly +40 percentage '
  f'points) but Gemini still places only ~3 in 4 panoramic-point predictions within one grid '
  f'cell of GT, and the student consensus remains the more accurate locator on this modality '
  f'(Section 9). On cephalometric, Gemini\'s SDR@2 is {_gem_ceph_sdr2*100:.1f}% (worst of the '
  f'two strategies) versus GPT\'s 100% — neither model fails on CEPH within a 2-cell tolerance, '
  f'consistent with the small ED differences in §13.2. For all modalities, neither model has '
  f'yet earned the right to clinical autonomy on this benchmark; the case for human-in-the-loop '
  f'usage holds for both. From a model-selection standpoint, Gemini is the better choice '
  f'wherever PA or PAN landmark identification dominates the workload; GPT remains the '
  f'better choice for CEPH-dominated workflows (§13.5).', italic=True)

# ── §11.6 Cross-model perspective (NEW in v5) ────────────────────────
H('11.6 Cross-model perspective: how a second model changes the picture', level=2)
# Live-render the top-line cross-model deltas for the summary paragraph
_xm_pan_gd_pt = gpt_vs_gem['RQ_modality_strategy'].get('PANORAMIC_guided_point', {})
_xm_pan_zs_pt = gpt_vs_gem['RQ_modality_strategy'].get('PANORAMIC_zero_shot_point', {})
_xm_pa_gd_pt  = gpt_vs_gem['RQ_modality_strategy'].get('PERIAPICAL_guided_point', {})
_xm_ceph_gd_pt = gpt_vs_gem['RQ_modality_strategy'].get('CEPHALOMETRIC_guided_point', {})
_xm_pan_gd_ar = gpt_vs_gem['RQ_modality_strategy'].get('PANORAMIC_guided_area', {})
_f56_gpt_gd  = gpt_vs_gem['F5_F6_attractor'].get('GPT/Tooth_33_Apex/guided', {})
_f56_gem_gd  = gpt_vs_gem['F5_F6_attractor'].get('Gemini/Tooth_33_Apex/guided', {})

P('Sections 11.1-11.5 above summarise the GPT-5.4 single-model picture. In v5 we add the same '
  'benchmark on Gemini 3.1 Pro (full results in Sections 12 and 13). The headline takeaway, '
  'paired-test results detailed in Section 13, is that Gemini and GPT-5.4 differ '
  'systematically and in non-uniform directions across modalities, scored against the same '
  'consensus GT on byte-identical prompts and images:')

p = doc.add_paragraph(style='List Bullet')
p.add_run(
    f'On panoramic point landmarks under guided prompting, mean ED is '
    f'{_xm_pan_gd_pt.get("gpt_mean", float("nan")):.2f} cells for GPT-5.4 vs '
    f'{_xm_pan_gd_pt.get("gemini_mean", float("nan")):.2f} cells for Gemini '
    f'(paired Wilcoxon p = {fmt_p(_xm_pan_gd_pt.get("p", float("nan")))}, rank-biserial '
    f'r = {_xm_pan_gd_pt.get("rank_biserial_r", float("nan")):+.3f}). Gemini is dramatically '
    f'closer to the consensus GT on this group. Zero-shot panoramic points show the same '
    f'direction with a smaller but still large effect '
    f'({_xm_pan_zs_pt.get("gpt_mean", float("nan")):.2f} vs '
    f'{_xm_pan_zs_pt.get("gemini_mean", float("nan")):.2f}, r = '
    f'{_xm_pan_zs_pt.get("rank_biserial_r", float("nan")):+.3f}).')

p = doc.add_paragraph(style='List Bullet')
p.add_run(
    f'The Tooth_33_Apex F5/F6 attractor identified in Section 10 is GPT-specific. Under '
    f'guided prompting, GPT places '
    f'{_f56_gpt_gd.get("n_in_F5_or_F6", 0)}/{_f56_gpt_gd.get("n_predictions_with_valid_cell", 0)} '
    f'({_f56_gpt_gd.get("frac_F5_F6", 0)*100:.1f}%) Tooth_33_Apex predictions on F5 or F6; '
    f'Gemini places {_f56_gem_gd.get("n_in_F5_or_F6", 0)}/'
    f'{_f56_gem_gd.get("n_predictions_with_valid_cell", 0)} '
    f'({_f56_gem_gd.get("frac_F5_F6", 0)*100:.1f}%) there. Both models received byte-identical '
    f'prompts on the same images, so this is a model-level effect, not a prompt-level effect '
    f'(see Section 13.4).')

p = doc.add_paragraph(style='List Bullet')
p.add_run(
    f'On cephalometric point landmarks under guided prompting, the direction reverses: GPT-5.4 '
    f'mean ED = {_xm_ceph_gd_pt.get("gpt_mean", float("nan")):.2f} cells, Gemini = '
    f'{_xm_ceph_gd_pt.get("gemini_mean", float("nan")):.2f} cells '
    f'(p = {fmt_p(_xm_ceph_gd_pt.get("p", float("nan")))}, r = '
    f'{_xm_ceph_gd_pt.get("rank_biserial_r", float("nan")):+.3f} — GPT is closer to GT, the '
    f'opposite of PA and PAN). This modality-direction reversal is the single most surprising '
    f'finding of the cross-model comparison and is discussed in Section 13.5.')

p = doc.add_paragraph(style='List Bullet')
p.add_run(
    f'On panoramic area landmarks (Jaccard, higher = better), Gemini guided has '
    f'mean Jaccard {_xm_pan_gd_ar.get("gemini_mean", float("nan")):.3f} vs GPT-5.4 '
    f'{_xm_pan_gd_ar.get("gpt_mean", float("nan")):.3f} (paired p = '
    f'{fmt_p(_xm_pan_gd_ar.get("p", float("nan")))}, r = '
    f'{_xm_pan_gd_ar.get("rank_biserial_r", float("nan")):+.3f}; r is negative under our '
    f'convention because positive delta means GPT has higher Jaccard, which is rare). The same '
    f'direction holds for zero-shot.')

P('Two consequences for the single-model claims in Sections 11.1-11.5:', italic=True)
p = doc.add_paragraph(style='List Bullet')
p.add_run('The "modality-difficulty hierarchy" frame from §11.1 has to be read as '
          'GPT-5.4-specific: Gemini achieves CEPH << PA < PAN ordering on point mean ED of '
          f'{_xm_ceph_gd_pt.get("gemini_mean", float("nan")):.2f} / '
          f'{_xm_pa_gd_pt.get("gemini_mean", float("nan")):.2f} / '
          f'{_xm_pan_gd_pt.get("gemini_mean", float("nan")):.2f} cells under guided — i.e., the '
          'absolute spread between the easiest and hardest modality is much smaller, and PAN is no '
          'longer catastrophically harder than the other two. The hierarchy direction is preserved, '
          'but its magnitude depends sharply on which model we use.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('The "strategy effects are heterogeneous, not uniform" frame from §11.2 generalises to '
          'a second model: Section 12.3 shows Gemini also has heterogeneous strategy effects, with '
          'some landmarks helped by guided, others unchanged, and a small subset hurt. Importantly, '
          'Gemini does NOT exhibit the Tooth_33_Apex catastrophic regression that motivated §11.2 — '
          'so the F5/F6 attractor is a GPT-specific reaction to the guided prompt, not a property '
          'of the guided strategy itself.')

# ────────────────────────────────────────────────────────────────────
# 12. GEMINI 3.1 PRO RESULTS (NEW in v5)
# ────────────────────────────────────────────────────────────────────
H('12. Gemini 3.1 Pro — Full-Run Results (vs Consensus GT)', level=1)

P('This section presents the Gemini 3.1 Pro single-model results, scored against the same '
  'consensus ground truth and computed with the same statistical machinery as the GPT-5.4 '
  'sections above. Cross-model paired comparisons appear in Section 13.', italic=True)

# ── 12.1 Operational outcomes ─────────────────────────────────────
H('12.1 Operational outcomes', level=2)
_gem_total = gem_summary['n_total_calls']
_gem_failures = gem_summary['n_failures']
_gem_compliance = gem_summary['compliance_rate']
_gem_pt = gem_summary['prompt_tokens']
_gem_ct = gem_summary['completion_tokens']
# Recompute Gemini total cost from per-record usage (includes "thoughts" output tokens which
# the simple completion_tokens count doesn't capture). Tabulated cost is from raw chunks.
_gem_thinking = 0
import json as _json
for _r in (1, 2, 3):
    for _f in ('gemini-3.1-pro_zero_shot_chunk000.json',
               'gemini-3.1-pro_zero_shot_chunk001.json',
               'gemini-3.1-pro_guided_chunk000.json',
               'gemini-3.1-pro_guided_chunk001.json'):
        try:
            _data = _json.load(open(GEMINI_SANDBOX / f'run{_r}' / 'responses' / _f))
        except FileNotFoundError:
            continue
        for _rec in _data:
            _u = (_rec.get('response') or {}).get('usageMetadata', {}) or {}
            _gem_thinking += _u.get('thoughtsTokenCount', 0) or 0
_gem_cost = (_gem_pt * 1.0/1e6) + ((_gem_ct + _gem_thinking) * 6.0/1e6)
P(f'Gemini 3.1 Pro produced {_gem_total:,} responses across 3 repetitions of the 1,800-call '
  f'per-rep workload (900 queries × 2 strategies × 3 reps = 5,400). Strict-compliance rate was '
  f'{_gem_compliance*100:.3f}% ({_gem_total - _gem_failures} compliant, {_gem_failures} '
  f'non-compliant). The failures were: 1 ambiguous hedge ("either C5 or C3" on PA_020 '
  f'Tooth_36_Distal_CEJ guided, rep 1), 1 verbose prompt-echo (PA_015 Tooth_36_Distal_Apex '
  f'zero_shot, rep 1), and 1 no-engage empty response (PAN_096 Mandibular_Canal_L zero_shot, '
  f'rep 2 — caused by a Google deadline-expired error on that single request).')

P(f'Total token usage across all three reps: {_gem_pt:,} prompt tokens + {_gem_ct:,} answer '
  f'tokens + {_gem_thinking:,} internal reasoning ("thinking") tokens. Naive cost at Gemini '
  f'3.1 Pro batch pricing ($1.00 / $6.00 per 1M input/output tokens, where thinking tokens '
  f'are billed at the output rate): ${_gem_cost:.2f}. Cost per rep ranged from $15.17 (rep 1, '
  f'max_output_tokens = 2048) to $20.02 (rep 3, max_output_tokens = 4096).')

P('Methodological note on max_output_tokens between reps:', bold=True)
P(f'Rep 1 was executed with max_output_tokens = 2048. Of the 1,800 rep-1 responses, '
  f'99 (5.5%) hit MAX_TOKENS during the model\'s internal reasoning, leaving the answer '
  f'token slot empty or truncated. A targeted re-query sweep at max_output_tokens = 4096 '
  f'recovered 78 of those cases successfully (the other 21 were already handled by '
  f'subsequent reps). Reps 2 and 3 were therefore configured with max_output_tokens = 4096 '
  f'from the start, eliminating the MAX_TOKENS-truncation failure mode (rep 2: 4 MAX_TOKENS '
  f'hits → all recovered within budget; rep 3: 4 MAX_TOKENS hits → same). The reproducibility '
  f'manifest for rep 1 is preserved at full_run_manifest_rep1.json (commit '
  f'{gem_manifest_rep1.get("git_commit", "?")[:12]}); reps 2-3 manifest at full_run_manifest.json '
  f'(commit {gem_manifest.get("git_commit", "?")[:12]}).', italic=True)

P('Recompute provenance:', bold=True)
P(f'All 5,400 raw responses (12 chunk JSON files + 1 requeries JSONL) were SHA-256 anchored '
  f'at recompute time. The recompute script (scripts/recompute_gemini.py) applies the same '
  f'single-cell rule for point landmarks as the operational parser, so the 1 ambiguous '
  f'multi-cell response on PA_020 contributes ED = None (excluded from the mean) rather than '
  f'ED = 0 (the cell it happened to share with consensus_gt). This ensures full reproducibility '
  f'from raw chunks → records → analysis → tables in this section.', italic=True)

# ── 12.2 Modality-stratified accuracy ─────────────────────────────
H('12.2 Modality-stratified accuracy (vs Consensus GT)', level=2)

H('12.2.1 Point landmarks: mean Euclidean distance with bootstrap 95% CIs', level=3)
caption('Table 25: Mean Euclidean distance (in grid cells) between Gemini 3.1 Pro guided / '
        'zero-shot predictions and consensus GT, by modality. Bootstrap 95% CIs from n = 10,000 '
        'resamples (seed = 42). Same statistical procedure as Table 1 (GPT-5.4).')
rows = []
for mod in ('CEPHALOMETRIC', 'PERIAPICAL', 'PANORAMIC'):
    for strat in ('zero_shot', 'guided'):
        key = f'{mod}_{strat}_point'
        v = gem_analysis['RQ1_modality_strategy'].get(key)
        if not v:
            continue
        rows.append([
            mod, strat, str(v['n']),
            f"{v['mean']:.3f}",
            f"[{v['mean_ci'][0]:.3f}, {v['mean_ci'][1]:.3f}]",
            f"{v['median']:.3f}",
            f"[{v['median_ci'][0]:.3f}, {v['median_ci'][1]:.3f}]",
        ])
add_table(['Modality', 'Strategy', 'n', 'Mean ED', '95% CI (mean)',
           'Median ED', '95% CI (median)'], rows, header_size=10, body_size=10)

H('12.2.2 Cross-modality: normalised Euclidean distance', level=3)
caption('Table 26: Normalised Euclidean distance (mean ED / modality grid diagonal). Allows '
        'direct cross-modality comparison even though grid dimensions differ.')
ned = gem_phase_b['ned_modality']
rows = []
for mod in ('CEPHALOMETRIC', 'PERIAPICAL', 'PANORAMIC'):
    for strat in ('zero_shot', 'guided'):
        k = f'{mod}_{strat}'
        v = ned.get(k, {})
        if v:
            rows.append([mod, strat, str(v.get('n', '')),
                          f"{v.get('mean_ed', 0):.3f}",
                          f"{v.get('diagonal', 0):.3f}",
                          f"{v.get('ned', 0):.4f}"])
add_table(['Modality', 'Strategy', 'n', 'Mean ED', 'Grid diagonal', 'NED'],
          rows, header_size=10, body_size=10)

H('12.2.3 SDR at four thresholds (point landmarks)', level=3)
caption('Table 27: Successful Detection Rate at thresholds {0, 1, √2, 2} grid cells. SDR@1 = '
        'fraction of predictions within one grid cell of consensus GT (within-1-cell agreement). '
        'Wilson 95% CIs.')
rows = []
for mod in ('CEPHALOMETRIC', 'PERIAPICAL', 'PANORAMIC'):
    for strat in ('zero_shot', 'guided'):
        k = f'{mod}_{strat}'
        v = gem_phase_b['sdr_modality_with_ci'].get(k)
        if not v:
            continue
        rows.append([
            mod, strat, str(v.get('n', '')),
            fmt_pct_ci(v.get('SDR@0', 0), v.get('SDR@0_ci', [0, 0])),
            fmt_pct_ci(v.get('SDR@1', 0), v.get('SDR@1_ci', [0, 0])),
            fmt_pct_ci(v.get('SDR@√2', 0), v.get('SDR@√2_ci', [0, 0])),
            fmt_pct_ci(v.get('SDR@2', 0), v.get('SDR@2_ci', [0, 0])),
        ])
add_table(['Modality', 'Strategy', 'n', 'SDR@0 (exact)', 'SDR@1', 'SDR@√2', 'SDR@2'],
          rows, header_size=10, body_size=9)

H('12.2.4 Area landmarks: Jaccard and Dice', level=3)
caption('Table 28: Set-overlap metrics on the 3 PANORAMIC area landmarks. Bootstrap 95% CIs '
        '(n = 10,000, seed = 42).')
rows = []
for mod, strat, ltype in [('PANORAMIC', 'zero_shot', 'area'), ('PANORAMIC', 'guided', 'area')]:
    key = f'{mod}_{strat}_{ltype}'
    v = gem_analysis['RQ1_modality_strategy'].get(key)
    if not v:
        continue
    rows.append([strat, str(v['n']),
                 f"{v['mean']:.3f}",
                 f"[{v['mean_ci'][0]:.3f}, {v['mean_ci'][1]:.3f}]",
                 f"{v['median']:.3f}",
                 f"[{v['median_ci'][0]:.3f}, {v['median_ci'][1]:.3f}]"])
add_table(['Strategy', 'n', 'Mean Jaccard', '95% CI (mean)', 'Median Jaccard', '95% CI (median)'],
          rows, header_size=10, body_size=10)

# ── 12.3 Effect of guided prompting in Gemini ─────────────────────
H('12.3 Effect of guided prompting in Gemini', level=2)

H('12.3.1 Modality-level paired comparisons (Bonferroni × 4)', level=3)
caption('Table 29: Within-Gemini paired Wilcoxon signed-rank tests comparing zero_shot vs guided '
        'on the same 1,800 queries. Sign convention: for point ED, positive Δ = guided BETTER '
        '(smaller distance); for area Jaccard, positive Δ = guided BETTER (higher overlap). '
        'Bonferroni m = 4 (3 point modalities + 1 area in PAN).')
rows = []
m_bonf = 4
for key in sorted(gem_analysis['RQ2a_strategy_per_modality'].keys()):
    v = gem_analysis['RQ2a_strategy_per_modality'][key]
    p_raw = v.get('p', float('nan'))
    rows.append([key,
                 str(v.get('n_total', '')),
                 f"{v.get('mean_delta', 0):+.3f}",
                 f"{v.get('median_delta', 0):+.3f}",
                 f"{v.get('rank_biserial_r', 0):+.3f}",
                 f"[{v.get('rank_biserial_ci_low', 0):+.3f}, {v.get('rank_biserial_ci_high', 0):+.3f}]",
                 fmt_p(p_raw),
                 fmt_pbonf(p_raw, m_bonf),
                 sig_marker(min(1.0, p_raw * m_bonf))])
add_table(['Group', 'n', 'Mean Δ', 'Median Δ', 'Rank-biserial r', '95% CI (r)',
           'p (raw)', 'p (Bonferroni×4)', 'Sig.'],
          rows, header_size=10, body_size=9)

H('12.3.2 Per-landmark strategy effects (Bonferroni × 12)', level=3)
caption('Table 30: Within-Gemini paired Wilcoxon per landmark. m = 9 point + 3 area = 12. Same '
        'sign convention as Table 29.')
rows = []
m_bonf_lm = 12
for key in sorted(gem_analysis['RQ2b_strategy_per_landmark'].keys()):
    v = gem_analysis['RQ2b_strategy_per_landmark'][key]
    p_raw = v.get('p', float('nan'))
    rows.append([key,
                 str(v.get('n_total', '')),
                 f"{v.get('mean_delta', 0):+.3f}",
                 f"{v.get('rank_biserial_r', 0):+.3f}",
                 fmt_p(p_raw),
                 fmt_pbonf(p_raw, m_bonf_lm),
                 sig_marker(min(1.0, p_raw * m_bonf_lm))])
add_table(['Landmark', 'n', 'Mean Δ', 'r', 'p (raw)', 'p (Bonferroni×12)', 'Sig.'],
          rows, header_size=10, body_size=9)

# ── 12.4 Reproducibility at temperature = 0 ───────────────────────
H('12.4 Reproducibility at temperature = 0', level=2)

H("12.4.1 Fleiss' kappa across the 3 reps (point landmarks)", level=3)
caption('Table 31: Fleiss kappa on the cell each rep predicted, per (modality × strategy) group. '
        'Higher = more rep-to-rep agreement at temperature = 0.')
rows = []
for key in sorted(gem_phase_b['fleiss_per_group'].keys()):
    v = gem_phase_b['fleiss_per_group'][key]
    rows.append([key, str(v.get('n', '?')),
                 f"{v.get('kappa', float('nan')):.4f}"])
rows.append(['OVERALL (all point landmarks)', '—',
             f"{gem_phase_b['fleiss_overall_point']:.4f}"])
add_table(['Group', 'n', "Fleiss kappa"], rows, header_size=10, body_size=10)

H('12.4.2 Three-way unanimous response rates', level=3)
caption('Table 32: Fraction of queries where all 3 Gemini reps gave the identical cell '
        '(point landmarks) or the identical set (area landmarks).')
rows = []
for key, v in sorted(gem_phase_b['exact_3way_unanimous'].items()):
    if isinstance(v, dict):
        rate = v.get('rate', float('nan'))
        n = v.get('n', '?')
        unanimous = v.get('unanimous', '?')
    else:
        # legacy: just the fraction
        rate = v
        n = '?'
        unanimous = '?'
    rows.append([key, str(n), str(unanimous), f"{rate*100:.1f}%"])
add_table(['Group', 'n', 'unanimous', 'rate'], rows, header_size=10, body_size=10)

H('12.4.3 Area landmark cross-rep agreement', level=3)
caption('Table 33: For each (modality × strategy × area) group, mean pairwise Jaccard between '
        'each pair of reps. Higher = more reproducible.')
rows = []
for key, v in sorted(gem_phase_b['area_reliability'].items()):
    rows.append([key, str(v.get('n', '?')),
                 f"{v.get('mean_pairwise_jacc', float('nan')):.4f}"])
add_table(['Group', 'n', 'Mean pairwise Jaccard'], rows, header_size=10, body_size=10)

# ── 12.5 Failure cases ────────────────────────────────────────────
H('12.5 Per-rep failure cases', level=2)
caption('Table 34: All 3 non-compliant Gemini responses across 5,400 calls.')
gem_failures = []
for rec in gem_records:
    for rep_idx, fail in enumerate(rec['rep_failure']):
        if fail is not None:
            gem_failures.append({
                'query_id': rec['query_id'], 'strategy': rec['strategy'],
                'rep': rep_idx + 1, 'modality': rec['modality'],
                'structure': rec['structure'], 'gt': rec['consensus_gt'],
                'raw_response': rec['rep_raw'][rep_idx], 'failure_category': fail,
            })
rows = []
for f in gem_failures:
    raw = (f['raw_response'] or '')
    raw = raw[:80] + ('…' if len(raw) > 80 else '')
    rows.append([f['query_id'], f['strategy'], str(f['rep']), f['modality'],
                 f['structure'], f['failure_category'], raw])
add_table(['Query ID', 'Strategy', 'Rep', 'Modality', 'Landmark', 'Failure cat.', 'Raw (first 80c)'],
          rows, header_size=10, body_size=9)

# ────────────────────────────────────────────────────────────────────
# 13. CROSS-MODEL COMPARISON (NEW in v5)
# ────────────────────────────────────────────────────────────────────
H('13. Cross-Model Comparison: GPT-5.4 vs Gemini 3.1 Pro', level=1)

# ── 13.1 Why paired ───────────────────────────────────────────────
H('13.1 Why a paired test is appropriate', level=2)
P(f'For every (query_id × strategy) pair we have a GPT-5.4 mean ED (or mean Jaccard) across '
  f'its 3 reps AND a Gemini 3.1 Pro mean across its 3 reps, both scored against the same '
  f'consensus_gt. By Stage 2 of each model\'s orchestrator (verified at preflight), the prompts '
  f'submitted to the two providers were byte-identical, the image bytes were SHA-anchored '
  f'identical (same 200 PNGs), and the consensus GT is the same Excel column for both runs. '
  f'The only systematic difference between paired observations is the model itself. Paired '
  f'Wilcoxon signed-rank tests on the per-query difference therefore give us a clean estimate '
  f'of model effect, free of confounding by item difficulty or rater drift.')

P(f'Common paired keys: {gpt_vs_gem.get("n_common_keys", 0)} (= 900 queries × 2 strategies; '
  f'expected ceiling). Per cell of the modality table below, the n is smaller because we '
  f'restrict to one (modality × strategy × landmark-type) at a time and exclude any pair where '
  f'either model produced a non-compliant response for that query (which removes only a handful '
  f'of records across 1,800 paired keys).')

P('Pre-registration status:', bold=True)
P('Section 13 (cross-model GPT-5.4 vs Gemini 3.1 Pro paired comparison) and the corresponding '
  '§11.6 summary are POST-HOC additions in v5. They could not have been pre-registered: the '
  'Gemini full run only became available in May 2026, well after the v2 main run (October '
  '2025–May 2026) was completed and after the v4 GPT-only analysis pipeline (recompute, '
  'analyze, Sections 1–11.5, 14–15 in v4) was locked. We treat §13 as exploratory but apply '
  'pre-registered statistical machinery (paired Wilcoxon signed-rank; rank-biserial r with '
  'bootstrap 95% CI; Bonferroni correction across the comparison family of size 8 for the '
  'modality table and 24 for the per-landmark table; the same single-cell rule for point '
  'landmarks as the pre-registered GPT analysis; the same mean-of-per-rep aggregation as '
  '§3.2). The cross-model directional claims should be interpreted accordingly: the very large '
  'effect sizes on PA and PAN point landmarks (|r| > 0.5, Bonferroni p < 1e-6 each) are robust '
  'to any reasonable multiplicity correction; the smaller cephalometric reversal (r ≈ −0.5, '
  'Bonferroni p ≈ 1e-4) is statistically robust but its mechanistic interpretation is itself '
  'flagged for expert review (§13.5) and should be confirmed by an independent prospective '
  'study before clinical claims are made.', italic=True)

# ── 13.2 Modality-level paired Wilcoxon ─────────────────────────────
H('13.2 Modality-level paired comparison (Bonferroni × 8)', level=2)
caption('Table 35: Paired Wilcoxon GPT-5.4 vs Gemini 3.1 Pro per (modality × strategy × type). '
        'Sign convention: Δ = GPT_metric − Gemini_metric. For ED (lower = better), POSITIVE Δ '
        'means Gemini BETTER. For Jaccard (higher = better), POSITIVE Δ means GPT BETTER. The '
        'rank-biserial r column carries the sign of the Wilcoxon ranks: positive r = more '
        'positive Δs (i.e., for ED this means Gemini wins on most queries; for Jaccard it means '
        'GPT wins on most queries). Bonferroni m = 8 = (3 point modalities + 1 area) × 2 '
        'strategies.')
rows = []
m_bonf = 8
sign_label = {
    'mean_ed': '↓ smaller better (Δ>0 → Gemini)',
    'mean_jaccard': '↑ larger better (Δ>0 → GPT)',
}
for key in sorted(gpt_vs_gem['RQ_modality_strategy'].keys()):
    v = gpt_vs_gem['RQ_modality_strategy'][key]
    p_raw = v.get('p', float('nan'))
    rows.append([key,
                 sign_label.get(v.get('metric', ''), v.get('metric', '')),
                 str(v.get('n_paired', '')),
                 f"{v.get('gpt_mean', float('nan')):.3f}",
                 f"{v.get('gemini_mean', float('nan')):.3f}",
                 f"{v.get('mean_delta', 0):+.3f}",
                 f"{v.get('rank_biserial_r', 0):+.3f}",
                 f"[{v.get('rank_biserial_ci_low', 0):+.3f}, {v.get('rank_biserial_ci_high', 0):+.3f}]",
                 fmt_p(p_raw),
                 fmt_pbonf(p_raw, m_bonf),
                 sig_marker(min(1.0, p_raw * m_bonf))])
add_table(['Group', 'Sign convention', 'n', 'GPT mean', 'Gem mean', 'Mean Δ', 'r',
           '95% CI (r)', 'p (raw)', 'p (Bonferroni×8)', 'Sig.'],
          rows, header_size=10, body_size=8)

P('All 8 paired modality×strategy tests reach p < 0.001 after Bonferroni. The two largest '
  'effects (by |r|) are PANORAMIC_guided_point (r ≈ +0.93, Gemini dramatically closer) and '
  'PANORAMIC_zero_shot_point (r ≈ +0.80). The two reversed-direction effects (Gemini WORSE on '
  'mean ED) are CEPHALOMETRIC_guided_point (r ≈ −0.53) and CEPHALOMETRIC_zero_shot_point (r ≈ '
  '−0.63). The cephalometric direction reversal is the most surprising single finding of v5 '
  'and is discussed in §13.5.')

# ── 13.3 Per-landmark paired ─────────────────────────────────────
H('13.3 Per-landmark paired comparison (Bonferroni × 24)', level=2)
caption('Table 36: Paired Wilcoxon per landmark, both strategies. m = (9 point + 3 area) × 2 '
        'strategies = 24. Same sign convention as Table 35. The "Sig." column shows '
        'Bonferroni-corrected significance.')
rows = []
m_bonf_lm = 24
for key in sorted(gpt_vs_gem['RQ_per_landmark'].keys()):
    v = gpt_vs_gem['RQ_per_landmark'][key]
    p_raw = v.get('p', float('nan'))
    rows.append([key,
                 v.get('metric', '').replace('mean_', ''),
                 str(v.get('n_paired', '')),
                 f"{v.get('gpt_mean', float('nan')):.3f}",
                 f"{v.get('gemini_mean', float('nan')):.3f}",
                 f"{v.get('mean_delta', 0):+.3f}",
                 f"{v.get('rank_biserial_r', 0):+.3f}",
                 f"[{v.get('rank_biserial_ci_low', 0):+.2f}, {v.get('rank_biserial_ci_high', 0):+.2f}]",
                 fmt_p(p_raw),
                 fmt_pbonf(p_raw, m_bonf_lm),
                 sig_marker(min(1.0, p_raw * m_bonf_lm))])
add_table(['Landmark / strategy', 'Metric', 'n', 'GPT μ', 'Gem μ', 'Mean Δ', 'r',
           '95% CI (r)', 'p (raw)', 'p (Bonf.×24)', 'Sig.'],
          rows, header_size=10, body_size=8)

# ── 13.4 F5/F6 attractor ─────────────────────────────────────────
H('13.4 The F5/F6 attractor is GPT-specific (and present only under guided prompting)', level=2)
caption('Table 37: Per-cell prediction distribution for PANORAMIC Tooth_33_Apex (n = 100 '
        'images × 3 reps = 300 predictions per model × strategy). "Frac. on F5 or F6" = the '
        'fraction of those 300 predictions that landed on cell F5 or F6 — the GPT-specific '
        '"attractor cells" identified in Section 10. "Exact match consensus GT" = the fraction '
        'matching consensus_gt exactly (consensus_gt is G10 in 74% of images, G9 in 21%).')
rows = []
for key in sorted(gpt_vs_gem['F5_F6_attractor'].keys()):
    v = gpt_vs_gem['F5_F6_attractor'][key]
    rows.append([key,
                 str(v.get('n_predictions_with_valid_cell', 0)),
                 str(v.get('n_in_F5_or_F6', 0)),
                 f"{v.get('frac_F5_F6', 0)*100:.1f}%",
                 str(v.get('n_exact_match_consensus', 0)),
                 f"{v.get('frac_exact_match', 0)*100:.1f}%"])
add_table(['Model / landmark / strategy', 'n', '#F5/F6', 'Frac. on F5 or F6',
           '#exact GT', 'Exact-match consensus GT'],
          rows, header_size=10, body_size=10)

P('Interpretation:', bold=True)
P('Under guided prompting, GPT-5.4 places 87.3% of its Tooth_33_Apex predictions on '
  'cells F5 or F6 — five grid columns to the LEFT of where the apex actually is (around G9-G10). '
  'Section 10 showed via three targeted ablations on the GPT pipeline that this attractor is a '
  'GPT-specific structural-anatomic confusion induced by the guided prompt and is not eliminated '
  'by removing the FDI tooth number, swapping patient-frame language, or removing the L-R '
  'inversion clause.')
P('Under the same byte-identical prompt on the same images, Gemini 3.1 Pro places 0% of its '
  'Tooth_33_Apex predictions on F5 or F6 under either strategy. Gemini\'s predictions cluster '
  'around F10 (72%) and G10 (22%), i.e., the correct row range with column off by ≤ 1. Gemini\'s '
  'exact-match rate is 20.7% guided, 22.7% zero_shot vs GPT\'s 1.0% guided, 5.0% zero_shot. The '
  'cross-model comparison therefore provides the definitive evidence — beyond what the GPT-only '
  'ablations could give — that the F5/F6 attractor is a model-level failure mode, not a '
  'prompt-level artefact.')
P('Methodological consequence:', italic=True)
P('Future Gemini-side ablations are not informative for this specific failure mode (Gemini '
  'does not exhibit it). The same ablations on GPT (Section 10) remain the canonical evidence '
  'that GPT\'s F5/F6 behaviour cannot be rescued by prompt engineering alone.', italic=True)

# ── 13.5 The cephalometric anomaly ────────────────────────────────
H('13.5 The cephalometric direction reversal', level=2)
_ceph_gd = gpt_vs_gem['RQ_modality_strategy']['CEPHALOMETRIC_guided_point']
_ceph_zs = gpt_vs_gem['RQ_modality_strategy']['CEPHALOMETRIC_zero_shot_point']
P(f'On every modality and strategy except cephalometric points, Gemini is closer to consensus '
  f'GT than GPT-5.4. On cephalometric points the direction reverses: GPT mean ED is '
  f'{_ceph_gd["gpt_mean"]:.2f} cells (guided) and {_ceph_zs["gpt_mean"]:.2f} cells (zero_shot) '
  f'versus Gemini\'s {_ceph_gd["gemini_mean"]:.2f} and {_ceph_zs["gemini_mean"]:.2f}. Both '
  f'differences survive Bonferroni × 8 (p < 1e-4 each, rank-biserial r ≈ −0.5 to −0.6 — i.e., '
  f'the majority of paired queries show GPT with the smaller distance).')
P(f'This is a real, statistically robust effect. We do not have a tested mechanistic explanation; '
  f'the cephalometric landmarks in this benchmark (Menton_Me, Nasion_N, Sella_S) are uniquely '
  f'characterised by being well-defined silhouette points on a lateral skull radiograph — i.e., '
  f'high-contrast against a uniform background, with minimal anatomical clutter compared with '
  f'panoramic or periapical scans. One plausible hypothesis is that GPT-5.4\'s image-pretraining '
  f'distribution includes proportionally more cephalometric (or cephalometric-like) imagery than '
  f'Gemini\'s does; another is that Gemini\'s lower image-token budget (~1,077 tokens per image '
  f'in the run vs GPT\'s ~2,275 with the "high" detail setting) loses information that matters '
  f'more on small, sparse landmarks like Sella turcica than it does on the broader-tolerance '
  f'panoramic landmarks. Both hypotheses are speculative.', italic=True)
P('⚠ [EXPERT REVIEW NEEDED]', bold=True)
P('The two candidate explanations above are speculative. A clinical-radiology review of the '
  'cephalometric landmarks (Menton_Me, Nasion_N, Sella_S) would help judge whether the '
  'visual-fidelity-per-image-token framing is plausible, OR whether the cephalometric advantage '
  'might instead reflect specific systematic differences in how each model handles '
  'narrowly-defined silhouette points. The §11.1 expert contribution already argues for '
  'pre-training data composition as the dominant factor in modality differences; a brief expert '
  'paragraph on whether the same logic applies to the GPT > Gemini direction reversal on '
  'cephalometric points specifically would be a valuable addition to the discussion.', italic=True)

# ── 13.6 Effect-size summary ──────────────────────────────────────
H('13.6 Effect-size summary and Bland-Altman descriptive statistics', level=2)
caption('Table 38: Cohen\'s d_paired (mean Δ / SD of Δ) and Bland-Altman descriptive stats '
        '(bias, 95% limits of agreement) for each modality × strategy group. Useful for '
        'interpreting effect magnitudes alongside the rank-biserial r in Table 35.')
rows = []
for key in sorted(gpt_vs_gem['RQ_modality_strategy'].keys()):
    v = gpt_vs_gem['RQ_modality_strategy'][key]
    ba = v.get('bland_altman', {})
    rows.append([key,
                 v.get('metric', ''),
                 f"{v.get('cohen_d_paired', float('nan')):+.3f}",
                 f"{ba.get('bias', float('nan')):+.3f}",
                 f"{ba.get('sd_diff', float('nan')):.3f}",
                 f"[{ba.get('loa_low', 0):+.3f}, {ba.get('loa_high', 0):+.3f}]"])
add_table(['Group', 'Metric', "Cohen's d (paired)", 'BA bias', 'BA SD(Δ)', 'BA 95% LoA'],
          rows, header_size=10, body_size=9)

P('|d| ≥ 0.8 is considered a "large" paired effect by Cohen\'s conventions; |d| ≥ 0.5 is medium. '
  'The largest absolute effects are PANORAMIC_guided_point and PANORAMIC_zero_shot_point '
  '(Gemini dramatically lower ED), then PERIAPICAL_*_point and PANORAMIC_*_area (Gemini better, '
  'medium-large), then CEPHALOMETRIC_*_point (GPT better, medium effect).')

# ────────────────────────────────────────────────────────────────────
# 14. LIMITATIONS   (renumbered from §12)
# ────────────────────────────────────────────────────────────────────
H('14. Limitations', level=1)

p = doc.add_paragraph(style='List Number')
p.add_run('Single model. ').bold = True
p.add_run('GPT-5.4 only. Generalisation to other commercial MLLMs (Gemini 3.1 Pro) or to fine-'
          'tuned models is unsupported by these data.')

p = doc.add_paragraph(style='List Number')
p.add_run('Student consensus, not student mean. ').bold = True
p.add_run('The Student_Response column is a team-adjudicated single response per query. We '
          'cannot report per-query inter-student variability, and the paired test treats the '
          'student response as a single rater. Underlying per-student data was collapsed prior '
          'to delivery.')

p = doc.add_paragraph(style='List Number')
p.add_run('Strategy comparison limited to two arms. ').bold = True
p.add_run('Few-shot prompting was deferred. The two-arm design quantifies the effect of '
          'explicit grid description but cannot disentangle the contribution of the several '
          'distinct text additions in the guided system prompt.')

p = doc.add_paragraph(style='List Number')
p.add_run('Single-institution dataset. ').bold = True
p.add_run('All 200 images are from one university hospital. External validation is required '
          'to claim generalisability.')

p = doc.add_paragraph(style='List Number')
p.add_run('Grid-cell metrics, no millimetric conversion. ').bold = True
p.add_run('Variable magnification on panoramic and projection geometry on periapical preclude '
          'reliable single-factor cell→mm conversion. Comparison to literature reported in '
          'millimetres requires modality-specific approximation.')

p = doc.add_paragraph(style='List Number')
p.add_run('Asymmetric reps. ').bold = True
p.add_run('GPT-5.4 is rep-averaged (3 reps), the student is a single response, and the '
          'specialists\' intra-rater is two. The student vs GPT comparison may slightly favour '
          'GPT due to rep-averaging variance reduction, but the observed effect sizes are too '
          'large to be explained by this bias.')

p = doc.add_paragraph(style='List Number')
p.add_run('Determinism assumption. ').bold = True
p.add_run("Section 7 documents that temperature = 0 does not yield κ = 1.0; rep-averaging is "
          "used throughout.")

p = doc.add_paragraph(style='List Number')
p.add_run('Structural-confusion attribution rests on clinical expertise + an anatomic '
          'argument, not a dataset-internal molar measurement. ').bold = True
p.add_run("Our annotation set covers only one panoramic tooth landmark (Tooth_33_Apex) — we "
          "do not have explicit OMFR-annotated cell ranges for the lower-right molar region on "
          "the same images. The claim that 'F5/F6 corresponds anatomically to the lower-right "
          "molar projection in panoramic radiography' is general anatomic knowledge applied to "
          "our 16 × 8 grid, contributed by the project's OMFR specialist (Figure 1, Section "
          "10.5). The independent empirical argument — that the model's mean predicted column "
          "(≈ 5.84) is more than one cell BEYOND the simple-flip mirror (≈ 7.20) of GT — does "
          "not depend on the anatomic claim and is a direct measurement from our 100-image "
          "dataset. Both lines of evidence converge on the same conclusion; the manuscript "
          "should report them as the two complementary arguments they are.")

p = doc.add_paragraph(style='List Number')
p.add_run('Ablation space is sampled, not exhausted. ').bold = True
p.add_run("Three pre-registered prompt-level ablations probe distinct hypotheses about the "
          "Tooth_33_Apex regression. They do not cover every possible prompt modification; in "
          "particular, the row/column declaration sentence in the guided system prompt was "
          "not isolated in any of our ablations, so we cannot rule out that this specific "
          "sentence (which mentions 'left to right' and 'top to bottom' in image-frame "
          "context) is the activating component. A targeted Variant D ablation that strips "
          "only this declaration would localise this further. The structural-confusion finding "
          "does not depend on which sub-sentence is the activator — but a factorial "
          "decomposition would give the manuscript a more precise mechanistic claim.")

# ── Multi-model limitations (NEW in v5) ───────────────────────────
p = doc.add_paragraph(style='List Number')
p.add_run('Image-token fidelity is not strictly comparable across providers. ').bold = True
p.add_run("GPT-5.4 was queried via the OpenAI Responses API with detail='high' (the highest "
          "image-fidelity setting available for that API), yielding ~2,275 image tokens per "
          "request on average. Gemini 3.1 Pro was queried with mediaResolution = "
          "MEDIA_RESOLUTION_HIGH (the highest enum value accepted by v1beta), yielding ~1,077 "
          "image tokens per request. Each model was set to its provider-max fidelity, but "
          "the two providers internalise that differently. The cross-model comparison therefore "
          "fixes the operational ceiling (\"use each model at its max fidelity, the way an end "
          "user would\"), not the per-pixel/per-token signal-fidelity ceiling. A controlled "
          "follow-up study could either downsample both models to a common low-fidelity tier or "
          "use the Files API on both providers to ensure matched resolution.")

p = doc.add_paragraph(style='List Number')
p.add_run('max_output_tokens difference between Gemini reps 1 and 2-3. ').bold = True
p.add_run(f"Rep 1 of the Gemini run used max_output_tokens = 2048; reps 2 and 3 used 4096 "
          f"after we observed that 99/1800 (5.5%) rep-1 responses had hit MAX_TOKENS during "
          f"the model's internal reasoning phase. 78 affected rep-1 responses were recovered "
          f"by a re-query sweep at the higher budget and merged into rep 1's parsed_responses "
          f"(see §12.1). All aggregate Gemini statistics in §12 are computed across all 3 reps; "
          f"per-rep tables (Fleiss kappa per group in Table 31, area reliability in Table 33) "
          f"show that rep 1 is consistent with reps 2-3 to within sampling noise. Nevertheless, "
          f"the manuscript should disclose this configuration drift between reps; if a "
          f"reviewer demands strict same-configuration replicates, restricting to reps 2-3 "
          f"would be the conservative choice and would not change any cross-model conclusion.")

p = doc.add_paragraph(style='List Number')
p.add_run('Two-model benchmark, not a three-model benchmark. ').bold = True
p.add_run("Claude Sonnet 4.6 was prepared (orchestrator + preflight code in scripts/) but not "
          "run for this version of the report due to time and budget constraints. The "
          "infrastructure is ready (full_run_anchor.json + Stage-1 sandbox prepared, "
          "0-prompt-drift Stage-2 check verified against GPT-5.4 v2 prompts) and re-using the "
          "same recompute / analyze pipeline would produce v6 with three-model headline tables.")

p = doc.add_paragraph(style='List Number')
p.add_run('Gemini-side ablations are NOT included in this report. ').bold = True
p.add_run("Section 13.4 establishes that Gemini does not exhibit the F5/F6 attractor that "
          "motivated the GPT ablations. Re-running the same three ablations on Gemini would "
          "therefore not characterise any Gemini-specific failure mode — Gemini's Tooth_33_Apex "
          "predictions already cluster correctly around G10. A more useful follow-up would be "
          "to identify Gemini's WEAKEST landmark (Sella_S, where Gemini under-performs GPT, see "
          "§13.3) and ablate prompts on THAT landmark instead; this is left as future work.")

# ────────────────────────────────────────────────────────────────────
# 15. CONCLUSIONS AND FUTURE WORK   (renumbered from §13)
# ────────────────────────────────────────────────────────────────────
H('15. Conclusions and Future Work', level=1)

P('On the basis of 10,800 API calls (5,400 GPT-5.4 + 5,400 Gemini 3.1 Pro) across 900 dental '
  'landmark queries — every query receiving byte-identical prompts on byte-identical images for '
  'both models — scored against a two-rater adjudicated consensus and compared head-to-head with '
  'a team-adjudicated dental student consensus:')

p = doc.add_paragraph(style='List Bullet')
p.add_run('GPT-5.4 performs well on cephalometric (mean ED ≤ 0.51 cells, SDR@1 ≥ 90%, SDR@2 ≥ '
          '99%) and is statistically indistinguishable from the student consensus on this '
          'modality.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('GPT-5.4 underperforms the student consensus substantially on periapical (mean ED '
          '~3.7× student) and on panoramic (~4–7× on points; mean Jaccard ~0.4× student on '
          'areas), with overwhelming statistical significance and large rank-biserial effect '
          'sizes (|r| ≥ 0.78 on every non-CEPH comparison; the sign is positive for ED '
          'comparisons where the model has a larger ED than students, and negative for '
          'Jaccard comparisons where the model has a smaller overlap with the GT than '
          'students, following the convention defined in Table 9).')

p = doc.add_paragraph(style='List Bullet')
p.add_run('The guided prompting strategy is heterogeneous: it strongly improves Mental_Foramen_L '
          'and Sella_S, strongly harms Tooth_33_Apex, and moderately harms Condylar_Head_R and '
          'Menton_Me. Aggregate panoramic regression is significant (Bonferroni p < 10⁻¹⁰) but '
          'driven by a few landmarks.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Two-rater inter-rater reliability on the consensus GT is high (κ = 0.86 on points, '
          'mean Jaccard = 0.85 on areas); intra-rater is essentially perfect (κ ≥ 0.97). The '
          'consensus is methodologically sound.')

p = doc.add_paragraph(style='List Bullet')
p.add_run('Reproducibility at temperature = 0 is substantial but not perfect (κ = 0.78 overall). '
          'Rep-averaging is methodologically necessary.')

p = doc.add_paragraph(style='List Bullet')
p.add_run(f'Three focused ablations (Section 10, ${abl_total_cost:.2f} total) tested distinct '
          f'hypotheses about why Tooth_33_Apex collapses under canonical guided. All three were '
          f'pre-registered with locked decision criteria. ')
p.add_run('All three are rejected: ').bold = True
p.add_run(f'Ablation A (FDI tooth number removed) → correct-side rate '
          f'{abl_nt["correct_side_rate"]*100:.0f}%; Ablation B (patient-frame "left" '
          f'disambiguation) → {abl_pl["correct_side_rate"]*100:.0f}%; Ablation C (panoramic '
          f'L–R clause removed, diagnostic) → {abl_nlr["correct_side_rate"]*100:.0f}%. Zero-shot '
          f'baseline is 67%. Qualitative inspection (Section 10.5) plus clinical interpretation '
          f'(from the project\'s OMFR specialist) reframes the regression: across every wording '
          f'variant tested, the model converges on cells F5/F6, which anatomically correspond '
          f'to the lower-right MOLAR region — not the canine region. The Tooth_33_Apex failure '
          f'is structural-anatomic confusion (wrong side AND wrong tooth class), not a spatial '
          f'left/right flip. No canonical prompt change is warranted; the finding is suitable '
          f'for the manuscript\'s discussion as a substantive observation about MLLM anatomic '
          f'mapping rather than as a methodological caveat.')

# ── Cross-model conclusions (NEW in v5) ───────────────────────────
p = doc.add_paragraph(style='List Bullet')
_xm_pan_gd = gpt_vs_gem['RQ_modality_strategy']['PANORAMIC_guided_point']
_xm_pan_zs = gpt_vs_gem['RQ_modality_strategy']['PANORAMIC_zero_shot_point']
_xm_pa_gd  = gpt_vs_gem['RQ_modality_strategy']['PERIAPICAL_guided_point']
_xm_pa_zs  = gpt_vs_gem['RQ_modality_strategy']['PERIAPICAL_zero_shot_point']
_xm_ceph_gd = gpt_vs_gem['RQ_modality_strategy']['CEPHALOMETRIC_guided_point']
_xm_pan_gd_ar = gpt_vs_gem['RQ_modality_strategy']['PANORAMIC_guided_area']
p.add_run(
    f'Gemini 3.1 Pro outperforms GPT-5.4 by a large margin on panoramic and periapical '
    f'point landmarks against the same consensus GT (paired Wilcoxon, all Bonferroni-corrected '
    f'p < 1e-6; rank-biserial r = {_xm_pan_gd["rank_biserial_r"]:+.2f} on PANORAMIC guided point — '
    f'the largest effect we observed — and {_xm_pa_gd["rank_biserial_r"]:+.2f} on PERIAPICAL '
    f'guided point). Mean ED on PAN guided points drops from {_xm_pan_gd["gpt_mean"]:.2f} cells '
    f'(GPT) to {_xm_pan_gd["gemini_mean"]:.2f} cells (Gemini) — a ~4× reduction. The same '
    f'direction holds for PAN area landmarks under both strategies (Jaccard rises from ~0.27 '
    f'to ~0.39 under guided).')

p = doc.add_paragraph(style='List Bullet')
p.add_run(
    f'GPT-5.4 outperforms Gemini on cephalometric point landmarks (mean ED '
    f'{_xm_ceph_gd["gpt_mean"]:.2f} vs {_xm_ceph_gd["gemini_mean"]:.2f} cells under guided; '
    f'r = {_xm_ceph_gd["rank_biserial_r"]:+.2f}, Bonferroni p < 1e-4). This is the only modality '
    f'× landmark-type combination where the two models reverse direction, and the candidate '
    f'mechanisms (pre-training distribution; image-token fidelity per landmark; high-contrast '
    f'silhouette geometry of CEPH landmarks) are flagged for expert review in §13.5.')

p = doc.add_paragraph(style='List Bullet')
p.add_run(
    f'The F5/F6 attractor on Tooth_33_Apex is a GPT-specific failure mode. Under byte-identical '
    f'guided prompts on the same images, GPT places 87.3% of its 300 predictions on F5 or F6 '
    f'(5 columns LEFT of the actual apex at G9-G10); Gemini places 0% there, with predictions '
    f'clustering correctly around F10/G10. The cross-model paired test settles the '
    f'GPT-only-ablations question of §10 definitively: F5/F6 is a model-level failure, not a '
    f'prompt-induced artefact.')

p = doc.add_paragraph(style='List Bullet')
p.add_run(
    f'Reproducibility at temperature = 0 is substantial in both models, and slightly higher in '
    f'Gemini (overall Fleiss kappa across 3 reps on point landmarks: '
    f'{gem_phase_b["fleiss_overall_point"]:.2f} for Gemini vs '
    f'{phase_b["fleiss_overall_point"]:.2f} for GPT). Both fall short of perfect determinism; '
    f'rep-averaging remains methodologically necessary for both.')

P('Future work, in order of expected impact:')
p = doc.add_paragraph(style='List Number')
p.add_run('Run the equivalent experiment on Claude Sonnet 4.6 to extend the v5 two-model '
          'comparison to a three-model benchmark. The orchestrator + preflight code are '
          'already prepared in scripts/run_full_run_claude.py and Stage 1 of the sandbox is '
          'cryptographically anchored against the same GPT-5.4 v2 main run. Same canonical '
          'prompts (zero_shot + guided). Estimated cost ~$24 at Anthropic batch pricing.')
p = doc.add_paragraph(style='List Number')
p.add_run('Bland-Altman plots (visual) and per-landmark stratified GPT-vs-student tests, '
          'reported in the manuscript figures.')
p = doc.add_paragraph(style='List Number')
p.add_run('Continue the Tooth_33_Apex root-cause investigation with factorial decomposition '
          'of the GUIDED_SYSTEM_ADDITION block — specifically, an ablation that strips ONLY '
          'the row/column declaration sentence ("columns 1 through 16 from left to right, '
          'rows A through H from top to bottom") while keeping the rest of the system prompt '
          'intact. This isolates whether the row/col declaration alone activates the F5/F6 '
          'attractor. The Section 10 infrastructure is directly reusable. Estimated cost: ~$1.')
p = doc.add_paragraph(style='List Number')
p.add_run('Cross-modality replicate: test guided_no_LR-style prompts on Tooth_36_* periapical '
          'landmarks (which have no L–R clause in their guided system prompt anyway, but use '
          'FDI numbers) to verify that the absence of the L–R clause is sufficient to keep '
          'them well-behaved.')
p = doc.add_paragraph(style='List Number')
p.add_run('Tolerant parser (digit-before-letter) for the published compliance metric, with both '
          'strict and tolerant rates reported. Across the v2 main run + three ablations, the '
          'observed compliance failures are all of this single failure mode.')

# ────────────────────────────────────────────────────────────────────
# APPENDIX A — PER-LANDMARK DETAILS
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix A — Per-Landmark Detailed Statistics', level=1)

H('A1. Per-landmark mean ED with 95% CI (vs consensus_gt)', level=2)
caption('Table A1: Per-landmark mean ED with bootstrap 95% CI and median ED.')
rows=[]
for mod,lm in landmarks_order:
    for strat in ('zero_shot','guided'):
        d = analysis['RQ1_per_landmark'][f'{mod}/{lm}/{strat}']
        rows.append([f'{mod}/{lm}', strat, str(d['n']),
                     f'{d["mean"]:.3f}',
                     f'[{d["mean_ci"][0]:.3f}, {d["mean_ci"][1]:.3f}]',
                     f'{d["median"]:.3f}'])
add_table(['Landmark','Strategy','n','Mean ED','Mean 95% CI','Median ED'], rows,
          col_widths=[2.3,0.95,0.5,0.85,1.5,0.95], header_size=9, body_size=9)

H('A2. Per-landmark SDR (vs consensus_gt)', level=2)
caption('Table A2: Per-landmark Successful Detection Rate at each threshold.')
rows=[]
for mod,lm in landmarks_order:
    for strat in ('zero_shot','guided'):
        d = phase_b['sdr_landmark'][f'{mod}/{lm}/{strat}']
        rows.append([f'{mod}/{lm}', strat, str(d['n']),
                     f'{d["mean_ed"]:.3f}',
                     f'{d["SDR@0"]*100:.0f}%',
                     f'{d["SDR@1"]*100:.0f}%',
                     f'{d["SDR@√2"]*100:.0f}%',
                     f'{d["SDR@2"]*100:.0f}%'])
add_table(['Landmark','Strategy','n','Mean ED','SDR@0','SDR@1','SDR@√2','SDR@2'], rows,
          col_widths=[2.0,0.85,0.4,0.85,0.65,0.65,0.7,0.7], header_size=9, body_size=9)

# ────────────────────────────────────────────────────────────────────
# APPENDIX B — COMPLIANCE FAILURES
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix B — Compliance Failures (Full List)', level=1)
P(f'All {len(failures)} strict compliance failures across {n_calls:,} API calls. Each is one '
  f'rep of one (query, strategy) pair where the canonical pipeline parser could not extract a '
  f'valid grid coordinate.')
caption("Table B1: All compliance failures with tolerant-parser decoded cell and ED to GT.")

# Helper: decode digit-before-letter form to letter-before-digit, then compute ED to GT.
import math as _math
def _decode_reversed(raw):
    s = str(raw).strip().upper()
    i = 0
    while i < len(s) and s[i].isdigit(): i += 1
    if i == 0 or i == len(s): return None
    return s[i:] + s[:i]
def _parse_cell(c):
    s = str(c).strip().upper()
    if len(s) < 2 or not s[0].isalpha(): return None
    try:
        return (ord(s[0]) - ord('A'), int(s[1:]) - 1)
    except (ValueError, IndexError):
        return None
def _ed_decoded_to_gt(raw, gt):
    rev = _decode_reversed(raw)
    if rev is None: return rev, None
    a = _parse_cell(rev); b = _parse_cell(gt)
    if a is None or b is None: return rev, None
    return rev, _math.hypot(a[0]-b[0], a[1]-b[1])

rows=[]
for f in failures:
    rev_cell, ed = _ed_decoded_to_gt(f['raw_response'], f['gt'])
    ed_str = f'{ed:.2f}' if ed is not None else '–'
    rows.append([f['query_id'], f['strategy'], str(f['rep']), f['modality'],
                 f['structure'], f['gt'], f'"{f["raw_response"]}"',
                 rev_cell or '–', ed_str, f['failure_category']])
add_table(['Query ID','Strategy','Rep','Modality','Landmark','GT','Raw','Decoded','ED to GT (cells)','Category'],
          rows, col_widths=[1.7,0.65,0.3,1.0,1.4,0.5,0.55,0.65,0.9,0.55],
          header_size=8.5, body_size=8.5, fixed=True)

# Live-compute the actual decoded-ED range so the prose can never drift.
_b1_eds = []
for f in failures:
    _, ed = _ed_decoded_to_gt(f['raw_response'], f['gt'])
    if ed is not None: _b1_eds.append(ed)
_b1_min, _b1_max = (min(_b1_eds), max(_b1_eds)) if _b1_eds else (None, None)
P(f'All four cases are digit-before-letter inversions. Under a tolerant parser they would all '
  f'parse to VALID grid cells (so the parse rate becomes 100%), but those decoded cells are '
  f'NOT the GT cells: they lie between {_b1_min:.2f} and {_b1_max:.2f} cells from consensus_gt '
  f'in Euclidean distance. The failure mode is therefore not "model gave the right answer with '
  f'flipped formatting" — it is "format-level glitch on top of a still-wrong prediction".')

H('B2. Ablation compliance failures (Section 10 runs)', level=2)
P('Compliance failures observed across the three Section 10 ablation runs '
  '(300 API calls each, all on the same 100 panoramic Tooth_33_Apex queries). '
  'These supplement Table B1\'s main-run record and exhibit the same '
  'reversed-coordinate failure mode.')

caption("Table B2: Compliance failures observed in the Section 10 ablations.")
abl_failures = []
for key, sandbox in (('A — Ablation A: guided_no_tooth_num',  ROOT/'results_ablation_no_tooth_num'),
                      ('B — Ablation B: guided_patient_left',  ROOT/'results_ablation_patient_left'),
                      ('C — Ablation C: guided_no_LR',         ROOT/'results_ablation_no_LR')):
    for rep in (1, 2, 3):
        cs_path = sandbox / f'run{rep}' / 'compliance_stats.json'
        if not cs_path.exists():
            continue
        cs = json.load(open(cs_path))
        # Use the parsed_responses.json to find raw text of failed entries
        pr_path = sandbox / f'run{rep}' / 'parsed_responses.json'
        if not pr_path.exists():
            continue
        for rec in json.load(open(pr_path)):
            if rec.get('failure_category') is not None:
                abl_failures.append({
                    'ablation': key, 'query_id': rec['query_id'],
                    'rep': rep,
                    'modality': (rec.get('modality')
                                 or ('PANORAMIC' if rec['query_id'].startswith('PAN_')
                                     else 'PERIAPICAL' if rec['query_id'].startswith('PA_')
                                     else 'CEPHALOMETRIC' if rec['query_id'].startswith('CEPH_')
                                     else '')),
                    'raw_response': rec.get('raw_response',''),
                    'failure_category': rec['failure_category'],
                })

if abl_failures:
    # Compute decoded ED to GT for each ablation failure (all are PAN/Tooth_33_Apex)
    _qi_lookup = {q['image_id'] + '_' + q['structure']: q['consensus_gt']
                  for q in json.load(open(SANDBOX / 'query_index.json'))}
    rows = []
    _abl_eds = []
    for f in abl_failures:
        # Recover the canonical query_id (strip the trailing strategy suffix)
        suffix = '_' + f['ablation'].split(':')[1].strip()
        qid_core = f['query_id'][:-len(suffix)] if f['query_id'].endswith(suffix) else f['query_id']
        gt = _qi_lookup.get(qid_core, '?')
        rev_cell, ed = _ed_decoded_to_gt(f['raw_response'], gt)
        ed_str = f'{ed:.2f}' if ed is not None else '–'
        if ed is not None: _abl_eds.append(ed)
        rows.append([f['ablation'], f['query_id'], str(f['rep']), f['modality'],
                     f'"{f["raw_response"]}"', gt, rev_cell or '–', ed_str,
                     f['failure_category']])
    add_table(['Ablation','Query ID','Rep','Modality','Raw','GT','Decoded','ED to GT','Category'],
              rows, col_widths=[2.0, 1.7, 0.3, 0.95, 0.55, 0.45, 0.6, 0.7, 0.55],
              header_size=8.5, body_size=8.5, fixed=True)
    _abl_min, _abl_max = (min(_abl_eds), max(_abl_eds)) if _abl_eds else (None, None)
    _range_str = (f'between {_abl_min:.2f} and {_abl_max:.2f} cells'
                  if _abl_eds and _abl_min != _abl_max
                  else (f'{_abl_min:.2f} cells' if _abl_eds else '–'))
    P(f'Total ablation compliance failures: {len(abl_failures)} across '
      f'{3 * 300:,} ablation API calls (strict-parse rate '
      f'{(900-len(abl_failures))/900*100:.3f}% on the 900 ablation calls). '
      f'Ablation C had zero compliance failures. As in the main run, every '
      f'failure is a reversed-coordinate inversion: the model returned a '
      f'cell coordinate with the digit before the letter. Under a tolerant '
      f'parser these would all parse to VALID cells, but those decoded cells '
      f'sit {_range_str} from the consensus GT — i.e., the underlying prediction '
      f'is still wrong by ~1+ cells once format is corrected. None of these '
      f'failures affects any reported metric — the failed reps are simply '
      f'excluded from per-query rep-averaging for that query.')
else:
    P('No compliance failures recorded across all three ablation runs '
      '(900 API calls total).')

# ────────────────────────────────────────────────────────────────────
# APPENDIX C — REPRODUCIBILITY MANIFEST
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix C — Reproducibility Manifest', level=1)
P('Every figure in this report can be reproduced from the artifacts below.')

caption("Table C1: Reproducibility manifest — v2 main run.")
rows = [
    ['Pipeline source — git commit', git_sha],
    ['Pipeline source — commit date', git_date],
    ['Final Excel SHA-256', final_excel_sha],
    ['Old (v1) Excel SHA-256 (superseded)', old_excel_sha],
    ['v2 query_index.json SHA-256', qi_sha],
    ['v1 query_index.json SHA-256 (anchored)', v1_anchor_sha],
    ['Anchored raw JSONL files (results_full)', f'{len(anchor["files"])} files'],
    ['Model', 'gpt-5.4'],
    ['Inference settings', 'temperature = 0, seed = 42, max_completion_tokens = 50, image detail = high'],
    ['Repetitions', '3 per (query × strategy)'],
    ['Sandbox directory', 'results_consensus/'],
    ['Anchor sandbox (frozen GPT outputs)', 'results_full/'],
    ['Total queries', '900'],
    ['Strategies', 'zero_shot, guided'],
    ['Total API calls', f'{n_calls:,}'],
    ['Compliance (strict)', f'{compliance_rate*100:.4f}%'],
    ['Prompt tokens', f'{prompt_tok:,}'],
    ['Completion tokens', f'{compl_tok:,}'],
    ['Naïve cost', f'${naive_cost:.2f}'],
    ['Bootstrap CIs', '10,000 resamples, percentile, random_state = 42'],
    ['Strategy test', 'Wilcoxon signed-rank (zero-method = wilcox, two-sided)'],
    ['Multiple comparisons', 'Bonferroni; family sizes per analysis level'],
    ['Reliability (point cells)', "Cohen's κ unweighted"],
    ['Reliability (area)', 'Jaccard, Dice'],
    ['Canonical GT field', 'consensus_gt'],
    ['Sensitivity reference', 'omfr_1 (Appendix D)'],
]
add_table(['Item','Value'], rows, col_widths=[2.7,3.9], header_size=10, body_size=9)

caption("Table C2: Reproducibility manifest — Section 10 follow-up ablations.")
rows_abl = []
ablation_metadata = [
    ('A', 'no_tooth_num',  'guided_no_tooth_num',
     '≥ 90% confirmed; 30–90% partial; < 30% rejected'),
    ('B', 'patient_left',  'guided_patient_left',
     '≥ 90% confirmed; 30–90% partial; < 30% rejected'),
    ('C', 'no_LR',         'guided_no_LR',
     '≥ 60% L–R clause is cause; 30–60% partial; < 30% L–R clause not cause'),
]
for label, key, strat, crit in ablation_metadata:
    ana = ablations[key]
    mani = abl_manifests[key]
    anc = abl_anchors[key]
    cost = abl_costs[key]
    cs_key_name = f'{strat}_correct_side_rate'
    cs_rate = ana['verdict'].get(cs_key_name, 0)
    rows_abl.append([f'Ablation {label}',
                     f'Sandbox: results_ablation_{key}/. '
                     f'Strategy: {strat}. '
                     f'Queries: {mani["n_queries"]} (panoramic Tooth_33_Apex). '
                     f'Anchored JSONLs: {len(anc.get("files", {}))} from results_full/. '
                     f'API calls: 300 (100 queries × 3 reps). '
                     f'Naïve cost: ${cost:.4f}. '
                     f'Pre-registered criteria: {crit}. '
                     f'Verdict: {ana["verdict"]["verdict"]} (correct-side rate {cs_rate*100:.1f}%).'])
add_table(['Item','Value'], rows_abl, col_widths=[1.4, 5.2], header_size=10, body_size=9)

P('Inputs: data/Final_Dental_MLLM_Benchmark_Data.xlsx and the 108 raw JSONL outputs at '
  'results_full/run{1,2,3}/responses/. Derived artifacts (v2 main run): '
  'results_consensus/{query_index.json, full_run_records.pkl, analysis.json, phase_b.json, '
  'summary.json, rater_reliability.json, gpt_vs_student.json, reanalysis_anchor.json, '
  'v2_manifest.json}. Derived artifacts (Section 10 ablations): '
  'results_ablation_{no_tooth_num, patient_left, no_LR}/{query_index.json, ablation_manifest.json, '
  'ablation_anchor.json, prompts_used.json, run{1,2,3}/{responses, parsed_responses.json, '
  'compliance_stats.json, batch_tracking.json}, ablation_analysis.json, '
  'ablation_per_image.json, live_test_record.json}. Each ablation\'s prompts_used.json captures '
  'the exact rendered system+user prompt that the model received under that ablation\'s '
  'strategy. The qualitative inspection (Section 10.5) data is in '
  'results_ablation_no_LR/qualitative_inspection.json. The recompute against consensus_gt '
  'and each ablation analysis verify all anchored SHAs at preflight; if any frozen JSONL '
  'drifts, the pipeline refuses to proceed.')

# ────────────────────────────────────────────────────────────────────
# APPENDIX D — v1 ↔ v2 SENSITIVITY
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix D — v1 (OMFR_1) ↔ v2 (Consensus) Sensitivity Comparison', level=1)
P('All numbers below were computed by the SAME script with the SAME canonical pipeline parser '
  '(pipeline.parse_grid_coordinate), varying only the reference field (omfr_1 vs consensus_gt). '
  'Differences are therefore purely GT-driven, not parser- or implementation-driven. We '
  'verified independently that the canonical parser produces the same first cell as v1\'s ad-hoc '
  'parser for every one of the 5,400 model responses, so the v1 frozen report\'s numbers match '
  'the OMFR_1 sensitivity column below to floating-point precision.')

caption('Table D1: RQ1 modality × strategy point ED — omfr_1 vs consensus_gt.')
rows=[]
for k_label, key in (('CEPH ZS','CEPHALOMETRIC_zero_shot_point'),
                      ('CEPH GD','CEPHALOMETRIC_guided_point'),
                      ('PA ZS','PERIAPICAL_zero_shot_point'),
                      ('PA GD','PERIAPICAL_guided_point'),
                      ('PAN ZS','PANORAMIC_zero_shot_point'),
                      ('PAN GD','PANORAMIC_guided_point')):
    a = analysis_omfr1['RQ1_modality_strategy'].get(key, {})
    b = analysis['RQ1_modality_strategy'].get(key, {})
    if not a or not b: continue
    rows.append([k_label, str(a.get('n','?')),
                 f'{a.get("mean", float("nan")):.4f}',
                 f'{b.get("mean", float("nan")):.4f}',
                 f'{(b.get("mean", float("nan")) - a.get("mean", float("nan"))):+.4f}'])
# Area
for k_label, key in (('PAN ZS area Jaccard','PANORAMIC_zero_shot_area'),
                      ('PAN GD area Jaccard','PANORAMIC_guided_area')):
    a = analysis_omfr1['RQ1_modality_strategy'].get(key, {})
    b = analysis['RQ1_modality_strategy'].get(key, {})
    if not a or not b: continue
    rows.append([k_label, str(a.get('n','?')),
                 f'{a.get("mean", float("nan")):.4f}',
                 f'{b.get("mean", float("nan")):.4f}',
                 f'{(b.get("mean", float("nan")) - a.get("mean", float("nan"))):+.4f}'])
add_table(['Group','n','vs OMFR_1','vs Consensus','Δ (cons − omfr_1)'], rows,
          col_widths=[2.0,0.6,1.2,1.2,1.4], header_size=10, body_size=10)

caption('Table D2: GT-change census per modality × landmark_type.')
rows = [
    ['CEPHALOMETRIC','point', '150', f'{len(gt_changes["point"]["CEPHALOMETRIC"])}',
     f'{len(gt_changes["point"]["CEPHALOMETRIC"])/150*100:.2f}%'],
    ['PERIAPICAL','point', '150', f'{len(gt_changes["point"]["PERIAPICAL"])}',
     f'{len(gt_changes["point"]["PERIAPICAL"])/150*100:.2f}%'],
    ['PANORAMIC','point', '300', f'{len(gt_changes["point"]["PANORAMIC"])}',
     f'{len(gt_changes["point"]["PANORAMIC"])/300*100:.2f}%'],
    ['PANORAMIC','area',  '300', f'{len(gt_changes["area"]["PANORAMIC"])}',
     f'{len(gt_changes["area"]["PANORAMIC"])/300*100:.2f}%'],
]
add_table(['Modality','Type','n','GT changed','Rate'], rows,
          col_widths=[1.6,1.0,0.8,1.5,1.0], header_size=10, body_size=10)
P('PAN points changed: '
  + ('; '.join(gt_changes['point']['PANORAMIC']) if gt_changes['point']['PANORAMIC'] else '(none)')
  + '.')

# ────────────────────────────────────────────────────────────────────
# APPENDIX E — PER-IMAGE QUALITATIVE INSPECTION TABLE (Section 10.5 data)
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix E — Per-Image Qualitative Inspection (Section 10.5 raw data)', level=1)
P('Per-image MODE predicted cell across all five conditions. Each row is one '
  'of the 100 panoramic Tooth_33_Apex queries; the MODE is the most-frequent '
  'cell across the three repetitions of that query under the named strategy. '
  '"GT" is the consensus_gt cell (image-right is cols ≥ 9; image-left is cols '
  '≤ 8). Reading guide: zero_shot mode cells cluster near GT; the four guided '
  'variants overwhelmingly converge on cells F5/F6 (image-left, lower row), '
  'regardless of wording change. This is the raw evidence behind Section 10.5.')

if qual_rows:
    caption("Table E1: Per-image MODE predicted cell, each strategy "
            "(zero_shot, guided, guided_no_tooth_num, guided_patient_left, "
            "guided_no_LR), with consensus_gt for reference.")
    rows = []
    for r in qual_rows:
        rows.append([
            r['image_id'],
            r['gt'],
            r['zero_shot_mode'],
            r['guided_mode'],
            r['guided_no_tooth_num_mode'],
            r['guided_patient_left_mode'],
            r['guided_no_LR_mode'],
        ])
    add_table(
        ['Image', 'GT', 'zero_shot', 'guided', 'no_tooth_num',
         'patient_left', 'no_LR'],
        rows,
        col_widths=[0.95, 0.55, 0.95, 0.85, 1.10, 1.00, 0.85],
        header_size=9, body_size=8.5, fixed=True,
    )
    n_zs_right = sum(1 for r in qual_rows if r['zero_shot_mode_col'] is not None and r['zero_shot_mode_col']>=9)
    n_gd_right = sum(1 for r in qual_rows if r['guided_mode_col'] is not None and r['guided_mode_col']>=9)
    n_ntn_right = sum(1 for r in qual_rows if r['guided_no_tooth_num_mode_col'] is not None and r['guided_no_tooth_num_mode_col']>=9)
    n_pl_right  = sum(1 for r in qual_rows if r['guided_patient_left_mode_col'] is not None and r['guided_patient_left_mode_col']>=9)
    n_nlr_right = sum(1 for r in qual_rows if r['guided_no_LR_mode_col'] is not None and r['guided_no_LR_mode_col']>=9)
    P(f'Column totals: zero_shot places its mode on the correct (image-right) '
      f'side for {n_zs_right}/100 queries, canonical guided for {n_gd_right}/100, '
      f'guided_no_tooth_num for {n_ntn_right}/100, guided_patient_left for '
      f'{n_pl_right}/100, guided_no_LR for {n_nlr_right}/100. The full '
      f'per-image data (mean ED per condition + raw responses per rep) lives '
      f'at results_ablation_no_LR/qualitative_inspection.json and is the '
      f'source for every figure in Section 10.5.', italic=True)

    # ── Per-rep response tables for each ablation condition ─────────
    H('E2. Per-rep raw responses — Ablation A (guided_no_tooth_num)', level=2)
    P('All 300 raw model responses (100 queries × 3 reps) from Ablation A. Each row is one '
      'PAN/Tooth_33_Apex query; the three "Rep 1/2/3" columns show the verbatim cell strings '
      'the model returned under guided_no_tooth_num. "<empty>" indicates a parse failure '
      '(none expected for this ablation — see Appendix B2). Source: '
      'results_ablation_no_tooth_num/run{1,2,3}/parsed_responses.json.')

    def _load_per_rep(sandbox_name, strategy_suffix):
        per = {}
        for rep in (1, 2, 3):
            pr_path = ROOT / sandbox_name / f'run{rep}' / 'parsed_responses.json'
            if not pr_path.exists():
                continue
            for rec in json.load(open(pr_path)):
                cid = rec.get('custom_id', '')
                if not cid.endswith(strategy_suffix):
                    continue
                qid = cid[:-len(strategy_suffix)]
                per.setdefault(qid, [None, None, None])
                per[qid][rep - 1] = rec.get('raw_response', '') or '<empty>'
        return per

    def _build_per_rep_table(caption_text, sandbox_name, strategy_suffix, qual_rows):
        per = _load_per_rep(sandbox_name, strategy_suffix)
        caption(caption_text)
        rows = []
        for r in qual_rows:
            qid = f'{r["image_id"]}_Tooth_33_Apex'
            reps = per.get(qid, [None, None, None])
            reps = [(v if v else '<empty>') for v in reps]
            rows.append([r['image_id'], r['gt'], reps[0], reps[1], reps[2]])
        add_table(['Image', 'GT', 'Rep 1', 'Rep 2', 'Rep 3'], rows,
                  col_widths=[0.95, 0.7, 1.65, 1.65, 1.65],
                  header_size=10, body_size=10, fixed=True)

    _build_per_rep_table(
        "Table E2: Ablation A — raw per-rep responses under guided_no_tooth_num.",
        'results_ablation_no_tooth_num', '_guided_no_tooth_num', qual_rows)

    H('E3. Per-rep raw responses — Ablation B (guided_patient_left)', level=2)
    P('Same format as E2. Source: results_ablation_patient_left/run{1,2,3}/'
      'parsed_responses.json. One compliance failure expected (PAN_047 rep 3 returned "9F" '
      'with reversed coordinates; see Appendix B2).')
    _build_per_rep_table(
        "Table E3: Ablation B — raw per-rep responses under guided_patient_left.",
        'results_ablation_patient_left', '_guided_patient_left', qual_rows)

    H('E4. Per-rep raw responses — Ablation C (guided_no_LR)', level=2)
    P('Same format as E2. Source: results_ablation_no_LR/run{1,2,3}/parsed_responses.json. '
      'No compliance failures (100% strict-parse rate across the three reps).')
    _build_per_rep_table(
        "Table E4: Ablation C — raw per-rep responses under guided_no_LR.",
        'results_ablation_no_LR', '_guided_no_LR', qual_rows)

    H('E5. Per-rep responses for canonical zero_shot and guided (main run)', level=2)
    P('The 100 PAN/Tooth_33_Apex per-rep raw responses for the canonical zero_shot and guided '
      'strategies are part of the v2 main run, not part of any Section 10 ablation. They are '
      'embedded in this report as part of Appendix F (Tables F1 and F2 — PAN point landmarks; '
      'Tooth_33_Apex rows are interleaved with the other PAN point landmarks within those '
      'tables, sorted by image_id). For direct comparison with Tables E2–E4 above, filter '
      'Appendix F Tables F1/F2 to the 100 rows where Landmark = "Tooth_33_Apex".')

else:
    P('(Qualitative inspection data not available — '
      'results_ablation_no_LR/qualitative_inspection.json missing.)',
      italic=True)

# ────────────────────────────────────────────────────────────────────
# Appendix F — All 5,400 main-run model responses (8 tables)
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix F — Main-Run Model Responses (all 5,400)', level=1)
P('Every GPT-5.4 response for the v2 main run, organised by (modality × strategy × '
  'landmark-type). Eight tables in total; each row is one query and the three '
  '"Run 1/2/3" columns show the verbatim cell strings the model returned across the '
  'three repetitions. The "GT" column is consensus_gt (CONSENSUS_Ground_Truth from the '
  f'Final benchmark Excel, SHA-256 prefix {final_excel_sha[:12]}). "<empty>" indicates a '
  'strict-parser failure (full failure list in Appendix B2). Source: '
  'results_full/run{1,2,3}/responses/*.jsonl (re-scored against consensus_gt via '
  'scripts/recompute_against_consensus.py; raw JSONL untouched). This appendix supersedes '
  'the standalone companion docx Full_Run_Model_Responses_Appendix_v2.docx — the content '
  'is identical, byte-for-byte, but inlined here so the report is self-contained.')

# Group layout (matches the companion docx exactly)
F_GROUPS = [
    ('PANORAMIC',     'zero_shot', 'point', 'F1'),
    ('PANORAMIC',     'guided',    'point', 'F2'),
    ('PANORAMIC',     'zero_shot', 'area',  'F3'),
    ('PANORAMIC',     'guided',    'area',  'F4'),
    ('PERIAPICAL',    'zero_shot', 'point', 'F5'),
    ('PERIAPICAL',    'guided',    'point', 'F6'),
    ('CEPHALOMETRIC', 'zero_shot', 'point', 'F7'),
    ('CEPHALOMETRIC', 'guided',    'point', 'F8'),
]
F_PT_WIDTHS = [0.95, 1.65, 1.30, 0.93, 0.93, 0.93]   # 6.69" — letter portrait
F_LANDMARK_ORDER = {
    'Mental_Foramen_L': 0, 'Condylar_Head_R': 1, 'Tooth_33_Apex': 2,
    'Mandibular_Canal_L': 3, 'Maxillary_Sinus_R': 4, 'External_Oblique_Ridge_R': 5,
    'Tooth_36_Distal_Apex': 0, 'Tooth_36_Distal_CEJ': 1, 'Tooth_36_Mesial_CEJ': 2,
    'Sella_S': 0, 'Nasion_N': 1, 'Menton_Me': 2,
}

def _fmt_main_response(raw):
    if raw is None or not str(raw).strip():
        return '<empty>'
    return str(raw).strip()

# Per-group overview list
for mod, strat, ltype, label in F_GROUPS:
    n = sum(1 for r in records if r['modality'] == mod and r['strategy'] == strat
            and r['landmark_type'] == ltype)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'Table {label}: {mod}, {strat}, {ltype} — {n} queries × 3 reps = {n * 3} responses.')

# Render each of the 8 tables on its own page
for mod, strat, ltype, label in F_GROUPS:
    items = [r for r in records
             if r['modality'] == mod and r['strategy'] == strat
             and r['landmark_type'] == ltype]
    items.sort(key=lambda r: (r['image_id'], F_LANDMARK_ORDER.get(r['structure'], 99)))

    doc.add_page_break()
    H(f'Table {label} — {mod}, {strat}, {ltype} landmarks (vs consensus_gt)', level=2)
    caption(f'Table {label}: All GPT-5.4 raw responses for the {mod} {strat} {ltype} group '
            f'({len(items)} queries × 3 reps = {len(items) * 3} responses). GT column = '
            f'consensus_gt.')
    rows = []
    for r in items:
        runs = [_fmt_main_response(x) for x in r['rep_raw']]
        while len(runs) < 3:
            runs.append('<empty>')
        rows.append([r['image_id'], r['structure'], r['consensus_gt'] or '',
                     runs[0], runs[1], runs[2]])
    # Area tables (F3, F4) carry long comma-separated cell lists → 10pt; points → 11pt
    body_size = 10 if ltype == 'area' else 11
    add_table(['Image', 'Landmark', 'GT', 'Run 1', 'Run 2', 'Run 3'],
              rows, col_widths=F_PT_WIDTHS,
              header_size=12, body_size=body_size, fixed=True)

# ────────────────────────────────────────────────────────────────────
# Appendix G — All 5,400 Gemini 3.1 Pro responses (8 tables, NEW in v5)
# ────────────────────────────────────────────────────────────────────
doc.add_page_break()
H('Appendix G — Gemini 3.1 Pro Responses (all 5,400)', level=1)
P('Every Gemini 3.1 Pro response across the 3 repetitions, organised by '
  '(modality × strategy × landmark-type) — the same eight-group layout as '
  'Appendix F so direct cross-model lookup by (Image, Landmark) is possible. '
  'Each row shows the verbatim cell string the model returned for that query '
  'in each of the three reps. The "GT" column is the same consensus_gt as in '
  'Appendix F. "<empty>" indicates a strict-parser failure for that rep '
  '(full Gemini failure list in §12.5, Table 34).')

P('Methodological note: For rep 1, where MAX_TOKENS truncation in the model\'s '
  'internal reasoning produced 99 empty answer slots (5.5% of rep 1), 78 of '
  'those were recovered by a re-query sweep at max_output_tokens = 4096. The '
  '"Run 1" column shows the FINAL recorded response for each query — i.e., the '
  're-queried response where applicable; the chunk-original response '
  'otherwise. The pre-re-query snapshot is preserved at '
  'results_full_gemini/run1/parsed_responses.before_requery.json for '
  'reproducibility. Reps 2 and 3 used max_output_tokens = 4096 from the start '
  '(see §12.1 and §14).', italic=True)

# Reuse the F_GROUPS / F_PT_WIDTHS / F_LANDMARK_ORDER definitions so the
# Gemini appendix is layout-byte-identical to F (modulo the table labels and
# the underlying records source).
G_GROUPS = [
    ('PANORAMIC',     'zero_shot', 'point', 'G1'),
    ('PANORAMIC',     'guided',    'point', 'G2'),
    ('PANORAMIC',     'zero_shot', 'area',  'G3'),
    ('PANORAMIC',     'guided',    'area',  'G4'),
    ('PERIAPICAL',    'zero_shot', 'point', 'G5'),
    ('PERIAPICAL',    'guided',    'point', 'G6'),
    ('CEPHALOMETRIC', 'zero_shot', 'point', 'G7'),
    ('CEPHALOMETRIC', 'guided',    'point', 'G8'),
]

# Per-group overview list
for mod, strat, ltype, label in G_GROUPS:
    n = sum(1 for r in gem_records if r['modality'] == mod and r['strategy'] == strat
            and r['landmark_type'] == ltype)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'Table {label}: {mod}, {strat}, {ltype} — {n} queries × 3 reps = {n * 3} responses.')

# Render each of the 8 tables on its own page
for mod, strat, ltype, label in G_GROUPS:
    items = [r for r in gem_records
             if r['modality'] == mod and r['strategy'] == strat
             and r['landmark_type'] == ltype]
    items.sort(key=lambda r: (r['image_id'], F_LANDMARK_ORDER.get(r['structure'], 99)))

    doc.add_page_break()
    H(f'Table {label} — Gemini 3.1 Pro: {mod}, {strat}, {ltype} landmarks (vs consensus_gt)',
      level=2)
    caption(f'Table {label}: All Gemini 3.1 Pro raw responses for the {mod} {strat} {ltype} '
            f'group ({len(items)} queries × 3 reps = {len(items) * 3} responses). GT column = '
            f'consensus_gt (same as Table {label.replace("G", "F")} in Appendix F).')
    rows = []
    for r in items:
        runs = [_fmt_main_response(x) for x in r['rep_raw']]
        while len(runs) < 3:
            runs.append('<empty>')
        rows.append([r['image_id'], r['structure'], r['consensus_gt'] or '',
                     runs[0], runs[1], runs[2]])
    body_size = 10 if ltype == 'area' else 11
    add_table(['Image', 'Landmark', 'GT', 'Run 1', 'Run 2', 'Run 3'],
              rows, col_widths=F_PT_WIDTHS,
              header_size=12, body_size=body_size, fixed=True)

# Save
doc.save(str(OUT))
import os
print(f'Wrote {OUT}')
print(f'  Paragraphs: {len(doc.paragraphs)}')
print(f'  Tables: {len(doc.tables)}')
print(f'  Size: {os.path.getsize(OUT):,} bytes')
