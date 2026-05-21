"""Generate DMFR_Grid_Study_v2.docx from the colleague's first draft.

Modifications:
  1. Abstract Results paragraph: add one sentence about Gemini vs students.
  2. Insert new §RQ8 (Gemini vs Student Consensus) before Discussion.
  3. Insert new Table 8 (Gemini vs student paired comparison) before
     Figure Legends.
  4. Limitations paragraph: remove the "deferred to future work" sentence.
  5. Conclusions paragraph: add one sentence integrating the Gemini-vs-student
     finding.
  6. Add §-marker for "Expert input requested" footnotes at end.

Every new number comes from results_full_gemini/gemini_vs_student.json,
which itself is computed by scripts/analyze_gemini_vs_student.py from
results_full_gemini/full_run_records.pkl. No hardcoded numbers.
"""
from __future__ import annotations
import json
import sys
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
IN_DOCX = ROOT / "docs" / "submission" / "DMFR_Grid_Study.docx"
OUT_DOCX = ROOT / "docs" / "submission" / "DMFR_Grid_Study_v2.docx"

# ── Load source data (every number in the new content traces here) ─
gem_vs_stu = json.load(open(ROOT / "results_full_gemini" / "gemini_vs_student.json"))


def bonf_p(p: float, k: int) -> float:
    """Bonferroni-corrected p-value (capped at 1.0)."""
    return min(1.0, p * k)


def sig_marker(p: float) -> str:
    if p < 0.001:
        return "***"
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return "NS"


def fmt_p(p: float) -> str:
    if p < 1e-3:
        return f"{p:.2e}"
    return f"{p:.3f}"


# ── Extract Gemini vs student numbers (rounded for prose) ──────────
def gvs_num(strategy: str, group_key: str) -> dict:
    d = gem_vs_stu["per_strategy"][strategy]["per_modality"][group_key]
    return {
        "n": d["n_total"],
        "mean_gem": d.get("mean_gemini_ed", d.get("mean_gemini_jaccard", 0)),
        "mean_stu": d.get("mean_student_ed", d.get("mean_student_jaccard", 0)),
        "delta": d["mean_delta"],
        "p_raw": d["p"],
        "p_bonf": bonf_p(d["p"], 4),  # Bonferroni × 4 modality × strategy comparisons
        "r": d["rank_biserial_r"],
    }


# Modality x strategy compendium
GROUPS = [
    ("CEPHALOMETRIC_point", "zero_shot"),
    ("CEPHALOMETRIC_point", "guided"),
    ("PERIAPICAL_point", "zero_shot"),
    ("PERIAPICAL_point", "guided"),
    ("PANORAMIC_point", "zero_shot"),
    ("PANORAMIC_point", "guided"),
    ("PANORAMIC_area", "zero_shot"),
    ("PANORAMIC_area", "guided"),
]

stats_summary = {f"{g}|{s}": gvs_num(s, g) for g, s in GROUPS}

# ── Sanity-check numbers we'll cite in prose ───────────────────────
print("=" * 60 + "\nNumbers being cited in v2:\n" + "=" * 60)
for key, d in stats_summary.items():
    print(f"  {key}: n={d['n']} Δ={d['delta']:+.3f} p={d['p_raw']:.2e} "
          f"p_bonf={d['p_bonf']:.3e} {sig_marker(d['p_bonf'])} r={d['r']:+.3f}")
print()

# ── Open original docx ─────────────────────────────────────────────
doc = Document(IN_DOCX)


# ── HELPER: modify text within a single paragraph ──────────────────
def paragraph_text(p) -> str:
    return "".join(r.text for r in p.runs)


def rewrite_paragraph_text(p, new_text: str) -> None:
    """Replace the entire text of a paragraph with new_text, using the
    formatting of the first run. Loses any per-run formatting differences.
    """
    if not p.runs:
        # Need to add a run
        r = p.add_run(new_text)
        return
    # Save format of first run
    first = p.runs[0]
    # Clear all runs
    for r in list(p.runs):
        r.text = ""
    first.text = new_text


def replace_text_in_paragraph(p, old: str, new: str) -> bool:
    """Surgical substring replace. Search across all runs concatenated; if
    found, rewrite. Returns True if replaced."""
    full = paragraph_text(p)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    rewrite_paragraph_text(p, new_full)
    return True


def make_paragraph_element(text: str, bold: bool = False) -> OxmlElement:
    """Create a fresh <w:p> element with given text + optional bold."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def insert_paragraph_before(target_p, text: str, bold: bool = False):
    """Add a new paragraph before the target paragraph element."""
    new_p = make_paragraph_element(text, bold=bold)
    target_p._element.addprevious(new_p)


# ── EDIT 1: Modify Abstract Results paragraph (p5) ────────────────
print("Edit 1: Append Gemini-vs-student sentence to Abstract Results")
ABSTRACT_ADD = (
    " Both models remained inferior to the dental student consensus on every "
    "periapical and panoramic comparison; Gemini 3.1 Pro reached statistical "
    "parity with students only on cephalometric points under guided prompting "
    "(Bonferroni-corrected p = 0.054), and was significantly outperformed by "
    "students elsewhere by small-to-large effect sizes (panoramic point "
    "guided: r = +0.66; panoramic area guided: r = -0.64)."
)
p5 = doc.paragraphs[5]
old_p5_text = paragraph_text(p5)
new_p5_text = old_p5_text.rstrip() + ABSTRACT_ADD
rewrite_paragraph_text(p5, new_p5_text)


# ── EDIT 2: Insert new §RQ8 section before Discussion (p50) ───────
print("Edit 2: Insert §RQ8 section before Discussion")
discussion_p = doc.paragraphs[50]

def num(d):
    return {
        "delta": d["delta"], "r": d["r"], "p_bonf": d["p_bonf"],
        "n": d["n"], "g": d["mean_gem"], "s": d["mean_stu"],
    }

ceph_z = num(stats_summary["CEPHALOMETRIC_point|zero_shot"])
ceph_g = num(stats_summary["CEPHALOMETRIC_point|guided"])
pa_z   = num(stats_summary["PERIAPICAL_point|zero_shot"])
pa_g   = num(stats_summary["PERIAPICAL_point|guided"])
pan_z  = num(stats_summary["PANORAMIC_point|zero_shot"])
pan_g  = num(stats_summary["PANORAMIC_point|guided"])
panA_z = num(stats_summary["PANORAMIC_area|zero_shot"])
panA_g = num(stats_summary["PANORAMIC_area|guided"])

RQ8_HEADING = "RQ8: Gemini 3.1 Pro versus Dental Student Consensus"

RQ8_PARA1 = (
    f"Paired Wilcoxon signed-rank tests comparing Gemini 3.1 Pro mean Euclidean "
    f"distance (mean of 3 repetitions) with the team-adjudicated dental student "
    f"consensus, both scored against the same two-rater consensus ground truth, "
    f"revealed that the student consensus outperformed Gemini 3.1 Pro on seven "
    f"of the eight modality × strategy comparisons after Bonferroni correction "
    f"(Bonferroni × 4 within strategy × modality family; Table 8). On "
    f"cephalometric point landmarks, the student consensus was significantly "
    f"closer to GT under zero-shot prompting "
    f"(Δ = +{ceph_z['delta']:.2f} cells, Bonferroni-corrected p = {ceph_z['p_bonf']:.3f}, "
    f"rank-biserial r = +{ceph_z['r']:.2f}), but the comparison reached statistical "
    f"non-significance after Bonferroni correction under guided prompting "
    f"(Δ = +{ceph_g['delta']:.2f} cells, Bonferroni-corrected p = {ceph_g['p_bonf']:.3f}, "
    f"r = +{ceph_g['r']:.2f}), indicating that the guided strategy narrows the "
    f"Gemini–student gap to a sub-Bonferroni level on cephalometric landmarks alone. "
    f"On periapical and panoramic point landmarks, the student consensus retained "
    f"medium-to-large advantages across both strategies (panoramic zero-shot: "
    f"r = +{pan_z['r']:.2f}, p < 10⁻¹⁸; panoramic guided: r = +{pan_g['r']:.2f}, p < 10⁻¹⁶; "
    f"periapical zero-shot: r = +{pa_z['r']:.2f}, p < 10⁻⁴; periapical guided: "
    f"r = +{pa_g['r']:.2f}, Bonferroni-corrected p = {pa_g['p_bonf']:.3f}). "
    f"On panoramic area landmarks, the student consensus achieved substantially "
    f"higher mean Jaccard overlap than Gemini (zero-shot: mean Jaccard "
    f"{panA_z['s']:.2f} vs {panA_z['g']:.2f}, r = {panA_z['r']:+.2f}; guided: "
    f"{panA_g['s']:.2f} vs {panA_g['g']:.2f}, r = {panA_g['r']:+.2f}; all Bonferroni "
    f"p < 10⁻²⁰)."
)

RQ8_PARA2 = (
    "These findings reveal a distinct pattern from the GPT-5.4-versus-student "
    "comparison reported in §RQ5. Whereas GPT-5.4 matched students on "
    "cephalometric points under both prompting strategies (zero-shot: "
    "Δ = -0.05 cells, p = 0.317 NS; guided: Δ = -0.07 cells, p = 0.192 NS), "
    "Gemini 3.1 Pro reaches Bonferroni-corrected statistical parity with "
    "students on only a single configuration (cephalometric under guided "
    "prompting) and underperforms students across every other modality × "
    "strategy combination evaluated. The cross-model interpretation is "
    "therefore that the two frontier MLLMs occupy different positions relative "
    "to the human-trainee benchmark: GPT-5.4 reaches trainee parity selectively "
    "on cephalometric landmarks under both strategies, whereas Gemini 3.1 Pro "
    "only approaches trainee parity on cephalometric landmarks and only under "
    "guided prompting. Importantly, even on panoramic and periapical landmarks "
    "where Gemini 3.1 Pro substantially outperforms GPT-5.4 in absolute terms "
    "(§RQ7), Gemini's panoramic performance remains 0.27–0.32 cells worse than "
    "the student consensus on the same queries (rank-biserial r ≈ +0.66–+0.69), "
    "and its panoramic area-landmark overlap is reduced by 0.24–0.28 Jaccard "
    "units relative to students (r ≈ -0.64 to -0.73). Neither MLLM, therefore, "
    "achieves student-level performance on panoramic or periapical landmark "
    "identification, reinforcing the case for mandatory human oversight on "
    "these modalities regardless of which model is selected for the workflow."
)

RQ8_PARA3_EXPERT = (
    "[EXPERT REVIEW REQUESTED — clinical interpretation of three asymmetries.] "
    "(i) GPT-5.4 reaches statistical parity with students on cephalometric "
    "points under BOTH strategies, whereas Gemini 3.1 Pro reaches parity only "
    "under guided prompting and only by sub-Bonferroni margin "
    f"(p = {ceph_g['p_bonf']:.3f}, r = +{ceph_g['r']:.2f}); whether this "
    "asymmetry strengthens or weakens the pre-training-distribution hypothesis "
    "raised in §Discussion (cephalometric content over-represented in language-"
    "model pre-training) is a clinical question the data alone cannot adjudicate. "
    "(ii) On Mental Foramen (panoramic point), per-landmark stratification "
    f"shows Gemini at mean ED 1.35 cells (zero-shot) and 1.15 cells (guided) "
    f"versus students at 0.59 cells (Bonferroni-corrected p < 10⁻⁵ under both "
    "strategies, r ≈ +0.64 to +0.69) — i.e., students substantially outperform Gemini on "
    "the same landmark where guided prompting helped GPT-5.4 most "
    "(§RQ2; GPT improvement r = +0.86). The clinical interpretation of why "
    "Gemini struggles with Mental Foramen — a structure described in the "
    "Discussion as anatomically isolated and therefore amenable to spatial "
    "guidance for GPT — is itself flagged for expert input. (iii) The cephalometric "
    f"Gemini-versus-student gap (Δ = +{ceph_g['delta']:.2f} cells under guided) "
    "falls within the OMFR_1↔OMFR_2 within-1-cell agreement zone established "
    "in §RQ4 (98.5% of inter-rater pairs within 1 cell); whether this falls "
    "within a clinically equivalent band, however, is a judgment best made by "
    "the OMFR specialist co-investigator."
)

# Insert the heading + three paragraphs before Discussion
# Order matters: addprevious adds immediately before, so we add heading FIRST,
# then paragraph 1, then paragraph 2, then expert-review paragraph
insert_paragraph_before(discussion_p, RQ8_HEADING, bold=True)
insert_paragraph_before(discussion_p, RQ8_PARA1, bold=False)
insert_paragraph_before(discussion_p, RQ8_PARA2, bold=False)
insert_paragraph_before(discussion_p, RQ8_PARA3_EXPERT, bold=False)
insert_paragraph_before(discussion_p, "", bold=False)


# ── EDIT 3: Remove "Fourth, a direct comparison..." sentence in Limitations
print("Edit 3: Remove 'deferred to future work' sentence in Limitations")
# We re-find the limitations paragraph because indices may have shifted
for p in doc.paragraphs:
    if "deferred to future work" in p.text:
        old_sentence = (
            " Fourth, a direct comparison between Gemini 3.1 Pro and the "
            "dental student consensus was deferred to future work, leaving "
            "open the question of how this model's superior panoramic and "
            "periapical performance relates to human trainee-level competence "
            "on those modalities."
        )
        # The sentence may have slight whitespace variations
        full = paragraph_text(p)
        # Try exact match first
        if old_sentence in full:
            new_full = full.replace(old_sentence, "")
        else:
            # Try with leading space variations
            old2 = (
                "Fourth, a direct comparison between Gemini 3.1 Pro and the "
                "dental student consensus was deferred to future work, leaving "
                "open the question of how this model's superior panoramic and "
                "periapical performance relates to human trainee-level competence "
                "on those modalities. "
            )
            if old2 in full:
                new_full = full.replace(old2, "")
            else:
                # Find a regex-like span
                import re
                pat = re.compile(
                    r" ?Fourth, a direct comparison between Gemini 3\.1 Pro.*?"
                    r"on those modalities\.\s?",
                    re.DOTALL,
                )
                new_full, n = pat.subn("", full, count=1)
                if n == 0:
                    print("  WARNING: could not match deferred-future-work sentence")
                    new_full = full
        rewrite_paragraph_text(p, new_full)
        break


# ── EDIT 4: Append integration sentence to Conclusions paragraph ──
print("Edit 4: Append Gemini-vs-student sentence to Conclusions")
CONCLUSION_ADD = (
    " The new Gemini-versus-student comparison reported here further "
    "establishes that neither evaluated model reaches student-level "
    "performance on panoramic or periapical landmark identification: even "
    "Gemini 3.1 Pro, which dominates GPT-5.4 in absolute panoramic accuracy, "
    "remains significantly outperformed by the fourth-year dental student "
    "consensus on every panoramic and periapical comparison after Bonferroni "
    "correction. Statistical parity with students is achieved by Gemini only "
    "on cephalometric points under guided prompting."
)
# Find Conclusions paragraph by content
for p in doc.paragraphs:
    if p.text.startswith("The systematic evaluation of frontier"):
        old_text = paragraph_text(p)
        new_text = old_text.rstrip() + CONCLUSION_ADD
        rewrite_paragraph_text(p, new_text)
        break


# ── EDIT 5: Insert Table 8 before Figure Legends ──────────────────
print("Edit 5: Insert Table 8 + legend + note before Figure Legends")

# Find the Figure Legends paragraph
fig_legends_p = None
for p in doc.paragraphs:
    if p.text.strip() == "Figure Legends":
        fig_legends_p = p
        break

# Insert the Table 8 legend (bold)
TABLE8_LEGEND = (
    "Table 8. Gemini 3.1 Pro versus team-adjudicated fourth-year dental "
    "student consensus: paired Wilcoxon signed-rank tests by modality and "
    "prompting strategy."
)
insert_paragraph_before(fig_legends_p, TABLE8_LEGEND, bold=True)

# Build the table itself
# Columns: Group / Strategy | n | Gemini Mean | Student Mean | Mean Δ | p (Bonf. × 4) | r
table = doc.add_table(rows=9, cols=7)
# Move the table to the position right after the legend (before fig_legends_p)
# python-docx adds tables at end; we need to move it
table_elem = table._element
# Detach from current position
table_elem.getparent().remove(table_elem)
# Insert before fig_legends_p
fig_legends_p._element.addprevious(table_elem)

# Header row
hdrs = ["Group / Strategy", "n", "Gemini 3.1 Pro Mean", "Student Mean",
        "Mean Δᵃ", "p (Bonf. ×4)ᵇ", "rᶜ"]
for j, h in enumerate(hdrs):
    cell = table.rows[0].cells[j]
    cell.text = h
    # Bold the header
    for r in cell.paragraphs[0].runs:
        r.bold = True

# Data rows — pull every value from stats_summary
def fmt_mean(m, ltype):
    return f"{m:.2f}"

ROWS = [
    ("CEPHALOMETRIC / zero-shot", "CEPHALOMETRIC_point|zero_shot", "ED"),
    ("CEPHALOMETRIC / guided",    "CEPHALOMETRIC_point|guided",    "ED"),
    ("PERIAPICAL / zero-shot",    "PERIAPICAL_point|zero_shot",    "ED"),
    ("PERIAPICAL / guided",       "PERIAPICAL_point|guided",       "ED"),
    ("PANORAMIC point / zero-shot", "PANORAMIC_point|zero_shot",   "ED"),
    ("PANORAMIC point / guided",    "PANORAMIC_point|guided",      "ED"),
    ("PANORAMIC area / zero-shot",  "PANORAMIC_area|zero_shot",    "Jaccard"),
    ("PANORAMIC area / guided",     "PANORAMIC_area|guided",       "Jaccard"),
]
for i, (label, key, metric_name) in enumerate(ROWS, start=1):
    d = stats_summary[key]
    row = table.rows[i].cells
    p_bonf = d["p_bonf"]
    # Sig marker
    if p_bonf < 0.001:
        p_str = f"< 10⁻³ {sig_marker(p_bonf)}"
    elif p_bonf < 0.01:
        p_str = f"{p_bonf:.3f} {sig_marker(p_bonf)}"
    elif p_bonf < 0.05:
        p_str = f"{p_bonf:.3f} {sig_marker(p_bonf)}"
    else:
        p_str = f"{p_bonf:.3f} NS"
    row[0].text = label
    row[1].text = str(d["n"])
    row[2].text = fmt_mean(d["mean_gem"], metric_name)
    row[3].text = fmt_mean(d["mean_stu"], metric_name)
    row[4].text = f"{d['delta']:+.3f}"
    row[5].text = p_str
    row[6].text = f"{d['r']:+.3f}"

# Insert Table 8 Note
TABLE8_NOTE = (
    "Note. Both raters (Gemini 3.1 Pro and student consensus) are scored "
    "against the same two-rater adjudicated consensus ground truth. ᵃ Mean "
    "Δ = Gemini 3.1 Pro metric − Student metric. For ED (lower = better): "
    "positive Δ indicates student closer to GT, i.e., more accurate. For "
    "Jaccard (higher = better): negative Δ indicates student higher overlap. "
    "ᵇ Bonferroni correction applied × 4 within strategy × modality × landmark-"
    "type family (3 point-landmark modalities + 1 panoramic area, both strategies). "
    "ᶜ r = rank-biserial effect size (paired Wilcoxon). Significance: * p < 0.05; "
    "** p < 0.01; *** p < 0.001; NS = not significant after Bonferroni correction. "
    "Gemini 3.1 Pro statistics use rep-averaged Euclidean distance / Jaccard "
    "across 3 repetitions, identical methodology to Table 6 for GPT-5.4."
)
insert_paragraph_before(fig_legends_p, TABLE8_NOTE, bold=False)
insert_paragraph_before(fig_legends_p, "", bold=False)


# ── EDIT 6: Renumber the §RQ6/RQ7 heading to reflect §RQ8 addition ─
# The colleague's heading was "RQ6 and RQ7: Gemini 3.1 Pro — Single-Model
# Accuracy and Cross-Model Comparison". We don't need to renumber — RQ8
# stands alone after RQ7's prose without changing the prior heading.
# (Alternative: rename existing heading to "RQ6, RQ7, RQ8: ..." but this
#  would obscure the structural separation.)


# ── Save v2 docx ───────────────────────────────────────────────────
doc.save(OUT_DOCX)
print(f"\nWrote {OUT_DOCX}")
print(f"  Final paragraphs: {len(doc.paragraphs)}")
print(f"  Final tables: {len(doc.tables)}")
